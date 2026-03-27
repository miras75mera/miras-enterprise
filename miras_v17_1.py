"""
MİRAS v17.0 Enterprise — Mera İhtiyaç ve Rasyonel Amenajman Sistemi
T.C. Ardahan Valiliği İl Tarım ve Orman Müdürlüğü
Geliştirici: Emre ÖZTÜRK — Ziraat Mühendisi
Kurulum: pip install google-generativeai pandas openpyxl reportlab bcrypt python-docx
Giriş  : admin / Admin123!
"""
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import tkinter.simpledialog
import sqlite3, pandas as pd, threading, shutil, os, json
import urllib.request, hashlib, logging, time, sys, traceback
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional

try:
    import bcrypt; BCRYPT_OK = True
except ImportError:
    BCRYPT_OK = False
try:
    import google.generativeai as genai; GEMINI_OK = True
except ImportError:
    GEMINI_OK = False
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors as rl_colors
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_OK = True
except ImportError:
    PDF_OK = False
try:
    import matplotlib
    MPL_OK = True
except ImportError:
    MPL_OK = False
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
try:
    from ctypes import windll; windll.shcore.SetProcessDpiAwareness(1)
except Exception:
    pass

logging.basicConfig(filename="miras_debug.log", level=logging.ERROR,
    format="%(asctime)s [%(levelname)s] %(message)s")

def _global_exc(t, v, tb):
    logging.error("".join(traceback.format_exception(t, v, tb)))
    try: messagebox.showerror("Sistem Hatası", f"Hata: {str(v)[:300]}\n\nDetay: miras_debug.log")
    except Exception: pass
sys.excepthook = _global_exc

# ─── SABİTLER ────────────────────────────────────────────────────────────────
PROG_ADI, VERSIYON = "MİRAS Enterprise", "v17.1"
MEVCUT_SURUM = "v17.1"  # <--- İŞTE BURAYA EKLİYORSUN
DEV_ADI = "Emre ÖZTÜRK"
DEV_UNVAN = "Ziraat Mühendisi"
DEV_TEL = "0 545 689 00 75"
DEV_MAIL = "miras75mera@gmail.com"
HAKLAR = f"© {datetime.now().year} {DEV_ADI} — Tüm Hakları Saklıdır."
CONFIG_FILE, API_KEY_FILE = "miras_config_v17.json", "miras_gemini_key.txt"
BACKUP_DIR = Path("miras_yedekler")
GUNLUK_YEM_KG, OTLATMA_GUN = 30, 135
PAYDA = GUNLUK_YEM_KG * OTLATMA_GUN
MAX_LOGIN_FAIL = 3
LOCKOUT_MIN = 15
SIFRE_OMUR_GUN = 180
SIFRE_UYARI_GUN = 7
PAGE_SIZE = 25
OTOSAVE_SEC = 600
DEV_WA = "905456890075"
GUNLUK_OT_BBHB = 12.5
GUNLUK_OT_KBHB = 1.25
GITHUB_REPO = "miras75mera/miras-enterprise"
MERA_VASIF = ["Zayıf","Orta","İyi","Çok İyi"]
OT_VERIM_KURU = {"Zayıf":45,"Orta":90,"İyi":135,"Çok İyi":180}

# Renkler
C_WHITE = "#FFFFFF"
C_DANGER = "#C0392B"
C_INFO = "#1A6FA8"
C_WARN = "#D68910"
C_DARK = "#2C3E50"
C_FOOTER = "#555555"

TEMALAR = {
    "Orman Yeşili": {"pri":"#1E5631","acc":"#2D8C55","bg":"#F0F4F2","side":"#163D22","side_text":"#D4E8DB","side_hover":"#1E5631","card":"#FFFFFF"},
    "Deniz Mavisi": {"pri":"#1A5276","acc":"#2980B9","bg":"#EBF5FB","side":"#143D5C","side_text":"#C8DFF0","side_hover":"#1A5276","card":"#FFFFFF"},
    "Koyu Bordo":   {"pri":"#641E16","acc":"#922B21","bg":"#FDEDEC","side":"#4A1610","side_text":"#E8C4BF","side_hover":"#641E16","card":"#FFFFFF"},
    "Koyu Gece":    {"pri":"#1A1A2E","acc":"#16213E","bg":"#0F0F1A","side":"#0D0D1A","side_text":"#8888AA","side_hover":"#1A1A2E","card":"#1F1F2E"},
}

HAYVAN_TURLERI = [
    ("Kültür Irkı Süt İneği",1.00),("Kültür Melezi Süt İneği",0.75),
    ("Yerli İnek",0.50),("Kültür Irkı Dana-Düve",0.60),
    ("Kültür Melezi Dana-Düve",0.45),("Yerli Dana-Düve",0.30),
    ("Koyun",0.10),("Keçi",0.08),("Manda (Erkek)",0.90),
    ("Manda (Dişi)",0.75),("Öküz",0.60),("Kuzu-Oğlak",0.04),
    ("Boğa",1.50),
]
TAHSIS_ASAMALARI = [
    "1-Dilekçe Alındı","2-Teknik İnceleme",
    "3-Komisyon Değerlendirmesi","4-Karar Yazıldı",
    "5-Tapu Müdürlüğüne Gönderildi","6-Tamamlandı",
]
IHALE_DURUMLARI = [
    "Başvuru Alındı","Evrak Kontrolü","Komisyon Değerlendirmesi",
    "İhale Yapıldı","Sözleşme Aşaması","Aktif","İptal Edildi","Tamamlandı",
]
ILCELER = ["Merkez","Göle","Çıldır","Posof","Hanak","Damal"]
CEZA_TURLERI = [
    "Kaçak Hayvan Otlatılması",
    "Otlatma Kapasitesi Üstünde Hayvan Sokulması",
    "Mera Alanında İzinsiz Yapılaşma",
    "Mera Alanını İzinsiz Sürme/Tahrip",
    "İzinsiz Ağaç Kesimi",
    "Diğer",
]
DB_PATH: Optional[str] = None

# ─── AÇIK RIZA METNİ ─────────────────────────────────────────────────────────
RIZA_METNI = """
MİRAS Enterprise — Kullanım Koşulları ve Açık Rıza Metni

1. Bu yazılım, T.C. Ardahan Valiliği İl Tarım ve Orman Müdürlüğü Çayır, Mera ve 
   Yem Bitkileri Şube Müdürlüğü bünyesinde kullanılmak üzere geliştirilmiştir.

2. Yazılımda yer alan yapay zekâ destekli özellikler (MERA AI) tavsiye niteliğinde 
   olup, üretilen evrak ve hesaplamaların doğruluğu kullanıcının sorumluluğundadır.

3. Otlatma kapasitesi hesaplamaları, ihale bedelleri ve idari para cezası tutarları 
   dahil tüm sayısal veriler kullanıcı tarafından kontrol edilmelidir.

4. Kullanıcı, sisteme girdiği verilerin doğruluğundan ve güncelliğinden sorumludur.

5. Sistem yedekleme hizmeti sunar ancak veri kaybına karşı ek önlemler kullanıcının 
   sorumluluğundadır.

6. Bu yazılımın kullanımından doğabilecek hata, zarar veya kayıplardan geliştirici 
   sorumlu tutulamaz.

7. Kullanıcı bu koşulları kabul ederek sisteme giriş yapar.

Geliştirici: Emre ÖZTÜRK — Ziraat Mühendisi
İletişim: miras75mera@gmail.com
"""

# ─── MEVZUAT REHBERİ ─────────────────────────────────────────────────────────
MEVZUAT_KARTLARI = {
    "Mera": {
        "baslik": "Mera Nedir?",
        "tanim": "Hayvanların otlatılması ve otundan yararlanılması için tahsis edilen veya kadimden beri bu amaçla kullanılan yerlerdir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/a",
        "detay": "Meralar Devletin hüküm ve tasarrufundadır. Özel mülkiyete geçirilemez, zaman aşımı uygulanamaz, sınırları daraltılamaz ve amacı dışında kullanılamaz."
    },
    "Yaylak": {
        "baslik": "Yaylak Nedir?",
        "tanim": "Çiftçilerin hayvanları ile birlikte yaz mevsiminde çıkarak otlatma ve hayvancılık yaptıkları yerlerdir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/b",
        "detay": "Yaylaklar da meralar gibi kamu malı niteliğinde olup tahsis amacı dışında kullanılamaz."
    },
    "Kışlak": {
        "baslik": "Kışlak Nedir?",
        "tanim": "Hayvanların kış mevsiminde barındırılması ve otundan yararlanılması için tahsis edilen veya kadimden beri bu amaçla kullanılan yerlerdir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/f",
        "detay": "Kışlaklar genellikle sıcak bölgelerde bulunur ve kış mevsiminde hayvanların korunması amacıyla kullanılır."
    },
    "Tespit": {
        "baslik": "Tespit Nedir?",
        "tanim": "Bir yerin mera, yaylak ve kışlak arazisi olup olmadığının resmi evrakla ve bilirkişi ifadeleri ile belgelendirilmesidir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/j",
        "detay": "Tespit işlemi, İl Mera Komisyonu tarafından gerçekleştirilir. Kadastro kayıtları, tapu sicili, mahkeme kararları ve bilirkişi beyanları değerlendirilir."
    },
    "Tahdit": {
        "baslik": "Tahdit Nedir?",
        "tanim": "Mera, yaylak ve kışlak arazisi olduğuna karar verilen yerlerin sınırlarının 1/5000 ölçekli haritalar üzerinde belirlenmesi ve arazi üzerinde kalıcı işaretlerle işaretlenmesidir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/i",
        "detay": "Tahdit işlemi teknik ekipler tarafından yapılır. Sonuçlar 30 gün askıda kalır, itirazlar 60 gün içinde karara bağlanır."
    },
    "Tahsis": {
        "baslik": "Tahsis Nedir?",
        "tanim": "Mera, yaylak ve kışlakların verimlilik ve sosyal adalet ilkelerine uygun şekilde düzenlenerek köy veya belediyeye bırakılmasıdır.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/ı",
        "detay": "Tahsis kararında otlatma kapasitesi, aile başına hayvan sayısı, sulama ve geçit yerleri belirtilir. Tahsis kararı Valilik onayına sunulur."
    },
    "BBHB": {
        "baslik": "Büyükbaş Hayvan Birimi (BBHB) Nedir?",
        "tanim": "Hayvan sayısının, 500 kg canlı ağırlığa çevrilerek ifade edilen standardize şeklidir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/n",
        "detay": "Hesaplama: Otlatma Kapasitesi = (Mera Alanı × Yem Verimi) / (Günlük Yem İhtiyacı × Otlatma Gün Sayısı). Ardahan için: 30 kg/gün × 135 gün = 4050 kg."
    },
    "Madde 14": {
        "baslik": "Madde 14 — Tahsis Amacı Değişikliği",
        "tanim": "Mera, yaylak ve kışlakların tahsis amacının enerji, turizm, kamu yatırımı gibi nedenlerle değiştirilmesidir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 14",
        "detay": "Tahsis amacı değişikliği için 20 yıllık ot geliri hesaplanır. İlgili Bakanlık talebi, Maliye ve Valilik uygun görüşü gerekir. Tescilden itibaren 2 yıl içinde yatırıma başlanmazsa iptal edilir."
    },
    "Otlatma Hakkı": {
        "baslik": "Otlatma Hakkı",
        "tanim": "Çiftçi ailelerinin meradan yararlanma hakkıdır. Tahsis yapılan köyde en az 6 aydır ikamet şartı aranır.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 22",
        "detay": "Otlatma hakkından fazla hayvan otlatılamaz. Otlatma hakkı başka çiftçi ailelerine devredilemez. İhtiyaç fazlası mera alanları kiralanabilir."
    },
    "İdari Para Cezası": {
        "baslik": "İdari Para Cezası",
        "tanim": "Mera Kanunu'na aykırı davranan kişilere uygulanan idari yaptırımdır.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 27",
        "detay": "Mera, yaylak, kışlakları izinsiz kullanan, tahrip eden, kapasiteden fazla hayvan otlatanlara idari para cezası uygulanır. Mükerrer ihlallerde ceza artırılır."
    },
    "Mera Komisyonu": {
        "baslik": "Mera Komisyonu",
        "tanim": "Vali yardımcısı başkanlığında 8 kişiden oluşan, mera tespit, tahdit ve tahsis işlemlerini yürüten komisyondur.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 6",
        "detay": "Üyeler: Vali Yardımcısı (Başkan), İl Tarım Müdürü, Ziraat Mühendisi, DSİ temsilcisi, Orman temsilcisi, Muhtar, Milli Emlak temsilcisi, Kadastro temsilcisi."
    },
    "Kiralama": {
        "baslik": "Mera Kiralama İhaleleri",
        "tanim": "İhtiyaç fazlası mera alanlarının 2886 sayılı Devlet İhale Kanunu'na göre kiralanmasıdır.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 12; Mera Yönetmeliği Madde 7, 13",
        "detay": "Kiralama bedeli: %25 Bakanlık hesabına, %75 köy hesabına yatırılır. Geçici teminat %30, kesin teminat %6 oranındadır. Damga vergisi ‰5.69 uygulanır."
    },
}

# ─── TC KİMLİK DOĞRULAMA ─────────────────────────────────────────────────────
def tc_dogrula(tc):
    """TC Kimlik numarası algoritma kontrolü"""
    if not tc or len(str(tc)) != 11:
        return False
    try:
        d = [int(c) for c in str(tc)]
    except ValueError:
        return False
    if d[0] == 0:
        return False
    # 10. hane kontrolü
    t1 = ((d[0]+d[2]+d[4]+d[6]+d[8])*7 - (d[1]+d[3]+d[5]+d[7])) % 10
    if t1 != d[9]:
        return False
    # 11. hane kontrolü
    t2 = sum(d[:10]) % 10
    if t2 != d[10]:
        return False
    return True

def telefon_format(tel):
    """Telefon numarası format kontrolü"""
    t = tel.replace(" ","").replace("-","").replace("(","").replace(")","")
    if t.startswith("+90"): t = t[3:]
    elif t.startswith("0"): t = t[1:]
    return len(t) == 10 and t.isdigit()

def tc_kontrol_ve_devam(tc, parent=None):
    """TC doğrula, geçersizse kullanıcıya sor — tüm modüllerde kullanılır"""
    if not tc: return True
    if tc_dogrula(tc): return True
    return messagebox.askyesno("TC Uyarı",f"'{tc}' geçerli bir TC numarası değil.\nDevam etmek istiyor musunuz?",parent=parent)

def muhammen_bedel_hesapla(alan_da, kuru_ot_verimi, otlatma_gun, kuru_ot_fiyati):
    """Muhammen bedel — Tahdit Raporu formülü"""
    try:
        alan=float(alan_da); kov=float(kuru_ot_verimi); ogs=float(otlatma_gun); kof=float(kuru_ot_fiyati)
        kapasite_bbhb = (alan * kov) / (ogs * GUNLUK_OT_BBHB)
        kapasite_kbhb = kapasite_bbhb * 10
        bedel = kapasite_bbhb * ogs * GUNLUK_OT_BBHB * kof
        return {"alan":alan,"kuru_ot":kov,"gun":ogs,"fiyat":kof,"bbhb":kapasite_bbhb,"kbhb":kapasite_kbhb,"bedel":bedel}
    except: return None

# ─── GÜVENLİK ────────────────────────────────────────────────────────────────
def hash_pw(p):
    if BCRYPT_OK: return bcrypt.hashpw(p.encode(), bcrypt.gensalt()).decode()
    return hashlib.sha256(p.encode()).hexdigest()

def verify_pw(p, h):
    if BCRYPT_OK and h.startswith("$2b$"): return bcrypt.checkpw(p.encode(), h.encode())
    return hashlib.sha256(p.encode()).hexdigest() == h

def strong_pw(p):
    if len(p)<6: return False,"En az 6 karakter olmalı."
    if not any(c.isupper() for c in p): return False,"Büyük harf gerekli."
    if not any(c.isdigit() for c in p): return False,"Rakam gerekli."
    return True,""

def db_log(kul, islem, detay=""):
    if not DB_PATH: return
    try:
        with sqlite3.connect(DB_PATH) as c:
            c.execute("INSERT INTO Loglar(tarih,kul,islem,detay)VALUES(?,?,?,?)",
                (datetime.now().strftime("%Y-%m-%d %H:%M:%S"),kul,islem,detay))
    except Exception as e: logging.error(f"db_log:{e}")

# ─── SESSION ─────────────────────────────────────────────────────────────────
class SessionManager:
    TIMEOUT=900
    def __init__(self):
        self._last=time.time(); self._lock=threading.Lock()
        self._cbs=[]; self._on=True
        threading.Thread(target=self._watch,daemon=True).start()
    def ping(self):
        with self._lock: self._last=time.time()
    def register(self,fn): self._cbs.append(fn)
    def stop(self): self._on=False
    def _watch(self):
        while self._on:
            time.sleep(30)
            with self._lock: exp=time.time()-self._last>self.TIMEOUT
            if exp:
                self._on=False
                for fn in self._cbs: fn()

# ─── YEDEKLEME ───────────────────────────────────────────────────────────────
class YedekYoneticisi:
    MAX=30
    def __init__(self,uid):
        self.uid=uid; BACKUP_DIR.mkdir(exist_ok=True)
        self.drive_path=self._load_drive_path()
        threading.Thread(target=self._oto,daemon=True).start()
        threading.Thread(target=self._otosave,daemon=True).start()
    def _load_drive_path(self):
        """Config'den Drive klasör yolunu oku"""
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE) as f: data=json.load(f)
                dp=data.get("drive_path","")
                if dp and os.path.isdir(dp): return dp
            except Exception: pass
        return ""
    def _save_drive_path(self,yol):
        """Config'e Drive klasör yolunu kaydet"""
        self.drive_path=yol
        data={}
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE) as f: data=json.load(f)
            except Exception: pass
        data["drive_path"]=yol
        with open(CONFIG_FILE,"w") as f: json.dump(data,f)
    def set_drive_path(self,yol):
        """Drive klasörünü ayarla ve kaydet"""
        if yol and os.path.isdir(yol):
            self._save_drive_path(yol); return True
        return False
    def _oto(self):
        while True:
            time.sleep(86400)
            try: self.al(otomatik=True)
            except Exception as e: logging.error(f"oto yedek:{e}")
    def _otosave(self):
        """10 dakikada bir otomatik yedek"""
        while True:
            time.sleep(OTOSAVE_SEC)
            try: self.al(otomatik=True)
            except Exception: pass
    def al(self,otomatik=False):
        if not DB_PATH: raise ValueError("DB seçilmemiş")
        ts=datetime.now().strftime("%Y%m%d_%H%M%S")
        e="oto" if otomatik else "manuel"
        h=BACKUP_DIR/f"Miras_{e}_{ts}.db"
        shutil.copy2(DB_PATH,h); self._temizle()
        # Otomatik Drive kopyalama
        drive_ok=False
        if self.drive_path and os.path.isdir(self.drive_path):
            try:
                drive_dosya=os.path.join(self.drive_path,f"Miras_{e}_{ts}.db")
                shutil.copy2(DB_PATH,drive_dosya)
                drive_ok=True
                # Drive'da eski yedekleri temizle (son 10 tane kalsın)
                self._drive_temizle()
            except Exception as ex:
                logging.error(f"Drive kopyalama hatası: {ex}")
        detay=str(h)
        if drive_ok: detay+=f" + Drive"
        db_log(self.uid,"Yedekleme",detay); return str(h)
    def _temizle(self):
        ys=sorted(BACKUP_DIR.glob("*.db"),key=lambda p:p.stat().st_mtime)
        for y in ys[:-self.MAX]: y.unlink(missing_ok=True)
    def _drive_temizle(self):
        """Drive klasöründeki eski yedekleri temizle (son 10 kalsın)"""
        if not self.drive_path or not os.path.isdir(self.drive_path): return
        try:
            drive_files=sorted(
                [f for f in Path(self.drive_path).glob("Miras_*.db")],
                key=lambda p:p.stat().st_mtime)
            for f in drive_files[:-10]: f.unlink(missing_ok=True)
        except Exception as e: logging.error(f"drive_temizle:{e}")
    def listele(self):
        ys=sorted(BACKUP_DIR.glob("*.db"),key=lambda p:p.stat().st_mtime,reverse=True)
        return [{"ad":p.name,"yol":str(p),"boyut":f"{p.stat().st_size/1024:.1f} KB",
                 "tarih":datetime.fromtimestamp(p.stat().st_mtime).strftime("%d.%m.%Y %H:%M")}
                for p in ys]
    def geri(self,yol):
        try: self.al(True); shutil.copy2(yol,DB_PATH); db_log(self.uid,"Geri Yükleme",yol); return True
        except Exception as e: logging.error(f"geri:{e}"); return False

# ─── GEMİNİ AI ───────────────────────────────────────────────────────────────
SISTEM_PROMPT=(
    "Sen T.C. Ardahan İl Tarım ve Orman Müdürlüğü Çayır Mera Şubesi için geliştirilmiş "
    "MERA AI yapay zeka asistanısın. Mera Kanunu(4342), BBHB hesaplamaları, "
    "ihale/tahsis/islah süreçleri ve tarım mevzuatı konularında Türkçe, teknik ve "
    "kurumsal dilde yardımcı olursun. Kendini her zaman MERA AI olarak tanıt."
)

class GeminiAsistan:
    def __init__(self):
        self.model=None; self.chat=None; self.hazir=False
        self.api_key=self._key()
        if self.api_key and GEMINI_OK:
            threading.Thread(target=self._baglan,daemon=True).start()
    def _key(self):
        k=os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
        if k: return k
        if os.path.exists(API_KEY_FILE):
            with open(API_KEY_FILE) as f: k=f.read().strip()
            if k: return k
        return None
    def _baglan(self):
        try:
            genai.configure(api_key=self.api_key)
            tercih=["gemini-2.0-flash","gemini-2.0-flash-lite","gemini-1.5-flash","gemini-1.5-flash-latest","gemini-1.5-pro"]
            try:
                mevcut=[m.name.replace("models/","") for m in genai.list_models() if "generateContent" in m.supported_generation_methods]
            except Exception: mevcut=tercih
            secilen=next((m for m in tercih if m in mevcut),None)
            if not secilen: secilen=next((m for m in mevcut if "flash" in m or "pro" in m),None)
            if not secilen: raise ValueError("Kullanılabilir model yok.")
            self.model=genai.GenerativeModel(secilen,system_instruction=SISTEM_PROMPT)
            self.chat=self.model.start_chat(history=[])
            self.hazir=True; logging.info(f"Gemini:{secilen}")
        except Exception as e: logging.error(f"Gemini:{e}"); self.hazir=False
    def key_kaydet(self,k):
        self.api_key=k.strip()
        with open(API_KEY_FILE,"w") as f: f.write(self.api_key)
        self.hazir=False
        threading.Thread(target=self._baglan,daemon=True).start()
    def yeni_chat(self):
        if self.model: self.chat=self.model.start_chat(history=[])
    def sor(self,mesaj):
        if not self.hazir:
            return ("⚠️ MERA AI aktif değil.\n"
                    "API anahtarı için: aistudio.google.com/app/apikey\n"
                    "MERA AI sekmesi → API Ayarları bölümünden anahtarınızı girin.")
        return self._guvenli_istek(lambda: self.chat.send_message(mesaj).text)
    def tek(self,mesaj):
        if not self.hazir or not self.model: return "⚠️ Gemini bağlı değil."
        return self._guvenli_istek(lambda: self.model.generate_content(mesaj).text)
    def _guvenli_istek(self, fn, max_deneme=3):
        """429 hatasında otomatik bekle ve tekrar dene"""
        import re
        for deneme in range(max_deneme):
            try:
                return fn()
            except Exception as e:
                hata_str=str(e)
                if "429" in hata_str or "quota" in hata_str.lower() or "rate" in hata_str.lower():
                    # Bekleme süresini hatadan çıkar
                    bekleme=45  # varsayılan
                    m=re.search(r'retry in (\d+)',hata_str)
                    if m: bekleme=int(m.group(1))+2
                    m2=re.search(r'retry_delay.*?seconds:\s*(\d+)',hata_str,re.DOTALL)
                    if m2: bekleme=int(m2.group(1))+2
                    if deneme < max_deneme-1:
                        logging.warning(f"Gemini 429 - {bekleme}s bekleniyor (deneme {deneme+1}/{max_deneme})")
                        time.sleep(bekleme)
                        continue
                    else:
                        # Son deneme de başarısız
                        if "per day" in hata_str.lower() or "perday" in hata_str.lower():
                            return ("⚠️ Günlük AI kullanım kotası doldu.\n\n"
                                    "Google ücretsiz plan günlük istek sınırı aşıldı.\n"
                                    "Yarın sıfırlanacak veya farklı bir API key deneyin.\n\n"
                                    "💡 Kota bilgisi: ai.google.dev/gemini-api/docs/rate-limits")
                        else:
                            return (f"⚠️ AI şu an yoğun — {max_deneme} kez denendi.\n\n"
                                    f"Dakikalık istek sınırı aşıldı.\n"
                                    f"Lütfen {bekleme} saniye sonra tekrar deneyin.\n\n"
                                    "💡 Kota bilgisi: ai.google.dev/gemini-api/docs/rate-limits")
                else:
                    logging.error(f"Gemini hata: {e}")
                    return f"❌ AI Hatası: {hata_str[:300]}"
        return "❌ Beklenmeyen hata oluştu."

# ─── PDF MOTORU ───────────────────────────────────────────────────────────────
_FN="Helvetica"; _FNB="Helvetica-Bold"

def _init_fonts():
    global _FN,_FNB
    if not PDF_OK: return
    for r,b in [
        ("C:/Windows/Fonts/arial.ttf","C:/Windows/Fonts/arialbd.ttf"),
        ("C:/Windows/Fonts/calibri.ttf","C:/Windows/Fonts/calibrib.ttf"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ]:
        if os.path.exists(r):
            try:
                pdfmetrics.registerFont(TTFont("TF",r))
                pdfmetrics.registerFont(TTFont("TFB",b if os.path.exists(b) else r))
                _FN="TF"; _FNB="TFB"; return
            except Exception: continue

def uret_pdf(dosya,veri,personel):
    if not PDF_OK: raise ImportError("reportlab kurulu degil")
    _init_fonts()
    fn,fnb=_FN,_FNB
    W,H=A4
    cv=rl_canvas.Canvas(dosya,pagesize=A4)
    def yaz(x,y,t,f=None,s=10,rk=None):
        cv.setFont(f or fn,s)
        cv.setFillColor(rk or rl_colors.black)
        cv.drawString(x*cm,y*cm,str(t))
    def yazm(y,t,f=None,s=10):
        cv.setFont(f or fn,s); cv.setFillColor(rl_colors.black)
        cv.drawCentredString(W/2,y*cm,str(t))
    def cizgi(y,sl=2.0,sr=19.0):
        cv.setStrokeColor(rl_colors.HexColor("#1E5631")); cv.setLineWidth(0.8)
        cv.line(sl*cm,y*cm,sr*cm,y*cm)

    # Başlık bandı
    cv.setFillColor(rl_colors.HexColor("#1E5631"))
    cv.rect(0,H-3.2*cm,W,3.2*cm,fill=1,stroke=0)
    # Alt dekoratif çizgi
    cv.setFillColor(rl_colors.HexColor("#2D8C55"))
    cv.rect(0,H-3.4*cm,W,0.2*cm,fill=1,stroke=0)
    cv.setFillColor(rl_colors.white)
    LOGO="bakanlik_logo.png"
    if os.path.exists(LOGO):
        try: cv.drawImage(LOGO,0.4*cm,H-3.0*cm,width=2.4*cm,height=2.4*cm,preserveAspectRatio=True,mask="auto")
        except Exception: pass
    cv.setFont(fnb,13); cv.drawCentredString(W/2,H-1.2*cm,"T.C. ARDAHAN VALİLİĞİ")
    cv.setFont(fn,10); cv.drawCentredString(W/2,H-1.8*cm,"İl Tarım ve Orman Müdürlüğü")
    cv.setFont(fn,9); cv.drawCentredString(W/2,H-2.4*cm,"Çayır, Mera ve Yem Bitkileri Şube Müdürlüğü")

    y=26.0
    yazm(y,"OTLATMA KAPASİTESİ RAPORU",fnb,12); y-=0.6
    yazm(y,f"ARDAHAN İLİ {veri['ilce'].upper()} İLÇESİ {veri['koy'].upper()} KÖYÜ",fn,10)
    y-=0.9; cizgi(y); y-=0.6
    yaz(2,y,f"Sayı : {veri['rapor_no']}",fn,10)
    yaz(13,y,f"Tarih: {datetime.now().strftime('%d.%m.%Y')}",fn,10)
    y-=0.5; yaz(2,y,f"Konu : {veri['ilce']} İlçesi {veri['koy']} Köyü Mera Otlatma Kapasitesi",fn,10)
    y-=0.5; yaz(2,y,f"Talep Eden: {veri['talep_eden']}   TC: {veri['tc']}",fn,10)
    y-=0.3; cizgi(y)

    ok=veri["ok_bbhb"]
    y-=0.7; yaz(2,y,"HESAPLAMA BİLGİLERİ",fnb,10); y-=0.5
    yaz(2.5,y,f"Mera Alanı: {veri['alan']:.0f} da  |  Yararlanılabilir Yeşil Yem: {veri['yem']:.0f} kg/da  |  Otlatma: {OTLATMA_GUN} gün  |  Günlük İhtiyaç: {GUNLUK_YEM_KG} kg",fn,9)
    y-=0.5
    yaz(2,y,f"OTLATMA KAPASİTESİ: {ok:.2f} BBHB   |   İşletme Sayısı: {veri['aktif']}   |   İşletme Başı: {ok/veri['aktif']:.2f} BBHB" if veri['aktif']>0 else f"OTLATMA KAPASİTESİ: {ok:.2f} BBHB",fnb,11,rl_colors.HexColor("#1E5631"))
    y-=0.3; cizgi(y)

    y-=0.6; yazm(y,"BBHB'YE GÖRE OTLATILABİLECEK HAYVAN SAYILARI",fnb,10); y-=0.4
    CW=[6.0,2.2,4.5,5.3]
    HDRS=["IRK VE TÜR","KATSAYI","TOPLAM HAYVAN SAYISI","İŞLETME BAŞINA"]
    cx=[2.0]
    for w in CW[:-1]: cx.append(cx[-1]+w)
    cv.setFillColor(rl_colors.HexColor("#1E5631"))
    cv.rect(2*cm,(y-0.45)*cm,sum(CW)*cm,0.5*cm,fill=1,stroke=0)
    cv.setFillColor(rl_colors.white); cv.setFont(fnb,8)
    for i,(h,x) in enumerate(zip(HDRS,cx)):
        cv.drawCentredString((x+CW[i]/2)*cm,(y-0.22)*cm,h)
    y-=0.5
    rh=0.42
    for idx,(tur,kat) in enumerate(HAYVAN_TURLERI):
        if y<4.5: cv.showPage(); y=27.0
        toplam=ok/kat if kat>0 else 0
        isb=toplam/veri["aktif"] if veri["aktif"]>0 else 0
        bg=rl_colors.HexColor("#EAF4EE") if idx%2==0 else rl_colors.white
        cv.setFillColor(bg)
        cv.rect(2*cm,(y-rh+0.1)*cm,sum(CW)*cm,rh*cm,fill=1,stroke=0)
        cv.setFillColor(rl_colors.black); cv.setFont(fn,8)
        for i,(txt,x) in enumerate(zip([tur,f"{kat:.2f}",f"{round(toplam)}",f"{round(isb)}"],cx)):
            cv.drawCentredString((x+CW[i]/2)*cm,(y-0.15)*cm,txt)
        cv.setStrokeColor(rl_colors.lightgrey)
        cv.rect(2*cm,(y-rh+0.1)*cm,sum(CW)*cm,rh*cm,fill=0,stroke=1)
        y-=rh
    y-=0.3; cizgi(y); y-=0.5
    if veri.get("aciklama"):
        yaz(2,y,f"Not: {veri['aciklama']}",fn,8)
    y-=1.0
    # Dekoratif imza bölümü
    if y<5.5: cv.showPage(); y=27.0
    cv.setFillColor(rl_colors.HexColor("#F0F4F2"))
    cv.rect(1.5*cm,(y-2.0)*cm,7.5*cm,2.2*cm,fill=1,stroke=0)
    cv.rect(12*cm,(y-2.0)*cm,7.5*cm,2.2*cm,fill=1,stroke=0)
    cv.setStrokeColor(rl_colors.HexColor("#1E5631")); cv.setLineWidth(1.5)
    cv.line(1.5*cm,(y+0.2)*cm,9*cm,(y+0.2)*cm)
    cv.line(12*cm,(y+0.2)*cm,19.5*cm,(y+0.2)*cm)
    yaz(2,y,personel["ad"],fnb,10); yaz(13,y,personel["sube_mudur"],fnb,10)
    y-=0.5; yaz(2,y,personel["unvan"],fn,9); yaz(13,y,"Şube Müdürü V.",fn,9)
    y-=0.5; yaz(2,y,"Hazırlayan",fn,8,rl_colors.HexColor("#888888")); yaz(13,y,"Onaylayan",fn,8,rl_colors.HexColor("#888888"))
    # Alt bilgi
    cv.setFillColor(rl_colors.HexColor("#1E5631"))
    cv.rect(0,0,W,1.2*cm,fill=1,stroke=0)
    cv.setFillColor(rl_colors.white); cv.setFont(fn,7)
    cv.drawCentredString(W/2,0.5*cm,f"MİRAS Enterprise {VERSIYON} — {HAKLAR}")
    cv.save()

# ─── WORD ŞABLON DOLDURMA ─────────────────────────────────────────────────────
def _yil():
    """Mevcut yılı döndür"""
    return str(datetime.now().year)

def word_katilim_evrak(dosya, veri):
    """İhale Katılım Evrakları Word şablonu — orijinal evrakla birebir"""
    if not DOCX_OK: raise ImportError("python-docx kurulu değil")
    doc = DocxDocument()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MERA KİRALAMA İHALELERİ KATILIM EVRAKLARI")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x1E,0x56,0x31)
    doc.add_paragraph()
    tbl = doc.add_table(rows=5, cols=2, style='Table Grid')
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    bilgiler = [
        ("ADI SOYADI:", veri.get("ad_soyad","")),
        ("T.C. NO:", veri.get("tc","")),
        ("İHALEYE GİRİLECEK MERA:", veri.get("mera","")),
        ("ADA/PARSEL:", veri.get("ada_parsel","")),
        ("TARİH:", veri.get("tarih", datetime.now().strftime("%d/%m/%Y"))),
    ]
    for i,(lbl,val) in enumerate(bilgiler):
        tbl.cell(i,0).text = lbl
        tbl.cell(i,1).text = str(val)
        for cell in tbl.rows[i].cells:
            for p2 in cell.paragraphs:
                for r2 in p2.runs: r2.font.size = Pt(11)
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("ALINAN EVRAKLAR (ZORUNLU OLANLAR)")
    r.bold = True; r.font.size = Pt(12)
    evraklar = [
        "Aile Nüfus Kayıt Örneği",
        "İkametgâh Belgesi",
        "Nüfus Cüzdanı Fotokopisi",
        "İki Adet Vesikalık Fotoğraf",
        "Savcılık İyi Hal Kâğıdı (Sabıka Kaydı)",
        "Geçici Teminatın Yatırıldığını Gösterir Dekont veya Teminat Mektubu",
        "Onaylı Hayvan Listesi (Büyükbaş Hayvan sayısı)",
        "Hayvancılık İşletme Tescil Belgesi",
        "520,00 ₺ Dosya Bedeli Makbuzu",
        "SGK ve Maliyeden Vadesi Geçmiş Borcu Olmadığına Dair Belge",
    ]
    for ev in evraklar:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"☐  {ev}").font.size = Pt(10)
    doc.add_paragraph(); doc.add_paragraph()
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.cell(0,0).text = "TESLİM ALAN"
    tbl2.cell(0,1).text = "TESLİM EDEN"
    for cell in tbl2.rows[0].cells:
        for p2 in cell.paragraphs:
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r2 in p2.runs: r2.bold = True
    doc.save(dosya)

def word_kiralama_sozlesme(dosya, veri):
    """Kiralama Sözleşmesi — 14 maddelik orijinal şablon birebir"""
    if not DOCX_OK: raise ImportError("python-docx kurulu değil")
    doc = DocxDocument()
    yil = _yil()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"4342 SAYILI MERA KANUNU VE 31 TEMMUZ 1998 TARİHLİ MERA YÖNETMELİĞİNİN\n7. VE 13. MADDESİ KAPSAMINDA KİRALANAN MERA-YAYLA ALANLARINA AİT\nKİRALAMA SÖZLEŞMESİ")
    r.bold = True; r.font.size = Pt(11)
    doc.add_paragraph()
    doc.add_paragraph("28 Şubat 1998 tarihli ve 4342 sayılı Mera Kanunu ve bu kanunun usul ve esaslarını "
        "belirleyen ve 31 Temmuz 1998 tarihli ve 23418 sayılı Resmi Gazete'de yayımlanan Mera "
        "Yönetmeliğinin 7. ve 13. Maddeleri kapsamında aşağıda belirtilen mera yaylak ve kışlakların "
        "kiralanmasına ilişkin hükümler taraflarca kabul edilmiştir.")
    doc.add_paragraph()
    p = doc.add_paragraph(); r = p.add_run("GENEL HÜKÜMLER"); r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph("Kiralanan Mera-Yayla Bulunduğu Yerin")
    doc.add_paragraph(f"İli\t\t: Ardahan")
    doc.add_paragraph(f"İlçesi\t\t: {veri.get('ilce','')}")
    doc.add_paragraph(f"Köyü\t\t: {veri.get('koy','')}")
    doc.add_paragraph()
    doc.add_paragraph("Kiralanan Mera-Yaylanın Özellikleri")
    doc.add_paragraph("Durumu ve sınıfı – İyi, kısmen orta")
    doc.add_paragraph(f"Otlatma kapasitesi {veri.get('kapasite','')} adet büyükbaş")
    doc.add_paragraph()
    doc.add_paragraph("Kiralayanın")
    doc.add_paragraph(f"Adı Soyadı: {veri.get('ad_soyad','')}")
    doc.add_paragraph(f"Doğum Yeri ve Tarihi: {veri.get('dogum_yeri','')}")
    doc.add_paragraph(f"Nüfusa Kayıtlı Olduğu Yer: Ardahan İli {veri.get('ilce','')} İlçesi {veri.get('koy','')} Köyü")
    doc.add_paragraph(f"İkametgahı: {veri.get('adres','')}")
    doc.add_paragraph(f"Tebligat Adresi: {veri.get('adres','')}")
    doc.add_paragraph()
    bas = veri.get('baslangic', f'01/06/{yil}')
    bit = veri.get('bitis', f'30/09/{yil}')
    doc.add_paragraph(f"3- Kiralama Süresi")
    doc.add_paragraph(f"Kiranın Başlama Tarihi\t\t: {bas}")
    doc.add_paragraph(f"Kiralamanın Sona Erme Tarihi\t: {bit} (4 Ay)")
    doc.add_paragraph()
    doc.add_paragraph(f"4- Otlatılacak Hayvanın Cinsi ve Miktarı: {veri.get('hayvan_bilgi','')}")
    doc.add_paragraph()
    doc.add_paragraph("5- Otlatma süresi")
    doc.add_paragraph(f"Otlatma başlangıcı\t\t: {bas}")
    doc.add_paragraph(f"Otlatmanın sona ermesi\t: {bit} (4 Ay)")
    maddeler = [
        f"6- Kiralanan yer için komisyonca belirlenen ve ilan edilen hayvan sayısından fazla hayvan otlatılmayacaktır. Kiralanan yer başka amaçla kullanılamaz. Kiralamaya ilişkin her türlü vergi ve resmi harçlar kiracıya aittir.",
        f"7- Kiracıya teslim edilen mera, yaylak ve kışlaklara 3. kişiler tarafından herhangi bir suretle yapılacak tecavüzleri, kiracı 7 gün içinde İl Tarım ve Orman Müdürlüğüne bildirmekle yükümlüdür.",
        f"8- Kiracılar, bölgelerinde huzursuzluk çıkartarak, ekili ve dikili alanlara zarar verdikleri tespit edildiğinde, Valilik onayı ile sözleşmeleri fesih edilir.",
        f"9- Kiracı bu hakkını devredemez, ortak alamaz, kiraya verilen mera, yaylak veya kışlağın sınırlarını daraltamaz, genişletemez, amacı dışında kullanamaz.",
        f"10- Kiracı, ihaleye girdiği sırada dosyasında kulak küpe numaralarını belirttiği hayvanların dışında hayvan meraya götüremez. Aksi halde her büyükbaş hayvan için 360,00 ₺, küçükbaş hayvan için 36,00 ₺ İdari Para Cezası uygulanır.",
        f"11- Komisyonun görevlendirdiği elemanların mera, yaylak ve kışlaklarda yapacakları çalışmaları kiracı tarafından hiçbir surette engellenemez.",
        f"12- Kiracı (Göçer), {yil}/01 sayılı Valilik Genel Emrinde belirtilen esaslara uymak ve yükümlülüklerini yerine getirmek zorundadır.",
        f"13- Göçerler kiraladıkları mera alanına getirecekleri hayvanları ekte yer alan komisyon kararında belirlenen yerlere kadar motorlu araçlarla taşıyacaklardır.",
        f"14- Yukarıda belirtilen maddelere uymayan kiracıların sözleşmesi tek taraflı olarak İl Mera Komisyonu kararı ile fesih edilir.",
    ]
    for m in maddeler:
        doc.add_paragraph(m)
    doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
    tbl = doc.add_table(rows=3, cols=2)
    tbl.cell(0,0).text = veri.get("komisyon_baskani","Semih CEMBEKLİ")
    tbl.cell(1,0).text = "Vali Yardımcısı"
    tbl.cell(2,0).text = "İl Mera Komisyonu Başkanı"
    tbl.cell(0,1).text = veri.get("ad_soyad","")
    tbl.cell(2,1).text = "Kiracı"
    doc.save(dosya)

def word_kiralama_sartname(dosya, veri):
    """Kiralama Şartnamesi — 22 maddelik orijinal şablon birebir"""
    if not DOCX_OK: raise ImportError("python-docx kurulu değil")
    doc = DocxDocument()
    yil = _yil()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MERA YAYLAK ve KIŞLAK ŞARTNAMESİ")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x1E,0x56,0x31)
    doc.add_paragraph()
    p = doc.add_paragraph(); r = p.add_run("GENEL ŞARTLAR"); r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph()
    ilce=veri.get('ilce',''); koy=veri.get('koy',''); ada=veri.get('ada_parsel','')
    kapasite=veri.get('kapasite','')
    bas=f"01.06.{yil}"; bit=f"30.09.{yil}"
    maddeler_genel = [
        f"Madde 1- Ardahan İli {ilce} İlçesi {koy} Köy kadastro alanı içerisinde yer alan {ada} numaralı mera-yaylak alanının 4342 sayılı Mera Kanunu ve Mera Yönetmeliğinin 7. ve 13. maddeleri kapsamında kiralanmasıdır.",
        f"Madde 2- Kiraya verilecek mera, yaylak ve kışlağın otlatma kapasitesi {kapasite} adet Büyükbaş hayvan ile sınırlıdır.",
        f"Madde 3- İhale Ardahan İli Mera Komisyonunca 2886 Sayılı Devlet İhale Kanununun 51. Maddesi (g) bendi ve 4342 sayılı Mera Kanunu ve bu kanunun usul ve esaslarını belirleyen Mera Yönetmeliğinin 7. ve 13. Maddeleri kapsamında yapılacaktır.",
        f"Madde 4- Mera yaylak ve kışlakların kiralama süresi 4 aydır. ({bas} – {bit})",
        f"Madde 5- Geçici teminat, tahmini bedelinin %30'udur. Defterdarlık Muhasebesinde İl Müdürlüğümüz (Ardahan İl Tarım ve Orman Müdürlüğü) hesabına yatırılacaktır. IBAN: TR 5100 0010 0100 0003 5015 4026",
        f"Madde 6- Kesin teminat, ihale bedelinin %6'sıdır. Sözleşmenin Noterde onaylanmasının ardından Defterdarlık Muhasebesinde İl Müdürlüğümüz hesabına yatırılacaktır.",
        f"Madde 7- İhale komisyonu gerekçesini kararda belirtmek suretiyle ihaleyi yapıp yapmamakta serbesttir.",
        f"Madde 8- İhale komisyonu tarafından alınan ihale kararları ita amirlerince karar tarihinden itibaren en geç 15 iş günü içinde onaylanır veya iptal edilir.",
        f"Madde 9- Kiracılık hakkı sona erdiğinde kiraya verilen mera, yaylak, kışlak, kiracı tarafından İl Tarım ve Orman Müdürlüğü'ne (Çayır, Mera ve Yem Bitkileri Şube Müdürlüğü) eksiksiz teslim edilecektir.",
        f"Madde 10- Kira süresi sona erdiği halde kiraya verilen yer kiracı tarafından komisyona müracaat ederek teslim edilmediği takdirde, her geçen gün için günlük kira bedelinin iki katı tazminat olarak tahsil edilir.",
        f"Madde 11- Kiraya verilen mera, yaylak ve umuma ait çayır kiracı tarafından korunacak, değerini düşürmeyecek, özelliğini ve verim gücünü bozmayacaktır.",
        f"Madde 12- Kiracı bu hakkını devredemez, ortak alamaz, kiraya verilen mera, yaylak veya kışlağın sınırlarını daraltamaz, genişletemez, amacı dışında kullanamaz.",
        f"Madde 13- Sözleşme süresinin bitimi veya süresinden evvel iptali halinde tebligatı müteakip 15 gün içerisinde tahliye edilir.",
        f"Madde 14- Göçerler görevlilerin denetiminde gerekli bilgi ve belgeleri göstermekle yükümlüdür.",
        f"Madde 15- Yukarıda yazılan hususlarla birlikte kiraya verilen mera, yaylak, kışlağın komisyonca belirleyeceği hususlara uyulacaktır.",
        f"Madde 16- Bu şartname Mera komisyonunun belirleyeceği diğer hususlarla birlikte sözleşmenin ekine teşkil edilir.",
        f"Madde 17- Geçici teminat kira bedelinin tamamının yatırılmasından sonra, kesin teminat ise kira süresi sonunda ilgiliye ödenecektir.",
        f"Madde 18- Gerektiğinde ihtilafların hal mercii Ardahan İcra Daireleri ve Mahkemeleridir.",
        f"Madde 19- Göçer kiracılar 1774 Sayılı Kimlik Bildirme Kanununa göre 15 gün içerisinde kiraladıkları meranın bağlı bulunduğu Jandarma karakoluna bildirimde bulunacaklardır.",
        f"Madde 20- Kiracılar, bölgelerinde huzursuzluk çıkartarak, ekili ve dikili alanlara zarar verdikleri tespit edildiğinde, Valilik onayı ile bölgeden çıkarılacaklardır.",
        f"Madde 21- İhaleye birden fazla katılım olması durumunda en yüksek 1. teklifi veren kişinin ihale yükümlülüklerini komisyonda belirtilen sürede yerine getirmemesi halinde 2. en yüksek teklifi veren kişiye hak tanınır.",
        f"Madde 22- İş bu şartname bu madde dâhil 22 (yirmi iki) maddedir.",
    ]
    for m in maddeler_genel:
        doc.add_paragraph(m)
    doc.add_paragraph()
    doc.add_paragraph("Yukarıdaki maddeleri kabul ediyorum.")
    doc.add_paragraph()
    doc.add_paragraph(f"Adı Soyadı: {veri.get('ad_soyad','')}")
    doc.add_paragraph("İmza:")
    doc.add_paragraph()
    # ÖZEL ŞARTLAR
    p = doc.add_paragraph(); r = p.add_run("ÖZEL ŞARTLAR"); r.bold = True; r.font.size = Pt(12)
    ozel_sartlar = [
        "1- Kiralanan yer için komisyonca belirlenen ve ilan edilen hayvan sayısından fazla hayvan otlatılmayacaktır. Kiralanan yer başka amaçla kullanılamaz.",
        "2- Kiralamaya ilişkin her türlü vergi ve resmi harçlar kiracıya aittir.",
        "3- Kiracıya teslim edilen mera, yaylak ve kışlaklara 3'üncü kişiler tarafından herhangi bir suretle yapılacak tecavüzleri, kiracı 7 gün içinde İl Tarım ve Orman Müdürlüğüne bildirmekle yükümlüdür.",
        "4- Mera, yaylak, kışlaktan faydalanacak sürü sahipleri Komisyonun/Jandarmanın vereceği kimlik belgelerini yanında bulunduracaklar.",
        "5- Komisyonun görevlendirdiği elemanlar ve İl/İlçe Tarım ve Orman Müdürlüğü personellerince kiralanan mera, yaylak ve kışlakta yapılacak çalışmalar engellenemez.",
        f"6- Kiraya verilen mera, yaylak, kışlağın 4 aylık kira süresi ({bas}-{bit}) dışında süre uzatımı talebinde bulunulamaz.",
        "7- Kiracı, kiralanan alan, alan üzerinde bulunan bina, koruyucu çit duvar ve buna benzer tesislere zarar veremez, yenilerini tesis edemez.",
        "8- Yapılan ihtara rağmen kiracı tarafından kira bedelinin süresinde ödenmemesi ya da sözleşmede belirtilen diğer hususlara uyulmaması halinde sözleşme fesih edilir.",
        "9- Kiracı tarafından ödenmeyen kira bedeli, 6183 sayılı Amme Alacaklarının Tahsil Usulü Hakkında Kanun hükümlerine göre tahsil edilir.",
        "10- Sözleşmenin hazırlanmasının ardından göçerlerin mera, yaylak ve kışlaklara çıkabilmesi için gerekli belgeler ilgili yerin Mülki Amirliklerince düzenlenecektir.",
        "11- İhaleyi kazanan kimse herhangi bir şekilde üçüncü şahıslara mera, yaylak veya kışlağı kiralayamaz ve amacı dışında kullanılamaz.",
        "12- Hayvanlar yolun gittiği yere kadar motorlu araçlar ile götürülecektir.",
        "13- Kiracı tespit edilen hayvan sayısı üzerinde hayvan getiremeyecektir. Aksi halde her büyükbaş hayvan için 360,00 ₺, küçükbaş hayvan için 36,00 ₺ İdari Para Cezası uygulanır.",
        f"14- {int(yil)-1} yılında mera, yaylak, kışlak ve umuma ait çayırları kiralayıp Jandarma tarafından, kiraladıkları yere kira şartnamesinde, sözleşmesinde ve Valilik Genel Emrinde belirtilen esaslara aykırı hareket edenlere yaptırım uygulanır.",
    ]
    for s in ozel_sartlar:
        doc.add_paragraph(s)
    doc.add_paragraph()
    doc.add_paragraph("Bu şartnamedeki yazılı hususları olduğu gibi kabul ve taahhüt ederim.")
    doc.add_paragraph()
    doc.add_paragraph(f"Adı Soyadı: {veri.get('ad_soyad','')}")
    doc.add_paragraph("İmza:")
    doc.add_paragraph(f"İmza Tarihi: {veri.get('tarih', datetime.now().strftime('%d/%m/%Y'))}")
    doc.add_paragraph()
    # ÖZEL HÜKÜMLER
    p = doc.add_paragraph(); r = p.add_run("ÖZEL HÜKÜMLER"); r.bold = True; r.font.size = Pt(12)
    ozel_hukumler = [
        "Kiracı, kiralanan alan, alan üzerinde bulunan bina, koruyucu çit, duvar ve buna benzer tesislere zarar veremez, yenilerini inşa edemez.",
        "Kiralama ücreti her yıl en geç otlatma mevsimine başlamadan peşin olarak Mera Özel Gelir Hesabına yatırılır.",
        "Bu alanları kiralayanlar, Komisyonlarca belirlenen ıslah, amenajman planlarını ve otlatma planlarına uymakla yükümlüdürler.",
        "Kiracı tarafından kira bedelinin yapılan ihtara rağmen süresinde ödenmemesi halinde sözleşme fesih edilir.",
        "Kiracı tarafından ödenmeyen kira bedeli, 6183 sayılı Amme Alacaklarını Tahsil Usulü Hakkında Kanun hükümlerine göre tahsil edilir.",
        "Kiracılar, Otlatma bedelinin tamamını otlatma izni verildiğini belirten sözleşmenin yapılması sırasında Mera Özel Gelir Hesabına yatırır.",
        "Komisyon gerekli gördüğü takdirde bölgelerinin özel şartlarına göre sözleşmeye özel hükümler eklemeye yetkilidirler.",
        "Bu sözleşmeyle Genel Şartlar, Özel Şartlar ve Özel Hükümleri olduğu gibi kabul ve taahhüt ederim.",
        "Göçerlerin yayla ve kışlaklarına hareket edecek olan büyükbaş hayvan sürülerine Şap, Şarbon ve LSD aşılanmış olması mecburidir.",
    ]
    for h in ozel_hukumler:
        doc.add_paragraph(h)
    doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
    tbl = doc.add_table(rows=3, cols=2)
    tbl.cell(0,0).text = veri.get("komisyon_baskani","Semih CEMBEKLİ")
    tbl.cell(1,0).text = "Vali Yardımcısı"
    tbl.cell(2,0).text = "İl Mera Komisyonu Başkanı"
    tbl.cell(0,1).text = veri.get("ad_soyad","")
    tbl.cell(2,1).text = "Kiralayan Şahıs"
    doc.save(dosya)

def word_idari_ceza(dosya, veri):
    """İdari Para Cezası Kararı Oluru — orijinal tablo formatında birebir"""
    if not DOCX_OK: raise ImportError("python-docx kurulu değil")
    doc = DocxDocument()
    yil = _yil()
    # Ana tablo — 26 satır, orijinal yapıda
    tbl = doc.add_table(rows=26, cols=4, style='Table Grid')
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    def _cell(r,c,txt,bold=False):
        cell=tbl.cell(r,c); cell.text=str(txt)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size=Pt(10)
                if bold: run.bold=True
    def _merge(r,c1,c2,txt,bold=False):
        merged=tbl.cell(r,c1).merge(tbl.cell(r,c2))
        merged.text=str(txt)
        for p in merged.paragraphs:
            for run in p.runs:
                run.font.size=Pt(10)
                if bold: run.bold=True
    # Başlıklar
    _merge(0,0,3,"İDARİ PARA CEZASI UYGULANAN ŞAHSIN",True)
    _merge(1,0,3,"KİMLİK BİLGİLERİ",True)
    # Kimlik bilgileri
    _merge(2,0,1,"Adı Soyadı",True); _merge(2,2,3,veri.get("ad_soyad",""))
    _merge(3,0,1,"T.C. Kimlik No",True); _merge(3,2,3,veri.get("tc",""))
    _merge(4,0,1,"Baba Adı",True); _merge(4,2,3,veri.get("baba_adi",""))
    _merge(5,0,1,"Doğum Yeri-Tarihi",True); _merge(5,2,3,veri.get("dogum",""))
    _cell(6,0,"Nüfusa Kayıtlı Olduğu",True); _cell(6,1,"İl",True); _merge(6,2,3,"Ardahan")
    _cell(7,0,"Nüfusa Kayıtlı Olduğu",True); _cell(7,1,"İlçe",True); _merge(7,2,3,veri.get("ilce","Merkez"))
    _merge(8,0,1,"Belge Türü",True); _cell(8,2,"Kimlik"); _cell(8,3,"")
    _merge(9,0,1,"Plaka No-Ehliyet Ruhsat No",True); _merge(9,2,3,"")
    _merge(10,0,1,"İkametgâh Adresi",True); _merge(10,2,3,veri.get("adres",""))
    # Ceza bilgileri
    _merge(11,0,3,"İDARİ PARA CEZASINA İLİŞKİN BİLGİLER",True)
    _merge(12,0,2,"Miktarı (Rakamla)",True); _cell(12,3,f"{veri.get('tutar','')} TL")
    _merge(13,0,2,"Ödeneceği Yer",True); _cell(13,3,"Ardahan Defterdarlığı")
    _merge(14,0,2,"Son Ödeme Tarihi",True); _cell(14,3,"Tebliğ tarihinden itibaren en geç 30 (otuz) gün içerisinde")
    # Yasal dayanak
    _merge(15,0,3,"(*) İDARİ PARA CEZASININ YASAL DAYANAĞI",True)
    dayanak = ("4342 Sayılı Mera Kanunu 26. ve 27. Maddeleri hükümlerine göre, "
        f"{yil}/01 Sayılı Valilik Genel Emrinin 18. Maddesi (h) bendi, "
        f"Çayır Mera ve Yem Bitkileri Şube Müdürlüğü ekiplerinin "
        f"{veri.get('tarih', datetime.now().strftime('%d/%m/%Y'))} tarihinde "
        f"{veri.get('mera','')} mera alanında yaptıkları denetim sonucunda "
        f"{veri.get('ad_soyad','')} isimli şahsın {veri.get('konu','')} "
        f"fiilinden dolayı İdari Para Cezası uygulanmasına karar verilmiştir.")
    _merge(16,0,3,dayanak)
    for i in range(17,20): _merge(i,0,3,"")
    not_txt = (f"Not: (*) {yil}/01 Sayılı Valilik Genel Emrinin 18. Maddesi (h) bendinde; "
        "meralara kapasitesinden fazla veya izinsiz hayvan sokan ya da meralardan izinsiz "
        "faydalanan kişilere 4342 sayılı Mera Kanununun 26. maddesinde belirtilen İPC uygulanır.")
    _merge(20,0,3,not_txt)
    _merge(21,0,3,"")
    # İtiraz bilgileri
    _merge(22,0,3,"KESİLEN İDARİ PARA CEZASINA",True)
    _merge(23,0,2,"İtiraz Merci",True); _cell(23,3,"Ardahan Sulh Ceza Hakimliği")
    _merge(24,0,2,"Son İtiraz Tarihi",True); _cell(24,3,"Tebliğ tarihinden itibaren 15 (onbeş) gün içerisinde")
    _merge(25,0,3,"Kararın tebliğ tarihinden itibaren 15 (onbeş) gün içerisinde yetkili Sulh Ceza Hakimliğine başvurulabilir. Süresinde başvurulmaması halinde karar kesinleşir.")
    # Açıklama paragrafı
    doc.add_paragraph()
    doc.add_paragraph(
        f"Yukarıda açık kimliği yazılı {veri.get('ad_soyad','')} hakkında "
        f"4342 Sayılı Mera Kanununun 26. Maddesi gereğince {veri.get('tutar','')} TL "
        f"İdari Para Cezası uygulanmasına karar verilmiştir.")
    doc.add_paragraph()
    # İl Müdürü tablosu
    tbl2 = doc.add_table(rows=2, cols=1)
    tbl2.cell(0,0).text = "Muhammet Fatih CİNEVİZ"
    tbl2.cell(1,0).text = "İl Müdürü"
    doc.add_paragraph()
    # OLUR tablosu
    tbl3 = doc.add_table(rows=6, cols=1)
    tbl3.cell(0,0).text = "OLUR"
    tbl3.cell(1,0).text = f"…../{datetime.now().strftime('%m')}/{yil}"
    tbl3.cell(2,0).text = veri.get("komisyon_baskani","Semih CEMBEKLİ")
    tbl3.cell(3,0).text = "Vali a."
    tbl3.cell(4,0).text = "Vali Yardımcısı"
    tbl3.cell(5,0).text = "Mera Komisyon Başkanı"
    doc.save(dosya)

def word_bilgi_notu(dosya, veri):
    """Bilgi Notu Word şablonu"""
    if not DOCX_OK: raise ImportError("python-docx kurulu değil")
    doc = DocxDocument()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BİLGİ NOTU"); r.bold = True; r.font.size = Pt(14)
    doc.add_paragraph()
    tbl = doc.add_table(rows=4, cols=2, style='Table Grid')
    bilgiler = [
        ("Sunulan Makam", veri.get("makam","")),
        ("Konu", veri.get("konu","")),
        ("Hazırlayan Birim", veri.get("birim","Çayır, Mera ve Yem Bitkileri Şube Müdürlüğü")),
        ("Hazırlama Tarihi", veri.get("tarih", datetime.now().strftime("%d.%m.%Y"))),
    ]
    for i,(lbl,val) in enumerate(bilgiler):
        tbl.cell(i,0).text = lbl
        tbl.cell(i,1).text = str(val)
        for cell in tbl.rows[i].cells:
            for p2 in cell.paragraphs:
                for r2 in p2.runs: r2.font.size = Pt(11)
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("İçerik:"); r.bold = True
    doc.add_paragraph(veri.get("icerik",""))
    doc.add_paragraph()
    doc.add_paragraph(f"Hazırlayan: {veri.get('hazirlayan','')}")
    doc.save(dosya)

# ─── İHALE BEDEL HESAPLAMA ───────────────────────────────────────────────────
def para_parse(metin):
    """Türk formatındaki para değerini float'a çevir: 52.507,18 → 52507.18"""
    if not metin: return 0.0
    t=str(metin).strip().replace("₺","").replace("TL","").strip()
    # Türk formatı: 52.507,18 → nokta binlik, virgül kuruş
    if "," in t and "." in t:
        t=t.replace(".","").replace(",",".")
    elif "," in t:
        t=t.replace(",",".")
    # sadece nokta varsa → zaten ondalık (1500.50) veya binlik (1.500)?
    # Eğer noktadan sonra 3+ hane varsa binlik ayracı
    elif "." in t:
        parcalar=t.split(".")
        if len(parcalar)==2 and len(parcalar[1])==3:
            t=t.replace(".","")  # binlik ayracı
    return float(t)

def para_format(sayi):
    """Float'ı Türk para formatına çevir: 52507.18 → 52.507,18 ₺"""
    try:
        s=f"{float(sayi):,.2f}"  # 52,507.18
        # İngilizce → Türkçe: virgül↔nokta swap
        s=s.replace(",","X").replace(".",",").replace("X",".")
        return f"{s} ₺"
    except (ValueError,TypeError):
        return "0,00 ₺"

def ihale_bedel_hesapla(bedel):
    """İhale bedelinden tüm kalemleri hesaplar"""
    try:
        b = para_parse(bedel)
        if b <= 0: return None
        return {
            "toplam": b,
            "bakanlik_25": b * 0.25,
            "koy_75": b * 0.75,
            "kesin_teminat_6": b * 0.06,
            "damga_vergisi": b * 5.69 / 1000,
        }
    except (ValueError, TypeError):
        return None

# ─── ARAYÜZ BİLEŞENLERİ ──────────────────────────────────────────────────────
class MBtn(tk.Button):
    def __init__(self,master,text,command=None,color="#2D8C55",**kw):
        kw.setdefault("pady",7); kw.setdefault("padx",14)
        super().__init__(master,text=text,command=command,bg=color,fg="white",
            font=("Segoe UI",10,"bold"),relief="flat",activebackground="#5DADE2",cursor="hand2",**kw)
        self._orig=text
    def loading(self,d):
        self.config(state="disabled" if d else "normal",text="⏳ İşleniyor..." if d else self._orig)

class StatKart(tk.Frame):
    def __init__(self,master,baslik,deger,renk="#1E5631",ikon="",**kw):
        super().__init__(master,bg=C_WHITE,relief="flat",highlightbackground="#D0DDD8",highlightthickness=1,**kw)
        tk.Label(self,text=ikon,font=("Segoe UI",22),bg=C_WHITE).pack(pady=(14,2))
        self._v=tk.Label(self,text=str(deger),font=("Segoe UI",26,"bold"),fg=renk,bg=C_WHITE); self._v.pack()
        tk.Label(self,text=baslik,font=("Segoe UI",9),fg="#666",bg=C_WHITE).pack(pady=(2,14))
    def set(self,v): self._v.config(text=str(v))

class StatusBar(tk.Label):
    def __init__(self,master,**kw):
        super().__init__(master,text="  Hazır.",anchor="w",bg="#D5E5DC",fg="#1E5631",font=("Segoe UI",9),padx=12,**kw)
    def set(self,msg): self.config(text=f"  {msg}"); self.update_idletasks()

class AramaFrame(tk.Frame):
    """Canlı arama — yazarken anında filtreler (debounce 200ms)"""
    def __init__(self,master,cb,ph="🔍 Ara...",**kw):
        super().__init__(master,**kw); self._cb=cb; self._timer=None
        self.var=tk.StringVar(); self.var.trace_add("write",self._on_change)
        tk.Label(self,text="🔍",font=("Segoe UI",11),bg=self["bg"]).pack(side="left",padx=(0,4))
        ttk.Entry(self,textvariable=self.var,width=30).pack(side="left")
        tk.Button(self,text="✕",command=lambda:self.var.set(""),relief="flat",cursor="hand2",bg=self["bg"]).pack(side="left",padx=4)
    def _on_change(self,*_):
        if self._timer: self.after_cancel(self._timer)
        self._timer=self.after(200,lambda:self._cb(self.var.get()))

# ─── VERİTABANI ───────────────────────────────────────────────────────────────
def init_db():
    if not DB_PATH: return
    with sqlite3.connect(DB_PATH) as c:
        c.executescript("""
        CREATE TABLE IF NOT EXISTS Kullanicilar(k_adi TEXT PRIMARY KEY,sifre TEXT NOT NULL,
            yetki TEXT NOT NULL DEFAULT 'Uzman',ad TEXT NOT NULL,unvan TEXT,
            aktif INTEGER NOT NULL DEFAULT 1,fail_count INTEGER NOT NULL DEFAULT 0,lockout_ts TEXT,
            riza_onay INTEGER DEFAULT 0);
        CREATE TABLE IF NOT EXISTS Ayarlar(k_adi TEXT PRIMARY KEY,tema TEXT DEFAULT 'Orman Yeşili',
            punto INTEGER DEFAULT 10,sube_mudur TEXT DEFAULT 'Leyla ARSLAN');
        CREATE TABLE IF NOT EXISTS Loglar(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tarih TEXT,kul TEXT,islem TEXT,detay TEXT);
        CREATE TABLE IF NOT EXISTS Duyurular(id INTEGER PRIMARY KEY AUTOINCREMENT,
            mesaj TEXT,tarih TEXT,gonderen TEXT);
        CREATE TABLE IF NOT EXISTS Mera_Varligi(koy TEXT PRIMARY KEY,ilce TEXT,alan REAL,yem REAL,turkvet_tarihi TEXT);
        CREATE TABLE IF NOT EXISTS Rapor_Gecmisi(rapor_no TEXT PRIMARY KEY,koy TEXT,
            talep_eden TEXT,tc TEXT,islem_tarihi TEXT,duzenleyen TEXT,aciklama TEXT);
        CREATE TABLE IF NOT EXISTS Ihaleler(id INTEGER PRIMARY KEY AUTOINCREMENT,
            koy TEXT,ilce TEXT,ad_soyad TEXT,tc TEXT,telefon TEXT,adres TEXT,
            bedel REAL,durum TEXT DEFAULT 'Başvuru Alındı',notlar TEXT,tarih TEXT,
            ada_parsel TEXT,kapasite TEXT);
        CREATE TABLE IF NOT EXISTS Ihale_Log(id INTEGER PRIMARY KEY AUTOINCREMENT,
            ihale_id INTEGER,tarih TEXT,personel TEXT,durum TEXT,not_icerik TEXT);
        CREATE TABLE IF NOT EXISTS Islah_Amenajman(id INTEGER PRIMARY KEY AUTOINCREMENT,
            koy TEXT,ilce TEXT,dilekce_tarihi TEXT,talep_eden TEXT,talep_alani REAL,
            talep_aciklama TEXT,verilen_alan REAL DEFAULT 0,verilmeme_neden TEXT,
            is_programi TEXT,durum TEXT DEFAULT 'Bekliyor',kapanma_tarihi TEXT);
        CREATE TABLE IF NOT EXISTS Tahsisler(id INTEGER PRIMARY KEY AUTOINCREMENT,
            koy TEXT,ilce TEXT,ada TEXT,parsel TEXT,kurum TEXT,amac TEXT,alan_ha REAL,
            asama TEXT DEFAULT '1-Dilekçe Alındı',durum TEXT DEFAULT 'Devam Ediyor',
            basvuru_t TEXT,sonuc_t TEXT,notlar TEXT,
            madde14_bent TEXT,ot_geliri REAL,tescil_tarihi TEXT);
        CREATE TABLE IF NOT EXISTS Tahsis_Log(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tahsis_id INTEGER,tarih TEXT,personel TEXT,asama TEXT,aciklama TEXT);
        CREATE TABLE IF NOT EXISTS Muhtarlar(id INTEGER PRIMARY KEY AUTOINCREMENT,
            ilce TEXT,koy TEXT,ad_soyad TEXT,telefon TEXT,email TEXT,notlar TEXT);
        CREATE TABLE IF NOT EXISTS Veri_Kayit(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tarih TEXT,kul TEXT,baslik TEXT,kategori TEXT,detay TEXT);
        CREATE TABLE IF NOT EXISTS Ajanda(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tarih TEXT,sure TEXT,baslik TEXT,tur TEXT DEFAULT 'Hatırlatıcı',
            durum TEXT DEFAULT 'Bekliyor',icerik TEXT,k_adi TEXT);
        CREATE TABLE IF NOT EXISTS Sifre_Talepleri(id INTEGER PRIMARY KEY AUTOINCREMENT,
            k_adi TEXT,tarih TEXT,durum TEXT DEFAULT 'Bekliyor');
        CREATE TABLE IF NOT EXISTS Kayit_Talepleri(id INTEGER PRIMARY KEY AUTOINCREMENT,
            ad TEXT,unvan TEXT,k_adi TEXT,tarih TEXT,durum TEXT DEFAULT 'Bekliyor');
        CREATE TABLE IF NOT EXISTS Sikayetler(id INTEGER PRIMARY KEY AUTOINCREMENT,
            koy TEXT,ilce TEXT,sikayet_eden TEXT,tur TEXT,aciklama TEXT,
            durum TEXT DEFAULT 'Yeni',tarih TEXT,kapanma TEXT,sonuc TEXT);
        CREATE TABLE IF NOT EXISTS Personel_Takvim(id INTEGER PRIMARY KEY AUTOINCREMENT,
            k_adi TEXT,ad TEXT,tarih TEXT,baslangic TEXT,bitis TEXT,
            tur TEXT DEFAULT 'Görev',aciklama TEXT);
        CREATE TABLE IF NOT EXISTS Son_Islemler(id INTEGER PRIMARY KEY AUTOINCREMENT,
            k_adi TEXT,modul TEXT,kayit_adi TEXT,tarih TEXT);
        CREATE TABLE IF NOT EXISTS Idari_Cezalar(id INTEGER PRIMARY KEY AUTOINCREMENT,
            ad_soyad TEXT,tc TEXT,il TEXT DEFAULT 'Ardahan',ilce TEXT,
            mera_koy TEXT,mera_ada_parsel TEXT,yil INTEGER,konu TEXT,
            hayvan_sayisi TEXT,ceza_miktari TEXT,ipc_tutari REAL,
            tarih TEXT,durum TEXT DEFAULT 'Uygulandı',notlar TEXT);
        CREATE TABLE IF NOT EXISTS Ihale_Yerleri(id INTEGER PRIMARY KEY AUTOINCREMENT,
            ilce TEXT,koy TEXT,ada TEXT,parsel TEXT,alan_da REAL,kapasite_bbhb REAL,
            tahmini_bedel REAL,vasif TEXT,durum TEXT DEFAULT 'Aktif',yil TEXT);
        CREATE TABLE IF NOT EXISTS Iletisim_Formu(id INTEGER PRIMARY KEY AUTOINCREMENT,
            gonderen TEXT,konu_tipi TEXT,mesaj TEXT,tarih TEXT,durum TEXT DEFAULT 'Yeni');
        CREATE TABLE IF NOT EXISTS Silme_Talepleri(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tablo TEXT,kayit_id INTEGER,kayit_ozet TEXT,talep_eden TEXT,tarih TEXT,
            durum TEXT DEFAULT 'Bekliyor');
        CREATE TABLE IF NOT EXISTS Dahili_Mesajlar(id INTEGER PRIMARY KEY AUTOINCREMENT,
            gonderen TEXT,alici TEXT,konu TEXT,mesaj TEXT,tarih TEXT,
            okundu INTEGER DEFAULT 0,tur TEXT DEFAULT 'Mesaj');
        CREATE TABLE IF NOT EXISTS Islah_Projeler(id INTEGER PRIMARY KEY AUTOINCREMENT,
            yil TEXT,ilce TEXT,koy TEXT,gubre_da REAL DEFAULT 0,mera_tohum_da REAL DEFAULT 0,
            yem_tohum_da REAL DEFAULT 0,golgelik INTEGER DEFAULT 0,sivat INTEGER DEFAULT 0,
            boru_m REAL DEFAULT 0,coban_evi INTEGER DEFAULT 0,tuzluk INTEGER DEFAULT 0,
            kasinka INTEGER DEFAULT 0,finansman TEXT,notlar TEXT,durum TEXT DEFAULT 'Tamamlandı');
        CREATE TABLE IF NOT EXISTS Iletisim_Bilgileri(anahtar TEXT PRIMARY KEY,deger TEXT);
        CREATE TABLE IF NOT EXISTS Loglar_Arsiv(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tarih TEXT,kul TEXT,islem TEXT,detay TEXT);
        """)
        # Migration — yeni sütunlar
        for tbl, col, tip in [
            ("Ihaleler","ada_parsel","TEXT"),("Ihaleler","kapasite","TEXT"),
            ("Ihaleler","ihale_yeri_id","INTEGER"),
            ("Tahsisler","madde14_bent","TEXT"),("Tahsisler","ot_geliri","REAL"),
            ("Tahsisler","tescil_tarihi","TEXT"),
            ("Kullanicilar","riza_onay","INTEGER DEFAULT 0"),
            ("Kullanicilar","sifre_tarih","TEXT"),
            ("Ayarlar","vali_yardimcisi","TEXT DEFAULT 'Semih CEMBEKLİ'"),
        ]:
            try: c.execute(f'ALTER TABLE {tbl} ADD COLUMN {col} {tip}')
            except Exception: pass
        # DB indexes — performans
        for idx in [
            "CREATE INDEX IF NOT EXISTS idx_log_tarih ON Loglar(tarih)",
            "CREATE INDEX IF NOT EXISTS idx_ihale_koy ON Ihaleler(koy)",
            "CREATE INDEX IF NOT EXISTS idx_ceza_tc ON Idari_Cezalar(tc)",
            "CREATE INDEX IF NOT EXISTS idx_mera_koy ON Mera_Varligi(koy)",
        ]:
            try: c.execute(idx)
            except Exception: pass
        # WAL mode — çok kullanıcılı erişim
        try: c.execute("PRAGMA journal_mode=WAL")
        except Exception: pass
        # Log arşivleme — 1 yıldan eski logları arşivle
        try:
            bir_yil_once=(datetime.now()-timedelta(days=365)).strftime("%Y-%m-%d")
            c.execute("INSERT INTO Loglar_Arsiv(tarih,kul,islem,detay) SELECT tarih,kul,islem,detay FROM Loglar WHERE tarih<?",(bir_yil_once,))
            c.execute("DELETE FROM Loglar WHERE tarih<?",(bir_yil_once,))
        except Exception: pass
        # İletişim bilgileri default
        for k,v in [("telefon",DEV_TEL),("email",DEV_MAIL),("whatsapp",DEV_WA)]:
            try: c.execute("INSERT OR IGNORE INTO Iletisim_Bilgileri(anahtar,deger)VALUES(?,?)",(k,v))
            except Exception: pass
        if c.execute("SELECT COUNT(*) FROM Kullanicilar").fetchone()[0]==0:
            c.execute("INSERT INTO Kullanicilar(k_adi,sifre,yetki,ad,unvan,riza_onay)VALUES(?,?,?,?,?,?)",
                ("admin",hash_pw("Admin123!"),"Admin","Sistem Yöneticisi","Ziraat Mühendisi",1))
            c.execute("INSERT INTO Kullanicilar(k_adi,sifre,yetki,ad,unvan,riza_onay)VALUES(?,?,?,?,?,?)",
                ("emre",hash_pw("Emre1234"),"Uzman","Emre ÖZTÜRK","Ziraat Mühendisi",1))

# ─── ANA UYGULAMA ─────────────────────────────────────────────────────────────
def guncelleme_kontrol_et():
    # GITHUB_REPO değişkeni kodunda "miras75mere/miras-enterprise" olarak tanımlı
    api_url = f"https://api.github.com/repos/{GITHUB_REPO}/releases/latest"
    try:
        req = urllib.request.Request(api_url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=10) as response:
            data = json.loads(response.read().decode())
            
        en_yeni_surum = data.get("tag_name")
        
        # Eğer GitHub'daki sürüm mevcut sürümden farklıysa güncelleme uyarısı ver
        if en_yeni_surum and en_yeni_surum != MEVCUT_SURUM:
            cevap = messagebox.askyesno("Güncelleme Bulundu", 
                                        f"MİRAS'ın yeni bir sürümü ({en_yeni_surum}) mevcut.\nŞimdi güncellensin mi?")
            if cevap:
                assets = data.get("assets", [])
                if assets:
                    # Yüklediğin .exe dosyasının indirme linkini al
                    indirme_linki = assets[0].get("browser_download_url")
                    guncellemeyi_indir_ve_kur(indirme_linki)
                else:
                    messagebox.showerror("Hata", "GitHub'da indirilecek .exe dosyası bulunamadı.")
    except Exception as e:
        print(f"Güncelleme kontrolü başarısız oldu: {e}")

def guncellemeyi_indir_ve_kur(indirme_linki):
    import subprocess
    # Eğer kod exe olarak çalıştırılıyorsa
    if getattr(sys, 'frozen', False): 
        exe_yolu = sys.executable
    else:
        messagebox.showinfo("Bilgi", "Uygulama şu an kaynak kodundan (.py) çalışıyor.\nOtomatik güncelleme işlemi yalnızca derlenmiş .exe versiyonunda aktiftir.")
        return

    eski_exe = exe_yolu + ".eski"
    yeni_exe_gecici = "guncel_miras_gecici.exe"

    try:
        # İnternetten yeni exe dosyasını indir
        req = urllib.request.Request(indirme_linki, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as response, open(yeni_exe_gecici, 'wb') as out_file:
            shutil.copyfileobj(response, out_file)
        
        # Windows'un dosya üzerine yazma kilidini aşmak için isimleri değiştir
        if os.path.exists(eski_exe):
            os.remove(eski_exe)
            
        os.rename(exe_yolu, eski_exe)
        os.rename(yeni_exe_gecici, exe_yolu)
        
        messagebox.showinfo("Başarılı", "Güncelleme tamamlandı. Yeni sürüm başlatılıyor...")
        
        # Yeni güncellenmiş uygulamayı aç ve mevcut açık olanı kapat
        subprocess.Popen([exe_yolu])
        sys.exit()
        
    except Exception as e:
        messagebox.showerror("Hata", f"Güncelleme sırasında bir hata oluştu:\n{e}")
        if os.path.exists(yeni_exe_gecici):
            os.remove(yeni_exe_gecici)
class MirasApp:
    MENU_ITEMS_UZMAN = [
        ("🏠","Dashboard","dash"),("🧮","Kapasite","kapasite"),
        ("⚖️","İhale Takip","ihale"),("📍","İhale Yerleri","ihale_yer"),
        ("🌱","Islah/Amenajman","islah"),
        ("📋","Tahsis/Md.14","tahsis"),("🚨","Şikayet","sikayet"),
        ("💰","İdari Ceza","ceza"),("📊","Veri Kayıt","veri"),
        ("👥","Muhtarlar","muhtar"),("📅","Ajanda","ajanda"),
        ("👨‍💼","Personel","personel"),
        ("📈","İstatistik","istatistik"),("📝","Evrak Üretici","evrak"),
        ("🌿","MERA AI","ai"),("📖","Mevzuat","mevzuat"),
        ("📞","İletişim","iletisim"),("⚙️","Ayarlar","ayarlar"),
    ]
    MENU_ITEMS_ADMIN = [("🛡️","Admin","admin")]
    MENU_ITEMS_IZLEYICI_HIDE = ["ihale","islah","tahsis","ceza","evrak"]

    def __init__(self,root):
        self.root=root; self.root.title(f"{PROG_ADI} {VERSIYON}"); self.root.geometry("520x380")
        self.ai=GeminiAsistan(); self.session=None; self.yedekci=None; self.status=None
        self.u_id=self.u_yetki=self.u_ad=self.u_unvan=None
        self.tema="Orman Yeşili"; self.punto=10; self.sube_mudur="Leyla ARSLAN"
        self._active_menu=None; self._content_frame=None; self._menu_btns={}
        BACKUP_DIR.mkdir(exist_ok=True)
        self._check_db()

    def gc(self,k): return TEMALAR.get(self.tema,TEMALAR["Orman Yeşili"]).get(k,"#1E5631")

    def _check_db(self):
        global DB_PATH
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE) as f: data=json.load(f)
                yol=data.get("db_path","")
                if yol and os.path.exists(yol):
                    DB_PATH=yol; init_db(); self._build_login(); return
                elif yol:
                    messagebox.showwarning("Veritabanı Bulunamadı",f"Kayıtlı DB bulunamadı:\n{yol}\nLütfen yeniden seçin.")
            except Exception as e: logging.error(f"config:{e}")
        self._build_db_setup()

    def _build_db_setup(self):
        self._clear(); self.root.geometry("560x420")
        f=tk.Frame(self.root,bg=C_WHITE,padx=40,pady=40); f.pack(fill="both",expand=True,padx=20,pady=20)
        tk.Label(f,text="🌿 MİRAS Ağ Kurulumu",font=("Segoe UI",17,"bold"),fg="#1E5631",bg=C_WHITE).pack(pady=(0,8))
        tk.Label(f,text="Ortak bir ağ klasörünü seçerek tüm personelin\naynı veritabanını kullanmasını sağlayın.",
            font=("Segoe UI",10),fg="#555",bg=C_WHITE,justify="center").pack(pady=(0,20))
        MBtn(f,"📂  Mevcut Veritabanını Seç",command=lambda:self._sec_db(False),color=C_INFO,width=36).pack(pady=8)
        MBtn(f,"➕  Yeni Veritabanı Oluştur",command=lambda:self._sec_db(True),color="#2D8C55",width=36).pack(pady=8)
        tk.Label(f,text="İpucu: \\\\sunucu\\tarim\\miras.db gibi ağ yolu kullanın",font=("Segoe UI",8),fg="#999",bg=C_WHITE).pack(pady=(16,0))

    def _sec_db(self,yeni):
        global DB_PATH
        if yeni: yol=filedialog.asksaveasfilename(defaultextension=".db",filetypes=[("SQLite","*.db")])
        else: yol=filedialog.askopenfilename(filetypes=[("SQLite","*.db")])
        if yol:
            DB_PATH=yol
            with open(CONFIG_FILE,"w") as f: json.dump({"db_path":DB_PATH},f)
            init_db(); self._build_login()

    def _build_login(self):
        self._clear(); self.root.geometry("460x700"); self.root.configure(bg="#F0F4F2")
        logo=tk.Frame(self.root,bg="#1E5631",height=140); logo.pack(fill="x")
        tk.Label(logo,text="🌿 MİRAS",font=("Segoe UI",30,"bold"),fg=C_WHITE,bg="#1E5631").pack(pady=(28,2))
        tk.Label(logo,text=f"Enterprise {VERSIYON}",font=("Segoe UI",10),fg="#8FCF9F",bg="#1E5631").pack()
        form=tk.Frame(self.root,bg=C_WHITE); form.pack(fill="both",expand=True,padx=50)
        for lbl,attr,show in [("Kullanıcı Adı","e_u",""),("Şifre","e_p","●")]:
            tk.Label(form,text=lbl,bg=C_WHITE,font=("Segoe UI",10,"bold"),fg="#444").pack(anchor="w",pady=(20 if "Ku" in lbl else 14,2))
            e=ttk.Entry(form,font=("Segoe UI",11),show=show); e.pack(fill="x"); setattr(self,attr,e)
        MBtn(form,"SİSTEME GİRİŞ YAP →",command=self._login,width=28).pack(pady=(28,8))
        self.lbl_hata=tk.Label(form,text="",fg=C_DANGER,bg=C_WHITE,font=("Segoe UI",9)); self.lbl_hata.pack()
        lf=tk.Frame(form,bg=C_WHITE); lf.pack(fill="x",pady=10)
        tk.Button(lf,text="🔑 Şifremi Unuttum",command=self._unuttum,fg="#1E5631",bg=C_WHITE,
            relief="flat",cursor="hand2",font=("Segoe UI",9)).pack(side="left")
        tk.Button(lf,text="📝 Kayıt Talebi",command=self._kayit_talep,fg=C_INFO,bg=C_WHITE,
            relief="flat",cursor="hand2",font=("Segoe UI",9)).pack(side="right")
        alt=tk.Frame(self.root,bg="#F0F4F2"); alt.pack(side="bottom",fill="x",pady=8)
        tk.Label(alt,text=f"{HAKLAR} | {PROG_ADI} {VERSIYON}",fg="#333333",bg="#F0F4F2",font=("Segoe UI",9,"bold")).pack()
        self.root.bind("<Return>",lambda e:self._login()); self.e_u.focus()

    def _unuttum(self):
        if not DB_PATH: messagebox.showerror("Hata","Veritabanı seçilmemiş."); return
        u=tkinter.simpledialog.askstring("Şifre Sıfırlama","Kullanıcı adınızı girin:")
        if u:
            with sqlite3.connect(DB_PATH) as c:
                ex=c.execute("SELECT COUNT(*) FROM Kullanicilar WHERE k_adi=?",(u,)).fetchone()[0]
            if ex:
                with sqlite3.connect(DB_PATH) as c:
                    c.execute("INSERT INTO Sifre_Talepleri(k_adi,tarih)VALUES(?,?)",(u,datetime.now().strftime("%Y-%m-%d")))
                messagebox.showinfo("Tamam","Şifre sıfırlama talebiniz Admin'e iletildi.")
            else: messagebox.showerror("Hata","Bu kullanıcı adı sistemde kayıtlı değil.")

    def _kayit_talep(self):
        if not DB_PATH: return
        win=tk.Toplevel(self.root); win.title("Kayıt Talebi"); win.geometry("340x320")
        win.configure(bg=C_WHITE); win.grab_set()
        tk.Label(win,text="Yeni Personel Kayıt Talebi",font=("Segoe UI",13,"bold"),bg=C_WHITE,fg="#1E5631").pack(pady=16)
        al={}
        for l in ["Ad Soyad:","Ünvan:","Kullanıcı Adı:"]:
            tk.Label(win,text=l,bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=24,pady=(8,0))
            e=ttk.Entry(win,width=32); e.pack(padx=24); al[l]=e
        def _g():
            k=al["Kullanıcı Adı:"].get().strip(); a=al["Ad Soyad:"].get().strip()
            if not k or not a: messagebox.showwarning("Uyarı","Tüm alanları doldurun.",parent=win); return
            with sqlite3.connect(DB_PATH) as c:
                c.execute("INSERT INTO Kayit_Talepleri(ad,unvan,k_adi,tarih)VALUES(?,?,?,?)",(a,al["Ünvan:"].get(),k,datetime.now().strftime("%Y-%m-%d")))
            messagebox.showinfo("Tamam","Talebiniz Admin'e iletildi.",parent=win); win.destroy()
        MBtn(win,"Gönder",command=_g,width=20).pack(pady=16)

    def _login(self):
        if not DB_PATH: messagebox.showerror("Hata","Veritabanı seçilmemiş."); return
        u=self.e_u.get().strip(); p=self.e_p.get()
        if not u or not p: self.lbl_hata.config(text="❌ Kullanıcı adı ve şifre gerekli."); return
        try:
            with sqlite3.connect(DB_PATH) as conn:
                row=conn.execute("SELECT k_adi,sifre,yetki,ad,unvan,aktif,fail_count,lockout_ts,riza_onay FROM Kullanicilar WHERE k_adi=?",(u,)).fetchone()
                ayr=conn.execute("SELECT tema,punto,sube_mudur FROM Ayarlar WHERE k_adi=?",(u,)).fetchone()
        except Exception as e:
            self.lbl_hata.config(text="❌ Veritabanı bağlantı hatası."); logging.error(f"login:{e}"); return
        if not row: self.lbl_hata.config(text="❌ Kullanıcı bulunamadı."); db_log(u,"Başarısız","Bilinmeyen"); return
        k_adi,hashed,yetki,ad,unvan,aktif,fail_c,lkts=row[:8]
        riza_onay=row[8] if len(row)>8 else 0
        if not aktif: self.lbl_hata.config(text="❌ Hesabınız devre dışı."); return
        if lkts:
            bitis=datetime.fromisoformat(lkts)
            if datetime.now()<bitis:
                kalan=int((bitis-datetime.now()).total_seconds()/60)
                self.lbl_hata.config(text=f"🔒 Hesap kilitli. {kalan+1} dk bekleyin."); return
            else:
                with sqlite3.connect(DB_PATH) as c: c.execute("UPDATE Kullanicilar SET fail_count=0,lockout_ts=NULL WHERE k_adi=?",(k_adi,))
                fail_c=0
        if verify_pw(p,hashed):
            with sqlite3.connect(DB_PATH) as c: c.execute("UPDATE Kullanicilar SET fail_count=0,lockout_ts=NULL WHERE k_adi=?",(k_adi,))
            db_log(k_adi,"Giriş","Başarılı")
            self.u_id=k_adi; self.u_yetki=yetki; self.u_ad=ad; self.u_unvan=unvan
            if ayr: self.tema,self.punto,self.sube_mudur=ayr[0] or "Orman Yeşili",ayr[1] or 10,ayr[2] or "Leyla ARSLAN"
            self.root.unbind("<Return>"); self.session=SessionManager()
            self.session.register(self._oturum_bitti); self.yedekci=YedekYoneticisi(self.u_id)
            # Rıza kontrolü
            if not riza_onay:
                self._riza_goster()
            else:
                self._sifre_yas_kontrol()
        else:
            self._login_fail(k_adi,fail_c)

    def _sifre_yas_kontrol(self):
        """Madde 13: Şifre ömrü 180 gün kontrolü"""
        try:
            with sqlite3.connect(DB_PATH) as c:
                row=c.execute("SELECT sifre_tarih FROM Kullanicilar WHERE k_adi=?",(self.u_id,)).fetchone()
            st=row[0] if row and row[0] else None
            if not st:
                with sqlite3.connect(DB_PATH) as c:
                    c.execute("UPDATE Kullanicilar SET sifre_tarih=? WHERE k_adi=?",(datetime.now().strftime("%Y-%m-%d"),self.u_id))
                self._build_app(); return
            sifre_t=datetime.strptime(st,"%Y-%m-%d")
            gecen=(datetime.now()-sifre_t).days
            kalan=SIFRE_OMUR_GUN-gecen
            if kalan<=0:
                messagebox.showwarning("🔒 Şifre Süresi Doldu",
                    f"Şifreniz {SIFRE_OMUR_GUN} günlük ömrünü tamamladı.\n"
                    "Devam etmek için şifrenizi değiştirmeniz gerekiyor.")
                self._zorla_sifre_degistir()
            elif kalan<=SIFRE_UYARI_GUN:
                messagebox.showwarning("⚠️ Şifre Uyarısı",
                    f"Şifrenizin süresi dolmak üzere!\n"
                    f"Kalan gün: {kalan}\n\n"
                    "Lütfen Ayarlar → Şifre sekmesinden şifrenizi değiştirin.")
                self._build_app()
            else:
                self._build_app()
        except Exception as e:
            logging.error(f"sifre_yas:{e}"); self._build_app()
    def _zorla_sifre_degistir(self):
        """Şifre süresi dolmuşsa zorunlu değiştirme ekranı"""
        win=tk.Toplevel(self.root); win.title("Şifre Değiştir"); win.geometry("420x320")
        win.configure(bg=C_WHITE); win.grab_set(); win.transient(self.root)
        win.protocol("WM_DELETE_WINDOW",lambda:None)
        tk.Label(win,text="🔒 Zorunlu Şifre Değişikliği",font=("Segoe UI",14,"bold"),fg=C_DANGER,bg=C_WHITE).pack(pady=16)
        tk.Label(win,text="Şifrenizin süresi doldu. Yeni şifre belirleyin.",font=("Segoe UI",10),fg="#555",bg=C_WHITE).pack(pady=(0,12))
        al={}
        for l,k in [("Mevcut Şifre:","eski"),("Yeni Şifre:","yeni"),("Tekrar:","tekrar")]:
            tk.Label(win,text=l,bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=30,pady=(6,0))
            e=ttk.Entry(win,show="●",width=30); e.pack(padx=30); al[k]=e
        def _k():
            eski=al["eski"].get(); yeni=al["yeni"].get(); tekrar=al["tekrar"].get()
            if yeni!=tekrar: messagebox.showerror("Hata","Şifreler eşleşmiyor.",parent=win); return
            ok,msg=strong_pw(yeni)
            if not ok: messagebox.showerror("Hata",msg,parent=win); return
            with sqlite3.connect(DB_PATH) as c:
                db_s=c.execute("SELECT sifre FROM Kullanicilar WHERE k_adi=?",(self.u_id,)).fetchone()[0]
            if not verify_pw(eski,db_s): messagebox.showerror("Hata","Mevcut şifre yanlış.",parent=win); return
            with sqlite3.connect(DB_PATH) as c:
                c.execute("UPDATE Kullanicilar SET sifre=?,sifre_tarih=? WHERE k_adi=?",(hash_pw(yeni),datetime.now().strftime("%Y-%m-%d"),self.u_id))
            db_log(self.u_id,"Şifre Değiştir","Zorunlu - süre doldu"); win.destroy(); self._build_app()
        MBtn(win,"🔒 Şifreyi Değiştir ve Giriş Yap",command=_k,color=C_DANGER,width=30).pack(pady=16)

    def _login_fail(self,k_adi,fail_c):
        """Login başarısız — ayrı metod olarak"""
        fail_c+=1; lockout=None
        kalan_hak=MAX_LOGIN_FAIL-fail_c
        if fail_c>=MAX_LOGIN_FAIL:
            lockout=(datetime.now()+timedelta(minutes=LOCKOUT_MIN)).isoformat()
            self.lbl_hata.config(text=f"🔒 {MAX_LOGIN_FAIL} hatalı giriş. {LOCKOUT_MIN} dk kilitlendi.")
        else:
            self.lbl_hata.config(text=f"❌ Hatalı şifre. ({fail_c}/{MAX_LOGIN_FAIL}) — {kalan_hak} hakkınız kaldı.")
        with sqlite3.connect(DB_PATH) as c: c.execute("UPDATE Kullanicilar SET fail_count=?,lockout_ts=? WHERE k_adi=?",(fail_c,lockout,k_adi))
        db_log(k_adi,"Başarısız",f"{fail_c}/{MAX_LOGIN_FAIL}"); self.e_p.delete(0,tk.END)

    def _riza_goster(self):
        """Açık rıza metni göster ve onay al"""
        win=tk.Toplevel(self.root); win.title("Kullanım Koşulları"); win.geometry("600x520")
        win.configure(bg=C_WHITE); win.grab_set(); win.transient(self.root)
        tk.Label(win,text="📋 Kullanım Koşulları ve Açık Rıza",font=("Segoe UI",14,"bold"),fg="#1E5631",bg=C_WHITE).pack(pady=14)
        txt=scrolledtext.ScrolledText(win,font=("Segoe UI",10),wrap="word",bg="#FAFAFA",relief="flat",
            highlightbackground="#D0DDD8",highlightthickness=1,padx=14,pady=12)
        txt.pack(fill="both",expand=True,padx=16,pady=8)
        txt.insert("1.0",RIZA_METNI); txt.config(state="disabled")
        var=tk.BooleanVar()
        ttk.Checkbutton(win,text="Yukarıdaki koşulları okudum ve kabul ediyorum.",variable=var).pack(pady=8)
        def _onayla():
            if not var.get():
                messagebox.showwarning("Uyarı","Koşulları kabul etmeniz gerekiyor.",parent=win); return
            with sqlite3.connect(DB_PATH) as c:
                c.execute("UPDATE Kullanicilar SET riza_onay=1 WHERE k_adi=?",(self.u_id,))
            db_log(self.u_id,"Rıza Onay","Kabul edildi"); win.destroy(); self._build_app()
        def _reddet():
            self.u_id=None; win.destroy(); self._build_login()
        bf=tk.Frame(win,bg=C_WHITE); bf.pack(pady=10)
        MBtn(bf,"✅ Kabul Ediyorum",command=_onayla,color="#2D8C55",width=20).pack(side="left",padx=8)
        MBtn(bf,"❌ Reddet",command=_reddet,color=C_DANGER,width=14).pack(side="left",padx=8)

    def _oturum_bitti(self):
        self.root.after(0,lambda:(messagebox.showwarning("Oturum Sona Erdi","15 dakika hareketsizlik nedeniyle oturum kapatıldı."),self._cikis()))

    def _cikis(self):
        # Çıkışta otomatik yedek al
        if self.yedekci:
            try: self.yedekci.al(otomatik=True)
            except Exception: pass
        if self.session: self.session.stop(); self.session=None
        self.u_id=self.u_yetki=self.u_ad=self.u_unvan=None
        self._build_login()

    def _clear(self):
        for w in self.root.winfo_children(): w.destroy()

    # ═══ SOL MENÜ (SIDEBAR) YAPISI ═══════════════════════════════════════════
    def _build_app(self):
        self._clear(); self.root.state("zoomed"); self.root.configure(bg=self.gc("bg"))
        self.root.bind_all("<Motion>",lambda e:self.session and self.session.ping())
        self.root.bind_all("<KeyPress>",lambda e:self.session and self.session.ping())

        # Ana container
        main=tk.Frame(self.root,bg=self.gc("bg")); main.pack(fill="both",expand=True)

        # ── Header ──
        h=tk.Frame(main,bg=self.gc("pri"),height=52); h.pack(fill="x",side="top"); h.pack_propagate(False)
        tk.Label(h,text="🌿 MİRAS",font=("Segoe UI",18,"bold"),fg=C_WHITE,bg=self.gc("pri")).pack(side="left",padx=18,pady=8)
        tk.Label(h,text=VERSIYON,font=("Segoe UI",8),fg="#8FCF9F",bg=self.gc("pri")).pack(side="left")
        right=tk.Frame(h,bg=self.gc("pri")); right.pack(side="right",padx=14)
        tk.Label(right,text=f"👤 {self.u_ad}  |  {self.u_unvan}  |  {self.u_yetki}",
            font=("Segoe UI",9),fg=C_WHITE,bg=self.gc("pri")).pack(side="left",padx=10)
        MBtn(right,"⏏ Çıkış",command=self._cikis,color=C_DANGER,pady=3,padx=8).pack(side="left")

        # ── Body: Sidebar + Content ──
        body=tk.Frame(main,bg=self.gc("bg")); body.pack(fill="both",expand=True)

        # Sidebar
        self._sidebar=tk.Frame(body,bg=self.gc("side"),width=200); self._sidebar.pack(side="left",fill="y")
        self._sidebar.pack_propagate(False)

        # Sidebar menü butonları
        self._menu_btns={}
        items=list(self.MENU_ITEMS_UZMAN)
        if self.u_yetki=="Admin": items+=self.MENU_ITEMS_ADMIN

        # Scrollable sidebar with visible scrollbar
        s_canvas=tk.Canvas(self._sidebar,bg=self.gc("side"),highlightthickness=0,width=180)
        s_scroll=ttk.Scrollbar(self._sidebar,orient="vertical",command=s_canvas.yview)
        s_canvas.configure(yscrollcommand=s_scroll.set)
        s_scroll.pack(side="right",fill="y")
        s_canvas.pack(side="left",fill="both",expand=True)
        s_inner=tk.Frame(s_canvas,bg=self.gc("side"))
        s_canvas.create_window((0,0),window=s_inner,anchor="nw",width=180)
        s_inner.bind("<Configure>",lambda e:s_canvas.configure(scrollregion=s_canvas.bbox("all")))
        # Mouse wheel scroll - platform uyumlu
        def _sidebar_scroll(event):
            if event.delta: s_canvas.yview_scroll(int(-1*(event.delta/120)),"units")
            elif event.num==4: s_canvas.yview_scroll(-1,"units")
            elif event.num==5: s_canvas.yview_scroll(1,"units")
        s_canvas.bind("<MouseWheel>",_sidebar_scroll)
        s_canvas.bind("<Button-4>",_sidebar_scroll)
        s_canvas.bind("<Button-5>",_sidebar_scroll)
        s_inner.bind("<MouseWheel>",_sidebar_scroll)
        s_inner.bind("<Button-4>",_sidebar_scroll)
        s_inner.bind("<Button-5>",_sidebar_scroll)

        for ikon,baslik,key in items:
            if self.u_yetki=="İzleyici" and key in self.MENU_ITEMS_IZLEYICI_HIDE:
                continue
            btn=tk.Button(s_inner,text=f" {ikon} {baslik}",anchor="w",
                font=("Segoe UI",9),fg=self.gc("side_text"),bg=self.gc("side"),
                relief="flat",cursor="hand2",padx=10,pady=6,
                activebackground=self.gc("side_hover"),activeforeground=C_WHITE,
                command=lambda k=key:self._menu_click(k))
            btn.pack(fill="x",pady=0)
            btn.bind("<Enter>",lambda e,b=btn:b.config(bg=self.gc("side_hover"),fg=C_WHITE))
            btn.bind("<Leave>",lambda e,b=btn,k=key:b.config(
                bg=self.gc("pri") if self._active_menu==k else self.gc("side"),
                fg=C_WHITE if self._active_menu==k else self.gc("side_text")))
            btn.bind("<MouseWheel>",_sidebar_scroll)
            btn.bind("<Button-4>",_sidebar_scroll)
            btn.bind("<Button-5>",_sidebar_scroll)
            self._menu_btns[key]=btn

        # Content area
        content_wrapper=tk.Frame(body,bg=self.gc("bg")); content_wrapper.pack(side="left",fill="both",expand=True)
        self._content_frame=tk.Frame(content_wrapper,bg=self.gc("bg")); self._content_frame.pack(fill="both",expand=True)

        # Footer
        ft=tk.Frame(content_wrapper,bg="#E8E8E8",height=28); ft.pack(fill="x",side="bottom"); ft.pack_propagate(False)
        tk.Label(ft,text=HAKLAR,font=("Segoe UI",8),fg=C_FOOTER,bg="#E8E8E8").pack(side="left",padx=12)
        self.status=StatusBar(ft); self.status.pack(side="right",fill="y")

        # İlk sayfa: Dashboard
        self._menu_click("dash")
        self._duyuru()
        self.root.after(400,self._hosgeldin)
        self.root.after(1600,self._ajanda_uyar)
        self.root.after(2800,self._yedek_uyari)
        self._klavye_kisayollari()
        self.root.after(5000,self._guncelleme_kontrol)

    def _guncelleme_kontrol(self):
        """Madde 21: GitHub'dan güncelleme kontrolü"""
        def _kontrol():
            try:
                import urllib.request, json as jn
                req=urllib.request.Request(f"https://api.github.com/repos/{GITHUB_REPO}/releases/latest",
                    headers={"User-Agent":"MirasEnterprise"})
                with urllib.request.urlopen(req,timeout=5) as resp:
                    data=jn.loads(resp.read())
                    son_versiyon=data.get("tag_name","")
                    if son_versiyon and son_versiyon!=VERSIYON:
                        self.root.after(0,lambda:self._guncelleme_bildir(son_versiyon,data.get("html_url","")))
            except Exception: pass
        threading.Thread(target=_kontrol,daemon=True).start()

    def _guncelleme_bildir(self,yeni_v,url):
        if messagebox.askyesno("🔄 Güncelleme Mevcut",
            f"Yeni sürüm: {yeni_v}\nMevcut: {VERSIYON}\n\nİndirmek ister misiniz?"):
            import webbrowser; webbrowser.open(url)

    def _menu_click(self,key):
        """Sidebar menü tıklama — ilgili sayfayı yükle"""
        # Aktif menü stilini güncelle
        for k,btn in self._menu_btns.items():
            if k==key:
                btn.config(bg=self.gc("pri"),fg=C_WHITE)
            else:
                btn.config(bg=self.gc("side"),fg=self.gc("side_text"))
        self._active_menu=key

        # İçerik alanını temizle
        for w in self._content_frame.winfo_children(): w.destroy()

        # Madde 12: Modül bazlı görev kontrolü
        # Kapasite raporu herkese açık, Admin her yere erişir
        MODUL_GOREV_MAP={"ihale":"İhale","ihale_yer":"İhale","sikayet":"Şikayet",
            "tahsis":"Tahsis","ceza":"Ceza","islah":"Islah"}
        if key in MODUL_GOREV_MAP and self.u_yetki!="Admin":
            try:
                with sqlite3.connect(DB_PATH) as c:
                    gorevler=c.execute("SELECT gorevler FROM Kullanicilar WHERE k_adi=?",(self.u_id,)).fetchone()
                gorev_str=gorevler[0] if gorevler and gorevler[0] else "*"
                if gorev_str!="*" and MODUL_GOREV_MAP[key] not in gorev_str:
                    tk.Label(self._content_frame,text=f"🔒 Bu modüle erişim yetkiniz bulunmuyor.\n\nGörevleriniz: {gorev_str}\n\nAdmin'den görev ataması isteyin.",
                        font=("Segoe UI",13),fg=C_DANGER,bg=self.gc("bg"),justify="center").pack(expand=True)
                    return
            except Exception: pass

        # İlgili sayfayı yükle
        builder=getattr(self,f"_t_{key}",None)
        if builder:
            try: builder(self._content_frame)
            except Exception as e:
                logging.error(f"Sayfa yükleme hatası ({key}): {e}")
                tk.Label(self._content_frame,text=f"❌ Sayfa yüklenirken hata oluştu:\n{e}",
                    font=("Segoe UI",12),fg=C_DANGER,bg=self.gc("bg")).pack(expand=True)

    def _hosgeldin(self):
        win=tk.Toplevel(self.root); win.overrideredirect(True); win.configure(bg=self.gc("pri"))
        w,h=440,110; rx=self.root.winfo_x()+(self.root.winfo_width()-w)//2; ry=self.root.winfo_y()+52
        win.geometry(f"{w}x{h}+{rx}+{ry}"); win.attributes("-alpha",0.93)
        tk.Label(win,text=f"Hoş Geldiniz, {self.u_ad}! 🌿",font=("Segoe UI",16,"bold"),fg=C_WHITE,bg=self.gc("pri")).pack(pady=(18,4))
        tk.Label(win,text=self.u_unvan,font=("Segoe UI",11),fg="#A9DFBF",bg=self.gc("pri")).pack()
        self.root.after(3200,win.destroy)

    def _duyuru(self):
        try:
            with sqlite3.connect(DB_PATH) as c:
                r=c.execute("SELECT mesaj,tarih,gonderen FROM Duyurular ORDER BY id DESC LIMIT 1").fetchone()
            if r: messagebox.showinfo("📢 MİRAS Duyurusu",f"{r[0]}\n\n— {r[2]} ({r[1]})")
        except Exception: pass

    def _tv(self,parent,cols,h=14):
        frame=tk.Frame(parent,bg=self.gc("bg")); frame.pack(fill="both",expand=True)
        tv=ttk.Treeview(frame,columns=[c[0] for c in cols],show="headings",height=h)
        for col,w,bas in cols:
            tv.heading(col,text=bas,command=lambda _c=col,_t=tv:self._sort(_t,_c))
            tv.column(col,width=w,anchor="center")
        sb=ttk.Scrollbar(frame,orient="vertical",command=tv.yview)
        tv.configure(yscrollcommand=sb.set)
        tv.pack(side="left",fill="both",expand=True); sb.pack(side="left",fill="y")
        return tv

    @staticmethod
    def _sort(tree,col):
        data=[(tree.set(k,col),k) for k in tree.get_children("")]
        try: data.sort(key=lambda t:float(t[0]))
        except ValueError: data.sort()
        for i,(_,k) in enumerate(data): tree.move(k,"",i)

    def _son_islem_kaydet(self,modul,kayit_adi):
        try:
            with sqlite3.connect(DB_PATH) as c:
                c.execute("INSERT INTO Son_Islemler(k_adi,modul,kayit_adi,tarih)VALUES(?,?,?,?)",
                    (self.u_id,modul,kayit_adi,datetime.now().strftime("%d.%m.%Y %H:%M")))
                c.execute("DELETE FROM Son_Islemler WHERE id NOT IN (SELECT id FROM Son_Islemler WHERE k_adi=? ORDER BY id DESC LIMIT 20)",(self.u_id,))
        except Exception: pass

    def _guvenli_sil(self, tablo, kayit_id, kayit_ozet, sil_sql, sil_params, yenile_fn):
        """Madde 2: Admin direkt siler, diğerleri talep oluşturur"""
        if self.u_yetki=="Admin":
            if messagebox.askyesno("Onay",f"'{kayit_ozet}' silinsin mi?"):
                with sqlite3.connect(DB_PATH) as c: c.execute(sil_sql, sil_params)
                db_log(self.u_id,"Silme",f"{tablo}:{kayit_ozet}"); yenile_fn()
        else:
            if messagebox.askyesno("Silme Talebi",f"'{kayit_ozet}' için Admin'e silme talebi gönderilsin mi?"):
                with sqlite3.connect(DB_PATH) as c:
                    c.execute("INSERT INTO Silme_Talepleri(tablo,kayit_id,kayit_ozet,talep_eden,tarih)VALUES(?,?,?,?,?)",
                        (tablo,kayit_id,kayit_ozet,self.u_ad,datetime.now().strftime("%d.%m.%Y %H:%M")))
                db_log(self.u_id,"Silme Talebi",f"{tablo}:{kayit_ozet}")
                messagebox.showinfo("Tamam","Silme talebiniz Admin'e iletildi.")

    def _yedek_uyari(self):
        try:
            yedekler=sorted(BACKUP_DIR.glob("*.db"),key=lambda p:p.stat().st_mtime)
            if not yedekler:
                messagebox.showwarning("💾 Yedek Uyarısı","Hiç yedek alınmamış!\nAyarlar → Yedekleme sekmesinden yedek alın."); return
            son=datetime.fromtimestamp(yedekler[-1].stat().st_mtime)
            fark=(datetime.now()-son).days
            if fark>=7: messagebox.showwarning("💾 Yedek Uyarısı",f"Son yedekten {fark} gün geçti!\nAyarlar → Yedekleme sekmesinden yedek alın.")
        except Exception as e: logging.error(f"yedek_uyari:{e}")

    def _ajanda_uyar(self):
        try:
            bugun=datetime.now().strftime("%d.%m.%Y")
            with sqlite3.connect(DB_PATH) as c:
                rows=c.execute("SELECT baslik,tarih,sure FROM Ajanda WHERE k_adi=? AND durum='Bekliyor' AND tarih<=? ORDER BY tarih,sure",(self.u_id,bugun)).fetchall()
            if rows:
                liste="\n".join(f"• {r[1]} {r[2]} — {r[0]}" for r in rows[:5])
                mesaj=f"{len(rows)} bekleyen etkinlik var:\n\n{liste}"
                if len(rows)>5: mesaj+=f"\n... ve {len(rows)-5} daha"
                messagebox.showinfo("📅 Ajanda Hatırlatıcı",mesaj)
        except Exception as e: logging.error(f"ajanda_uyar:{e}")
        # Tahsis süre uyarısı
        self._tahsis_sure_uyar()

    def _tahsis_sure_uyar(self):
        """Madde 11: Tahsis süre takibi — son 7 gün uyarı"""
        try:
            with sqlite3.connect(DB_PATH) as c:
                rows=c.execute("SELECT id,koy,sure_bitis,sure_tipi FROM Tahsisler WHERE durum='Devam Ediyor' AND sure_bitis IS NOT NULL AND sure_bitis!=''").fetchall()
            uyarilar=[]
            for r in rows:
                try:
                    bitis=datetime.strptime(r[2],"%d.%m.%Y")
                    kalan=(bitis-datetime.now()).days
                    if kalan<=7 and kalan>=0:
                        uyarilar.append(f"⚠️ {r[1]} — {r[3]}: {kalan} gün kaldı! ({r[2]})")
                    elif kalan<0:
                        uyarilar.append(f"🔴 {r[1]} — {r[3]}: Süre {abs(kalan)} gün önce doldu!")
                except Exception: pass
            if uyarilar:
                messagebox.showwarning("📋 Tahsis Süre Uyarısı","\n".join(uyarilar))
        except Exception as e: logging.error(f"tahsis_sure:{e}")

    def _klavye_kisayollari(self):
        self.root.bind_all("<F5>",lambda e:self._menu_click("dash"))
        self.root.bind_all("<Control-n>",lambda e:self._menu_click("kapasite"))
        self.root.bind_all("<Control-h>",lambda e:self._koy_gecmis())
        self.root.bind_all("<F1>",lambda e:messagebox.showinfo("⌨️ Klavye Kısayolları",
            "F5        → Dashboard Yenile\nCtrl+N    → Kapasite\nCtrl+H    → Köy Geçmişi\nF1        → Bu pencere"))

    # ═══ 1. DASHBOARD ════════════════════════════════════════════════════════
    def _t_dash(self,p):
        bg=self.gc("bg"); f=tk.Frame(p,bg=bg); f.pack(fill="both",expand=True,padx=20,pady=16)
        tk.Label(f,text="📊 Genel Durum Paneli",font=("Segoe UI",16,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")

        # Bildirim çubuğu
        self._bildirim_frame=tk.Frame(f,bg="#FEF9E7",highlightbackground="#F0C040",highlightthickness=1)
        self._bildirim_frame.pack(fill="x",pady=(8,4))
        self.lbl_bildirim=tk.Label(self._bildirim_frame,text="📢 Bildirimler yükleniyor...",bg="#FEF9E7",
            font=("Segoe UI",9),fg="#7D6608",padx=10,pady=6)
        self.lbl_bildirim.pack(anchor="w")

        # Hızlı erişim butonları
        hf=tk.Frame(f,bg=bg); hf.pack(fill="x",pady=(4,8))
        tk.Label(hf,text="⚡ Hızlı Erişim:",bg=bg,font=("Segoe UI",9,"bold"),fg="#666").pack(side="left",padx=(0,8))
        for txt,key,renk in [("🧮 Yeni Rapor","kapasite",self.gc("acc")),("⚖️ Yeni İhale","ihale",C_INFO),
                              ("🚨 Yeni Şikayet","sikayet",C_WARN),("💰 Yeni Ceza","ceza",C_DANGER)]:
            if self.u_yetki!="İzleyici" or key not in self.MENU_ITEMS_IZLEYICI_HIDE:
                tk.Button(hf,text=txt,command=lambda k=key:self._menu_click(k),bg=renk,fg=C_WHITE,
                    font=("Segoe UI",8,"bold"),relief="flat",cursor="hand2",padx=10,pady=4).pack(side="left",padx=4)

        kf=tk.Frame(f,bg=bg); kf.pack(fill="x",pady=8)
        self.dk1=StatKart(kf,"Sistemdeki Köy","—",self.gc("pri"),"🏘️")
        self.dk2=StatKart(kf,"Düzenlenen Rapor","—",C_INFO,"📄")
        self.dk3=StatKart(kf,"Bekleyen Islah","—",C_DANGER,"🌱")
        self.dk4=StatKart(kf,"Aktif İhale","—",C_WARN,"⚖️")
        self.dk5=StatKart(kf,"Bekleyen Tahsis","—","#8E44AD","📋")
        for k in (self.dk1,self.dk2,self.dk3,self.dk4,self.dk5): k.pack(side="left",fill="both",expand=True,padx=4)
        lf=tk.LabelFrame(f,text="  Son İşlem Kayıtları  ",bg=C_WHITE,font=("Segoe UI",10,"bold"),fg=self.gc("pri"))
        lf.pack(fill="both",expand=True,pady=(0,8))
        self.tv_log=self._tv(lf,[("t",150,"Zaman"),("k",120,"Personel"),("i",140,"İşlem"),("d",500,"Detay")],10)
        sf=tk.LabelFrame(f,text="  ⏱ Son İşlemlerim  ",bg=C_WHITE,font=("Segoe UI",10,"bold"),fg=self.gc("pri"))
        sf.pack(fill="x",pady=(0,8))
        self.tv_son=self._tv(sf,[("modul",120,"Modül"),("kayit",280,"Kayıt"),("tarih",160,"Tarih")],4)
        MBtn(f,"🔄  Paneli Yenile",command=self._dash_yenile,color=self.gc("acc"),width=18).pack(pady=6)
        self._dash_yenile()

    def _dash_yenile(self):
        if not DB_PATH: return
        try:
            with sqlite3.connect(DB_PATH) as c:
                self.dk1.set(c.execute("SELECT COUNT(*) FROM Mera_Varligi").fetchone()[0])
                self.dk2.set(c.execute("SELECT COUNT(*) FROM Rapor_Gecmisi").fetchone()[0])
                self.dk3.set(c.execute("SELECT COUNT(*) FROM Islah_Amenajman WHERE durum='Bekliyor'").fetchone()[0])
                ihale_c=c.execute("SELECT COUNT(*) FROM Ihaleler WHERE durum NOT IN('Tamamlandı','İptal Edildi')").fetchone()[0]
                self.dk4.set(ihale_c)
                tahsis_c=c.execute("SELECT COUNT(*) FROM Tahsisler WHERE durum='Devam Ediyor'").fetchone()[0]
                self.dk5.set(tahsis_c)
                # Bildirimler
                sik_c=c.execute("SELECT COUNT(*) FROM Sikayetler WHERE durum='Yeni'").fetchone()[0]
                bugun=datetime.now().strftime("%d.%m.%Y")
                aj_c=c.execute("SELECT COUNT(*) FROM Ajanda WHERE k_adi=? AND durum='Bekliyor' AND tarih<=?",(self.u_id,bugun)).fetchone()[0]
                bildirimler=[]
                if ihale_c>0: bildirimler.append(f"⚖️ {ihale_c} aktif ihale")
                if sik_c>0: bildirimler.append(f"🚨 {sik_c} yeni şikayet")
                if aj_c>0: bildirimler.append(f"📅 {aj_c} bekleyen etkinlik")
                if tahsis_c>0: bildirimler.append(f"📋 {tahsis_c} devam eden tahsis")
                if bildirimler:
                    self.lbl_bildirim.config(text="📢 " + "  |  ".join(bildirimler))
                else:
                    self.lbl_bildirim.config(text="✅ Bekleyen bildirim yok.")

                self.tv_log.delete(*self.tv_log.get_children())
                log_q="SELECT tarih,kul,islem,detay FROM Loglar ORDER BY id DESC LIMIT 30"
                if self.u_yetki!="Admin":
                    log_q=f"SELECT tarih,kul,islem,detay FROM Loglar WHERE kul='{self.u_id}' ORDER BY id DESC LIMIT 30"
                for r in c.execute(log_q).fetchall():
                    self.tv_log.insert("","end",values=r)
                self.tv_son.delete(*self.tv_son.get_children())
                for r in c.execute("SELECT modul,kayit_adi,tarih FROM Son_Islemler WHERE k_adi=? ORDER BY id DESC LIMIT 20",(self.u_id,)).fetchall():
                    self.tv_son.insert("","end",values=r)
            if self.status: self.status.set(f"✔  Panel güncellendi — {datetime.now().strftime('%H:%M:%S')}")
        except Exception as e: logging.error(f"dash:{e}")

    # ═══ 2. KAPASİTE ════════════════════════════════════════════════════════
    def _t_kapasite(self,p):
        bg=self.gc("bg")
        cw=tk.Canvas(p,bg=bg,highlightthickness=0); vsb=ttk.Scrollbar(p,orient="vertical",command=cw.yview)
        cw.configure(yscrollcommand=vsb.set); vsb.pack(side="right",fill="y"); cw.pack(side="left",fill="both",expand=True)
        ana=tk.Frame(cw,bg=bg); wid=cw.create_window((0,0),window=ana,anchor="nw")
        cw.bind("<Configure>",lambda e:cw.itemconfig(wid,width=e.width))
        ana.bind("<Configure>",lambda e:cw.configure(scrollregion=cw.bbox("all")))
        px=20
        top=tk.LabelFrame(ana,text="  Mera Verisi & Rapor Hesaplama  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=14,pady=10)
        top.pack(fill="x",padx=px,pady=(14,6))
        r1=tk.Frame(top,bg=bg); r1.pack(fill="x",pady=4)
        if self.u_yetki!="İzleyici":
            MBtn(r1,"📥  Excel'den Mera Verisi Yükle",command=self._excel_mera,color=C_INFO,width=26).pack(side="left",padx=(0,14))
            MBtn(r1,"➕  Manuel Köy Ekle/Güncelle",command=self._koy_manuel,color="#8E44AD",width=26).pack(side="left")
        r2=tk.Frame(top,bg=bg); r2.pack(fill="x",pady=8)
        for lbl,attr,tip,w in [("Köy Seç:","kap_koy","combo",22),("Aktif İşletme:","kap_aktif","entry",8),
               ("Talep Eden:","kap_talep","entry",20),("TC No:","kap_tc","entry",14)]:
            tk.Label(r2,text=lbl,bg=bg,font=("Segoe UI",10)).pack(side="left",padx=(0,4))
            w_obj=ttk.Combobox(r2,width=w,state="readonly") if tip=="combo" else ttk.Entry(r2,width=w)
            w_obj.pack(side="left",padx=(0,14)); setattr(self,attr,w_obj)
        r3=tk.Frame(top,bg=bg); r3.pack(fill="x",pady=4)
        tk.Label(r3,text="Ek Açıklama:",bg=bg,font=("Segoe UI",10)).pack(side="left",padx=(0,6))
        self.kap_acik=ttk.Entry(r3,width=50); self.kap_acik.pack(side="left")
        if self.u_yetki!="İzleyici":
            MBtn(top,"🧮  Hesapla ve Kurumsal PDF Üret",command=self._hesapla_pdf,color=self.gc("acc")).pack(anchor="w",pady=10)
        kv=tk.LabelFrame(ana,text="  Kayıtlı Mera Verileri  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"))
        kv.pack(fill="x",padx=px,pady=(0,8))
        self.tv_mera=self._tv(kv,[("koy",130,"Köy"),("ilce",90,"İlçe"),("alan",90,"Alan (da)"),("yem",110,"Yem (kg/da)"),("ok",110,"OK (BBHB)")],6)
        rf=tk.LabelFrame(ana,text="  Düzenlenen Raporlar  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"))
        rf.pack(fill="both",expand=True,padx=px,pady=(0,14))
        AramaFrame(rf,self._filtre_rap,bg=bg).pack(anchor="w",padx=10,pady=6)
        self.tv_rap=self._tv(rf,[("no",80,"Rapor No"),("koy",110,"Köy"),("talep",150,"Talep Eden"),
            ("tc",110,"TC No"),("tarih",100,"Tarih"),("per",120,"Düzenleyen")],8)
        bf=tk.Frame(rf,bg=bg); bf.pack(pady=6)
        MBtn(bf,"🔄 Yenile",command=self._yenile_kap,color=C_INFO,width=14).pack(side="left",padx=6)
        if self.u_yetki!="İzleyici":
            MBtn(bf,"🗑 Seçili Raporu Sil",command=self._sil_rap,color=C_DANGER,width=18).pack(side="left",padx=6)
        self._yenile_kap()

    def _yenile_kap(self):
        if not DB_PATH: return
        try:
            with sqlite3.connect(DB_PATH) as c:
                koyler=[r[0] for r in c.execute("SELECT koy FROM Mera_Varligi ORDER BY koy")]
                self.kap_koy["values"]=koyler
                self.tv_mera.delete(*self.tv_mera.get_children())
                for r in c.execute("SELECT koy,ilce,alan,yem FROM Mera_Varligi ORDER BY koy"):
                    ok=(r[2]*r[3])/PAYDA
                    self.tv_mera.insert("","end",values=(r[0],r[1],f"{r[2]:.0f}",f"{r[3]:.0f}",f"{ok:.2f}"))
                self._all_rap=c.execute("SELECT rapor_no,koy,talep_eden,tc,islem_tarihi,duzenleyen FROM Rapor_Gecmisi ORDER BY rowid DESC").fetchall()
            self._filtre_rap("")
        except Exception as e: logging.error(f"yenile_kap:{e}")

    def _filtre_rap(self,a):
        self.tv_rap.delete(*self.tv_rap.get_children()); a=a.lower()
        for r in getattr(self,"_all_rap",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            self.tv_rap.insert("","end",values=r)

    def _excel_mera(self):
        yol=filedialog.askopenfilename(filetypes=[("Excel","*.xlsx *.xls")])
        if not yol: return
        try:
            df=pd.read_excel(yol)
            if len(df.columns)<4: messagebox.showerror("Hata","4 sütun gerekli: koy|ilce|alan_da|yem_kg_da"); return
            with sqlite3.connect(DB_PATH) as c:
                kayit=0
                for _,r in df.iterrows():
                    v=r.tolist()
                    c.execute("INSERT OR REPLACE INTO Mera_Varligi(koy,ilce,alan,yem)VALUES(?,?,?,?)",(str(v[0]),str(v[1]),float(v[2]),float(v[3])))
                    kayit+=1
            db_log(self.u_id,"Excel Yükleme",f"{kayit} köy"); self._yenile_kap()
            messagebox.showinfo("Tamam",f"{kayit} köy verisi sisteme yüklendi.")
        except Exception as e: messagebox.showerror("Excel Hatası",str(e))

    def _koy_manuel(self):
        win=tk.Toplevel(self.root); win.title("Köy Mera Verisi"); win.geometry("380x300"); win.configure(bg=C_WHITE); win.grab_set()
        tk.Label(win,text="Köy Mera Verisi Ekle/Güncelle",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=14)
        al={}
        for lbl in ["Köy Adı:","İlçe:","Mera Alanı (da):","Yem (kg/da):"]:
            tk.Label(win,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=24,pady=(6,0))
            e=ttk.Entry(win,width=32); e.pack(padx=24); al[lbl]=e
        def _k():
            try:
                koy=al["Köy Adı:"].get().strip(); ilce=al["İlçe:"].get().strip()
                alan=float(al["Mera Alanı (da):"].get()); yem=float(al["Yem (kg/da):"].get())
                if not koy or not ilce: raise ValueError("Köy ve İlçe zorunlu.")
                with sqlite3.connect(DB_PATH) as c: c.execute("INSERT OR REPLACE INTO Mera_Varligi(koy,ilce,alan,yem)VALUES(?,?,?,?)",(koy,ilce,alan,yem))
                db_log(self.u_id,"Mera Verisi",f"{koy}"); self._yenile_kap(); messagebox.showinfo("Tamam",f"'{koy}' kaydedildi.",parent=win); win.destroy()
            except ValueError as e: messagebox.showerror("Hata",str(e),parent=win)
        MBtn(win,"Kaydet",command=_k,width=20).pack(pady=16)

    def _hesapla_pdf(self):
        koy=self.kap_koy.get(); talep=self.kap_talep.get().strip(); tc=self.kap_tc.get().strip()
        try: aktif=int(self.kap_aktif.get())
        except ValueError: messagebox.showerror("Hata","Aktif işletme sayısı sayı olmalı."); return
        if not all([koy,talep,tc]): messagebox.showerror("Eksik","Köy, talep eden ve TC zorunlu."); return
        if aktif<=0: messagebox.showerror("Hata","Aktif işletme > 0 olmalı."); return
        # TC doğrulama
        if not tc_dogrula(tc):
            if not messagebox.askyesno("TC Uyarı",f"'{tc}' geçerli bir TC numarası değil.\nDevam etmek istiyor musunuz?"): return
        try:
            with sqlite3.connect(DB_PATH) as c:
                mera=c.execute("SELECT alan,yem,ilce FROM Mera_Varligi WHERE koy=?",(koy,)).fetchone()
        except Exception as e: messagebox.showerror("DB Hatası",str(e)); return
        if not mera: messagebox.showwarning("Eksik",f"'{koy}' mera verisi yok."); return
        alan,yem,ilce=mera; ok=(alan*yem)/PAYDA
        try:
            with sqlite3.connect(DB_PATH) as c:
                yil=datetime.now().year
                son=c.execute("SELECT rapor_no FROM Rapor_Gecmisi WHERE rapor_no LIKE ? ORDER BY rowid DESC LIMIT 1",(f"{yil}-%",)).fetchone()
                sira=(int(son[0].split("-")[1])+1) if son else 1
                rap_no=f"{yil}-{sira:03d}"
                c.execute("INSERT INTO Rapor_Gecmisi(rapor_no,koy,talep_eden,tc,islem_tarihi,duzenleyen,aciklama)VALUES(?,?,?,?,?,?,?)",
                    (rap_no,koy,talep,tc,datetime.now().strftime("%d.%m.%Y"),self.u_ad,self.kap_acik.get()))
        except Exception as e: messagebox.showerror("DB Hatası",str(e)); return
        db_log(self.u_id,"Kapasite Raporu",f"{koy} {rap_no} OK:{ok:.2f}BBHB")
        self._son_islem_kaydet("Rapor",f"{koy} — {rap_no}")
        if not PDF_OK:
            messagebox.showwarning("PDF Yok",f"reportlab kurulu değil.\n\n{koy}: {ok:.2f} BBHB")
            self._yenile_kap(); return
        dosya=filedialog.asksaveasfilename(defaultextension=".pdf",initialfile=f"Rapor_{koy}_{rap_no}.pdf",filetypes=[("PDF","*.pdf")])
        if not dosya: self._yenile_kap(); return
        try:
            uret_pdf(dosya,{"ilce":ilce,"koy":koy,"alan":alan,"yem":yem,"aktif":aktif,"ok_bbhb":ok,"rapor_no":rap_no,"talep_eden":talep,"tc":tc,"aciklama":self.kap_acik.get()},
                {"ad":self.u_ad,"unvan":self.u_unvan or "Ziraat Mühendisi","sube_mudur":self.sube_mudur})
            self._yenile_kap()
            messagebox.showinfo("Rapor Hazır",f"✅ Rapor No: {rap_no}\n{koy}\nOtlatma Kapasitesi: {ok:.2f} BBHB\nİşletme Başı: {ok/aktif:.2f} BBHB\n\nPDF: {dosya}")
        except Exception as e: messagebox.showerror("PDF Hatası",str(e)); logging.error(f"pdf:{e}")

    def _sil_rap(self):
        sel=self.tv_rap.selection()
        if not sel: messagebox.showwarning("Seçim","Raporu seçin."); return
        no=self.tv_rap.item(sel[0])["values"][0]
        self._guvenli_sil("Rapor_Gecmisi",no,f"Rapor {no}","DELETE FROM Rapor_Gecmisi WHERE rapor_no=?",(no,),self._yenile_kap)

    # ═══ 3. İHALE TAKİP ═════════════════════════════════════════════════════
    def _t_ihale(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2); bf2=ttk.Frame(nb2)
        nb2.add(lf,text="  📋 İhale Listesi  "); nb2.add(nf,text="  ➕ Yeni Başvuru  "); nb2.add(bf2,text="  💰 Bedel Hesaplama  ")
        # Liste
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,self._filtre_ihale,bg=bg).pack(anchor="w",pady=(0,6))
        df2=tk.Frame(ll,bg=bg); df2.pack(fill="x",pady=(0,6))
        tk.Label(df2,text="Durum:",bg=bg,font=("Segoe UI",10)).pack(side="left",padx=(0,6))
        self.ih_df=ttk.Combobox(df2,values=["Tümü"]+IHALE_DURUMLARI,state="readonly",width=22)
        self.ih_df.set("Tümü"); self.ih_df.bind("<<ComboboxSelected>>",lambda e:self._filtre_ihale(getattr(self,"_ih_ara","")))
        self.ih_df.pack(side="left")
        self.tv_ihale=self._tv(ll,[("id",40,"#"),("koy",100,"Köy"),("ilce",70,"İlçe"),("ad",120,"Başvuran"),
            ("tc",100,"TC"),("tel",100,"Telefon"),("bedel",80,"Bedel"),("durum",120,"Durum"),("tarih",90,"Tarih")],14)
        self.tv_ihale.tag_configure("aktif",background="#EAF4EE")
        self.tv_ihale.tag_configure("tamamlandi",background="#EBF5FB")
        self.tv_ihale.tag_configure("iptal",background="#FDEDEC")
        btnf=tk.Frame(ll,bg=bg); btnf.pack(pady=8)
        MBtn(btnf,"🔄 Yenile",command=self._yenile_ihale,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(btnf,"✏️ Durum Güncelle",command=self._ihale_durum,color=C_WARN,width=18).pack(side="left",padx=4)
        MBtn(btnf,"📋 Süreç",command=self._ihale_surec,color="#8E44AD",width=12).pack(side="left",padx=4)
        MBtn(btnf,"📄 Evrak Üret",command=self._ihale_evrak_uret,color=self.gc("acc"),width=14).pack(side="left",padx=4)
        MBtn(btnf,"🗑 Sil",command=self._ihale_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        # Yeni başvuru
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=16)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.45,anchor="center",width=640,height=560)
        tk.Label(card,text="Yeni İhale Başvurusu",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=4,pady=(18,14))
        self.ih_g={}
        for lbl,key,tip,r,col in [("Köy:","koy","entry",1,0),("İlçe:","ilce","combo",1,2),("Ad Soyad:","ad","entry",2,0),
            ("TC No:","tc","entry",2,2),("Telefon:","tel","entry",3,0),("Adres:","adres","entry",3,2),
            ("Teklif Bedeli (₺):","bedel","entry",4,0),("Ada/Parsel:","ada_parsel","entry",4,2),
            ("Otlatma Kapasitesi:","kapasite","entry",5,0)]:
            tk.Label(card,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).grid(row=r,column=col,padx=(20,4),pady=8,sticky="e")
            w_obj=ttk.Combobox(card,values=ILCELER,state="readonly",width=18) if tip=="combo" else ttk.Entry(card,width=20)
            w_obj.grid(row=r,column=col+1,padx=(0,14),pady=8,sticky="w"); self.ih_g[key]=w_obj
        tk.Label(card,text="Notlar:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=6,column=0,padx=(20,4),pady=8,sticky="e")
        self.ih_g["notlar"]=tk.Text(card,width=46,height=3,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=6,pady=4)
        self.ih_g["notlar"].grid(row=6,column=1,columnspan=3,padx=(0,20),pady=8,sticky="w")
        MBtn(card,"✅  Başvuruyu Kaydet",command=self._ihale_kaydet,color=self.gc("acc"),width=34).grid(row=7,column=0,columnspan=4,pady=18)
        # Bedel hesaplama
        bh=tk.Frame(bf2,bg=bg); bh.pack(fill="both",expand=True,padx=40,pady=30)
        tk.Label(bh,text="💰 İhale Bedel Hesaplayıcı",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")
        tk.Label(bh,text="İhale bedelini girin, tüm kalemler otomatik hesaplansın.",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,14))
        ef=tk.Frame(bh,bg=bg); ef.pack(anchor="w")
        tk.Label(ef,text="İhale Bedeli (₺):",bg=bg,font=("Segoe UI",11,"bold")).pack(side="left",padx=(0,8))
        self.e_bedel_hesap=ttk.Entry(ef,width=18,font=("Segoe UI",12)); self.e_bedel_hesap.pack(side="left",padx=(0,12))
        MBtn(ef,"Hesapla",command=self._bedel_hesapla,color=self.gc("acc"),width=12).pack(side="left")
        self.lbl_bedel_sonuc=tk.Label(bh,text="",bg=bg,font=("Segoe UI",11),fg="#333",justify="left")
        self.lbl_bedel_sonuc.pack(anchor="w",pady=16)
        self._yenile_ihale()

    def _bedel_hesapla(self):
        sonuc=ihale_bedel_hesapla(self.e_bedel_hesap.get())
        if not sonuc: messagebox.showerror("Hata","Geçerli bir bedel girin.\nÖrnek: 52.507,18 veya 52507.18"); return
        self.lbl_bedel_sonuc.config(text=
            f"📊 İhale Bedeli: {para_format(sonuc['toplam'])}\n\n"
            f"💰 Bakanlık Hesabına (%25): {para_format(sonuc['bakanlik_25'])}\n"
            f"🏘️ Köy Hesabına (%75): {para_format(sonuc['koy_75'])}\n"
            f"🔐 Kesin Teminat (%6): {para_format(sonuc['kesin_teminat_6'])}\n"
            f"📋 Damga Vergisi (‰5.69): {para_format(sonuc['damga_vergisi'])}\n"
            f"{'─'*40}\n"
            f"IBAN (Bakanlık): TR 5100 0010 0100 0003 5015 4026\n"
            f"IBAN (Emanet): TR 85 0001 0001 4900 0010 0059 60")

    def _yenile_ihale(self):
        if not DB_PATH: return
        try:
            with sqlite3.connect(DB_PATH) as c:
                self._all_ihale=c.execute("SELECT id,koy,ilce,ad_soyad,tc,telefon,bedel,durum,tarih FROM Ihaleler ORDER BY id DESC").fetchall()
            self._ih_ara=getattr(self,"_ih_ara",""); self._filtre_ihale(self._ih_ara)
        except Exception as e: logging.error(f"yenile_ihale:{e}")

    def _filtre_ihale(self,a):
        self._ih_ara=a.lower(); filtre=self.ih_df.get()
        self.tv_ihale.delete(*self.tv_ihale.get_children())
        for r in getattr(self,"_all_ihale",[]):
            if filtre!="Tümü" and r[7]!=filtre: continue
            if self._ih_ara and self._ih_ara not in " ".join(str(x) for x in r).lower(): continue
            tag="iptal" if "İptal" in str(r[7]) else "tamamlandi" if "Tamamlandı" in str(r[7]) else "aktif"
            # Bedel sütununu Türk formatında göster
            row=list(r)
            try: row[6]=para_format(row[6])
            except Exception: pass
            self.tv_ihale.insert("","end",values=row,tags=(tag,))

    def _ihale_kaydet(self):
        try:
            bedel=para_parse(self.ih_g["bedel"].get())
            koy=self.ih_g["koy"].get().strip()
            if not koy: raise ValueError("Köy adı boş bırakılamaz.")
            tc=self.ih_g["tc"].get().strip()
            if tc and not tc_dogrula(tc):
                if not messagebox.askyesno("TC Uyarı","TC numarası geçersiz. Devam?"): return
            notlar=self.ih_g["notlar"].get("1.0",tk.END).strip()
            with sqlite3.connect(DB_PATH) as conn:
                cur=conn.cursor()
                cur.execute("INSERT INTO Ihaleler(koy,ilce,ad_soyad,tc,telefon,adres,bedel,durum,notlar,tarih,ada_parsel,kapasite)VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                    (koy,self.ih_g["ilce"].get(),self.ih_g["ad"].get(),tc,
                     self.ih_g["tel"].get(),self.ih_g["adres"].get(),bedel,"Başvuru Alındı",notlar,
                     datetime.now().strftime("%d.%m.%Y"),self.ih_g["ada_parsel"].get(),self.ih_g["kapasite"].get()))
                ih_id=cur.lastrowid
                cur.execute("INSERT INTO Ihale_Log(ihale_id,tarih,personel,durum,not_icerik)VALUES(?,?,?,?,?)",
                    (ih_id,datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad,"Başvuru Alındı","Yeni başvuru."))
            db_log(self.u_id,"Yeni İhale",f"{koy} {bedel}₺")
            self._son_islem_kaydet("İhale",koy)
            for key,w in self.ih_g.items():
                if isinstance(w,(ttk.Entry,ttk.Combobox)): w.set("") if isinstance(w,ttk.Combobox) else w.delete(0,tk.END)
                elif isinstance(w,tk.Text): w.delete("1.0",tk.END)
            self._yenile_ihale(); messagebox.showinfo("Tamam",f"'{koy}' ihale başvurusu kaydedildi.")
        except ValueError as e: messagebox.showerror("Hata",str(e))
        except Exception as e: messagebox.showerror("DB Hatası",str(e))

    def _ihale_durum(self):
        sel=self.tv_ihale.selection()
        if not sel: messagebox.showwarning("Seçim","İhale seçin."); return
        ih_id=self.tv_ihale.item(sel[0])["values"][0]
        win=tk.Toplevel(self.root); win.title("Durum Güncelle"); win.geometry("400x280"); win.configure(bg=C_WHITE); win.grab_set()
        tk.Label(win,text="Yeni Durum:",bg=C_WHITE,font=("Segoe UI",11,"bold")).pack(pady=(16,4))
        cb=ttk.Combobox(win,values=IHALE_DURUMLARI,state="readonly",width=30); cb.pack(padx=20)
        tk.Label(win,text="Not:",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=20,pady=(10,2))
        not_e=tk.Text(win,height=4,width=42,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        not_e.pack(padx=20)
        def _k():
            yeni=cb.get(); acik=not_e.get("1.0",tk.END).strip()
            if not yeni: messagebox.showwarning("Uyarı","Durum seçin.",parent=win); return
            with sqlite3.connect(DB_PATH) as c:
                c.execute("UPDATE Ihaleler SET durum=? WHERE id=?",(yeni,ih_id))
                c.execute("INSERT INTO Ihale_Log(ihale_id,tarih,personel,durum,not_icerik)VALUES(?,?,?,?,?)",
                    (ih_id,datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad,yeni,acik))
            db_log(self.u_id,"İhale Durum",f"ID:{ih_id}→{yeni}"); self._yenile_ihale(); win.destroy()
        MBtn(win,"Güncelle",command=_k,width=18).pack(pady=12)

    def _ihale_surec(self):
        sel=self.tv_ihale.selection()
        if not sel: messagebox.showwarning("Seçim","İhale seçin."); return
        vals=self.tv_ihale.item(sel[0])["values"]; ih_id=vals[0]; koy=vals[1]
        win=tk.Toplevel(self.root); win.title(f"Süreç — {koy}"); win.geometry("700x400"); win.configure(bg=C_WHITE)
        tk.Label(win,text=f"İhale Süreç Geçmişi: {koy}",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=12)
        tv=ttk.Treeview(win,columns=("t","p","d","n"),show="headings",height=14)
        for col,w,bas in [("t",140,"Tarih"),("p",130,"Personel"),("d",130,"Durum"),("n",280,"Not")]:
            tv.heading(col,text=bas); tv.column(col,width=w)
        tv.pack(fill="both",expand=True,padx=12,pady=8)
        try:
            with sqlite3.connect(DB_PATH) as c:
                for r in c.execute("SELECT tarih,personel,durum,not_icerik FROM Ihale_Log WHERE ihale_id=? ORDER BY id",(ih_id,)).fetchall():
                    tv.insert("","end",values=r)
        except Exception as e: logging.error(f"surec:{e}")

    def _ihale_evrak_uret(self):
        """Seçili ihale için Word evrak üret"""
        sel=self.tv_ihale.selection()
        if not sel: messagebox.showwarning("Seçim","İhale seçin."); return
        vals=self.tv_ihale.item(sel[0])["values"]
        ih_id=vals[0]
        if not DOCX_OK: messagebox.showerror("Hata","python-docx kurulu değil.\npip install python-docx"); return
        try:
            with sqlite3.connect(DB_PATH) as c:
                r=c.execute("SELECT koy,ilce,ad_soyad,tc,telefon,adres,bedel,ada_parsel,kapasite FROM Ihaleler WHERE id=?",(ih_id,)).fetchone()
        except Exception: return
        if not r: return
        yil=_yil()
        veri={"koy":r[0],"ilce":r[1],"ad_soyad":r[2],"tc":r[3],"telefon":r[4],"adres":r[5],
              "bedel":r[6],"ada_parsel":r[7],"kapasite":r[8],"mera":f"{r[1]}/{r[0]}",
              "tarih":datetime.now().strftime("%d/%m/%Y"),
              "komisyon_baskani":"Semih CEMBEKLİ",
              "baslangic":f"01/06/{yil}","bitis":f"30/09/{yil}",
              "dogum_yeri":"","hayvan_bilgi":""}
        win=tk.Toplevel(self.root); win.title("Evrak Seç"); win.geometry("380x320"); win.configure(bg=C_WHITE); win.grab_set()
        tk.Label(win,text="Hangi evrakı üretmek istiyorsunuz?",font=("Segoe UI",12,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=14)
        def _uret(tip):
            dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile=f"{tip}_{r[0]}_{r[2]}.docx",filetypes=[("Word","*.docx")])
            if not dosya: return
            try:
                if tip=="Katılım Evrakları": word_katilim_evrak(dosya,veri)
                elif tip=="Kiralama Sözleşmesi": word_kiralama_sozlesme(dosya,veri)
                elif tip=="Kiralama Şartnamesi": word_kiralama_sartname(dosya,veri)
                messagebox.showinfo("Tamam",f"Evrak üretildi:\n{dosya}"); win.destroy()
            except Exception as e: messagebox.showerror("Hata",str(e))
        MBtn(win,"📋 Katılım Evrakları",command=lambda:_uret("Katılım Evrakları"),color=self.gc("acc"),width=30).pack(pady=6)
        MBtn(win,"📝 Kiralama Sözleşmesi (14 madde)",command=lambda:_uret("Kiralama Sözleşmesi"),color=C_INFO,width=30).pack(pady=6)
        MBtn(win,"📄 Kiralama Şartnamesi (22 madde)",command=lambda:_uret("Kiralama Şartnamesi"),color="#8E44AD",width=30).pack(pady=6)

    def _ihale_sil(self):
        sel=self.tv_ihale.selection()
        if not sel: return
        ih_id=self.tv_ihale.item(sel[0])["values"][0]; koy=self.tv_ihale.item(sel[0])["values"][1]
        self._guvenli_sil("Ihaleler",ih_id,f"İhale: {koy}","DELETE FROM Ihaleler WHERE id=?",(ih_id,),self._yenile_ihale)

    # ═══ 4. ISLAH ════════════════════════════════════════════════════════════
    def _t_islah(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2); pf=ttk.Frame(nb2)
        nb2.add(lf,text="  📋 Islah Listesi  "); nb2.add(nf,text="  ➕ Yeni Dilekçe  "); nb2.add(pf,text="  📊 Proje Takip Cetveli  ")
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,self._filtre_islah,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_islah=self._tv(ll,[("id",40,"#"),("koy",110,"Köy"),("ilce",70,"İlçe"),("talep",120,"Talep Eden"),
            ("dt",90,"Dilekçe T."),("talan",70,"Talep (da)"),("valan",70,"Verilen (da)"),("durum",100,"Durum")],12)
        self.tv_islah.tag_configure("bekliyor",background="#FEF9E7")
        self.tv_islah.tag_configure("verildi",background="#EAFAF1")
        self.tv_islah.tag_configure("verilmedi",background="#FDEDEC")
        bf=tk.Frame(ll,bg=bg); bf.pack(pady=8)
        MBtn(bf,"🔄 Yenile",command=self._yenile_islah,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(bf,"✏️ Sonuç Gir",command=self._islah_sonuc,color=C_WARN,width=16).pack(side="left",padx=4)
        MBtn(bf,"🗑 Sil",command=self._islah_sil,color=C_DANGER,width=12).pack(side="left",padx=4)
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=16)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.48,anchor="center",width=640,height=480)
        tk.Label(card,text="Yeni Islah/Amenajman Dilekçesi",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=4,pady=(18,14))
        self.isl_g={}
        for lbl,key,tip,r,col in [("Köy:","koy","entry",1,0),("İlçe:","ilce","combo",1,2),
            ("Talep Eden:","talep","entry",2,0),("Dilekçe Tarihi:","dt","entry",2,2),
            ("Talep Alanı (da):","alan","entry",3,0),("İş Programı:","ip","entry",3,2)]:
            tk.Label(card,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).grid(row=r,column=col,padx=(20,6),pady=8,sticky="e")
            w_obj=ttk.Combobox(card,values=ILCELER,state="readonly",width=18) if tip=="combo" else ttk.Entry(card,width=20)
            if key=="dt": w_obj.insert(0,datetime.now().strftime("%d.%m.%Y"))
            w_obj.grid(row=r,column=col+1,padx=(0,14),pady=8,sticky="w"); self.isl_g[key]=w_obj
        tk.Label(card,text="Açıklama:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=4,column=0,padx=(20,6),pady=8,sticky="ne")
        self.isl_g["acik"]=tk.Text(card,width=46,height=4,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        self.isl_g["acik"].grid(row=4,column=1,columnspan=3,padx=(0,20),pady=8,sticky="w")
        MBtn(card,"✅  Dilekçeyi Kaydet",command=self._islah_kaydet,color=self.gc("acc"),width=34).grid(row=5,column=0,columnspan=4,pady=18)
        # Proje Takip Cetveli
        pp=tk.Frame(pf,bg=bg); pp.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(pp,text="📊 Çayır Mera Islah ve Amenajman Projeleri — Proje Takip Cetveli",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,8))
        # Yıl seçimi
        yf2=tk.Frame(pp,bg=bg); yf2.pack(fill="x",pady=(0,6))
        tk.Label(yf2,text="Yıl:",bg=bg,font=("Segoe UI",10,"bold")).pack(side="left")
        self.prj_yil=ttk.Combobox(yf2,values=[str(y) for y in range(2020,2031)],state="readonly",width=8)
        self.prj_yil.set(_yil()); self.prj_yil.pack(side="left",padx=6)
        self.prj_yil.bind("<<ComboboxSelected>>",lambda e:self._yenile_proje())
        AramaFrame(yf2,self._filtre_proje,bg=bg).pack(side="left",padx=10)
        self.tv_prj=self._tv(pp,[("id",35,"#"),("ilce",70,"İlçe"),("koy",100,"Köy"),
            ("gub",65,"Gübre"),("mt",65,"M.Tohum"),("yt",65,"Y.Tohum"),
            ("gol",50,"Gölg."),("siv",50,"Sıvat"),("bor",55,"Boru"),
            ("cob",50,"Ç.Evi"),("tuz",50,"Tuzlk"),("kas",50,"K.Kzğ"),
            ("fin",100,"Finans"),("dur",80,"Durum")],10)
        bf2=tk.Frame(pp,bg=bg); bf2.pack(pady=6)
        MBtn(bf2,"➕ Yeni Proje",command=self._proje_ekle,color=self.gc("acc"),width=14).pack(side="left",padx=4)
        MBtn(bf2,"🔄 Yenile",command=self._yenile_proje,color=C_INFO,width=12).pack(side="left",padx=4)
        MBtn(bf2,"📥 Excel Import",command=self._proje_import,color="#8E44AD",width=14).pack(side="left",padx=4)
        MBtn(bf2,"📊 Excel Export",command=self._proje_export,color=self.gc("pri"),width=14).pack(side="left",padx=4)
        MBtn(bf2,"🗑 Sil",command=self._proje_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        self._yenile_islah(); self._yenile_proje()

    def _yenile_islah(self):
        if not DB_PATH: return
        try:
            with sqlite3.connect(DB_PATH) as c:
                self._all_islah=c.execute("SELECT id,koy,ilce,talep_eden,dilekce_tarihi,talep_alani,verilen_alan,durum FROM Islah_Amenajman ORDER BY id DESC").fetchall()
            self._filtre_islah("")
        except Exception as e: logging.error(f"yenile_islah:{e}")

    def _filtre_islah(self,a):
        self.tv_islah.delete(*self.tv_islah.get_children()); a=a.lower()
        for r in getattr(self,"_all_islah",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            d=str(r[7]); tag="bekliyor" if d=="Bekliyor" else "verildi" if "Verildi" in d else "verilmedi"
            self.tv_islah.insert("","end",values=r,tags=(tag,))

    def _islah_kaydet(self):
        try:
            koy=self.isl_g["koy"].get().strip(); talep=self.isl_g["talep"].get().strip()
            if not koy or not talep: raise ValueError("Köy ve talep eden zorunlu.")
            alan=float(self.isl_g["alan"].get() or 0)
            acik=self.isl_g["acik"].get("1.0",tk.END).strip()
            with sqlite3.connect(DB_PATH) as c:
                c.execute("INSERT INTO Islah_Amenajman(koy,ilce,dilekce_tarihi,talep_eden,talep_alani,talep_aciklama,is_programi,durum)VALUES(?,?,?,?,?,?,?,?)",
                    (koy,self.isl_g["ilce"].get(),self.isl_g["dt"].get(),talep,alan,acik,self.isl_g["ip"].get(),"Bekliyor"))
            db_log(self.u_id,"Islah Dilekçe",koy); self._son_islem_kaydet("Islah",koy)
            for key,w in self.isl_g.items():
                if isinstance(w,ttk.Entry): w.delete(0,tk.END)
                elif isinstance(w,ttk.Combobox): w.set("")
                elif isinstance(w,tk.Text): w.delete("1.0",tk.END)
            self.isl_g["dt"].insert(0,datetime.now().strftime("%d.%m.%Y"))
            self._yenile_islah(); messagebox.showinfo("Tamam",f"'{koy}' dilekçesi kaydedildi.")
        except ValueError as e: messagebox.showerror("Hata",str(e))

    def _islah_sonuc(self):
        sel=self.tv_islah.selection()
        if not sel: messagebox.showwarning("Seçim","Kayıt seçin."); return
        is_id=self.tv_islah.item(sel[0])["values"][0]; koy=self.tv_islah.item(sel[0])["values"][1]
        win=tk.Toplevel(self.root); win.title(f"Sonuç — {koy}"); win.geometry("460x340"); win.configure(bg=C_WHITE); win.grab_set()
        tk.Label(win,text=f"Islah Sonucu: {koy}",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=14)
        tk.Label(win,text="Durum:",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=20)
        cb=ttk.Combobox(win,values=["Bekliyor","Islah Yapılıyor","Islah Verildi","Islah Verilmedi","Kısmi Verildi"],state="readonly",width=30)
        cb.pack(padx=20,pady=4)
        tk.Label(win,text="Verilen Alan (da):",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=20)
        e_v=ttk.Entry(win,width=20); e_v.insert(0,"0"); e_v.pack(padx=20,pady=4,anchor="w")
        tk.Label(win,text="Açıklama:",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=20)
        txt=tk.Text(win,height=3,width=44,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        txt.pack(padx=20,pady=4)
        def _k():
            try:
                v=float(e_v.get() or 0); n=txt.get("1.0",tk.END).strip(); d=cb.get() or "Bekliyor"
                with sqlite3.connect(DB_PATH) as c:
                    c.execute("UPDATE Islah_Amenajman SET verilen_alan=?,verilmeme_neden=?,durum=?,kapanma_tarihi=? WHERE id=?",
                        (v,n,d,datetime.now().strftime("%d.%m.%Y"),is_id))
                db_log(self.u_id,"Islah Sonuç",f"ID:{is_id} {d}"); self._yenile_islah(); win.destroy()
            except ValueError as e: messagebox.showerror("Hata",str(e),parent=win)
        MBtn(win,"Sonucu Kaydet",command=_k,width=20).pack(pady=10)

    def _islah_sil(self):
        sel=self.tv_islah.selection()
        if not sel: return
        is_id=self.tv_islah.item(sel[0])["values"][0]; koy=self.tv_islah.item(sel[0])["values"][1]
        self._guvenli_sil("Islah_Amenajman",is_id,f"Islah: {koy}","DELETE FROM Islah_Amenajman WHERE id=?",(is_id,),self._yenile_islah)

    # ─── PROJE TAKİP CETVELİ ─────────────────────────────────────────────────
    ISLAH_FAALIYETLER = ["Gübreleme (da)","Mera Tohum (da)","Yem Bitkisi Tohum (da)",
        "Gölgelik (adet)","Sıvat (adet)","Kangal Boru (m)","Çoban Evi (adet)","Tuzluk (adet)","Kaşınma Kazığı (adet)"]
    FINANSMAN = ["DAP","AKAK","Bakanlık","İl Özel İdaresi","Köylere Hizmet Götürme","Diğer"]

    def _yenile_proje(self):
        if not DB_PATH: return
        yil=self.prj_yil.get() if hasattr(self,"prj_yil") else _yil()
        try:
            with sqlite3.connect(DB_PATH) as c:
                self._all_prj=c.execute("SELECT id,ilce,koy,gubre_da,mera_tohum_da,yem_tohum_da,golgelik,sivat,boru_m,coban_evi,tuzluk,kasinka,finansman,durum FROM Islah_Projeler WHERE yil=? ORDER BY ilce,koy",(yil,)).fetchall()
            self._filtre_proje("")
        except Exception as e: logging.error(f"proje:{e}")

    def _filtre_proje(self,a):
        if not hasattr(self,"tv_prj"): return
        self.tv_prj.delete(*self.tv_prj.get_children()); a=a.lower()
        for r in getattr(self,"_all_prj",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            self.tv_prj.insert("","end",values=r)

    def _proje_ekle(self):
        win=tk.Toplevel(self.root); win.title("Yeni Proje Kaydı"); win.geometry("560x520"); win.configure(bg=C_WHITE); win.grab_set()
        tk.Label(win,text="📊 Yeni Islah Projesi",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=14)
        pg={}
        for lbl,key in [("İlçe:","ilce"),("Köy:","koy")]:
            f=tk.Frame(win,bg=C_WHITE); f.pack(fill="x",padx=20,pady=4)
            tk.Label(f,text=lbl,bg=C_WHITE,font=("Segoe UI",10),width=18,anchor="w").pack(side="left")
            if key=="ilce":
                w=ttk.Combobox(f,values=ILCELER,state="readonly",width=20)
            else:
                w=ttk.Entry(f,width=22)
            w.pack(side="left"); pg[key]=w
        tk.Label(win,text="── Harcama Faaliyetleri ──",font=("Segoe UI",10,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=(10,4))
        faaliyet_keys=["gubre","mt","yt","gol","siv","bor","cob","tuz","kas"]
        faaliyet_labels=["Gübreleme (da)","Mera Tohum (da)","Yem Bitkisi Tohum (da)","Gölgelik (adet)","Sıvat (adet)","Kangal Boru (m)","Çoban Evi (adet)","Tuzluk (adet)","Kaşınma Kazığı (adet)"]
        for lbl,key in zip(faaliyet_labels,faaliyet_keys):
            f=tk.Frame(win,bg=C_WHITE); f.pack(fill="x",padx=20,pady=2)
            tk.Label(f,text=lbl+":",bg=C_WHITE,font=("Segoe UI",9),width=22,anchor="w").pack(side="left")
            e=ttk.Entry(f,width=12); e.insert(0,"0"); e.pack(side="left"); pg[key]=e
        f2=tk.Frame(win,bg=C_WHITE); f2.pack(fill="x",padx=20,pady=6)
        tk.Label(f2,text="Finansman:",bg=C_WHITE,font=("Segoe UI",10),width=18,anchor="w").pack(side="left")
        pg["fin"]=ttk.Combobox(f2,values=self.FINANSMAN,state="readonly",width=20); pg["fin"].pack(side="left")
        def _kaydet():
            koy=pg["koy"].get().strip()
            if not koy: messagebox.showerror("Hata","Köy zorunlu.",parent=win); return
            try:
                with sqlite3.connect(DB_PATH) as c:
                    c.execute("INSERT INTO Islah_Projeler(yil,ilce,koy,gubre_da,mera_tohum_da,yem_tohum_da,golgelik,sivat,boru_m,coban_evi,tuzluk,kasinka,finansman)VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",
                        (self.prj_yil.get(),pg["ilce"].get(),koy,
                         float(pg["gubre"].get() or 0),float(pg["mt"].get() or 0),float(pg["yt"].get() or 0),
                         int(pg["gol"].get() or 0),int(pg["siv"].get() or 0),float(pg["bor"].get() or 0),
                         int(pg["cob"].get() or 0),int(pg["tuz"].get() or 0),int(pg["kas"].get() or 0),pg["fin"].get()))
                db_log(self.u_id,"Proje Ekle",koy); win.destroy(); self._yenile_proje()
            except Exception as e: messagebox.showerror("Hata",str(e),parent=win)
        MBtn(win,"✅ Kaydet",command=_kaydet,color=self.gc("acc"),width=24).pack(pady=12)

    def _proje_import(self):
        yol=filedialog.askopenfilename(filetypes=[("Excel","*.xlsx")])
        if not yol: return
        try:
            df=pd.read_excel(yol)
            n=0
            with sqlite3.connect(DB_PATH) as c:
                for _,r in df.iterrows():
                    v=r.tolist()
                    if len(v)>=5:
                        c.execute("INSERT INTO Islah_Projeler(yil,ilce,koy,gubre_da,mera_tohum_da,yem_tohum_da,golgelik,sivat,boru_m,coban_evi,tuzluk,kasinka,finansman)VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",
                            (self.prj_yil.get(),str(v[0] if len(v)>0 else ""),str(v[1] if len(v)>1 else ""),
                             float(v[2] or 0) if len(v)>2 else 0,float(v[3] or 0) if len(v)>3 else 0,
                             float(v[4] or 0) if len(v)>4 else 0,int(v[5] or 0) if len(v)>5 else 0,
                             int(v[6] or 0) if len(v)>6 else 0,float(v[7] or 0) if len(v)>7 else 0,
                             int(v[8] or 0) if len(v)>8 else 0,int(v[9] or 0) if len(v)>9 else 0,
                             int(v[10] or 0) if len(v)>10 else 0,str(v[11]) if len(v)>11 else ""))
                        n+=1
            self._yenile_proje(); messagebox.showinfo("Tamam",f"{n} proje kaydı yüklendi.")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _proje_export(self):
        if not getattr(self,"_all_prj",None): return
        yol=filedialog.asksaveasfilename(defaultextension=".xlsx",filetypes=[("Excel","*.xlsx")])
        if yol:
            try:
                pd.DataFrame(self._all_prj,columns=["ID","İlçe","Köy","Gübreleme","M.Tohum","Y.Tohum","Gölgelik","Sıvat","Boru(m)","Ç.Evi","Tuzluk","K.Kazığı","Finansman","Durum"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"Dışa aktarıldı:\n{yol}")
            except Exception as e: messagebox.showerror("Hata",str(e))

    def _proje_sil(self):
        sel=self.tv_prj.selection()
        if not sel: return
        pid=self.tv_prj.item(sel[0])["values"][0]; koy=self.tv_prj.item(sel[0])["values"][2]
        self._guvenli_sil("Islah_Projeler",pid,f"Proje: {koy}","DELETE FROM Islah_Projeler WHERE id=?",(pid,),self._yenile_proje)

    # ═══ 5. TAHSİS / MADDE 14 ═══════════════════════════════════════════════
    def _t_tahsis(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2)
        nb2.add(lf,text="  📋 Tahsis Listesi  "); nb2.add(nf,text="  ➕ Yeni Tahsis/Md.14  ")
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,lambda a:None,bg=bg).pack(anchor="w",pady=(0,6))
        # Aşama göstergesi
        af=tk.Frame(ll,bg=bg); af.pack(fill="x",pady=(0,8))
        tk.Label(af,text="Aşama Durumu:",bg=bg,font=("Segoe UI",9,"bold")).pack(side="left",padx=(0,8))
        self._al=[]
        for a in TAHSIS_ASAMALARI:
            l=tk.Label(af,text=f" {a} ",bg="#E0E0E0",fg="#555",font=("Segoe UI",8),padx=6,pady=3)
            l.pack(side="left",padx=2); self._al.append(l)
        self.tv_tahsis=self._tv(ll,[("id",40,"#"),("koy",100,"Köy"),("ilce",70,"İlçe"),("ada",60,"Ada"),
            ("parsel",60,"Parsel"),("kurum",120,"Kurum"),("amac",120,"Amaç"),("alan",70,"Alan (ha)"),
            ("asama",140,"Aşama"),("durum",90,"Durum"),("bt",90,"Başvuru T.")],12)
        self.tv_tahsis.bind("<<TreeviewSelect>>",self._tahsis_sec)
        self.tv_tahsis.tag_configure("devam",background="#EBF5FB")
        self.tv_tahsis.tag_configure("tamam",background="#EAFAF1")
        self.tv_tahsis.tag_configure("red",background="#FDEDEC")
        bf=tk.Frame(ll,bg=bg); bf.pack(pady=8)
        MBtn(bf,"🔄 Yenile",command=self._yenile_tahsis,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(bf,"▶️ Sonraki Aşama",command=self._tahsis_ileri,color=self.gc("acc"),width=16).pack(side="left",padx=4)
        MBtn(bf,"❌ Reddet/Kapat",command=self._tahsis_red,color=C_DANGER,width=14).pack(side="left",padx=4)
        MBtn(bf,"📋 Süreç",command=self._tahsis_gecmis,color="#8E44AD",width=10).pack(side="left",padx=4)
        # Yeni
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=16)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.45,anchor="center",width=680,height=540)
        tk.Label(card,text="Yeni Tahsis / Madde 14 Başvurusu",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=4,pady=(18,14))
        self.tah_g={}
        for lbl,key,tip,r,col in [("Köy:","koy","entry",1,0),("İlçe:","ilce","combo",1,2),("Ada No:","ada","entry",2,0),
            ("Parsel:","parsel","entry",2,2),("Kurum:","kurum","entry",3,0),("Alan (ha):","alan","entry",3,2),
            ("Tahsis Amacı:","amac","entry",4,0),("Md.14 Bent:","madde14","md14combo",4,2)]:
            tk.Label(card,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).grid(row=r,column=col,padx=(20,6),pady=8,sticky="e")
            if tip=="combo": w_obj=ttk.Combobox(card,values=ILCELER,state="readonly",width=18)
            elif tip=="md14combo": w_obj=ttk.Combobox(card,values=["—","(a) Maden/Petrol","(b) Turizm","(c) Kamu Yatırımı","(d) Yerleşim/Muhafaza","(e) Köy Kanunu","(f) Güvenlik","(g) Doğal Afet"],state="readonly",width=22)
            else: w_obj=ttk.Entry(card,width=20)
            w_obj.grid(row=r,column=col+1,padx=(0,14),pady=8,sticky="w"); self.tah_g[key]=w_obj
        tk.Label(card,text="20 Yıllık Ot Geliri (₺):",bg=C_WHITE,font=("Segoe UI",10)).grid(row=5,column=0,padx=(20,6),pady=8,sticky="e")
        self.tah_g["ot_geliri"]=ttk.Entry(card,width=20); self.tah_g["ot_geliri"].grid(row=5,column=1,padx=(0,14),pady=8,sticky="w")
        tk.Label(card,text="Süre Bitiş Tarihi:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=5,column=2,padx=(20,6),pady=8,sticky="e")
        self.tah_g["sure_bitis"]=ttk.Entry(card,width=14); self.tah_g["sure_bitis"].grid(row=5,column=3,padx=(0,14),pady=8,sticky="w")
        tk.Label(card,text="Süre Tipi:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=6,column=0,padx=(20,6),pady=8,sticky="e")
        self.tah_g["sure_tipi"]=ttk.Combobox(card,values=["Askı İlanı (30 gün)","İtiraz Süresi (60 gün)","Tescil Süresi","Yatırım Başlama (2 yıl)","Diğer"],state="readonly",width=24)
        self.tah_g["sure_tipi"].grid(row=6,column=1,padx=(0,14),pady=8,sticky="w")
        tk.Label(card,text="Notlar:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=7,column=0,padx=(20,6),pady=8,sticky="ne")
        self.tah_g["notlar"]=tk.Text(card,width=46,height=3,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        self.tah_g["notlar"].grid(row=7,column=1,columnspan=3,padx=(0,20),pady=8,sticky="w")
        MBtn(card,"✅  Başvuruyu Kaydet",command=self._tahsis_kaydet,color=self.gc("acc"),width=34).grid(row=8,column=0,columnspan=4,pady=18)
        self._yenile_tahsis()

    def _yenile_tahsis(self):
        if not DB_PATH: return
        try:
            with sqlite3.connect(DB_PATH) as c:
                rows=c.execute("SELECT id,koy,ilce,ada,parsel,kurum,amac,alan_ha,asama,durum,basvuru_t FROM Tahsisler ORDER BY id DESC").fetchall()
            self.tv_tahsis.delete(*self.tv_tahsis.get_children())
            for r in rows:
                d=str(r[9]); tag="tamam" if d=="Tamamlandı" else "red" if d=="Reddedildi" else "devam"
                self.tv_tahsis.insert("","end",values=r,tags=(tag,))
        except Exception as e: logging.error(f"yenile_tahsis:{e}")

    def _tahsis_sec(self,event=None):
        sel=self.tv_tahsis.selection()
        if not sel: return
        asama=str(self.tv_tahsis.item(sel[0])["values"][8])
        mi=next((i for i,a in enumerate(TAHSIS_ASAMALARI) if a==asama),0)
        for i,l in enumerate(self._al):
            if i<mi: l.config(bg="#A9DFBF",fg="#1E5631")
            elif i==mi: l.config(bg=self.gc("pri"),fg=C_WHITE)
            else: l.config(bg="#E0E0E0",fg="#555")

    def _tahsis_kaydet(self):
        try:
            koy=self.tah_g["koy"].get().strip()
            if not koy: raise ValueError("Köy adı zorunlu.")
            alan=float(self.tah_g["alan"].get() or 0)
            ot_geliri=float(self.tah_g["ot_geliri"].get() or 0)
            notlar=self.tah_g["notlar"].get("1.0",tk.END).strip()
            md14=self.tah_g["madde14"].get()
            with sqlite3.connect(DB_PATH) as conn:
                cur=conn.cursor()
                sure_bitis=self.tah_g["sure_bitis"].get().strip() if "sure_bitis" in self.tah_g else ""
                sure_tipi=self.tah_g["sure_tipi"].get() if "sure_tipi" in self.tah_g else ""
                cur.execute("INSERT INTO Tahsisler(koy,ilce,ada,parsel,kurum,amac,alan_ha,asama,durum,basvuru_t,notlar,madde14_bent,ot_geliri,sure_bitis,sure_tipi)VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (koy,self.tah_g["ilce"].get(),self.tah_g["ada"].get(),self.tah_g["parsel"].get(),
                     self.tah_g["kurum"].get(),self.tah_g["amac"].get(),alan,TAHSIS_ASAMALARI[0],"Devam Ediyor",
                     datetime.now().strftime("%d.%m.%Y"),notlar,md14 if md14!="—" else "",ot_geliri,sure_bitis,sure_tipi))
                tid=cur.lastrowid
                cur.execute("INSERT INTO Tahsis_Log(tahsis_id,tarih,personel,asama,aciklama)VALUES(?,?,?,?,?)",
                    (tid,datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad,TAHSIS_ASAMALARI[0],"Başvuru alındı."))
            db_log(self.u_id,"Tahsis Başvuru",koy); self._son_islem_kaydet("Tahsis",koy)
            for k,w in self.tah_g.items():
                if isinstance(w,ttk.Entry): w.delete(0,tk.END)
                elif isinstance(w,ttk.Combobox): w.set("")
                elif isinstance(w,tk.Text): w.delete("1.0",tk.END)
            self._yenile_tahsis(); messagebox.showinfo("Tamam",f"'{koy}' tahsis başvurusu kaydedildi.")
        except ValueError as e: messagebox.showerror("Hata",str(e))

    def _tahsis_ileri(self):
        sel=self.tv_tahsis.selection()
        if not sel: messagebox.showwarning("Seçim","Tahsis seçin."); return
        vals=self.tv_tahsis.item(sel[0])["values"]; tid=vals[0]; koy=vals[1]; asama=str(vals[8])
        if asama==TAHSIS_ASAMALARI[-1]: messagebox.showinfo("Bilgi","Son aşamada."); return
        mi=next((i for i,a in enumerate(TAHSIS_ASAMALARI) if a==asama),0)
        sonraki=TAHSIS_ASAMALARI[mi+1]
        acik=tkinter.simpledialog.askstring("Not",f"'{koy}' → '{sonraki}'\nAçıklama:") or ""
        with sqlite3.connect(DB_PATH) as c:
            c.execute("UPDATE Tahsisler SET asama=? WHERE id=?",(sonraki,tid))
            if sonraki==TAHSIS_ASAMALARI[-1]:
                c.execute("UPDATE Tahsisler SET durum='Tamamlandı',sonuc_t=? WHERE id=?",(datetime.now().strftime("%d.%m.%Y"),tid))
            c.execute("INSERT INTO Tahsis_Log(tahsis_id,tarih,personel,asama,aciklama)VALUES(?,?,?,?,?)",
                (tid,datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad,sonraki,acik or "—"))
        db_log(self.u_id,"Tahsis Aşama",f"ID:{tid}→{sonraki}"); self._yenile_tahsis()

    def _tahsis_red(self):
        sel=self.tv_tahsis.selection()
        if not sel: return
        tid=self.tv_tahsis.item(sel[0])["values"][0]; koy=self.tv_tahsis.item(sel[0])["values"][1]
        n=tkinter.simpledialog.askstring("Red",f"'{koy}' red/kapatma nedeni:")
        if n is not None:
            with sqlite3.connect(DB_PATH) as c:
                c.execute("UPDATE Tahsisler SET durum='Reddedildi',sonuc_t=? WHERE id=?",(datetime.now().strftime("%d.%m.%Y"),tid))
                c.execute("INSERT INTO Tahsis_Log(tahsis_id,tarih,personel,asama,aciklama)VALUES(?,?,?,?,?)",
                    (tid,datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad,"REDDEDİLDİ",n))
            db_log(self.u_id,"Tahsis Red",f"ID:{tid}"); self._yenile_tahsis()

    def _tahsis_gecmis(self):
        sel=self.tv_tahsis.selection()
        if not sel: return
        tid=self.tv_tahsis.item(sel[0])["values"][0]; koy=self.tv_tahsis.item(sel[0])["values"][1]
        win=tk.Toplevel(self.root); win.title(f"Süreç — {koy}"); win.geometry("640x360"); win.configure(bg=C_WHITE)
        tk.Label(win,text=f"Tahsis Süreç Geçmişi: {koy}",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=12)
        tv=ttk.Treeview(win,columns=("t","p","a","n"),show="headings",height=14)
        for col,w,bas in [("t",140,"Tarih"),("p",130,"Personel"),("a",150,"Aşama"),("n",280,"Açıklama")]:
            tv.heading(col,text=bas); tv.column(col,width=w)
        tv.pack(fill="both",expand=True,padx=12,pady=8)
        try:
            with sqlite3.connect(DB_PATH) as c:
                for r in c.execute("SELECT tarih,personel,asama,aciklama FROM Tahsis_Log WHERE tahsis_id=? ORDER BY id",(tid,)).fetchall():
                    tv.insert("","end",values=r)
        except Exception as e: logging.error(f"tahsis_gecmis:{e}")

    # ═══ 6. ŞİKAYET TAKİP ═══════════════════════════════════════════════════
    def _t_sikayet(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2)
        nb2.add(lf,text="  📋 Şikayet Listesi  "); nb2.add(nf,text="  ➕ Yeni Şikayet  ")
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,self._filtre_sik,bg=bg).pack(anchor="w",pady=(0,6))
        df2=tk.Frame(ll,bg=bg); df2.pack(fill="x",pady=(0,6))
        tk.Label(df2,text="Durum:",bg=bg,font=("Segoe UI",10)).pack(side="left",padx=(0,6))
        self.sik_df=ttk.Combobox(df2,values=["Tümü","Yeni","İncelemede","Sonuçlandı","Kapatıldı"],state="readonly",width=16)
        self.sik_df.set("Tümü"); self.sik_df.bind("<<ComboboxSelected>>",lambda e:self._filtre_sik(""))
        self.sik_df.pack(side="left")
        self.tv_sik=self._tv(ll,[("id",40,"#"),("koy",110,"Köy"),("ilce",70,"İlçe"),
            ("sikayet_eden",130,"Şikayet Eden"),("tur",120,"Tür"),("tarih",90,"Tarih"),("durum",110,"Durum")],14)
        self.tv_sik.tag_configure("yeni",background="#FEF9E7")
        self.tv_sik.tag_configure("sonuclandi",background="#EAFAF1")
        bf=tk.Frame(ll,bg=bg); bf.pack(pady=8)
        MBtn(bf,"🔄 Yenile",command=self._yenile_sik,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(bf,"✏️ Güncelle",command=self._sik_guncelle,color=C_WARN,width=14).pack(side="left",padx=4)
        MBtn(bf,"🗑 Sil",command=self._sik_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        # Yeni
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=16)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.46,anchor="center",width=620,height=480)
        tk.Label(card,text="Yeni Şikayet / İhbar",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=4,pady=(18,14))
        self.sik_g={}
        for lbl,key,tip,r,col in [("Köy:","koy","entry",1,0),("İlçe:","ilce","combo",1,2),
            ("Şikayet Eden:","sikayet_eden","entry",2,0),("Tarih:","tarih","entry",2,2),
            ("Tür:","tur","siktip",3,0)]:
            tk.Label(card,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).grid(row=r,column=col,padx=(20,6),pady=8,sticky="e")
            if tip=="combo": w_obj=ttk.Combobox(card,values=ILCELER,state="readonly",width=18)
            elif tip=="siktip": w_obj=ttk.Combobox(card,state="readonly",width=28,values=["Kaçak Otlatma","Mera İşgali","İzinsiz Yapılaşma","Mera Tahrip","Hayvan Zararı","Diğer"])
            else: w_obj=ttk.Entry(card,width=20)
            if key=="tarih" and isinstance(w_obj,ttk.Entry): w_obj.insert(0,datetime.now().strftime("%d.%m.%Y"))
            w_obj.grid(row=r,column=col+1,padx=(0,14),pady=8,sticky="w"); self.sik_g[key]=w_obj
        tk.Label(card,text="Açıklama:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=4,column=0,padx=(20,6),pady=8,sticky="ne")
        self.sik_g["aciklama"]=tk.Text(card,width=46,height=5,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        self.sik_g["aciklama"].grid(row=4,column=1,columnspan=3,padx=(0,20),pady=8,sticky="w")
        MBtn(card,"✅  Kaydet",command=self._sik_kaydet,color=self.gc("acc"),width=34).grid(row=5,column=0,columnspan=4,pady=18)
        self._yenile_sik()

    def _yenile_sik(self):
        if not DB_PATH: return
        try:
            with sqlite3.connect(DB_PATH) as c:
                self._all_sik=c.execute("SELECT id,koy,ilce,sikayet_eden,tur,tarih,durum FROM Sikayetler ORDER BY id DESC").fetchall()
            self._filtre_sik("")
        except Exception as e: logging.error(f"yenile_sik:{e}")

    def _filtre_sik(self,a):
        self.tv_sik.delete(*self.tv_sik.get_children()); a=a.lower()
        filtre=self.sik_df.get()
        for r in getattr(self,"_all_sik",[]):
            if filtre!="Tümü" and r[6]!=filtre: continue
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            tag="sonuclandi" if r[6] in ("Sonuçlandı","Kapatıldı") else "yeni"
            self.tv_sik.insert("","end",values=r,tags=(tag,))

    def _sik_kaydet(self):
        try:
            koy=self.sik_g["koy"].get().strip()
            if not koy: raise ValueError("Köy adı zorunlu.")
            acik=self.sik_g["aciklama"].get("1.0",tk.END).strip()
            with sqlite3.connect(DB_PATH) as c:
                c.execute("INSERT INTO Sikayetler(koy,ilce,sikayet_eden,tur,aciklama,durum,tarih)VALUES(?,?,?,?,?,?,?)",
                    (koy,self.sik_g["ilce"].get(),self.sik_g["sikayet_eden"].get(),self.sik_g["tur"].get(),acik,"Yeni",self.sik_g["tarih"].get()))
            db_log(self.u_id,"Şikayet Kayıt",koy); self._son_islem_kaydet("Şikayet",koy)
            for k,w in self.sik_g.items():
                if isinstance(w,ttk.Entry): w.delete(0,tk.END)
                elif isinstance(w,ttk.Combobox): w.set("")
                elif isinstance(w,tk.Text): w.delete("1.0",tk.END)
            self.sik_g["tarih"].insert(0,datetime.now().strftime("%d.%m.%Y"))
            self._yenile_sik(); messagebox.showinfo("Tamam",f"'{koy}' şikayet kaydedildi.")
        except ValueError as e: messagebox.showerror("Hata",str(e))

    def _sik_guncelle(self):
        sel=self.tv_sik.selection()
        if not sel: messagebox.showwarning("Seçim","Şikayet seçin."); return
        sid=self.tv_sik.item(sel[0])["values"][0]
        win=tk.Toplevel(self.root); win.title("Güncelle"); win.geometry("420x260"); win.configure(bg=C_WHITE); win.grab_set()
        tk.Label(win,text="Yeni Durum:",bg=C_WHITE,font=("Segoe UI",11,"bold")).pack(pady=(16,4))
        cb=ttk.Combobox(win,values=["Yeni","İncelemede","Sonuçlandı","Kapatıldı"],state="readonly",width=28); cb.pack(padx=20)
        tk.Label(win,text="Sonuç:",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=20,pady=(10,2))
        txt=tk.Text(win,height=3,width=44,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6); txt.pack(padx=20)
        def _k():
            yeni=cb.get(); sonuc=txt.get("1.0",tk.END).strip()
            if not yeni: messagebox.showwarning("Uyarı","Durum seçin.",parent=win); return
            kapanma=datetime.now().strftime("%d.%m.%Y") if yeni in ["Sonuçlandı","Kapatıldı"] else None
            with sqlite3.connect(DB_PATH) as c:
                c.execute("UPDATE Sikayetler SET durum=?,sonuc=?,kapanma=? WHERE id=?",(yeni,sonuc,kapanma,sid))
            db_log(self.u_id,"Şikayet Güncelle",f"ID:{sid}→{yeni}"); self._yenile_sik(); win.destroy()
        MBtn(win,"Güncelle",command=_k,width=18).pack(pady=12)

    def _sik_sil(self):
        sel=self.tv_sik.selection()
        if not sel: return
        sid=self.tv_sik.item(sel[0])["values"][0]
        if messagebox.askyesno("Onay","Silinsin mi?"):
            with sqlite3.connect(DB_PATH) as c: c.execute("DELETE FROM Sikayetler WHERE id=?",(sid,))
            self._yenile_sik()

    # ═══ 7. İDARİ PARA CEZASI ════════════════════════════════════════════════
    def _t_ceza(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2); mf=ttk.Frame(nb2)
        nb2.add(lf,text="  📋 Ceza Listesi  "); nb2.add(nf,text="  ➕ Yeni Ceza  "); nb2.add(mf,text="  🔁 Mükerrer Cezalar  ")
        # Liste
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,self._filtre_ceza,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_ceza=self._tv(ll,[("id",40,"#"),("ad",130,"Ad Soyad"),("tc",100,"TC No"),
            ("ilce",70,"İlçe"),("mera",120,"Mera"),("yil",50,"Yıl"),("konu",180,"Konu"),
            ("tutar",90,"Tutar (₺)"),("tarih",90,"Tarih")],14)
        self.tv_ceza.tag_configure("mukerrer",background="#FDEDEC")
        bf=tk.Frame(ll,bg=bg); bf.pack(pady=8)
        MBtn(bf,"🔄 Yenile",command=self._yenile_ceza,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(bf,"📄 Evrak Üret",command=self._ceza_evrak,color=self.gc("acc"),width=14).pack(side="left",padx=4)
        MBtn(bf,"📊 Excel Export",command=self._ceza_excel,color="#8E44AD",width=14).pack(side="left",padx=4)
        MBtn(bf,"🗑 Sil",command=self._ceza_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        # Yeni ceza
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=16)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.45,anchor="center",width=660,height=540)
        tk.Label(card,text="Yeni İdari Para Cezası Kaydı",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=4,pady=(18,14))
        self.cz_g={}
        for lbl,key,tip,r,col in [("Ad Soyad:","ad","entry",1,0),("TC No:","tc","entry",1,2),
            ("İlçe:","ilce","combo",2,0),("Mera Köy:","mera_koy","entry",2,2),
            ("Ada/Parsel:","ada_parsel","entry",3,0),("Yıl:","yil","entry",3,2),
            ("Konu:","konu","ceza_tip",4,0),("Hayvan Sayısı:","hayvan","entry",4,2),
            ("Ceza Miktarı (₺/baş):","birim","entry",5,0),("İPC Tutarı (₺):","tutar","entry",5,2)]:
            tk.Label(card,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).grid(row=r,column=col,padx=(20,6),pady=8,sticky="e")
            if tip=="combo": w_obj=ttk.Combobox(card,values=ILCELER,state="readonly",width=18)
            elif tip=="ceza_tip": w_obj=ttk.Combobox(card,values=CEZA_TURLERI,state="readonly",width=30)
            else: w_obj=ttk.Entry(card,width=20)
            if key=="yil": w_obj.insert(0,str(datetime.now().year))
            w_obj.grid(row=r,column=col+1,padx=(0,14),pady=8,sticky="w"); self.cz_g[key]=w_obj
        tk.Label(card,text="Notlar:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=6,column=0,padx=(20,6),pady=8,sticky="ne")
        self.cz_g["notlar"]=tk.Text(card,width=46,height=2,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        self.cz_g["notlar"].grid(row=6,column=1,columnspan=3,padx=(0,20),pady=8,sticky="w")
        MBtn(card,"✅  Cezayı Kaydet",command=self._ceza_kaydet,color=C_DANGER,width=34).grid(row=7,column=0,columnspan=4,pady=18)
        # Mükerrer
        mm=tk.Frame(mf,bg=bg); mm.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(mm,text="🔁 Birden Fazla Ceza Uygulanan Şahıslar",font=("Segoe UI",13,"bold"),fg=C_DANGER,bg=bg).pack(anchor="w",pady=(0,10))
        self.tv_mukerrer=self._tv(mm,[("ad",160,"Ad Soyad"),("tc",110,"TC No"),("sayi",60,"Adet"),("konular",250,"Konular"),("tarihler",250,"Ceza Tarihleri")],14)
        MBtn(mm,"🔄 Yenile",command=self._yenile_mukerrer,color=C_INFO,width=14).pack(pady=8)
        self._yenile_ceza()
        self._yenile_mukerrer()

    def _yenile_ceza(self):
        if not DB_PATH: return
        try:
            with sqlite3.connect(DB_PATH) as c:
                self._all_ceza=c.execute("SELECT id,ad_soyad,tc,ilce,mera_koy,yil,konu,ipc_tutari,tarih FROM Idari_Cezalar ORDER BY id DESC").fetchall()
                # Mükerrer TC'leri bul
                mukerrer_tc=set()
                for r in c.execute("SELECT tc FROM Idari_Cezalar GROUP BY tc HAVING COUNT(*)>1").fetchall():
                    mukerrer_tc.add(str(r[0]))
                self._mukerrer_tc=mukerrer_tc
            self._filtre_ceza("")
        except Exception as e: logging.error(f"yenile_ceza:{e}")

    def _filtre_ceza(self,a):
        self.tv_ceza.delete(*self.tv_ceza.get_children()); a=a.lower()
        for r in getattr(self,"_all_ceza",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            tag="mukerrer" if str(r[2]) in getattr(self,"_mukerrer_tc",set()) else ""
            # Tutar sütununu Türk formatında göster
            row=list(r)
            try: row[7]=para_format(row[7])
            except Exception: pass
            self.tv_ceza.insert("","end",values=row,tags=(tag,) if tag else ())

    def _yenile_mukerrer(self):
        if not DB_PATH: return
        try:
            self.tv_mukerrer.delete(*self.tv_mukerrer.get_children())
            with sqlite3.connect(DB_PATH) as c:
                for r in c.execute("""SELECT ad_soyad,tc,COUNT(*) as cnt,
                    GROUP_CONCAT(DISTINCT konu) as konular,
                    GROUP_CONCAT(tarih,' / ') as tarihler
                    FROM Idari_Cezalar GROUP BY tc HAVING cnt>1 ORDER BY cnt DESC""").fetchall():
                    self.tv_mukerrer.insert("","end",values=r)
        except Exception as e: logging.error(f"mukerrer:{e}")

    def _ceza_kaydet(self):
        try:
            ad=self.cz_g["ad"].get().strip(); tc=self.cz_g["tc"].get().strip()
            if not ad: raise ValueError("Ad soyad zorunlu.")
            if tc and not tc_dogrula(tc):
                if not messagebox.askyesno("TC Uyarı","TC numarası geçersiz. Devam?"): return
            tutar=para_parse(self.cz_g["tutar"].get())
            notlar=self.cz_g["notlar"].get("1.0",tk.END).strip()
            # Mükerrer kontrolü
            with sqlite3.connect(DB_PATH) as c:
                onceki=c.execute("SELECT COUNT(*) FROM Idari_Cezalar WHERE tc=?",(tc,)).fetchone()[0]
                if onceki>0:
                    messagebox.showwarning("⚠️ Mükerrer Ceza",f"Bu kişiye daha önce {onceki} adet ceza uygulanmış!\nTC: {tc}")
                c.execute("INSERT INTO Idari_Cezalar(ad_soyad,tc,ilce,mera_koy,mera_ada_parsel,yil,konu,hayvan_sayisi,ceza_miktari,ipc_tutari,tarih,notlar)VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                    (ad,tc,self.cz_g["ilce"].get(),self.cz_g["mera_koy"].get(),self.cz_g["ada_parsel"].get(),
                     int(self.cz_g["yil"].get() or datetime.now().year),self.cz_g["konu"].get(),
                     self.cz_g["hayvan"].get(),self.cz_g["birim"].get(),tutar,
                     datetime.now().strftime("%d.%m.%Y"),notlar))
            db_log(self.u_id,"İdari Ceza",f"{ad} {tutar}₺"); self._son_islem_kaydet("Ceza",ad)
            for k,w in self.cz_g.items():
                if isinstance(w,ttk.Entry): w.delete(0,tk.END)
                elif isinstance(w,ttk.Combobox): w.set("")
                elif isinstance(w,tk.Text): w.delete("1.0",tk.END)
            self.cz_g["yil"].insert(0,str(datetime.now().year))
            self._yenile_ceza(); self._yenile_mukerrer()
            messagebox.showinfo("Tamam",f"'{ad}' ceza kaydı oluşturuldu.")
        except ValueError as e: messagebox.showerror("Hata",str(e))

    def _ceza_evrak(self):
        sel=self.tv_ceza.selection()
        if not sel: messagebox.showwarning("Seçim","Ceza seçin."); return
        if not DOCX_OK: messagebox.showerror("Hata","python-docx kurulu değil."); return
        cid=self.tv_ceza.item(sel[0])["values"][0]
        try:
            with sqlite3.connect(DB_PATH) as c:
                r=c.execute("SELECT ad_soyad,tc,ilce,mera_koy,konu,ipc_tutari,tarih FROM Idari_Cezalar WHERE id=?",(cid,)).fetchone()
        except Exception: return
        dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile=f"Ceza_{r[0]}.docx",filetypes=[("Word","*.docx")])
        if not dosya: return
        try:
            word_idari_ceza(dosya,{"ad_soyad":r[0],"tc":r[1],"mera":f"{r[2]}/{r[3]}","konu":r[4],"tutar":r[5],"tarih":r[6],"baba_adi":"","dogum":"","adres":""})
            messagebox.showinfo("Tamam",f"Evrak üretildi:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _ceza_excel(self):
        if not getattr(self,"_all_ceza",None): messagebox.showwarning("Veri Yok","Kayıt yok."); return
        yol=filedialog.asksaveasfilename(defaultextension=".xlsx",filetypes=[("Excel","*.xlsx")])
        if yol:
            try:
                pd.DataFrame(self._all_ceza,columns=["ID","Ad Soyad","TC","İlçe","Mera","Yıl","Konu","Tutar","Tarih"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"Dışa aktarıldı:\n{yol}")
            except Exception as e: messagebox.showerror("Hata",str(e))

    def _ceza_sil(self):
        sel=self.tv_ceza.selection()
        if not sel: return
        cid=self.tv_ceza.item(sel[0])["values"][0]; ad=self.tv_ceza.item(sel[0])["values"][1]
        self._guvenli_sil("Idari_Cezalar",cid,f"Ceza: {ad}","DELETE FROM Idari_Cezalar WHERE id=?",(cid,),lambda:(self._yenile_ceza(),self._yenile_mukerrer()))

    # ═══ 8. VERİ KAYIT ══════════════════════════════════════════════════════
    def _t_veri(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2)
        nb2.add(lf,text="  📋 Kayıtlar  "); nb2.add(nf,text="  ➕ Yeni Kayıt  ")
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,self._filtre_veri,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_veri=self._tv(ll,[("id",40,"#"),("t",120,"Tarih"),("k",110,"Personel"),("kat",110,"Kategori"),("bas",200,"Başlık"),("det",350,"Detay")],16)
        bf=tk.Frame(ll,bg=bg); bf.pack(pady=6)
        MBtn(bf,"🔄 Yenile",command=self._yenile_veri,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(bf,"📊 Excel Export",command=self._veri_excel,color=self.gc("acc"),width=16).pack(side="left",padx=4)
        if self.u_yetki!="İzleyici":
            MBtn(bf,"🗑 Sil",command=self._veri_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=20)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.45,anchor="center",width=560,height=380)
        tk.Label(card,text="Yeni Veri Kaydı",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=2,pady=(18,14))
        tk.Label(card,text="Başlık:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=1,column=0,padx=(28,6),pady=8,sticky="e")
        self.veri_bas=ttk.Entry(card,width=38); self.veri_bas.grid(row=1,column=1,padx=(0,28),sticky="w")
        tk.Label(card,text="Kategori:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=2,column=0,padx=(28,6),pady=8,sticky="e")
        self.veri_kat=ttk.Combobox(card,width=28,state="readonly",values=["Genel","Mera","Hayvancılık","İdari","Teknik","Diğer"])
        self.veri_kat.set("Genel"); self.veri_kat.grid(row=2,column=1,padx=(0,28),sticky="w")
        tk.Label(card,text="Detay:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=3,column=0,padx=(28,6),pady=8,sticky="ne")
        self.veri_det=tk.Text(card,width=36,height=6,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        self.veri_det.grid(row=3,column=1,padx=(0,28),pady=8,sticky="w")
        MBtn(card,"💾  Kaydet",command=self._veri_kaydet,color=self.gc("acc"),width=28).grid(row=4,column=0,columnspan=2,pady=14)
        self._yenile_veri()

    def _yenile_veri(self):
        if not DB_PATH: return
        try:
            with sqlite3.connect(DB_PATH) as c: self._all_veri=c.execute("SELECT id,tarih,kul,kategori,baslik,detay FROM Veri_Kayit ORDER BY id DESC").fetchall()
            self._filtre_veri("")
        except Exception as e: logging.error(f"yenile_veri:{e}")

    def _filtre_veri(self,a):
        self.tv_veri.delete(*self.tv_veri.get_children()); a=a.lower()
        for r in getattr(self,"_all_veri",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            self.tv_veri.insert("","end",values=r)

    def _veri_kaydet(self):
        bas=self.veri_bas.get().strip()
        if not bas: messagebox.showwarning("Uyarı","Başlık zorunlu."); return
        det=self.veri_det.get("1.0",tk.END).strip()
        with sqlite3.connect(DB_PATH) as c:
            c.execute("INSERT INTO Veri_Kayit(tarih,kul,baslik,kategori,detay)VALUES(?,?,?,?,?)",
                (datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad,bas,self.veri_kat.get(),det))
        db_log(self.u_id,"Veri Kayıt",bas); self.veri_bas.delete(0,tk.END); self.veri_det.delete("1.0",tk.END); self._yenile_veri()

    def _veri_excel(self):
        if not getattr(self,"_all_veri",None): messagebox.showwarning("Veri Yok","Kayıt yok."); return
        yol=filedialog.asksaveasfilename(defaultextension=".xlsx",filetypes=[("Excel","*.xlsx")])
        if yol:
            try:
                pd.DataFrame(self._all_veri,columns=["ID","Tarih","Personel","Kategori","Başlık","Detay"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"Dışa aktarıldı:\n{yol}")
            except Exception as e: messagebox.showerror("Hata",str(e))

    def _veri_sil(self):
        sel=self.tv_veri.selection()
        if not sel: return
        vid=self.tv_veri.item(sel[0])["values"][0]
        if messagebox.askyesno("Onay","Silinsin mi?"):
            with sqlite3.connect(DB_PATH) as c: c.execute("DELETE FROM Veri_Kayit WHERE id=?",(vid,))
            self._yenile_veri()

    # ═══ 9. MUHTARLAR ════════════════════════════════════════════════════════
    def _t_muhtar(self,p):
        bg=self.gc("bg"); f=tk.Frame(p,bg=bg); f.pack(fill="both",expand=True,padx=16,pady=14)
        tk.Label(f,text="👥 Muhtarlar Rehberi",font=("Segoe UI",15,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        if self.u_yetki!="İzleyici":
            form=tk.LabelFrame(f,text="  Yeni Muhtar  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=10,pady=8)
            form.pack(fill="x",pady=(0,10))
            self.muh_g={}
            for i,(lbl,key,w) in enumerate([("İlçe:","ilce",12),("Köy:","koy",16),("Ad Soyad:","ad",20),("Telefon:","tel",14),("E-posta:","email",20)]):
                tk.Label(form,text=lbl,bg=bg,font=("Segoe UI",10)).grid(row=0,column=i*2,padx=(0,4),sticky="w")
                w_obj=ttk.Combobox(form,values=ILCELER,state="readonly",width=w) if lbl=="İlçe:" else ttk.Entry(form,width=w)
                w_obj.grid(row=0,column=i*2+1,padx=(0,10)); self.muh_g[key]=w_obj
            MBtn(form,"➕ Ekle",command=self._muh_ekle).grid(row=0,column=10,padx=8)
        AramaFrame(f,self._filtre_muh,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_muh=self._tv(f,[("id",40,"#"),("ilce",90,"İlçe"),("koy",140,"Köy"),("ad",160,"Ad Soyad"),("tel",130,"Telefon"),("email",180,"E-posta")],18)
        bf=tk.Frame(f,bg=bg); bf.pack(pady=6)
        MBtn(bf,"🔄 Yenile",command=self._yenile_muh,color=C_INFO,width=14).pack(side="left",padx=4)
        if self.u_yetki!="İzleyici":
            MBtn(bf,"📥 Excel Import",command=self._muh_import,color="#8E44AD",width=14).pack(side="left",padx=4)
            MBtn(bf,"📊 Excel Export",command=self._muh_export,color=self.gc("acc"),width=14).pack(side="left",padx=4)
            MBtn(bf,"🗑 Sil",command=self._muh_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        self._tum_muh=[]; self._yenile_muh()

    def _yenile_muh(self):
        if not DB_PATH: return
        try:
            with sqlite3.connect(DB_PATH) as c: self._tum_muh=c.execute("SELECT id,ilce,koy,ad_soyad,telefon,email FROM Muhtarlar ORDER BY ilce,koy").fetchall()
            self._filtre_muh("")
        except Exception as e: logging.error(f"yenile_muh:{e}")

    def _filtre_muh(self,a):
        self.tv_muh.delete(*self.tv_muh.get_children()); a=a.lower()
        for r in self._tum_muh:
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            self.tv_muh.insert("","end",values=r)

    def _muh_ekle(self):
        koy=self.muh_g["koy"].get().strip(); ad=self.muh_g["ad"].get().strip()
        if not koy or not ad: messagebox.showwarning("Uyarı","Köy ve ad zorunlu."); return
        with sqlite3.connect(DB_PATH) as c: c.execute("INSERT INTO Muhtarlar(ilce,koy,ad_soyad,telefon,email)VALUES(?,?,?,?,?)",(self.muh_g["ilce"].get(),koy,ad,self.muh_g["tel"].get(),self.muh_g["email"].get()))
        db_log(self.u_id,"Muhtar Ekle",koy)
        for k,w in self.muh_g.items():
            if isinstance(w,ttk.Entry): w.delete(0,tk.END)
            elif isinstance(w,ttk.Combobox): w.set("")
        self._yenile_muh()

    def _muh_sil(self):
        sel=self.tv_muh.selection()
        if not sel: return
        mid=self.tv_muh.item(sel[0])["values"][0]
        if messagebox.askyesno("Onay","Silinsin mi?"):
            with sqlite3.connect(DB_PATH) as c: c.execute("DELETE FROM Muhtarlar WHERE id=?",(mid,))
            self._yenile_muh()

    def _muh_import(self):
        """Excel'den muhtar bilgilerini toplu yükle"""
        yol=filedialog.askopenfilename(filetypes=[("Excel","*.xlsx *.xls")])
        if not yol: return
        try:
            df=pd.read_excel(yol)
            if len(df.columns)<3:
                messagebox.showerror("Hata","En az 3 sütun gerekli: İlçe | Köy | Ad Soyad\nOpsiyonel: Telefon | E-posta"); return
            with sqlite3.connect(DB_PATH) as c:
                kayit=0
                for _,r in df.iterrows():
                    v=r.tolist()
                    ilce=str(v[0]) if len(v)>0 else ""
                    koy=str(v[1]) if len(v)>1 else ""
                    ad=str(v[2]) if len(v)>2 else ""
                    tel=str(v[3]) if len(v)>3 else ""
                    email=str(v[4]) if len(v)>4 else ""
                    if koy and ad:
                        c.execute("INSERT INTO Muhtarlar(ilce,koy,ad_soyad,telefon,email)VALUES(?,?,?,?,?)",(ilce,koy,ad,tel,email))
                        kayit+=1
            db_log(self.u_id,"Muhtar Import",f"{kayit} kayıt"); self._yenile_muh()
            messagebox.showinfo("Tamam",f"{kayit} muhtar bilgisi yüklendi.")
        except Exception as e: messagebox.showerror("Excel Hatası",str(e))

    def _muh_export(self):
        """Muhtar bilgilerini Excel'e aktar"""
        if not self._tum_muh: messagebox.showwarning("Veri Yok","Kayıt yok."); return
        yol=filedialog.asksaveasfilename(defaultextension=".xlsx",initialfile="Muhtarlar.xlsx",filetypes=[("Excel","*.xlsx")])
        if yol:
            try:
                pd.DataFrame(self._tum_muh,columns=["ID","İlçe","Köy","Ad Soyad","Telefon","E-posta"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"Dışa aktarıldı:\n{yol}")
            except Exception as e: messagebox.showerror("Hata",str(e))

    # ═══ 10. AJANDA ══════════════════════════════════════════════════════════
    def _t_ajanda(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2)
        nb2.add(lf,text="  📅 Tüm Notlar  "); nb2.add(nf,text="  ➕ Yeni Hatırlatıcı  ")
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,self._filtre_aj,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_aj=self._tv(ll,[("id",40,"#"),("t",100,"Tarih"),("s",70,"Saat"),("bas",190,"Başlık"),("tur",90,"Tür"),("dur",90,"Durum"),("ic",280,"İçerik")],14)
        self.tv_aj.tag_configure("bekliyor",background="#FEF9E7")
        self.tv_aj.tag_configure("tamamlandi",background="#EAFAF1")
        self.tv_aj.tag_configure("gecmis",background="#FDEDEC")
        bf=tk.Frame(ll,bg=bg); bf.pack(pady=8)
        MBtn(bf,"🔄 Yenile",command=self._yenile_aj,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(bf,"✅ Tamamlandı",command=lambda:self._aj_dur("Tamamlandı"),color=self.gc("acc"),width=16).pack(side="left",padx=4)
        MBtn(bf,"🗑 Sil",command=self._aj_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=20)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.45,anchor="center",width=560,height=460)
        tk.Label(card,text="Yeni Ajanda Notu",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=2,pady=(18,14))
        self.aj_g={}
        for i,(l,k,d) in enumerate([("Başlangıç Tarihi:","t",datetime.now().strftime("%d.%m.%Y")),
                                     ("Bitiş Tarihi:","t_bitis",""),
                                     ("Saat:","s","09:00"),("Başlık:","bas","")],1):
            tk.Label(card,text=l,bg=C_WHITE,font=("Segoe UI",10)).grid(row=i,column=0,padx=(28,6),pady=8,sticky="e")
            e=ttk.Entry(card,width=28); e.insert(0,d); e.grid(row=i,column=1,padx=(0,28),sticky="w"); self.aj_g[k]=e
        tk.Label(card,text="",bg=C_WHITE,fg="#888",font=("Segoe UI",8)).grid(row=2,column=1,sticky="w",padx=(0,28))
        lbl_hint=tk.Label(card,text="💡 Boş bırakırsanız tek gün kaydedilir.\nÖrn: 5 günlük izin için bitiş tarihi girin.",bg=C_WHITE,fg="#888",font=("Segoe UI",8),justify="left")
        lbl_hint.grid(row=2,column=1,sticky="se",padx=(0,28))
        tk.Label(card,text="Tür:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=5,column=0,padx=(28,6),pady=8,sticky="e")
        self.aj_g["tur"]=ttk.Combobox(card,values=["Hatırlatıcı","Toplantı","Denetim","Randevu","İzin","Saha Ziyareti","Diğer"],state="readonly",width=26)
        self.aj_g["tur"].set("Hatırlatıcı"); self.aj_g["tur"].grid(row=5,column=1,padx=(0,28),sticky="w")
        tk.Label(card,text="İçerik:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=6,column=0,padx=(28,6),pady=8,sticky="ne")
        self.aj_g["ic"]=tk.Text(card,width=28,height=4,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        self.aj_g["ic"].grid(row=6,column=1,padx=(0,28),pady=8,sticky="w")
        MBtn(card,"💾  Kaydet",command=self._aj_kaydet,color=self.gc("acc"),width=28).grid(row=7,column=0,columnspan=2,pady=16)
        self._yenile_aj()

    def _yenile_aj(self):
        if not DB_PATH: return
        try:
            with sqlite3.connect(DB_PATH) as c: self._all_aj=c.execute("SELECT id,tarih,sure,baslik,tur,durum,icerik FROM Ajanda WHERE k_adi=? ORDER BY tarih DESC,sure",(self.u_id,)).fetchall()
            self._filtre_aj("")
        except Exception as e: logging.error(f"yenile_aj:{e}")

    def _filtre_aj(self,a):
        self.tv_aj.delete(*self.tv_aj.get_children()); a=a.lower(); bugun=datetime.now().strftime("%d.%m.%Y")
        for r in getattr(self,"_all_aj",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            d=str(r[5]); tag="tamamlandi" if d=="Tamamlandı" else "gecmis" if r[1]<bugun else "bekliyor"
            self.tv_aj.insert("","end",values=r,tags=(tag,))

    def _aj_kaydet(self):
        bas=self.aj_g["bas"].get().strip()
        if not bas: messagebox.showwarning("Uyarı","Başlık zorunlu."); return
        ic=self.aj_g["ic"].get("1.0",tk.END).strip()
        tarih_bas=self.aj_g["t"].get().strip()
        tarih_bit=self.aj_g["t_bitis"].get().strip()
        saat=self.aj_g["s"].get()
        tur=self.aj_g["tur"].get()
        if tarih_bit:
            # Tarih aralığı — her gün için ayrı kayıt
            try:
                d_bas=datetime.strptime(tarih_bas,"%d.%m.%Y")
                d_bit=datetime.strptime(tarih_bit,"%d.%m.%Y")
                if d_bit<d_bas:
                    messagebox.showerror("Hata","Bitiş tarihi başlangıçtan önce olamaz."); return
                gun_sayisi=(d_bit-d_bas).days+1
                if gun_sayisi>60:
                    messagebox.showerror("Hata","En fazla 60 günlük aralık girilebilir."); return
                with sqlite3.connect(DB_PATH) as c:
                    for i in range(gun_sayisi):
                        gun=(d_bas+timedelta(days=i)).strftime("%d.%m.%Y")
                        etiket=f"{bas} ({i+1}/{gun_sayisi})" if gun_sayisi>1 else bas
                        c.execute("INSERT INTO Ajanda(tarih,sure,baslik,tur,durum,icerik,k_adi)VALUES(?,?,?,?,?,?,?)",
                            (gun,saat,etiket,tur,"Bekliyor",ic,self.u_id))
                self._yenile_aj()
                messagebox.showinfo("Tamam",f"{gun_sayisi} günlük kayıt oluşturuldu.\n{tarih_bas} — {tarih_bit}")
            except ValueError:
                messagebox.showerror("Hata","Tarih formatı hatalı.\nDoğru format: 01.04.2026"); return
        else:
            # Tek gün
            with sqlite3.connect(DB_PATH) as c:
                c.execute("INSERT INTO Ajanda(tarih,sure,baslik,tur,durum,icerik,k_adi)VALUES(?,?,?,?,?,?,?)",
                    (tarih_bas,saat,bas,tur,"Bekliyor",ic,self.u_id))
            self._yenile_aj()

    def _aj_dur(self,d):
        sel=self.tv_aj.selection()
        if not sel: return
        aid=self.tv_aj.item(sel[0])["values"][0]
        with sqlite3.connect(DB_PATH) as c: c.execute("UPDATE Ajanda SET durum=? WHERE id=?",(d,aid))
        self._yenile_aj()

    def _aj_sil(self):
        sel=self.tv_aj.selection()
        if not sel: return
        aid=self.tv_aj.item(sel[0])["values"][0]
        if messagebox.askyesno("Onay","Silinsin mi?"):
            with sqlite3.connect(DB_PATH) as c: c.execute("DELETE FROM Ajanda WHERE id=?",(aid,))
            self._yenile_aj()

    # ═══ 11. PERSONEL TAKVİMİ ════════════════════════════════════════════════
    def _t_personel(self,p):
        bg=self.gc("bg"); f=tk.Frame(p,bg=bg); f.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(f,text="👨‍💼 Personel Takvimi",font=("Segoe UI",15,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        self.tv_personel=self._tv(f,[("id",40,"#"),("ad",150,"Personel"),("tarih",90,"Tarih"),("bas",70,"Başlangıç"),("bit",70,"Bitiş"),("tur",100,"Tür"),("acik",260,"Açıklama")],16)
        bf=tk.Frame(f,bg=bg); bf.pack(pady=8)
        MBtn(bf,"🔄 Yenile",command=self._yenile_personel,color=C_INFO,width=14).pack(side="left",padx=4)
        if self.u_yetki!="İzleyici":
            MBtn(bf,"➕ Yeni Görev",command=self._personel_ekle_popup,color=self.gc("acc"),width=14).pack(side="left",padx=4)
            MBtn(bf,"🗑 Sil",command=self._personel_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        self._yenile_personel()

    def _yenile_personel(self):
        if not DB_PATH: return
        try:
            with sqlite3.connect(DB_PATH) as c:
                rows=c.execute("SELECT id,ad,tarih,baslangic,bitis,tur,aciklama FROM Personel_Takvim ORDER BY tarih DESC").fetchall()
            self.tv_personel.delete(*self.tv_personel.get_children())
            for r in rows: self.tv_personel.insert("","end",values=r)
        except Exception as e: logging.error(f"yenile_personel:{e}")

    def _personel_ekle_popup(self):
        win=tk.Toplevel(self.root); win.title("Yeni Görev/İzin"); win.geometry("440x380"); win.configure(bg=C_WHITE); win.grab_set()
        tk.Label(win,text="Yeni Görev/İzin/Toplantı",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=14)
        al={}
        for l,k,d in [("Personel Adı:","ad",""),("Tarih:","tarih",datetime.now().strftime("%d.%m.%Y")),
                       ("Başlangıç:","bas","08:00"),("Bitiş:","bit","17:00")]:
            tk.Label(win,text=l,bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=24,pady=(6,0))
            e=ttk.Entry(win,width=28); e.insert(0,d); e.pack(padx=24); al[k]=e
        tk.Label(win,text="Tür:",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=24,pady=(6,0))
        al["tur"]=ttk.Combobox(win,values=["Görev","İzin","Toplantı","Denetim","Saha Ziyareti","Eğitim"],state="readonly",width=26)
        al["tur"].set("Görev"); al["tur"].pack(padx=24)
        tk.Label(win,text="Açıklama:",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=24,pady=(6,0))
        al["acik"]=tk.Text(win,width=30,height=3,font=("Segoe UI",10)); al["acik"].pack(padx=24)
        def _k():
            ad=al["ad"].get().strip()
            if not ad: messagebox.showwarning("Uyarı","Personel adı zorunlu.",parent=win); return
            acik=al["acik"].get("1.0",tk.END).strip()
            with sqlite3.connect(DB_PATH) as c:
                c.execute("INSERT INTO Personel_Takvim(k_adi,ad,tarih,baslangic,bitis,tur,aciklama)VALUES(?,?,?,?,?,?,?)",
                    (self.u_id,ad,al["tarih"].get(),al["bas"].get(),al["bit"].get(),al["tur"].get(),acik))
            self._yenile_personel(); win.destroy()
        MBtn(win,"💾 Kaydet",command=_k,width=20).pack(pady=12)

    def _personel_sil(self):
        sel=self.tv_personel.selection()
        if not sel: return
        pid=self.tv_personel.item(sel[0])["values"][0]
        if messagebox.askyesno("Onay","Silinsin mi?"):
            with sqlite3.connect(DB_PATH) as c: c.execute("DELETE FROM Personel_Takvim WHERE id=?",(pid,))
            self._yenile_personel()

    # ═══ 12. HARİTA ══════════════════════════════════════════════════════════
    # ═══ İHALE YERLERİ HAVUZU + MUHAMMEN BEDEL ════════════════════════════════
    def _t_ihale_yer(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2); bf=ttk.Frame(nb2)
        nb2.add(lf,text="  📍 İhale Yerleri  "); nb2.add(nf,text="  ➕ Yeni Yer Ekle  "); nb2.add(bf,text="  💰 Muhammen Bedel  ")
        # Liste
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(ll,text="📍 İhaleye Çıkan Mera Yerleri Havuzu",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,8))
        tk.Label(ll,text="Sezon başında bir kere girilir. İhale başvurusunda buradan seçilir.",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,8))
        AramaFrame(ll,self._filtre_iyer,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_iyer=self._tv(ll,[("id",40,"#"),("ilce",80,"İlçe"),("koy",120,"Köy"),("ada",60,"Ada"),
            ("parsel",60,"Parsel"),("alan",80,"Alan(da)"),("kap",90,"Kapasite"),("bedel",110,"Tah.Bedel"),("vasif",80,"Vasıf"),("yil",50,"Yıl")],12)
        bff=tk.Frame(ll,bg=bg); bff.pack(pady=8)
        MBtn(bff,"🔄 Yenile",command=self._yenile_iyer,color=C_INFO,width=14).pack(side="left",padx=4)
        if self.u_yetki!="İzleyici":
            MBtn(bff,"📥 Excel Import",command=self._iyer_import,color="#8E44AD",width=14).pack(side="left",padx=4)
            MBtn(bff,"📊 Excel Export",command=self._iyer_export,color=self.gc("acc"),width=14).pack(side="left",padx=4)
            MBtn(bff,"🗑 Sil",command=self._iyer_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        # Yeni yer
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=16)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.45,anchor="center",width=600,height=420)
        tk.Label(card,text="Yeni İhale Yeri Ekle",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=4,pady=(18,14))
        self.iy_g={}
        for lbl,key,tip,r,col in [("İlçe:","ilce","combo",1,0),("Köy:","koy","entry",1,2),
            ("Ada:","ada","entry",2,0),("Parsel:","parsel","entry",2,2),
            ("Alan (da):","alan","entry",3,0),("Vasıf:","vasif","vasif",3,2),
            ("Yıl:","yil","entry",4,0)]:
            tk.Label(card,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).grid(row=r,column=col,padx=(20,6),pady=8,sticky="e")
            if tip=="combo": w_obj=ttk.Combobox(card,values=ILCELER,state="readonly",width=18)
            elif tip=="vasif": w_obj=ttk.Combobox(card,values=MERA_VASIF,state="readonly",width=18)
            else: w_obj=ttk.Entry(card,width=20)
            if key=="yil" and isinstance(w_obj,ttk.Entry): w_obj.insert(0,_yil())
            w_obj.grid(row=r,column=col+1,padx=(0,14),pady=8,sticky="w"); self.iy_g[key]=w_obj
        MBtn(card,"✅ Kaydet",command=self._iyer_kaydet,color=self.gc("acc"),width=30).grid(row=5,column=0,columnspan=4,pady=18)
        # Muhammen Bedel
        bh=tk.Frame(bf,bg=bg); bh.pack(fill="both",expand=True,padx=30,pady=20)
        tk.Label(bh,text="💰 Muhammen Bedel Hesaplayıcı (Tahdit Raporu)",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")
        tk.Label(bh,text="Formül: Kapasite(BBHB) = Alan × Kuru Ot Verimi / (Otlatma Gün × 12,5 kg)\nMuhammen Bedel = Kapasite × Otlatma Gün × 12,5 × Kuru Ot Fiyatı",
            font=("Segoe UI",9),fg="#666",bg=bg,justify="left").pack(anchor="w",pady=(4,14))
        mf=tk.Frame(bh,bg=bg); mf.pack(fill="x")
        self.mb_g={}
        for i,(l,k,d,w) in enumerate([("Mera Alanı (da):","alan","",14),("Vasıf:","vasif","",14),
            ("Otlatma Gün:","gun","",8),("Kuru Ot Fiyatı (₺/kg):","fiyat","",10)]):
            tk.Label(mf,text=l,bg=bg,font=("Segoe UI",10)).grid(row=0,column=i*2,padx=(0,4),sticky="w")
            if k=="vasif":
                w_obj=ttk.Combobox(mf,values=MERA_VASIF,state="readonly",width=w)
                w_obj.bind("<<ComboboxSelected>>",self._mb_vasif_sec)
            else: w_obj=ttk.Entry(mf,width=w)
            w_obj.grid(row=0,column=i*2+1,padx=(0,10)); self.mb_g[k]=w_obj
        tk.Label(bh,text="Kuru Ot Verimi (kg/da):",bg=bg,font=("Segoe UI",10,"bold")).pack(anchor="w",pady=(10,2))
        self.mb_kov=tk.Label(bh,text="— vasıf seçin —",bg=bg,font=("Segoe UI",12),fg=self.gc("pri"))
        self.mb_kov.pack(anchor="w")
        MBtn(bh,"💰 Hesapla",command=self._mb_hesapla,color=self.gc("acc"),width=18).pack(anchor="w",pady=10)
        self.lbl_mb_sonuc=tk.Label(bh,text="",bg=bg,font=("Segoe UI",11),fg="#333",justify="left")
        self.lbl_mb_sonuc.pack(anchor="w",pady=8)
        self._yenile_iyer()

    def _mb_vasif_sec(self,event=None):
        v=self.mb_g["vasif"].get()
        kov=OT_VERIM_KURU.get(v,"—")
        self.mb_kov.config(text=f"{kov} kg/da (Ardahan 500-650mm yağış kuşağı)")

    def _mb_hesapla(self):
        v=self.mb_g["vasif"].get()
        if not v: messagebox.showerror("Hata","Vasıf seçin."); return
        kov=OT_VERIM_KURU.get(v)
        sonuc=muhammen_bedel_hesapla(self.mb_g["alan"].get(),kov,self.mb_g["gun"].get(),self.mb_g["fiyat"].get())
        if not sonuc: messagebox.showerror("Hata","Tüm alanları doğru doldurun."); return
        self.lbl_mb_sonuc.config(text=
            f"📊 Mera Alanı: {sonuc['alan']:,.0f} da | Kuru Ot: {sonuc['kuru_ot']:.0f} kg/da\n"
            f"⏱ Otlatma: {sonuc['gun']:.0f} gün | Fiyat: {sonuc['fiyat']:.2f} ₺/kg\n\n"
            f"🐄 Otlatma Kapasitesi: {sonuc['bbhb']:,.1f} BBHB ({sonuc['kbhb']:,.0f} KBHB)\n\n"
            f"💰 MUHAMMEN BEDEL: {para_format(sonuc['bedel'])}\n"
            f"{'─'*40}\n"
            f"IBAN (Emanet): TR 85 0001 0001 4900 0010 0059 60")

    def _yenile_iyer(self):
        if not DB_PATH: return
        try:
            with sqlite3.connect(DB_PATH) as c:
                self._all_iyer=c.execute("SELECT id,ilce,koy,ada,parsel,alan_da,kapasite_bbhb,tahmini_bedel,vasif,yil FROM Ihale_Yerleri ORDER BY yil DESC,ilce,koy").fetchall()
            self._filtre_iyer("")
        except Exception as e: logging.error(f"yenile_iyer:{e}")

    def _filtre_iyer(self,a):
        self.tv_iyer.delete(*self.tv_iyer.get_children()); a=a.lower()
        for r in getattr(self,"_all_iyer",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            row=list(r)
            try: row[7]=para_format(row[7])
            except: pass
            self.tv_iyer.insert("","end",values=row)

    def _iyer_kaydet(self):
        try:
            koy=self.iy_g["koy"].get().strip()
            if not koy: raise ValueError("Köy zorunlu.")
            alan=float(self.iy_g["alan"].get() or 0)
            vasif=self.iy_g["vasif"].get()
            kov=OT_VERIM_KURU.get(vasif,90)
            kap=(alan*kov)/(135*GUNLUK_OT_BBHB) if alan>0 else 0
            with sqlite3.connect(DB_PATH) as c:
                c.execute("INSERT INTO Ihale_Yerleri(ilce,koy,ada,parsel,alan_da,kapasite_bbhb,vasif,yil)VALUES(?,?,?,?,?,?,?,?)",
                    (self.iy_g["ilce"].get(),koy,self.iy_g["ada"].get(),self.iy_g["parsel"].get(),alan,kap,vasif,self.iy_g["yil"].get()))
            db_log(self.u_id,"İhale Yeri Ekle",koy)
            for k,w in self.iy_g.items():
                if isinstance(w,ttk.Entry): w.delete(0,tk.END)
                elif isinstance(w,ttk.Combobox): w.set("")
            self.iy_g["yil"].insert(0,_yil()); self._yenile_iyer()
            messagebox.showinfo("Tamam",f"'{koy}' ihale yeri kaydedildi.\nKapasite: {kap:.1f} BBHB")
        except ValueError as e: messagebox.showerror("Hata",str(e))

    def _iyer_import(self):
        yol=filedialog.askopenfilename(filetypes=[("Excel","*.xlsx")])
        if not yol: return
        try:
            df=pd.read_excel(yol)
            with sqlite3.connect(DB_PATH) as c:
                n=0
                for _,r in df.iterrows():
                    v=r.tolist()
                    if len(v)>=5:
                        c.execute("INSERT INTO Ihale_Yerleri(ilce,koy,ada,parsel,alan_da,vasif,yil)VALUES(?,?,?,?,?,?,?)",
                            (str(v[0]),str(v[1]),str(v[2]),str(v[3]),float(v[4] or 0),str(v[5]) if len(v)>5 else "",_yil()))
                        n+=1
            self._yenile_iyer(); messagebox.showinfo("Tamam",f"{n} ihale yeri yüklendi.")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _iyer_export(self):
        if not getattr(self,"_all_iyer",None): return
        yol=filedialog.asksaveasfilename(defaultextension=".xlsx",filetypes=[("Excel","*.xlsx")])
        if yol:
            try:
                pd.DataFrame(self._all_iyer,columns=["ID","İlçe","Köy","Ada","Parsel","Alan(da)","Kapasite","Tah.Bedel","Vasıf","Yıl"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"Dışa aktarıldı:\n{yol}")
            except Exception as e: messagebox.showerror("Hata",str(e))

    def _iyer_sil(self):
        sel=self.tv_iyer.selection()
        if not sel: return
        if messagebox.askyesno("Onay","Silinsin mi?"):
            with sqlite3.connect(DB_PATH) as c: c.execute("DELETE FROM Ihale_Yerleri WHERE id=?",(self.tv_iyer.item(sel[0])["values"][0],))
            self._yenile_iyer()

    # ═══ 13. İSTATİSTİK + İHALE/CEZA ÖZETİ ═════════════════════════════════
    def _t_istatistik(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        gf=ttk.Frame(nb2); ihf=ttk.Frame(nb2); czf=ttk.Frame(nb2)
        nb2.add(gf,text="  📈 Genel  "); nb2.add(ihf,text="  ⚖️ İhale Özeti  "); nb2.add(czf,text="  💰 Ceza Özeti  ")
        # Genel
        f=tk.Frame(gf,bg=bg); f.pack(fill="both",expand=True,padx=14,pady=12)
        kf=tk.Frame(f,bg=bg); kf.pack(fill="x",pady=(0,10))
        self.ist_k={}
        for bas,key,renk,ikon in [("Toplam Köy","koy",self.gc("pri"),"🏘️"),("Toplam Rapor","rapor",C_INFO,"📄"),
            ("Aktif İhale","ihale",C_WARN,"⚖️"),("İdari Ceza","ceza",C_DANGER,"💰")]:
            k=StatKart(kf,bas,"—",renk,ikon); k.pack(side="left",fill="both",expand=True,padx=4)
            self.ist_k[key]=k
        if MPL_OK:
            self.ist_canvas_frame=tk.Frame(f,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
            self.ist_canvas_frame.pack(fill="both",expand=True)
        MBtn(f,"📊 Yenile",command=self._ist_ozet,color=self.gc("acc"),width=14).pack(pady=6)
        # İhale Özeti (yıl bazlı)
        ihh=tk.Frame(ihf,bg=bg); ihh.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(ihh,text="⚖️ İhale Özeti (Yıl Bazlı)",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        self.tv_ih_ozet=self._tv(ihh,[("yil",60,"Yıl"),("toplam",70,"Toplam"),("aktif",70,"Aktif"),
            ("tamam",80,"Tamamlanan"),("gelir",120,"Toplam Gelir"),("bakanlik",120,"Bakanlık(%25)"),
            ("koy",120,"Köy(%75)"),("damga",100,"Damga V.")],10)
        MBtn(ihh,"🔄 Yenile",command=self._ihale_ozet_yenile,color=C_INFO,width=14).pack(pady=6)
        # Ceza Özeti (yıl bazlı)
        czz=tk.Frame(czf,bg=bg); czz.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(czz,text="💰 Ceza Özeti (Yıl Bazlı)",font=("Segoe UI",14,"bold"),fg=C_DANGER,bg=bg).pack(anchor="w",pady=(0,10))
        self.tv_cz_ozet=self._tv(czz,[("yil",60,"Yıl"),("sayi",70,"Kişi"),("toplam",120,"Toplam Tutar"),
            ("ilce",140,"İlçe Dağılımı"),("konu",200,"Konu Dağılımı"),("mukerrer",80,"Mükerrer")],10)
        MBtn(czz,"🔄 Yenile",command=self._ceza_ozet_yenile,color=C_INFO,width=14).pack(pady=6)
        self._ist_ozet(); self._ihale_ozet_yenile(); self._ceza_ozet_yenile()

    def _ihale_ozet_yenile(self):
        if not DB_PATH: return
        self.tv_ih_ozet.delete(*self.tv_ih_ozet.get_children())
        try:
            with sqlite3.connect(DB_PATH) as c:
                yillar=c.execute("SELECT DISTINCT substr(tarih,-4) as y FROM Ihaleler ORDER BY y DESC").fetchall()
                for (yil,) in yillar:
                    if not yil: continue
                    toplam=c.execute("SELECT COUNT(*) FROM Ihaleler WHERE tarih LIKE ?",(f"%{yil}",)).fetchone()[0]
                    aktif=c.execute("SELECT COUNT(*) FROM Ihaleler WHERE tarih LIKE ? AND durum NOT IN('Tamamlandı','İptal Edildi')",(f"%{yil}",)).fetchone()[0]
                    tamam=c.execute("SELECT COUNT(*) FROM Ihaleler WHERE tarih LIKE ? AND durum='Tamamlandı'",(f"%{yil}",)).fetchone()[0]
                    gelir=c.execute("SELECT COALESCE(SUM(bedel),0) FROM Ihaleler WHERE tarih LIKE ? AND durum='Tamamlandı'",(f"%{yil}",)).fetchone()[0]
                    self.tv_ih_ozet.insert("","end",values=(yil,toplam,aktif,tamam,
                        para_format(gelir),para_format(gelir*0.25),para_format(gelir*0.75),para_format(gelir*5.69/1000)))
        except Exception as e: logging.error(f"ihale_ozet:{e}")

    def _ceza_ozet_yenile(self):
        if not DB_PATH: return
        self.tv_cz_ozet.delete(*self.tv_cz_ozet.get_children())
        try:
            with sqlite3.connect(DB_PATH) as c:
                yillar=c.execute("SELECT DISTINCT yil FROM Idari_Cezalar ORDER BY yil DESC").fetchall()
                for (yil,) in yillar:
                    if not yil: continue
                    sayi=c.execute("SELECT COUNT(DISTINCT tc) FROM Idari_Cezalar WHERE yil=?",(yil,)).fetchone()[0]
                    toplam=c.execute("SELECT COALESCE(SUM(ipc_tutari),0) FROM Idari_Cezalar WHERE yil=?",(yil,)).fetchone()[0]
                    ilceler=c.execute("SELECT ilce||'('||COUNT(*)||')' FROM Idari_Cezalar WHERE yil=? GROUP BY ilce",(yil,)).fetchall()
                    konular=c.execute("SELECT konu||'('||COUNT(*)||')' FROM Idari_Cezalar WHERE yil=? GROUP BY konu",(yil,)).fetchall()
                    mukerrer=c.execute("SELECT COUNT(*) FROM (SELECT tc FROM Idari_Cezalar WHERE yil=? GROUP BY tc HAVING COUNT(*)>1)",(yil,)).fetchone()[0]
                    self.tv_cz_ozet.insert("","end",values=(yil,sayi,para_format(toplam),
                        ", ".join(r[0] for r in ilceler),", ".join(r[0] for r in konular),mukerrer))
        except Exception as e: logging.error(f"ceza_ozet:{e}")

    def _ist_ozet(self):
        if not DB_PATH or not MPL_OK: return
        try:
            with sqlite3.connect(DB_PATH) as c:
                veri={
                    "koy":c.execute("SELECT COUNT(*) FROM Mera_Varligi").fetchone()[0],
                    "rapor":c.execute("SELECT COUNT(*) FROM Rapor_Gecmisi").fetchone()[0],
                    "ihale":c.execute("SELECT COUNT(*) FROM Ihaleler WHERE durum NOT IN('Tamamlandı','İptal Edildi')").fetchone()[0],
                    "ceza":c.execute("SELECT COUNT(*) FROM Idari_Cezalar").fetchone()[0],
                }
        except Exception: return
        for key,k in self.ist_k.items(): k.set(veri.get(key,"—"))
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        for w in self.ist_canvas_frame.winfo_children(): w.destroy()
        fig,axes=plt.subplots(1,2,figsize=(12,4),facecolor=self.gc("bg"))
        fig.suptitle(f"MİRAS Genel Özet — {datetime.now().strftime('%d.%m.%Y')}",fontsize=12,fontweight="bold",color=self.gc("pri"))
        ax1=axes[0]
        labels=["Köy","Rapor","İhale","Ceza"]
        vals=[veri.get("koy",0),veri.get("rapor",0),veri.get("ihale",0),veri.get("ceza",0)]
        renkler=["#1E5631","#2980B9","#D68910","#C0392B"]
        ax1.bar(labels,[max(1,v) for v in vals],color=renkler)
        ax1.set_title("Genel Durum",color=self.gc("pri"),fontsize=10)
        ax2=axes[1]
        try:
            with sqlite3.connect(DB_PATH) as c:
                rows=c.execute("SELECT ilce,COUNT(*) FROM Mera_Varligi GROUP BY ilce ORDER BY COUNT(*) DESC").fetchall()
            if rows:
                ilceler=[r[0] for r in rows]; sayilar=[r[1] for r in rows]
                ax2.barh(ilceler,sayilar,color=self.gc("acc"))
                ax2.set_title("İlçe Bazlı Köy Dağılımı",color=self.gc("pri"),fontsize=10)
        except Exception: pass
        fig.tight_layout()
        canvas=FigureCanvasTkAgg(fig,master=self.ist_canvas_frame)
        canvas.draw(); canvas.get_tk_widget().pack(fill="both",expand=True)
        plt.close(fig)

    # ═══ 14. EVRAK ÜRETİCİ ══════════════════════════════════════════════════
    def _t_evrak(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        bf=ttk.Frame(nb2); kf=ttk.Frame(nb2); cf=ttk.Frame(nb2)
        nb2.add(bf,text="  📋 Bilgi Notu  "); nb2.add(kf,text="  📑 Katılım Evrakları  "); nb2.add(cf,text="  💰 Ceza Oluru  ")
        # Bilgi Notu
        bn=tk.Frame(bf,bg=bg); bn.pack(fill="both",expand=True,padx=20,pady=16)
        tk.Label(bn,text="📋 Bilgi Notu Oluşturucu",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        tk.Label(bn,text="Bilgi notunun bilgilerini girin, Word dosyası oluşturulsun.",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,8))
        self.bn_g={}
        for l,k,d in [("Sunulan Makam:","makam",""),("Konu:","konu",""),("İçerik:","icerik","")]:
            tk.Label(bn,text=l,bg=bg,font=("Segoe UI",10)).pack(anchor="w",pady=(6,0))
            if k=="icerik":
                w_obj=tk.Text(bn,height=10,font=("Segoe UI",10),wrap="word",relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=10,pady=8)
            else:
                w_obj=ttk.Entry(bn,width=60); w_obj.insert(0,d)
            w_obj.pack(fill="x",pady=2); self.bn_g[k]=w_obj
        MBtn(bn,"📄 Word Olarak Kaydet",command=self._bilgi_notu_kaydet,color=self.gc("acc"),width=24).pack(anchor="w",pady=12)
        # Katılım Evrakları
        ke=tk.Frame(kf,bg=bg); ke.pack(fill="both",expand=True,padx=20,pady=16)
        tk.Label(ke,text="📑 İhale Katılım Evrakları Şablonu",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        tk.Label(ke,text="Başvuran bilgilerini girin, katılım evrakları Word dosyası oluşturulsun.",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,8))
        self.ke_g={}
        for l,k in [("Ad Soyad:","ad_soyad"),("TC No:","tc"),("Mera Adı:","mera"),("Ada/Parsel:","ada_parsel")]:
            tk.Label(ke,text=l,bg=bg,font=("Segoe UI",10)).pack(anchor="w",pady=(6,0))
            e=ttk.Entry(ke,width=40); e.pack(anchor="w"); self.ke_g[k]=e
        MBtn(ke,"📄 Word Olarak Kaydet",command=self._katilim_evrak_kaydet,color=self.gc("acc"),width=24).pack(anchor="w",pady=12)
        # Ceza Oluru
        co=tk.Frame(cf,bg=bg); co.pack(fill="both",expand=True,padx=20,pady=16)
        tk.Label(co,text="💰 İdari Para Cezası Oluru Şablonu",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        tk.Label(co,text="Ceza bilgilerini girin, olur Word dosyası oluşturulsun.",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,8))
        self.co_g={}
        for l,k in [("Ad Soyad:","ad_soyad"),("TC No:","tc"),("Mera/Köy:","mera"),("İhlal Konusu:","konu"),("Ceza Tutarı (₺):","tutar")]:
            tk.Label(co,text=l,bg=bg,font=("Segoe UI",10)).pack(anchor="w",pady=(6,0))
            e=ttk.Entry(co,width=40); e.pack(anchor="w"); self.co_g[k]=e
        MBtn(co,"📄 Word Olarak Kaydet",command=self._ceza_oluru_kaydet,color=C_DANGER,width=24).pack(anchor="w",pady=12)
        tk.Label(p,text="💡 AI ile otomatik evrak üretmek için: Sidebar → 🌿 MERA AI → AI Evrak Üret sekmesi",
            font=("Segoe UI",9),fg="#888",bg=bg).pack(anchor="w",padx=14,pady=(0,4))

    def _katilim_evrak_kaydet(self):
        if not DOCX_OK: messagebox.showerror("Hata","python-docx kurulu değil."); return
        dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile="Katilim_Evraklari.docx",filetypes=[("Word","*.docx")])
        if not dosya: return
        try:
            veri={k:e.get() for k,e in self.ke_g.items()}
            word_katilim_evrak(dosya,veri)
            messagebox.showinfo("Tamam",f"Evrak üretildi:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _ceza_oluru_kaydet(self):
        if not DOCX_OK: messagebox.showerror("Hata","python-docx kurulu değil."); return
        dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile="Ceza_Oluru.docx",filetypes=[("Word","*.docx")])
        if not dosya: return
        try:
            veri={k:e.get() for k,e in self.co_g.items()}
            veri["tarih"]=datetime.now().strftime("%d.%m.%Y"); veri["baba_adi"]=""; veri["dogum"]=""; veri["adres"]=""
            word_idari_ceza(dosya,veri)
            messagebox.showinfo("Tamam",f"Evrak üretildi:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _evrak_uret_ai(self):
        self.btn_ev.loading(True)
        prompt=(f"T.C. Ardahan Valiliği İl Tarım ve Orman Müdürlüğü adına resmi {self.ev_tur.get()} hazırla.\n"
               f"KONU: {self.ev_kon.get().strip()}\nTARİH: {datetime.now().strftime('%d.%m.%Y')}\nHAZIRLAYAN: {self.u_ad}\n"
               f"{'TALİMAT: '+self.ev_tal.get().strip() if self.ev_tal.get().strip() else ''}\n"
               "Türk resmi yazışma kurallarına uy.")
        def _g(t): self.ev_yan.config(state="normal"); self.ev_yan.delete("1.0",tk.END); self.ev_yan.insert(tk.END,t); self.ev_yan.config(state="disabled")
        _g("⏳ MERA AI hazırlıyor...\nKota aşımında otomatik yeniden denenecek, lütfen bekleyin.")
        def _bg():
            r=self.ai.tek(prompt)
            self.root.after(0,lambda:(_g(r),self.btn_ev.loading(False)))
        threading.Thread(target=_bg,daemon=True).start()

    def _bilgi_notu_kaydet(self):
        if not DOCX_OK: messagebox.showerror("Hata","python-docx kurulu değil."); return
        dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile="Bilgi_Notu.docx",filetypes=[("Word","*.docx")])
        if not dosya: return
        icerik=self.bn_g["icerik"].get("1.0",tk.END).strip() if isinstance(self.bn_g["icerik"],tk.Text) else self.bn_g["icerik"].get()
        try:
            word_bilgi_notu(dosya,{"makam":self.bn_g["makam"].get(),"konu":self.bn_g["konu"].get(),"icerik":icerik,"hazirlayan":self.u_ad})
            messagebox.showinfo("Tamam",f"Bilgi notu üretildi:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    # ═══ 15. MERA AI ═════════════════════════════════════════════════════════
    def _t_ai(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        chat_f=ttk.Frame(nb2); evrak_f=ttk.Frame(nb2); key_f=ttk.Frame(nb2)
        nb2.add(chat_f,text="  💬 Sohbet  "); nb2.add(evrak_f,text="  📝 AI Evrak Üret  ")
        if self.u_yetki=="Admin":
            nb2.add(key_f,text="  🔑 API Ayarları  ")
        # ── Durum bandı ──
        durum_renk="#2D8C55" if self.ai.hazir else "#C0392B"
        durum_txt="✅ MERA AI Aktif — Gemini bağlı" if self.ai.hazir else "❌ MERA AI Pasif — API Key gerekli"
        # ── 1) SOHBET ──
        cc=tk.Frame(chat_f,bg=bg); cc.pack(fill="both",expand=True,padx=14,pady=12)
        df=tk.Frame(cc,bg=durum_renk,height=30); df.pack(fill="x",pady=(0,8)); df.pack_propagate(False)
        tk.Label(df,text=durum_txt,fg=C_WHITE,bg=durum_renk,font=("Segoe UI",9,"bold")).pack(side="left",padx=10)
        if not self.ai.hazir and self.u_yetki=="Admin":
            tk.Button(df,text="🔑 API Key Gir →",command=lambda:nb2.select(key_f),fg=C_WHITE,bg=durum_renk,relief="flat",font=("Segoe UI",9,"bold"),cursor="hand2").pack(side="right",padx=10)
        self.chat_alan=scrolledtext.ScrolledText(cc,font=("Segoe UI",10),wrap="word",bg=C_WHITE,relief="flat",bd=1,
            highlightbackground="#D0DDD8",highlightthickness=1,padx=14,pady=12,state="disabled")
        self.chat_alan.pack(fill="both",expand=True)
        self.chat_alan.tag_configure("siz",foreground="#1A5276",font=("Segoe UI",10,"bold"))
        self.chat_alan.tag_configure("ai",foreground=self.gc("pri"))
        self.chat_alan.tag_configure("sis",foreground="#999",font=("Segoe UI",8,"italic"))
        hf=tk.Frame(cc,bg=bg); hf.pack(fill="x",pady=6)
        for s in ["Otlatma kapasitesi?","Tespit tahdit nedir?","Madde 14 süreci?","İhale şartları?","BBHB hesaplama?"]:
            tk.Button(hf,text=s,command=lambda q=s:self._chat_hizli(q),bg="#E8F5E9",fg=self.gc("pri"),font=("Segoe UI",8),relief="flat",cursor="hand2",padx=8,pady=3).pack(side="left",padx=4)
        gf=tk.Frame(cc,bg=bg); gf.pack(fill="x",pady=(6,0))
        self.chat_txt=tk.Text(gf,height=3,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#A9DFBF",highlightthickness=1,padx=10,pady=8)
        self.chat_txt.pack(side="left",fill="x",expand=True,padx=(0,10))
        self.chat_txt.bind("<Return>",lambda e:(self._chat_gonder(),"break")[1])
        sg=tk.Frame(gf,bg=bg); sg.pack(side="left")
        self.btn_cg=MBtn(sg,"➤ Gönder",command=self._chat_gonder,color=self.gc("pri"),width=12); self.btn_cg.pack(pady=2)
        MBtn(sg,"🗑 Temizle",command=self._chat_temizle,color="#888",width=12).pack(pady=2)
        self._chat_yaz("sis","🌿 MERA AI'ye Hoş Geldiniz!\nMera mevzuatı, BBHB hesaplama, ihale ve tahsis süreçleri hakkında sorular sorabilirsiniz.\n"+"─"*50+"\n")
        if not self.ai.hazir:
            self._chat_yaz("sis","⚠️ AI bağlı değil. Üstteki '🔑 API Key Gir →' butonuna tıklayın.\n"
                "veya MERA AI → API Ayarları sekmesinden key girin.\n"
                "Ücretsiz key almak için: aistudio.google.com/app/apikey\n"+"─"*50+"\n")
        # ── 2) AI EVRAK ÜRET ──
        ee=tk.Frame(evrak_f,bg=bg); ee.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(ee,text="✨ Yapay Zeka ile Otomatik Evrak Üretici",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,8))
        tk.Label(ee,text="Evrak türünü ve konuyu seçin, MERA AI resmi formatta hazırlasın.",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,10))
        ctrl=tk.Frame(ee,bg=bg); ctrl.pack(fill="x",pady=(0,10))
        for i,(l,attr,vals,d,w) in enumerate([
            ("Evrak Türü:","ev_tur",["Bilgi Notu","Makam Oluru","Mera Tespit Tutanağı","Otlatma Yasağı Kararı","Islah Proje Talebi","Şikayet Değerlendirme","İhale Duyurusu"],"Bilgi Notu",24),
            ("Konu/Yer:","ev_kon",None,"",28),("Ek Talimat:","ev_tal",None,"",20)]):
            tk.Label(ctrl,text=l,bg=bg,font=("Segoe UI",10)).grid(row=0,column=i*2,padx=(0,4),sticky="w")
            w_obj=(ttk.Combobox(ctrl,values=vals,state="readonly",width=w) if vals else ttk.Entry(ctrl,width=w))
            if vals: w_obj.set(d)
            w_obj.grid(row=0,column=i*2+1,padx=(0,10)); setattr(self,attr,w_obj)
        self.btn_ev=MBtn(ctrl,"✨ MERA AI Üret",command=self._evrak_uret_ai,color=self.gc("pri")); self.btn_ev.grid(row=0,column=6,padx=10)
        self.ev_yan=scrolledtext.ScrolledText(ee,font=("Segoe UI",10),wrap="word",bg="#F9F9FF",relief="flat",bd=1,highlightbackground="#C8B4E8",highlightthickness=1,padx=14,pady=12,state="disabled")
        self.ev_yan.pack(fill="both",expand=True)
        bf_ev=tk.Frame(ee,bg=bg); bf_ev.pack(pady=6)
        MBtn(bf_ev,"📋 Kopyala",command=lambda:(self.root.clipboard_clear(),self.root.clipboard_append(self.ev_yan.get("1.0",tk.END))),color="#555",width=14).pack(side="left",padx=4)
        MBtn(bf_ev,"📄 Word Olarak Kaydet",command=self._ai_evrak_word,color=self.gc("acc"),width=20).pack(side="left",padx=4)
        # ── 3) API AYARLARI (sadece Admin) ──
        if self.u_yetki=="Admin":
            aa=tk.Frame(key_f,bg=bg); aa.pack(fill="both",expand=True,padx=40,pady=30)
            tk.Label(aa,text="🔑 MERA AI — Gemini API Anahtarı",font=("Segoe UI",16,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")
            tk.Label(aa,text="Google AI Studio'dan ücretsiz API key alın ve aşağıya yapıştırın.",font=("Segoe UI",10),fg="#555",bg=bg).pack(anchor="w",pady=(6,4))
            # Adım adım rehber
            rehber_f=tk.Frame(aa,bg="#F0F4F2",highlightbackground="#D0DDD8",highlightthickness=1,padx=16,pady=12)
            rehber_f.pack(fill="x",pady=(8,16))
            tk.Label(rehber_f,text="📌 API Key Nasıl Alınır?",font=("Segoe UI",11,"bold"),fg=self.gc("pri"),bg="#F0F4F2").pack(anchor="w")
            for adim in ["1️⃣  aistudio.google.com/app/apikey adresine gidin",
                         "2️⃣  Google hesabınızla giriş yapın",
                         "3️⃣  'Create API Key' butonuna tıklayın",
                         "4️⃣  Oluşan anahtarı kopyalayıp aşağıya yapıştırın"]:
                tk.Label(rehber_f,text=adim,font=("Segoe UI",10),fg="#333",bg="#F0F4F2").pack(anchor="w",pady=2)
            tk.Label(aa,text="API Anahtarı:",bg=bg,font=("Segoe UI",11,"bold")).pack(anchor="w",pady=(0,4))
            self.ai_key_e=ttk.Entry(aa,width=60,show="•",font=("Segoe UI",11))
            if self.ai.api_key: self.ai_key_e.insert(0,self.ai.api_key)
            self.ai_key_e.pack(anchor="w",pady=6,fill="x")
            # Durum göstergesi
            dur_f=tk.Frame(aa,bg=bg); dur_f.pack(anchor="w",pady=(4,12))
            self.lbl_ai_durum=tk.Label(dur_f,text=durum_txt,font=("Segoe UI",10,"bold"),fg=durum_renk,bg=bg)
            self.lbl_ai_durum.pack(side="left")
            def _kk():
                k=self.ai_key_e.get().strip()
                if not k: messagebox.showwarning("Uyarı","API anahtarı boş."); return
                self.ai.key_kaydet(k)
                self.lbl_ai_durum.config(text="⏳ Bağlanıyor...",fg="#D68910")
                def _kontrol():
                    if self.ai.hazir:
                        self.lbl_ai_durum.config(text="✅ MERA AI Aktif — Gemini bağlı!",fg="#2D8C55")
                        messagebox.showinfo("MERA AI","✅ Başarıyla aktifleştirildi!\n\nArtık Sohbet ve AI Evrak Üret sekmelerini kullanabilirsiniz.")
                    else:
                        self.lbl_ai_durum.config(text="❌ Bağlantı kurulamadı — anahtarı kontrol edin",fg="#C0392B")
                        messagebox.showwarning("MERA AI","Bağlantı kurulamadı.\nAPI anahtarını kontrol edin.")
                self.root.after(3000,_kontrol)
            MBtn(aa,"✅ Aktifleştir",command=_kk,color=self.gc("pri"),width=30).pack(anchor="w",pady=4)
            tk.Label(aa,text="💡 API anahtarı miras_gemini_key.txt dosyasına kaydedilir.\nTüm kullanıcılar aynı anahtarı kullanır.",font=("Segoe UI",9),fg="#888",bg=bg).pack(anchor="w",pady=(12,0))

    def _ai_evrak_word(self):
        """AI tarafından üretilen evrakı Word olarak kaydet"""
        icerik=self.ev_yan.get("1.0",tk.END).strip()
        if not icerik or "⏳" in icerik: messagebox.showwarning("Uyarı","Önce evrak üretin."); return
        if not DOCX_OK: messagebox.showerror("Hata","python-docx kurulu değil."); return
        dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile=f"AI_Evrak_{datetime.now().strftime('%Y%m%d')}.docx",filetypes=[("Word","*.docx")])
        if not dosya: return
        try:
            doc=DocxDocument()
            for satir in icerik.split("\n"):
                if satir.strip(): doc.add_paragraph(satir)
            doc.save(dosya)
            messagebox.showinfo("Tamam",f"Evrak kaydedildi:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _chat_yaz(self,tag,metin):
        self.chat_alan.config(state="normal"); self.chat_alan.insert(tk.END,metin+"\n",tag)
        self.chat_alan.see(tk.END); self.chat_alan.config(state="disabled")

    def _chat_hizli(self,s): self.chat_txt.delete("1.0",tk.END); self.chat_txt.insert("1.0",s); self._chat_gonder()

    def _chat_gonder(self):
        m=self.chat_txt.get("1.0",tk.END).strip()
        if not m: return
        self.chat_txt.delete("1.0",tk.END); self._chat_yaz("siz",f"👤 Siz:\n{m}"); self.btn_cg.loading(True)
        self._chat_yaz("sis","⏳ Yanıt bekleniyor... (kota aşımında otomatik yeniden denenecek)")
        def _bg():
            r=self.ai.sor(m)
            self.root.after(0,lambda:(self._chat_yaz("ai",f"🌿 MERA AI:\n{r}\n{'─'*40}"),self.btn_cg.loading(False)))
        threading.Thread(target=_bg,daemon=True).start()

    def _chat_temizle(self):
        self.ai.yeni_chat()
        self.chat_alan.config(state="normal"); self.chat_alan.delete("1.0",tk.END); self.chat_alan.config(state="disabled")

    # ═══ 16. MEVZUAT REHBERİ ═════════════════════════════════════════════════
    def _t_mevzuat(self,p):
        bg=self.gc("bg"); f=tk.Frame(p,bg=bg); f.pack(fill="both",expand=True,padx=20,pady=16)
        tk.Label(f,text="📖 Mevzuat Rehberi & Bilgi Kartları",font=("Segoe UI",15,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,4))
        tk.Label(f,text="4342 sayılı Mera Kanunu ve Mera Yönetmeliği kapsamında temel kavramlar",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,12))
        # Scrollable kartlar
        cw=tk.Canvas(f,bg=bg,highlightthickness=0); vsb=ttk.Scrollbar(f,orient="vertical",command=cw.yview)
        cw.configure(yscrollcommand=vsb.set); vsb.pack(side="right",fill="y"); cw.pack(side="left",fill="both",expand=True)
        inner=tk.Frame(cw,bg=bg); wid=cw.create_window((0,0),window=inner,anchor="nw")
        cw.bind("<Configure>",lambda e:cw.itemconfig(wid,width=e.width))
        inner.bind("<Configure>",lambda e:cw.configure(scrollregion=cw.bbox("all")))
        for key,kart in MEVZUAT_KARTLARI.items():
            card=tk.Frame(inner,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1,padx=16,pady=12)
            card.pack(fill="x",pady=6,padx=4)
            tk.Label(card,text=kart["baslik"],font=("Segoe UI",12,"bold"),fg=self.gc("pri"),bg=C_WHITE,anchor="w").pack(fill="x")
            tk.Label(card,text=kart["tanim"],font=("Segoe UI",10),fg="#333",bg=C_WHITE,wraplength=800,justify="left",anchor="w").pack(fill="x",pady=(4,2))
            tk.Label(card,text=f"📌 {kart['kanun']}",font=("Segoe UI",9,"italic"),fg=C_INFO,bg=C_WHITE,anchor="w").pack(fill="x")
            tk.Label(card,text=kart["detay"],font=("Segoe UI",9),fg="#555",bg=C_WHITE,wraplength=800,justify="left",anchor="w").pack(fill="x",pady=(4,0))

    # ═══ 17. İLETİŞİM ════════════════════════════════════════════════════════
    def _t_iletisim(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        bf=ttk.Frame(nb2); ff=ttk.Frame(nb2)
        nb2.add(bf,text="  📞 Bilgiler  "); nb2.add(ff,text="  📝 İletişim Formu  ")
        if self.u_yetki=="Admin":
            af=ttk.Frame(nb2); nb2.add(af,text="  ⚙️ Düzenle  ")
        # Bilgiler
        bi=tk.Frame(bf,bg=bg); bi.pack(fill="both",expand=True,padx=40,pady=30)
        tk.Label(bi,text="📞 Geliştirici İletişim",font=("Segoe UI",18,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,20))
        card=tk.Frame(bi,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1,padx=30,pady=24); card.pack(fill="x")
        # DB'den iletişim bilgilerini çek
        try:
            with sqlite3.connect(DB_PATH) as c:
                ilet={r[0]:r[1] for r in c.execute("SELECT anahtar,deger FROM Iletisim_Bilgileri")}
        except: ilet={"telefon":DEV_TEL,"email":DEV_MAIL,"whatsapp":DEV_WA}
        for lbl,val,ikon in [("Ad Soyad",DEV_ADI,"👤"),("Ünvan",DEV_UNVAN,"🎓"),
                              ("E-posta",ilet.get("email",DEV_MAIL),"📧"),
                              ("Telefon",ilet.get("telefon",DEV_TEL),"📱"),
                              ("Program",f"{PROG_ADI} {VERSIYON}","💻")]:
            rf=tk.Frame(card,bg=C_WHITE); rf.pack(fill="x",pady=6)
            tk.Label(rf,text=f"{ikon}  {lbl}:",font=("Segoe UI",11,"bold"),fg="#333",bg=C_WHITE,width=14,anchor="w").pack(side="left")
            tk.Label(rf,text=val,font=("Segoe UI",11),fg=self.gc("pri"),bg=C_WHITE).pack(side="left",padx=8)
        # WhatsApp butonu
        wa=ilet.get("whatsapp",DEV_WA)
        def _wa():
            if messagebox.askyesno("WhatsApp","Sayfadan ayrılıp WhatsApp Web açılacak.\nDevam?"):
                import webbrowser; webbrowser.open(f"https://wa.me/{wa}")
        MBtn(bi,"💬 WhatsApp ile Mesaj Gönder",command=_wa,color="#25D366",width=28).pack(anchor="w",pady=16)
        tk.Label(bi,text="⚠️ Hata bildirimi için ekran görüntüsünü ve miras_debug.log dosyasını gönderin.",
            font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w")
        # İletişim formu
        fm=tk.Frame(ff,bg=bg); fm.pack(fill="both",expand=True,padx=30,pady=20)
        tk.Label(fm,text="📝 İstek / Şikayet / Bilgi Edinme Formu",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        tk.Label(fm,text=f"Gönderen: {self.u_ad}",font=("Segoe UI",10),fg="#555",bg=bg).pack(anchor="w",pady=(0,8))
        tk.Label(fm,text="Konu:",bg=bg,font=("Segoe UI",10,"bold")).pack(anchor="w")
        self.il_konu=ttk.Combobox(fm,values=["İstek","Şikayet","Bilgi Edinme","Öneri","Hata Bildirimi"],state="readonly",width=28)
        self.il_konu.set("İstek"); self.il_konu.pack(anchor="w",pady=(0,10))
        tk.Label(fm,text="Mesaj:",bg=bg,font=("Segoe UI",10,"bold")).pack(anchor="w")
        self.il_mesaj=tk.Text(fm,height=8,font=("Segoe UI",10),wrap="word",relief="flat",bd=1,
            highlightbackground="#D0DDD8",highlightthickness=1,padx=10,pady=8)
        self.il_mesaj.pack(fill="x",pady=(0,10))
        MBtn(fm,"📨 Gönder",command=self._iletisim_gonder,color=self.gc("acc"),width=20).pack(anchor="w")
        # Admin düzenle
        if self.u_yetki=="Admin":
            ae=tk.Frame(af,bg=bg); ae.pack(fill="both",expand=True,padx=40,pady=30)
            tk.Label(ae,text="⚙️ İletişim Bilgilerini Düzenle",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,14))
            self.il_ed={}
            for l,k,d in [("Telefon:","telefon",ilet.get("telefon","")),("E-posta:","email",ilet.get("email","")),("WhatsApp No:","whatsapp",ilet.get("whatsapp",""))]:
                tk.Label(ae,text=l,bg=bg,font=("Segoe UI",10,"bold")).pack(anchor="w",pady=(8,0))
                e=ttk.Entry(ae,width=40); e.insert(0,d); e.pack(anchor="w"); self.il_ed[k]=e
            MBtn(ae,"💾 Kaydet",command=self._iletisim_kaydet,color=C_INFO,width=20).pack(anchor="w",pady=14)

    def _iletisim_gonder(self):
        mesaj=self.il_mesaj.get("1.0",tk.END).strip()
        if not mesaj: messagebox.showwarning("Uyarı","Mesaj boş."); return
        with sqlite3.connect(DB_PATH) as c:
            c.execute("INSERT INTO Iletisim_Formu(gonderen,konu_tipi,mesaj,tarih)VALUES(?,?,?,?)",
                (self.u_ad,self.il_konu.get(),mesaj,datetime.now().strftime("%d.%m.%Y %H:%M")))
        db_log(self.u_id,"İletişim Formu",self.il_konu.get())
        self.il_mesaj.delete("1.0",tk.END)
        messagebox.showinfo("Tamam","Mesajınız Admin'e iletildi.")

    def _iletisim_kaydet(self):
        with sqlite3.connect(DB_PATH) as c:
            for k,e in self.il_ed.items():
                c.execute("INSERT OR REPLACE INTO Iletisim_Bilgileri(anahtar,deger)VALUES(?,?)",(k,e.get()))
        messagebox.showinfo("Tamam","İletişim bilgileri güncellendi.")

    # ═══ 18. AYARLAR ═════════════════════════════════════════════════════════
    def _t_ayarlar(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        gf=ttk.Frame(nb2); sf=ttk.Frame(nb2); yf=ttk.Frame(nb2)
        nb2.add(gf,text="  🎨 Görünüm  "); nb2.add(sf,text="  🔑 Şifre  "); nb2.add(yf,text="  💾 Yedekleme  ")
        # Görünüm
        gg=tk.Frame(gf,bg=bg); gg.pack(fill="both",expand=True,padx=40,pady=30)
        tk.Label(gg,text="Tema:",bg=bg,font=("Segoe UI",11,"bold")).pack(anchor="w")
        self.cb_tema=ttk.Combobox(gg,values=list(TEMALAR.keys()),state="readonly"); self.cb_tema.set(self.tema); self.cb_tema.pack(anchor="w",pady=(4,16))
        tk.Label(gg,text="Yazı Boyutu:",bg=bg,font=("Segoe UI",11,"bold")).pack(anchor="w")
        self.cb_pt=ttk.Combobox(gg,values=[8,9,10,11,12,14],state="readonly",width=6); self.cb_pt.set(self.punto); self.cb_pt.pack(anchor="w",pady=(4,16))
        if self.u_yetki=="Admin":
            tk.Label(gg,text="Şube Müdürü V. (PDF):",bg=bg,font=("Segoe UI",11,"bold")).pack(anchor="w")
            self.e_sm=ttk.Entry(gg,width=28); self.e_sm.insert(0,self.sube_mudur); self.e_sm.pack(anchor="w",pady=(4,16))
            tk.Label(gg,text="Vali Yardımcısı (Evrak):",bg=bg,font=("Segoe UI",11,"bold")).pack(anchor="w")
            self.e_vy=ttk.Entry(gg,width=28); self.e_vy.insert(0,getattr(self,"vali_yardimcisi","Semih CEMBEKLİ")); self.e_vy.pack(anchor="w",pady=(4,20))
        def _kg():
            sm=self.e_sm.get().strip() if self.u_yetki=="Admin" and hasattr(self,"e_sm") else self.sube_mudur
            vy=self.e_vy.get().strip() if self.u_yetki=="Admin" and hasattr(self,"e_vy") else getattr(self,"vali_yardimcisi","Semih CEMBEKLİ")
            with sqlite3.connect(DB_PATH) as c: c.execute("INSERT OR REPLACE INTO Ayarlar(k_adi,tema,punto,sube_mudur,vali_yardimcisi)VALUES(?,?,?,?,?)",(self.u_id,self.cb_tema.get(),int(self.cb_pt.get()),sm,vy))
            messagebox.showinfo("Kaydedildi","Bir sonraki girişte aktif olur.")
        MBtn(gg,"💾 Kaydet",command=_kg,color=C_INFO,width=22).pack(anchor="w")
        # Şifre
        ss=tk.Frame(sf,bg=bg); ss.pack(fill="both",expand=True,padx=40,pady=30)
        tk.Label(ss,text="Şifre Değiştir",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,16))
        self.sw={}
        for l,k in [("Mevcut Şifre:","eski"),("Yeni Şifre:","yeni"),("Tekrar:","tekrar")]:
            tk.Label(ss,text=l,bg=bg,font=("Segoe UI",10)).pack(anchor="w",pady=(8,2))
            e=ttk.Entry(ss,show="●",width=30); e.pack(anchor="w"); self.sw[k]=e
        def _sk():
            eski=self.sw["eski"].get(); yeni=self.sw["yeni"].get(); tekrar=self.sw["tekrar"].get()
            if yeni!=tekrar: messagebox.showerror("Hata","Şifreler eşleşmiyor."); return
            ok,msg=strong_pw(yeni)
            if not ok: messagebox.showerror("Hata",msg); return
            with sqlite3.connect(DB_PATH) as c:
                db_s=c.execute("SELECT sifre FROM Kullanicilar WHERE k_adi=?",(self.u_id,)).fetchone()[0]
            if not verify_pw(eski,db_s): messagebox.showerror("Hata","Mevcut şifre yanlış."); return
            with sqlite3.connect(DB_PATH) as c: c.execute("UPDATE Kullanicilar SET sifre=?,sifre_tarih=? WHERE k_adi=?",(hash_pw(yeni),datetime.now().strftime("%Y-%m-%d"),self.u_id))
            db_log(self.u_id,"Şifre Değiştir","Başarılı")
            for e in self.sw.values(): e.delete(0,tk.END)
            messagebox.showinfo("Tamam","Şifreniz değiştirildi.")
        MBtn(ss,"🔒 Güncelle",command=_sk,color=C_DANGER,width=22).pack(anchor="w",pady=16)
        # Şifre ömrü bilgisi
        try:
            with sqlite3.connect(DB_PATH) as c:
                st=c.execute("SELECT sifre_tarih FROM Kullanicilar WHERE k_adi=?",(self.u_id,)).fetchone()
            if st and st[0]:
                gecen=(datetime.now()-datetime.strptime(st[0],"%Y-%m-%d")).days
                kalan=max(0,SIFRE_OMUR_GUN-gecen)
                renk="#2D8C55" if kalan>SIFRE_UYARI_GUN else "#D68910" if kalan>0 else "#C0392B"
                tk.Label(ss,text=f"🔑 Şifre ömrü: {kalan} gün kaldı ({SIFRE_OMUR_GUN} günden {gecen}. gün)",
                    font=("Segoe UI",10,"bold"),fg=renk,bg=bg).pack(anchor="w",pady=(0,8))
        except Exception: pass
        # Giriş Geçmişi (son 10)
        tk.Label(ss,text="📋 Son 10 Giriş Denemesi",font=("Segoe UI",11,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(16,6))
        tv_giris=ttk.Treeview(ss,columns=("t","i","d"),show="headings",height=6)
        for col,w,bas in [("t",150,"Tarih"),("i",100,"Sonuç"),("d",200,"Detay")]:
            tv_giris.heading(col,text=bas); tv_giris.column(col,width=w)
        tv_giris.tag_configure("basarili",background="#EAFAF1")
        tv_giris.tag_configure("basarisiz",background="#FDEDEC")
        tv_giris.tag_configure("kilitli",background="#FEF9E7")
        tv_giris.pack(fill="x",pady=4)
        try:
            with sqlite3.connect(DB_PATH) as c:
                for r in c.execute("SELECT tarih,islem,detay FROM Loglar WHERE kul=? AND islem IN('Giriş','Başarısız') ORDER BY id DESC LIMIT 10",(self.u_id,)).fetchall():
                    tag="basarili" if "Giriş" in r[1] else "kilitli" if "Kilit" in str(r[2]) else "basarisiz"
                    tv_giris.insert("","end",values=r,tags=(tag,))
        except Exception: pass
        # Yedekleme
        yy=tk.Frame(yf,bg=bg); yy.pack(fill="both",expand=True,padx=20,pady=16)
        tk.Label(yy,text="💾 Yedekleme & Geri Yükleme",font=("Segoe UI",15,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")
        tk.Label(yy,text=f"Yerel Yedek: {BACKUP_DIR.absolute()}",font=("Segoe UI",9),fg="#888",bg=bg).pack(anchor="w",pady=(0,4))
        # Drive klasör bilgisi
        drive_f=tk.Frame(yy,bg=bg); drive_f.pack(fill="x",pady=(0,10))
        dp=self.yedekci.drive_path
        drive_durum="✅ "+dp if dp else "❌ Ayarlanmamış — aşağıdan klasör seçin"
        self.lbl_drive_yol=tk.Label(drive_f,text=f"☁️ Google Drive: {drive_durum}",font=("Segoe UI",9,"bold"),fg="#1A5276" if dp else "#C0392B",bg=bg)
        self.lbl_drive_yol.pack(side="left")
        if dp:
            tk.Label(drive_f,text="  (Her yedeklemede otomatik kopyalanır)",font=("Segoe UI",8),fg="#2D8C55",bg=bg).pack(side="left")
        bff=tk.Frame(yy,bg=bg); bff.pack(fill="x",pady=(0,12))
        MBtn(bff,"💾 Şimdi Yedekle",command=self._yedek_al,width=18).pack(side="left",padx=6)
        MBtn(bff,"🔄 Listeyi Yenile",command=self._yedek_yenile,color=C_INFO,width=16).pack(side="left",padx=6)
        MBtn(bff,"⏮ Yedeğe Dön",command=self._yedek_geri,color=C_DANGER,width=16).pack(side="left",padx=6)
        MBtn(bff,"☁️ Drive Klasörü Seç",command=self._drive_klasor_sec,color="#F39C12",width=20).pack(side="left",padx=6)
        MBtn(bff,"☁️ Şimdi Drive'a Kopyala",command=self._drive_kopyala,color="#1A5276",width=22).pack(side="left",padx=6)
        self.tv_yedek=self._tv(yy,[("ad",280,"Dosya"),("tarih",160,"Tarih"),("boyut",100,"Boyut"),("tur",100,"Tür")],12)
        self._yedek_yenile()

    def _yedek_al(self):
        try:
            yol=self.yedekci.al(False); self._yedek_yenile()
            drive_msg=""
            if self.yedekci.drive_path:
                drive_msg=f"\n☁️ Drive'a da kopyalandı:\n{self.yedekci.drive_path}"
            messagebox.showinfo("Tamam",f"💾 Yerel yedek:\n{yol}{drive_msg}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _yedek_yenile(self):
        self.tv_yedek.delete(*self.tv_yedek.get_children())
        for y in self.yedekci.listele():
            tur="🤖 Oto" if "oto" in y["ad"] else "✋ Manuel"
            self.tv_yedek.insert("","end",values=(y["ad"],y["tarih"],y["boyut"],tur),tags=(y["yol"],))

    def _yedek_geri(self):
        sel=self.tv_yedek.selection()
        if not sel: messagebox.showwarning("Seçim","Yedek seçin."); return
        yol=self.tv_yedek.item(sel[0])["tags"][0]
        if messagebox.askyesno("Onay",f"Mevcut veri ÜZERİNE yazılacak!\nDevam?"):
            if self.yedekci.geri(yol): messagebox.showinfo("Tamam","Geri yükleme başarılı.")
            else: messagebox.showerror("Hata","Geri yükleme başarısız.")

    def _drive_klasor_sec(self):
        """Google Drive sync klasörünü seç ve kaydet"""
        mevcut=self.yedekci.drive_path or ""
        hedef=filedialog.askdirectory(title="Google Drive Klasörünü Seçin (örn: G:\\My Drive\\Miras_Yedek)",initialdir=mevcut or None)
        if hedef:
            if self.yedekci.set_drive_path(hedef):
                db_log(self.u_id,"Drive Klasör Ayarla",hedef)
                messagebox.showinfo("Tamam",
                    f"☁️ Drive klasörü kaydedildi:\n{hedef}\n\n"
                    "Artık her yedeklemede (manuel + otomatik günlük)\n"
                    "veritabanı bu klasöre de otomatik kopyalanacak.\n\n"
                    "Google Drive Desktop açık olduğu sürece\n"
                    "dosyalar otomatik buluta sync olur.")
                # UI'ı yenile
                self._menu_click("ayarlar")
            else:
                messagebox.showerror("Hata","Geçersiz klasör yolu.")

    def _drive_kopyala(self):
        """Manuel olarak şu anda Drive'a kopyala"""
        if not self.yedekci.drive_path:
            messagebox.showwarning("Drive Ayarlanmamış",
                "Önce '☁️ Drive Klasörü Seç' butonuyla\n"
                "Google Drive sync klasörünüzü seçin.\n\n"
                "Örnek: G:\\My Drive\\Miras_Yedek")
            return
        if not os.path.isdir(self.yedekci.drive_path):
            messagebox.showerror("Hata",f"Drive klasörü bulunamadı:\n{self.yedekci.drive_path}\n\nKlasörü yeniden seçin.")
            return
        try:
            ts=datetime.now().strftime("%Y%m%d_%H%M%S")
            hedef_dosya=os.path.join(self.yedekci.drive_path,f"Miras_manuel_{ts}.db")
            shutil.copy2(DB_PATH,hedef_dosya)
            self.yedekci._drive_temizle()
            db_log(self.u_id,"Drive Kopyala",hedef_dosya)
            messagebox.showinfo("Tamam",f"☁️ Drive'a kopyalandı!\n{hedef_dosya}\n\nGoogle Drive Desktop sync edecektir.")
        except Exception as e: messagebox.showerror("Hata",str(e))

    # ═══ 19. ADMİN ═══════════════════════════════════════════════════════════
    def _t_admin(self,p):
        bg=self.gc("bg"); f=tk.Frame(p,bg=bg); f.pack(fill="both",expand=True,padx=20,pady=16)
        # Duyuru
        df=tk.LabelFrame(f,text="  📢 Duyuru  ",bg=C_WHITE,padx=10,pady=8); df.pack(fill="x",pady=5)
        self.e_dy=ttk.Entry(df,width=60); self.e_dy.pack(side="left",padx=10)
        def _dy():
            if self.e_dy.get():
                with sqlite3.connect(DB_PATH) as c: c.execute("INSERT INTO Duyurular(mesaj,tarih,gonderen)VALUES(?,?,?)",(self.e_dy.get(),datetime.now().strftime("%d.%m.%Y"),self.u_ad))
                messagebox.showinfo("Tamam","Duyuru eklendi."); self.e_dy.delete(0,tk.END)
        MBtn(df,"📢 Gönder",command=_dy,color=C_DANGER).pack(side="left")
        # Kullanıcılar
        tk.Label(f,text="Kullanıcı Yönetimi",font=("Segoe UI",12,"bold"),bg=bg,fg=self.gc("pri")).pack(anchor="w",pady=(14,6))
        self.tv_kul=self._tv(f,[("k",120,"Kullanıcı"),("ad",160,"Ad"),("unvan",160,"Ünvan"),("y",90,"Yetki"),("aktif",70,"Aktif"),("hata",60,"Hata")],8)
        bff=tk.Frame(f,bg=bg); bff.pack(fill="x",pady=(0,10))
        for txt,fn,renk in [("🔄 Yenile",self._kul_yenile,"#2D8C55"),("🔒 Kilit Kaldır",self._kul_kilit,C_WARN),
            ("❌ Pasif",self._kul_pasif,C_DANGER),("✅ Aktif",self._kul_aktif,"#2D8C55"),
            ("🔑 Şifre Sıfırla",self._kul_sifre,C_INFO),("✏️ K.Adı Değiştir",self._kul_adi_degistir,"#8E44AD")]:
            MBtn(bff,txt,command=fn,color=renk,width=16).pack(side="left",padx=3)
        # Yeni kullanıcı
        nf=tk.LabelFrame(f,text="  ➕ Yeni Kullanıcı  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=14,pady=10); nf.pack(fill="x",pady=6)
        self.kul_g={}
        for i,(l,k) in enumerate([("K.Adı","ka"),("Ad Soyad","ad"),("Ünvan","un"),("Şifre","si")]):
            tk.Label(nf,text=l,bg=bg,font=("Segoe UI",10)).grid(row=0,column=i*2,padx=(0,4))
            e=ttk.Entry(nf,width=14,show="*" if k=="si" else ""); e.grid(row=0,column=i*2+1,padx=(0,10)); self.kul_g[k]=e
        tk.Label(nf,text="Yetki",bg=bg,font=("Segoe UI",10)).grid(row=0,column=8)
        self.kul_g["y"]=ttk.Combobox(nf,values=["Uzman","Admin","İzleyici"],state="readonly",width=10); self.kul_g["y"].set("Uzman"); self.kul_g["y"].grid(row=0,column=9,padx=(0,10))
        MBtn(nf,"➕ Ekle",command=self._kul_ekle).grid(row=0,column=10)
        # Log filtresi (sadece admin)
        lff=tk.LabelFrame(f,text="  📋 Log Filtreleme  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=10,pady=8); lff.pack(fill="both",expand=True,pady=6)
        AramaFrame(lff,self._filtre_admin_log,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_admin_log=self._tv(lff,[("t",150,"Zaman"),("k",100,"Personel"),("i",120,"İşlem"),("d",400,"Detay")],10)
        # Talepler
        tk.Label(f,text="Şifre / Kayıt Talepleri",font=("Segoe UI",11,"bold"),bg=bg).pack(anchor="w",pady=(8,4))
        self.tv_st=self._tv(f,[("id",40,"#"),("ka",120,"Kullanıcı"),("t",100,"Tarih"),("d",100,"Durum")],3)
        self.tv_kt=self._tv(f,[("id",40,"#"),("ad",140,"Ad"),("un",140,"Ünvan"),("ka",110,"K.Adı"),("t",100,"Tarih")],3)
        bf2=tk.Frame(f,bg=bg); bf2.pack(pady=4)
        MBtn(bf2,"✅ Şifre Onayla",command=self._sifre_onayla,color=C_INFO,width=18).pack(side="left",padx=6)
        MBtn(bf2,"✅ Kayıt Onayla",command=self._kayit_onayla,color="#2D8C55",width=18).pack(side="left",padx=6)
        # Silme Talepleri
        tk.Label(f,text="🗑 Silme Talepleri",font=("Segoe UI",11,"bold"),bg=bg,fg=C_DANGER).pack(anchor="w",pady=(8,4))
        self.tv_silt=self._tv(f,[("id",40,"#"),("tablo",100,"Modül"),("ozet",200,"Kayıt"),("talep",120,"Talep Eden"),("t",120,"Tarih")],3)
        bf3=tk.Frame(f,bg=bg); bf3.pack(pady=4)
        MBtn(bf3,"✅ Onayla (Sil)",command=self._silme_onayla,color=C_DANGER,width=16).pack(side="left",padx=6)
        MBtn(bf3,"❌ Reddet",command=self._silme_reddet,color="#888",width=14).pack(side="left",padx=6)
        # İletişim Talepleri
        tk.Label(f,text="📝 İletişim Formları",font=("Segoe UI",11,"bold"),bg=bg,fg=C_INFO).pack(anchor="w",pady=(8,4))
        self.tv_ilet=self._tv(f,[("id",40,"#"),("gon",120,"Gönderen"),("konu",100,"Konu"),("mesaj",280,"Mesaj"),("t",120,"Tarih")],3)
        MBtn(f,"✅ Okundu",command=self._ilet_okundu,color=C_INFO,width=14).pack(anchor="w",padx=6,pady=4)
        self._kul_yenile()

    def _kul_yenile(self):
        self.tv_kul.delete(*self.tv_kul.get_children())
        self.tv_st.delete(*self.tv_st.get_children()); self.tv_kt.delete(*self.tv_kt.get_children())
        try:
            with sqlite3.connect(DB_PATH) as c:
                for r in c.execute("SELECT k_adi,ad,unvan,yetki,aktif,fail_count FROM Kullanicilar"):
                    self.tv_kul.insert("","end",values=(r[0],r[1],r[2],r[3],"✔" if r[4] else "✘",r[5]))
                for r in c.execute("SELECT id,k_adi,tarih,durum FROM Sifre_Talepleri WHERE durum='Bekliyor'"):
                    self.tv_st.insert("","end",values=r)
                for r in c.execute("SELECT id,ad,unvan,k_adi,tarih FROM Kayit_Talepleri WHERE durum='Bekliyor'"):
                    self.tv_kt.insert("","end",values=r)
                # Loglar
                self._all_admin_log=c.execute("SELECT tarih,kul,islem,detay FROM Loglar ORDER BY id DESC LIMIT 200").fetchall()
                # Silme talepleri
                if hasattr(self,'tv_silt'):
                    self.tv_silt.delete(*self.tv_silt.get_children())
                    for r in c.execute("SELECT id,tablo,kayit_ozet,talep_eden,tarih FROM Silme_Talepleri WHERE durum='Bekliyor'"):
                        self.tv_silt.insert("","end",values=r)
                # İletişim formları
                if hasattr(self,'tv_ilet'):
                    self.tv_ilet.delete(*self.tv_ilet.get_children())
                    for r in c.execute("SELECT id,gonderen,konu_tipi,mesaj,tarih FROM Iletisim_Formu WHERE durum='Yeni' ORDER BY id DESC"):
                        self.tv_ilet.insert("","end",values=r)
            self._filtre_admin_log("")
        except Exception as e: logging.error(f"kul_yenile:{e}")

    def _filtre_admin_log(self,a):
        self.tv_admin_log.delete(*self.tv_admin_log.get_children()); a=a.lower()
        for r in getattr(self,"_all_admin_log",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            self.tv_admin_log.insert("","end",values=r)

    def _sec_kul(self):
        sel=self.tv_kul.selection()
        if not sel: messagebox.showwarning("Seçim","Kullanıcı seçin."); return None
        return self.tv_kul.item(sel[0])["values"][0]

    def _kul_kilit(self):
        k=self._sec_kul()
        if k:
            with sqlite3.connect(DB_PATH) as c: c.execute("UPDATE Kullanicilar SET fail_count=0,lockout_ts=NULL WHERE k_adi=?",(k,))
            db_log(self.u_id,"Kilit Kaldır",k); self._kul_yenile(); messagebox.showinfo("Tamam",f"'{k}' kilidi kaldırıldı.")

    def _kul_pasif(self):
        k=self._sec_kul()
        if not k or k==self.u_id: messagebox.showwarning("Uyarı","Kendinizi pasifleştiremezsiniz."); return
        if k:
            with sqlite3.connect(DB_PATH) as c: c.execute("UPDATE Kullanicilar SET aktif=0 WHERE k_adi=?",(k,))
            db_log(self.u_id,"Pasifleştir",k); self._kul_yenile()

    def _kul_aktif(self):
        k=self._sec_kul()
        if k:
            with sqlite3.connect(DB_PATH) as c: c.execute("UPDATE Kullanicilar SET aktif=1 WHERE k_adi=?",(k,))
            self._kul_yenile()

    def _kul_sifre(self):
        k=self._sec_kul()
        if not k: return
        y=tkinter.simpledialog.askstring("Şifre",f"'{k}' için yeni şifre:",show="*")
        if y:
            ok,msg=strong_pw(y)
            if not ok: messagebox.showerror("Hata",msg); return
            with sqlite3.connect(DB_PATH) as c: c.execute("UPDATE Kullanicilar SET sifre=? WHERE k_adi=?",(hash_pw(y),k))
            db_log(self.u_id,"Şifre Sıfırla",k); messagebox.showinfo("Tamam","Şifre güncellendi.")

    def _kul_adi_degistir(self):
        k=self._sec_kul()
        if not k: return
        yeni=tkinter.simpledialog.askstring("K.Adı Değiştir",f"'{k}' için yeni kullanıcı adı:")
        if yeni and yeni.strip():
            yeni=yeni.strip()
            try:
                with sqlite3.connect(DB_PATH) as c:
                    c.execute("UPDATE Kullanicilar SET k_adi=? WHERE k_adi=?",(yeni,k))
                    c.execute("UPDATE Ayarlar SET k_adi=? WHERE k_adi=?",(yeni,k))
                db_log(self.u_id,"K.Adı Değiştir",f"{k}→{yeni}"); self._kul_yenile()
                messagebox.showinfo("Tamam",f"'{k}' → '{yeni}' olarak değiştirildi.")
            except sqlite3.IntegrityError: messagebox.showerror("Hata","Bu kullanıcı adı zaten alınmış.")

    def _kul_ekle(self):
        try:
            ka=self.kul_g["ka"].get().strip(); ad=self.kul_g["ad"].get().strip(); si=self.kul_g["si"].get()
            if not all([ka,ad,si]): raise ValueError("Tüm alanlar zorunlu.")
            ok,msg=strong_pw(si)
            if not ok: raise ValueError(msg)
            with sqlite3.connect(DB_PATH) as c: c.execute("INSERT INTO Kullanicilar(k_adi,sifre,yetki,ad,unvan)VALUES(?,?,?,?,?)",(ka,hash_pw(si),self.kul_g["y"].get(),ad,self.kul_g["un"].get()))
            db_log(self.u_id,"Kullanıcı Ekle",ka); self._kul_yenile()
            for e in self.kul_g.values():
                if isinstance(e,ttk.Entry): e.delete(0,tk.END)
            messagebox.showinfo("Tamam",f"'{ka}' oluşturuldu.")
        except ValueError as e: messagebox.showerror("Hata",str(e))
        except sqlite3.IntegrityError: messagebox.showerror("Hata","K.Adı alınmış.")

    def _sifre_onayla(self):
        sel=self.tv_st.selection()
        if not sel: return
        r=self.tv_st.item(sel[0])["values"]; sid=r[0]; ku=r[1]
        y=tkinter.simpledialog.askstring("Onay",f"'{ku}' için yeni şifre:",show="*")
        if y:
            ok,msg=strong_pw(y)
            if not ok: messagebox.showerror("Hata",msg); return
            with sqlite3.connect(DB_PATH) as c:
                c.execute("UPDATE Kullanicilar SET sifre=? WHERE k_adi=?",(hash_pw(y),ku))
                c.execute("UPDATE Sifre_Talepleri SET durum='Onaylandı' WHERE id=?",(sid,))
            self._kul_yenile()

    def _kayit_onayla(self):
        sel=self.tv_kt.selection()
        if not sel: return
        r=self.tv_kt.item(sel[0])["values"]; kid=r[0]
        y=tkinter.simpledialog.askstring("Onay",f"'{r[3]}' için şifre:",show="*")
        if y:
            ok,msg=strong_pw(y)
            if not ok: messagebox.showerror("Hata",msg); return
            try:
                with sqlite3.connect(DB_PATH) as c:
                    c.execute("INSERT INTO Kullanicilar(k_adi,sifre,yetki,ad,unvan)VALUES(?,?,?,?,?)",(r[3],hash_pw(y),"Uzman",r[1],r[2]))
                    c.execute("UPDATE Kayit_Talepleri SET durum='Onaylandı' WHERE id=?",(kid,))
                messagebox.showinfo("Tamam","Kullanıcı eklendi.")
            except sqlite3.IntegrityError: messagebox.showerror("Hata","K.Adı alınmış.")
            self._kul_yenile()

    def _silme_onayla(self):
        sel=self.tv_silt.selection()
        if not sel: return
        r=self.tv_silt.item(sel[0])["values"]; sid=r[0]; tablo=r[1]; ozet=r[2]
        if messagebox.askyesno("Silme Onayı",f"'{ozet}' kalıcı olarak silinecek.\nOnaylıyor musunuz?"):
            # Tablodan sil
            tablo_map={"Rapor_Gecmisi":"rapor_no","Ihaleler":"id","Idari_Cezalar":"id",
                       "Islah_Amenajman":"id","Tahsisler":"id","Sikayetler":"id","Veri_Kayit":"id"}
            pk=tablo_map.get(tablo,"id")
            try:
                with sqlite3.connect(DB_PATH) as c:
                    kayit_id=c.execute("SELECT kayit_id FROM Silme_Talepleri WHERE id=?",(sid,)).fetchone()[0]
                    c.execute(f"DELETE FROM {tablo} WHERE {pk}=?",(kayit_id,))
                    c.execute("UPDATE Silme_Talepleri SET durum='Onaylandı' WHERE id=?",(sid,))
                db_log(self.u_id,"Silme Onay",ozet); self._kul_yenile()
            except Exception as e: messagebox.showerror("Hata",str(e))

    def _silme_reddet(self):
        sel=self.tv_silt.selection()
        if not sel: return
        sid=self.tv_silt.item(sel[0])["values"][0]
        with sqlite3.connect(DB_PATH) as c: c.execute("UPDATE Silme_Talepleri SET durum='Reddedildi' WHERE id=?",(sid,))
        self._kul_yenile()

    def _ilet_okundu(self):
        sel=self.tv_ilet.selection()
        if not sel: return
        iid=self.tv_ilet.item(sel[0])["values"][0]
        with sqlite3.connect(DB_PATH) as c: c.execute("UPDATE Iletisim_Formu SET durum='Okundu' WHERE id=?",(iid,))
        self._kul_yenile()

    # ═══ KÖY KARTI ═════════════════════════════════════════════════════════════
    def _koy_gecmis(self):
        koy=tkinter.simpledialog.askstring("Köy Kartı","Köy adını girin:")
        if not koy: return
        self._son_islem_kaydet("Köy Kartı",koy)
        win=tk.Toplevel(self.root); win.title(f"📂 {koy} — Köy Kartı"); win.geometry("960x700"); win.configure(bg="#F0F4F2")
        # Üst başlık bandı
        hdr=tk.Frame(win,bg=self.gc("pri"),height=60); hdr.pack(fill="x"); hdr.pack_propagate(False)
        tk.Label(hdr,text=f"🏘️ {koy.upper()} KÖYÜ",font=("Segoe UI",18,"bold"),fg=C_WHITE,bg=self.gc("pri")).pack(side="left",padx=20,pady=10)
        # İlçe ve alan bilgisi
        try:
            with sqlite3.connect(DB_PATH) as c:
                mera=c.execute("SELECT ilce,alan,yem FROM Mera_Varligi WHERE LOWER(koy)=LOWER(?)",(koy,)).fetchone()
        except: mera=None
        if mera:
            tk.Label(hdr,text=f"İlçe: {mera[0]}  |  Mera: {mera[1]:.0f} da",font=("Segoe UI",11),fg="#A9DFBF",bg=self.gc("pri")).pack(side="right",padx=20)
        # Mini kartlar
        kf=tk.Frame(win,bg="#F0F4F2"); kf.pack(fill="x",padx=16,pady=10)
        sayilar={}
        try:
            with sqlite3.connect(DB_PATH) as c:
                sayilar["rapor"]=c.execute("SELECT COUNT(*) FROM Rapor_Gecmisi WHERE LOWER(koy)=LOWER(?)",(koy,)).fetchone()[0]
                sayilar["ihale"]=c.execute("SELECT COUNT(*) FROM Ihaleler WHERE LOWER(koy)=LOWER(?)",(koy,)).fetchone()[0]
                sayilar["islah"]=c.execute("SELECT COUNT(*) FROM Islah_Amenajman WHERE LOWER(koy)=LOWER(?)",(koy,)).fetchone()[0]
                sayilar["tahsis"]=c.execute("SELECT COUNT(*) FROM Tahsisler WHERE LOWER(koy)=LOWER(?)",(koy,)).fetchone()[0]
                sayilar["sikayet"]=c.execute("SELECT COUNT(*) FROM Sikayetler WHERE LOWER(koy)=LOWER(?)",(koy,)).fetchone()[0]
                sayilar["ceza"]=c.execute("SELECT COUNT(*) FROM Idari_Cezalar WHERE LOWER(mera_koy)=LOWER(?)",(koy,)).fetchone()[0]
        except: pass
        for bas,key,ikon,renk in [("Rapor","rapor","📄",C_INFO),("İhale","ihale","⚖️",C_WARN),
            ("Islah","islah","🌱","#2D8C55"),("Tahsis","tahsis","📋","#8E44AD"),
            ("Şikayet","sikayet","🚨",C_DANGER),("Ceza","ceza","💰","#C0392B")]:
            cf=tk.Frame(kf,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1,padx=14,pady=8)
            cf.pack(side="left",fill="both",expand=True,padx=4)
            tk.Label(cf,text=ikon,font=("Segoe UI",18),bg=C_WHITE).pack()
            tk.Label(cf,text=str(sayilar.get(key,0)),font=("Segoe UI",20,"bold"),fg=renk,bg=C_WHITE).pack()
            tk.Label(cf,text=bas,font=("Segoe UI",8),fg="#666",bg=C_WHITE).pack()
        # Detay notebook
        nb3=ttk.Notebook(win); nb3.pack(fill="both",expand=True,padx=10,pady=6)
        bolumler=[
            ("📄 Raporlar","SELECT rapor_no,islem_tarihi,talep_eden,tc,duzenleyen FROM Rapor_Gecmisi WHERE LOWER(koy)=LOWER(?)",[("no",90,"No"),("t",100,"Tarih"),("talep",150,"Talep"),("tc",110,"TC"),("per",130,"Personel")]),
            ("⚖️ İhaleler","SELECT id,tarih,ad_soyad,bedel,durum FROM Ihaleler WHERE LOWER(koy)=LOWER(?)",[("id",40,"#"),("t",90,"Tarih"),("ad",150,"Başvuran"),("bedel",90,"Bedel"),("d",120,"Durum")]),
            ("🌱 Islah","SELECT id,dilekce_tarihi,talep_eden,talep_alani,durum FROM Islah_Amenajman WHERE LOWER(koy)=LOWER(?)",[("id",40,"#"),("t",90,"Tarih"),("talep",140,"Talep"),("alan",80,"Alan"),("d",110,"Durum")]),
            ("📋 Tahsis","SELECT id,basvuru_t,kurum,amac,asama,durum FROM Tahsisler WHERE LOWER(koy)=LOWER(?)",[("id",40,"#"),("t",90,"Tarih"),("k",130,"Kurum"),("a",130,"Amaç"),("as",120,"Aşama"),("d",90,"Durum")]),
            ("🚨 Şikayet","SELECT id,tarih,sikayet_eden,tur,durum FROM Sikayetler WHERE LOWER(koy)=LOWER(?)",[("id",40,"#"),("t",90,"Tarih"),("s",130,"Şikayet Eden"),("tur",120,"Tür"),("d",110,"Durum")]),
            ("💰 Ceza","SELECT id,tarih,ad_soyad,konu,ipc_tutari FROM Idari_Cezalar WHERE LOWER(mera_koy)=LOWER(?)",[("id",40,"#"),("t",90,"Tarih"),("ad",130,"Ad Soyad"),("konu",180,"Konu"),("tutar",90,"Tutar")]),
        ]
        for tab_ad,sql,cols in bolumler:
            tf=ttk.Frame(nb3); nb3.add(tf,text=f"  {tab_ad}  ")
            tv=self._tv(tf,cols,12)
            try:
                with sqlite3.connect(DB_PATH) as c:
                    for r in c.execute(sql,(koy,)).fetchall(): tv.insert("","end",values=r)
            except Exception: pass
        # Son işlemler zaman çizelgesi
        tlf=ttk.Frame(nb3); nb3.add(tlf,text="  ⏱ Zaman Çizelgesi  ")
        tv_zaman=self._tv(tlf,[("t",140,"Tarih"),("k",120,"Personel"),("i",140,"İşlem"),("d",400,"Detay")],14)
        try:
            with sqlite3.connect(DB_PATH) as c:
                for r in c.execute("SELECT tarih,kul,islem,detay FROM Loglar WHERE LOWER(detay) LIKE ? ORDER BY id DESC LIMIT 30",(f"%{koy.lower()}%",)).fetchall():
                    tv_zaman.insert("","end",values=r)
        except Exception: pass

# ─── GİRİŞ NOKTASI ───────────────────────────────────────────────────────────
if __name__ == "__main__":
    # --- YENİ EKLENEN KISIM BAŞLANGICI (Eski dosyaları temizleme) ---
    # Eğer program güncellenmişse, arkada kalan ".eski" uzantılı çöp dosyayı siler
    if getattr(sys, 'frozen', False):
        eski_dosya = sys.executable + ".eski"
        if os.path.exists(eski_dosya):
            try:
                os.remove(eski_dosya)
            except:
                pass
    # --- YENİ EKLENEN KISIM BİTİŞİ ---

    app = tk.Tk()
    
    # --- YENİ EKLENEN KISIM BAŞLANGICI (Güncellemeyi tetikleme) ---
    # Arayüz açıldıktan 2 saniye (2000 milisaniye) sonra güncellemeyi kontrol et
    # after() kullanmak, programın açılış hızını yavaşlatmaz/dondurmaz
    app.after(2000, guncelleme_kontrol_et)
    # --- YENİ EKLENEN KISIM BİTİŞİ ---

    uygulama = MirasApp(app)
    app.mainloop()