"""
MİRAS v17.0 Enterprise — Mera İhtiyaç ve Rasyonel Amenajman Sistemi
T.C. Ardahan Valiliği İl Tarım ve Orman Müdürlüğü
Geliştirici: Emre ÖZTÜRK — Ziraat Mühendisi
Kurulum: pip install google-generativeai pandas openpyxl reportlab bcrypt python-docx
Giriş  : admin / Admin123!
"""
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import tkinter.simpledialog
import sqlite3, pandas as pd, threading, shutil, os, json
import urllib.request, hashlib, logging, time, sys, traceback
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional

try:
    import bcrypt; BCRYPT_OK = True
except ImportError:
    BCRYPT_OK = False
try:
    import google.generativeai as genai; GEMINI_OK = True
except ImportError:
    GEMINI_OK = False
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors as rl_colors
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_OK = True
except ImportError:
    PDF_OK = False
try:
    import matplotlib
    MPL_OK = True
except ImportError:
    MPL_OK = False
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
try:
    from ctypes import windll; windll.shcore.SetProcessDpiAwareness(1)
except Exception:
    pass

logging.basicConfig(filename="miras_debug.log", level=logging.ERROR,
    format="%(asctime)s [%(levelname)s] %(message)s")

def _global_exc(t, v, tb):
    logging.error("".join(traceback.format_exception(t, v, tb)))
    hata_str=str(v)
    try:
        if "database is locked" in hata_str.lower() or "locked" in hata_str.lower():
            messagebox.showwarning("⏳ Veritabanı Meşgul",
                "Başka bir personel şu an kayıt yapıyor.\n"
                "Birkaç saniye bekleyip tekrar deneyin.\n\n"
                "İpucu: Aynı anda iki kişi kayıt yaparsa bu uyarı çıkabilir.")
        else:
            messagebox.showerror("Sistem Hatası", f"Hata: {hata_str[:300]}\n\nDetay: miras_debug.log")
    except Exception: pass
sys.excepthook = _global_exc

# ─── SABİTLER ────────────────────────────────────────────────────────────────
PROG_ADI, VERSIYON = "MİRAS Enterprise", "v17.5"
DEV_ADI = "Emre ÖZTÜRK"
DEV_UNVAN = "Ziraat Mühendisi"
DEV_TEL = "0 545 689 00 75"
DEV_MAIL = "miras75mera@gmail.com"
HAKLAR = f"© {datetime.now().year} {DEV_ADI} — Tüm Hakları Saklıdır."
CONFIG_FILE, API_KEY_FILE = "miras_config_v17.json", "miras_gemini_key.txt"
BACKUP_DIR = Path("miras_yedekler")
GUNLUK_YEM_KG = 30            # Kapasite raporu: Bir BBHB'nin günlük YEŞİL YEM ihtiyacı (kg)
OTLATMA_GUN = 135             # Otlatma gün sayısı
GUNLUK_OT_BBHB = 12.5        # Muhammen bedel: Bir BBHB'nin günlük KURU OT ihtiyacı (kg)
# İKİ FARKLI FORMÜL VAR:
# Kapasite: OK = Alan × Yeşil Yem / (30 × 135)      → PAYDA_KAP
# Muhammen: OK = Alan × Kuru Ot / (135 × 12.5)      → PAYDA_MUH
PAYDA = GUNLUK_YEM_KG * OTLATMA_GUN  # 30 × 135 = 4050 (Kapasite Raporu için)
MAX_LOGIN_FAIL = 3
LOCKOUT_MIN = 15
SIFRE_OMUR_GUN = 180
SIFRE_UYARI_GUN = 7
# PAGE_SIZE = 25  # şu an kullanılmıyor
OTOSAVE_SEC = 600
SESSION_TIMEOUT_MIN = 30  # 30 dk hareketsizlik → otomatik çıkış
DEV_WA = "905456890075"
GUNLUK_OT_BBHB = 12.5
GUNLUK_OT_KBHB = 1.25
GITHUB_REPO = "miras75mera/miras-enterprise"
MERA_VASIF = ["Zayıf","Orta","İyi","Çok İyi"]
OT_VERIM_KURU = {"Zayıf":45,"Orta":90,"İyi":135,"Çok İyi":180}

# ─── MERKEZİ KONFİGÜRASYON (hardcoded değerler burada toplanır) ────────────
KURUM_IL = "Ardahan"
KURUM_ADI = "T.C. Ardahan Valiliği İl Tarım ve Orman Müdürlüğü"
KURUM_SUBE = "Çayır, Mera ve Yem Bitkileri Şube Müdürlüğü"
OTLATMA_BASLANGIC = "01/06"  # gün/ay
OTLATMA_BITIS = "30/09"      # gün/ay
OTLATMA_AY = 4
CEZA_BB_TL = 451.50   # 2026 varsayılan — Ayarlar'dan güncellenir
CEZA_KB_TL = 45.15    # 2026 varsayılan — Ayarlar'dan güncellenir

# Valilik Genel Emri yıl bazlı ceza geçmişi
GENEL_EMRI_GECMIS = {
    2025: {"sayi":"2025/01","bb":360.00,"kb":36.00,"madde":"18. Madde (h) bendi"},
    2026: {"sayi":"2026/01","bb":451.50,"kb":45.15,"madde":"18. Madde 9. fıkrası"},
}
DOSYA_BEDELI_TL = 520.00
IBAN_BAKANLIK = "TR 5100 0010 0100 0003 5015 4026"
IBAN_EMANET = "TR 85 0001 0001 4900 0010 0059 60"
ODEME_YERI = "Ardahan Defterdarlığı"
ITIRAZ_MERCI = "Ardahan Sulh Ceza Hakimliği"

# Renkler
C_WHITE = "#FFFFFF"
C_DANGER = "#C0392B"
C_INFO = "#1A6FA8"
C_WARN = "#D68910"
# C_DARK = "#2C3E50"  # şu an kullanılmıyor
C_FOOTER = "#555555"

TEMALAR = {
    "Orman Yeşili": {"pri":"#1E5631","acc":"#2D8C55","bg":"#F0F4F2","side":"#163D22","side_text":"#D4E8DB","side_hover":"#1E5631","card":"#FFFFFF"},
    "Deniz Mavisi": {"pri":"#1A5276","acc":"#2980B9","bg":"#EBF5FB","side":"#143D5C","side_text":"#C8DFF0","side_hover":"#1A5276","card":"#FFFFFF"},
    "Koyu Bordo":   {"pri":"#641E16","acc":"#922B21","bg":"#FDEDEC","side":"#4A1610","side_text":"#E8C4BF","side_hover":"#641E16","card":"#FFFFFF"},
    "Koyu Gece":    {"pri":"#1A1A2E","acc":"#16213E","bg":"#0F0F1A","side":"#0D0D1A","side_text":"#8888AA","side_hover":"#1A1A2E","card":"#1F1F2E"},
}

HAYVAN_TURLERI = [
    ("Kültür Irkı Süt İneği",1.00),("Kültür Melezi Süt İneği",0.75),
    ("Yerli İnek",0.50),("Kültür Irkı Dana-Düve",0.60),
    ("Kültür Melezi Dana-Düve",0.45),("Yerli Dana-Düve",0.30),
    ("Koyun",0.10),("Keçi",0.08),("Manda (Erkek)",0.90),
    ("Manda (Dişi)",0.75),("Öküz",0.60),("Kuzu-Oğlak",0.04),
    ("Boğa",1.50),
]
TAHSIS_ASAMALARI = [
    "1-Dilekçe Alındı","2-Teknik İnceleme",
    "3-Komisyon Değerlendirmesi","4-Karar Yazıldı",
    "5-Tapu Müdürlüğüne Gönderildi","6-Tamamlandı",
]
IHALE_DURUMLARI = [
    "Başvuru Alındı","Evrak Kontrolü","Komisyon Değerlendirmesi",
    "İhale Yapıldı","Sözleşme Aşaması","Aktif","İptal Edildi","Tamamlandı",
]
ILCELER = ["Merkez","Göle","Çıldır","Posof","Hanak","Damal"]
CEZA_TURLERI = [
    "Kaçak Hayvan Otlatılması",
    "Otlatma Kapasitesi Üstünde Hayvan Sokulması",
    "Mera Alanında İzinsiz Yapılaşma",
    "Mera Alanını İzinsiz Sürme/Tahrip",
    "İzinsiz Ağaç Kesimi",
    "Diğer",
]
DB_PATH: Optional[str] = None

# ─── AÇIK RIZA METNİ ─────────────────────────────────────────────────────────
RIZA_METNI = """
MİRAS Enterprise — Kullanım Koşulları ve Açık Rıza Metni

1. Bu yazılım, T.C. Ardahan Valiliği İl Tarım ve Orman Müdürlüğü Çayır, Mera ve 
   Yem Bitkileri Şube Müdürlüğü bünyesinde kullanılmak üzere geliştirilmiştir.

2. Yazılımda yer alan yapay zekâ destekli özellikler (MERA AI) tavsiye niteliğinde 
   olup, üretilen evrak ve hesaplamaların doğruluğu kullanıcının sorumluluğundadır.

3. Otlatma kapasitesi hesaplamaları, ihale bedelleri ve idari para cezası tutarları 
   dahil tüm sayısal veriler kullanıcı tarafından kontrol edilmelidir.

4. Kullanıcı, sisteme girdiği verilerin doğruluğundan ve güncelliğinden sorumludur.

5. Sistem yedekleme hizmeti sunar ancak veri kaybına karşı ek önlemler kullanıcının 
   sorumluluğundadır.

6. Bu yazılımın kullanımından doğabilecek hata, zarar veya kayıplardan geliştirici 
   sorumlu tutulamaz.

7. Kullanıcı bu koşulları kabul ederek sisteme giriş yapar.

Geliştirici: Emre ÖZTÜRK — Ziraat Mühendisi
İletişim: miras75mera@gmail.com
"""

# ─── MEVZUAT REHBERİ ─────────────────────────────────────────────────────────
MEVZUAT_KARTLARI = {
    "Mera": {
        "baslik": "Mera Nedir?",
        "tanim": "Hayvanların otlatılması ve otundan yararlanılması için tahsis edilen veya kadimden beri bu amaçla kullanılan yerlerdir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/a",
        "detay": "Meralar Devletin hüküm ve tasarrufundadır. Özel mülkiyete geçirilemez, zaman aşımı uygulanamaz, sınırları daraltılamaz ve amacı dışında kullanılamaz."
    },
    "Yaylak": {
        "baslik": "Yaylak Nedir?",
        "tanim": "Çiftçilerin hayvanları ile birlikte yaz mevsiminde çıkarak otlatma ve hayvancılık yaptıkları yerlerdir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/b",
        "detay": "Yaylaklar da meralar gibi kamu malı niteliğinde olup tahsis amacı dışında kullanılamaz."
    },
    "Kışlak": {
        "baslik": "Kışlak Nedir?",
        "tanim": "Hayvanların kış mevsiminde barındırılması ve otundan yararlanılması için tahsis edilen veya kadimden beri bu amaçla kullanılan yerlerdir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/f",
        "detay": "Kışlaklar genellikle sıcak bölgelerde bulunur ve kış mevsiminde hayvanların korunması amacıyla kullanılır."
    },
    "Tespit": {
        "baslik": "Tespit Nedir?",
        "tanim": "Bir yerin mera, yaylak ve kışlak arazisi olup olmadığının resmi evrakla ve bilirkişi ifadeleri ile belgelendirilmesidir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/j",
        "detay": "Tespit işlemi, İl Mera Komisyonu tarafından gerçekleştirilir. Kadastro kayıtları, tapu sicili, mahkeme kararları ve bilirkişi beyanları değerlendirilir."
    },
    "Tahdit": {
        "baslik": "Tahdit Nedir?",
        "tanim": "Mera, yaylak ve kışlak arazisi olduğuna karar verilen yerlerin sınırlarının 1/5000 ölçekli haritalar üzerinde belirlenmesi ve arazi üzerinde kalıcı işaretlerle işaretlenmesidir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/i",
        "detay": "Tahdit işlemi teknik ekipler tarafından yapılır. Sonuçlar 30 gün askıda kalır, itirazlar 60 gün içinde karara bağlanır."
    },
    "Tahsis": {
        "baslik": "Tahsis Nedir?",
        "tanim": "Mera, yaylak ve kışlakların verimlilik ve sosyal adalet ilkelerine uygun şekilde düzenlenerek köy veya belediyeye bırakılmasıdır.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/ı",
        "detay": "Tahsis kararında otlatma kapasitesi, aile başına hayvan sayısı, sulama ve geçit yerleri belirtilir. Tahsis kararı Valilik onayına sunulur."
    },
    "BBHB": {
        "baslik": "Büyükbaş Hayvan Birimi (BBHB) Nedir?",
        "tanim": "Hayvan sayısının, 500 kg canlı ağırlığa çevrilerek ifade edilen standardize şeklidir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 3/n",
        "detay": "Hesaplama: Otlatma Kapasitesi = (Mera Alanı × Yem Verimi) / (Günlük Yem İhtiyacı × Otlatma Gün Sayısı). Ardahan için: 30 kg/gün × 135 gün = 4050 kg."
    },
    "Madde 14": {
        "baslik": "Madde 14 — Tahsis Amacı Değişikliği",
        "tanim": "Mera, yaylak ve kışlakların tahsis amacının enerji, turizm, kamu yatırımı gibi nedenlerle değiştirilmesidir.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 14",
        "detay": "Tahsis amacı değişikliği için 20 yıllık ot geliri hesaplanır. İlgili Bakanlık talebi, Maliye ve Valilik uygun görüşü gerekir. Tescilden itibaren 2 yıl içinde yatırıma başlanmazsa iptal edilir."
    },
    "Otlatma Hakkı": {
        "baslik": "Otlatma Hakkı",
        "tanim": "Çiftçi ailelerinin meradan yararlanma hakkıdır. Tahsis yapılan köyde en az 6 aydır ikamet şartı aranır.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 22",
        "detay": "Otlatma hakkından fazla hayvan otlatılamaz. Otlatma hakkı başka çiftçi ailelerine devredilemez. İhtiyaç fazlası mera alanları kiralanabilir."
    },
    "İdari Para Cezası": {
        "baslik": "İdari Para Cezası",
        "tanim": "Mera Kanunu'na aykırı davranan kişilere uygulanan idari yaptırımdır.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 27",
        "detay": "Mera, yaylak, kışlakları izinsiz kullanan, tahrip eden, kapasiteden fazla hayvan otlatanlara idari para cezası uygulanır. Mükerrer ihlallerde ceza artırılır."
    },
    "Mera Komisyonu": {
        "baslik": "Mera Komisyonu",
        "tanim": "Vali yardımcısı başkanlığında 8 kişiden oluşan, mera tespit, tahdit ve tahsis işlemlerini yürüten komisyondur.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 6",
        "detay": "Üyeler: Vali Yardımcısı (Başkan), İl Tarım Müdürü, Ziraat Mühendisi, DSİ temsilcisi, Orman temsilcisi, Muhtar, Milli Emlak temsilcisi, Kadastro temsilcisi."
    },
    "Kiralama": {
        "baslik": "Mera Kiralama İhaleleri",
        "tanim": "İhtiyaç fazlası mera alanlarının 2886 sayılı Devlet İhale Kanunu'na göre kiralanmasıdır.",
        "kanun": "4342 sayılı Mera Kanunu, Madde 12; Mera Yönetmeliği Madde 7, 13",
        "detay": "Kiralama bedeli: %25 Bakanlık hesabına, %75 köy hesabına yatırılır. Geçici teminat %30, kesin teminat %6 oranındadır. Damga vergisi ‰5.69 uygulanır."
    },
}

# ─── TC KİMLİK DOĞRULAMA ─────────────────────────────────────────────────────
def tc_dogrula(tc):
    """TC Kimlik numarası algoritma kontrolü"""
    if not tc or len(str(tc)) != 11:
        return False
    try:
        d = [int(c) for c in str(tc)]
    except ValueError:
        return False
    if d[0] == 0:
        return False
    # 10. hane kontrolü
    t1 = ((d[0]+d[2]+d[4]+d[6]+d[8])*7 - (d[1]+d[3]+d[5]+d[7])) % 10
    if t1 != d[9]:
        return False
    # 11. hane kontrolü
    t2 = sum(d[:10]) % 10
    if t2 != d[10]:
        return False
    return True

def telefon_format(tel):
    """Telefon numarası format kontrolü"""
    t = tel.replace(" ","").replace("-","").replace("(","").replace(")","")
    if t.startswith("+90"): t = t[3:]
    elif t.startswith("0"): t = t[1:]
    return len(t) == 10 and t.isdigit()

def tc_kontrol_ve_devam(tc, parent=None):
    """TC doğrula, geçersizse kullanıcıya sor — tüm modüllerde kullanılır"""
    if not tc: return True
    if tc_dogrula(tc): return True
    return messagebox.askyesno("TC Uyarı",f"'{tc}' geçerli bir TC numarası değil.\nDevam etmek istiyor musunuz?",parent=parent)

def veri_dogrula(alanlar):
    """Veri doğrulama kuralları. alanlar: [(ad, deger, tip), ...] tip: 'str','float+','int+','tc'"""
    for ad,deger,tip in alanlar:
        if tip=="str" and (not deger or not str(deger).strip()):
            messagebox.showerror("Veri Hatası",f"'{ad}' alanı boş bırakılamaz."); return False
        elif tip=="float+":
            try:
                v=float(deger or 0)
                if v<0: messagebox.showerror("Veri Hatası",f"'{ad}' negatif olamaz."); return False
            except ValueError:
                messagebox.showerror("Veri Hatası",f"'{ad}' geçerli bir sayı değil."); return False
        elif tip=="int+":
            try:
                v=int(deger or 0)
                if v<0: messagebox.showerror("Veri Hatası",f"'{ad}' negatif olamaz."); return False
            except ValueError:
                messagebox.showerror("Veri Hatası",f"'{ad}' geçerli bir tam sayı değil."); return False
        elif tip=="float_pos":
            try:
                v=float(deger or 0)
                if v<=0: messagebox.showerror("Veri Hatası",f"'{ad}' sıfırdan büyük olmalı."); return False
            except ValueError:
                messagebox.showerror("Veri Hatası",f"'{ad}' geçerli bir sayı değil."); return False
    return True

def form_temizle(alanlar_dict):
    """Form alanlarını temizle — tekrar eden kodu azaltır"""
    for k,w in alanlar_dict.items():
        try:
            if isinstance(w,ttk.Entry): w.delete(0,tk.END)
            elif isinstance(w,ttk.Combobox): w.set("")
            elif isinstance(w,tk.Text): w.delete("1.0",tk.END)
        except Exception: pass


def sayi_yaziya(sayi):
    """Sayıyı Türkçe yazıya çevir: 30000 → otuzbintürklirası"""
    birler=["","bir","iki","üç","dört","beş","altı","yedi","sekiz","dokuz"]
    onlar=["","on","yirmi","otuz","kırk","elli","altmış","yetmiş","seksen","doksan"]
    sayi=int(sayi)
    if sayi==0: return "sıfır"
    sonuc=""
    if sayi>=1000000:
        m=sayi//1000000
        sonuc+=("" if m==1 else birler[m])+"milyon"
        sayi%=1000000
    if sayi>=1000:
        b=sayi//1000
        if b==1: sonuc+="bin"
        else:
            y=b//100; k=(b%100)//10; br=b%10
            if y>0: sonuc+=("" if y==1 else birler[y])+"yüz"
            if k>0: sonuc+=onlar[k]
            if br>0: sonuc+=birler[br]
            sonuc+="bin"
        sayi%=1000
    if sayi>=100:
        y=sayi//100
        sonuc+=("" if y==1 else birler[y])+"yüz"
        sayi%=100
    if sayi>=10:
        sonuc+=onlar[sayi//10]
        sayi%=10
    if sayi>0:
        sonuc+=birler[sayi]
    return sonuc+"türklirası"

def sanitize(metin, max_uzunluk=500):
    """Kullanıcı girişini temizle — SQL/XSS koruması"""
    if not metin: return ""
    metin = str(metin).strip()
    # Tehlikeli karakterleri temizle
    metin = metin.replace("\x00", "")  # null byte
    # Uzunluk limiti
    if len(metin) > max_uzunluk:
        metin = metin[:max_uzunluk]
    return metin

def oto_dosya_no(prefix="CEZA", tablo="Idari_Cezalar"):
    """Otomatik dosya/rapor numarası üret — YYYY-PREFIX-NNN formatında"""
    yil = int(_yil())
    no_pattern = f"{yil}-{prefix}-"
    try:
        with db_baglan() as c:
            # Tablodaki son kaydın sayısını bul
            count = c.execute(f"SELECT COUNT(*) FROM {tablo} WHERE yil=?", (yil,)).fetchone()[0]
            return f"{no_pattern}{count+1:03d}"
    except Exception:
        return f"{no_pattern}001"

def genel_emri_ceza(yil=None):
    """Belirtilen yılın Valilik Genel Emri ceza miktarlarını DB'den oku"""
    yil = yil or int(_yil())
    if DB_PATH:
        try:
            with db_baglan() as c:
                r = c.execute("SELECT sayi,bb_ceza,kb_ceza,madde FROM Genel_Emri WHERE yil=?",(yil,)).fetchone()
                if r: return {"yil":yil,"sayi":r[0],"bb":r[1],"kb":r[2],"madde":r[3]}
        except Exception: pass
    # DB'de yoksa sabit değerlerden al
    if yil in GENEL_EMRI_GECMIS:
        g = GENEL_EMRI_GECMIS[yil]
        return {"yil":yil,"sayi":g["sayi"],"bb":g["bb"],"kb":g["kb"],"madde":g["madde"]}
    # Hiç yoksa en son bilinen yılı kullan
    return {"yil":yil,"sayi":f"{yil}/01","bb":CEZA_BB_TL,"kb":CEZA_KB_TL,"madde":"18. Madde"}

def muhammen_bedel_hesapla(alan_da, kuru_ot_verimi, otlatma_gun, kuru_ot_fiyati):
    """Muhammen bedel — Tahdit Raporu formülü"""
    try:
        alan=para_parse(alan_da); kov=para_parse(kuru_ot_verimi); ogs=para_parse(otlatma_gun); kof=para_parse(kuru_ot_fiyati)
        kapasite_bbhb = (alan * kov) / (ogs * GUNLUK_OT_BBHB)
        kapasite_kbhb = kapasite_bbhb * 10
        bedel = kapasite_bbhb * ogs * GUNLUK_OT_BBHB * kof
        return {"alan":alan,"kuru_ot":kov,"gun":ogs,"fiyat":kof,"bbhb":kapasite_bbhb,"kbhb":kapasite_kbhb,"bedel":bedel}
    except Exception: return None

TAD_KATSAYILAR = {
    "14/a - Madencilik":3,"14/a - Petrol":1,"14/a - Jeotermal Kaynak":1,
    "14/b - Turizm Yatırımı":2,"14/c - Kamu Yatırımı":1,
    "14/d - Köy Yerleşim":1,"14/d - İmar Planı":2,"14/d - Toprak Muhafaza":0,
    "14/d - Gen Kaynakları Koruma":1,"14/d - Milli Park":1,"14/d - Kültürel Varlık":1,
    "14/d - Akarsu Düzenleme":1,"14/d - Su Ürünleri":1,"14/d - Termal Tarım":1,
    "14/e - Köy Kanunu 13-14":0,"14/f - Güvenlik/Olağanüstü Hal":0,
    "14/g - Doğal Afet Yerleşim":0,"14/ğ - Elektrik (RES/GES/HES)":3,
    "14/ğ - Doğal Gaz":1,"14/ğ - Petrol Piyasası":1,
    "14/h - Jeotermal Sera":1,"14/ı - Kentsel Dönüşüm":2,
    "14/i - Endüstri/Teknoloji/OSB":0,"14/j - Elektronik Haberleşme":1,
}

def tad_ot_bedeli_hesapla(alan_da, vasif, kuru_ot_fiyati, faaliyet_konu):
    """TAD 20 yıllık ot bedeli — Bakanlık 31.07.2025 formülü"""
    try:
        alan=para_parse(alan_da); kof=para_parse(kuru_ot_fiyati)
        vi=MERA_VASIF.index(vasif)
        ust_vasif=MERA_VASIF[min(vi+1,len(MERA_VASIF)-1)]
        kov_ust=OT_VERIM_KURU[ust_vasif]
        katsayi=TAD_KATSAYILAR.get(faaliyet_konu,1)
        yillik_uretim = alan * kov_ust
        yillik_gelir = yillik_uretim * kof
        ot_bedeli_20 = yillik_gelir * 20
        tad_ucreti = ot_bedeli_20 * katsayi
        return {"alan":alan,"vasif":vasif,"ust_vasif":ust_vasif,"kov":OT_VERIM_KURU[vasif],
                "kov_ust":kov_ust,"fiyat":kof,"faaliyet":faaliyet_konu,"katsayi":katsayi,
                "yillik_uretim":yillik_uretim,"yillik_gelir":yillik_gelir,
                "ot_bedeli_20":ot_bedeli_20,"tad_ucreti":tad_ucreti}
    except Exception: return None
def hash_pw(p):
    if BCRYPT_OK: return bcrypt.hashpw(p.encode(), bcrypt.gensalt()).decode()
    return hashlib.sha256(p.encode()).hexdigest()

def verify_pw(p, h):
    if BCRYPT_OK and h.startswith("$2b$"): return bcrypt.checkpw(p.encode(), h.encode())
    return hashlib.sha256(p.encode()).hexdigest() == h

def strong_pw(p):
    if len(p)<6: return False,"En az 6 karakter olmalı."
    if not any(c.isupper() for c in p): return False,"Büyük harf gerekli."
    if not any(c.isdigit() for c in p): return False,"Rakam gerekli."
    return True,""

_db_lock = threading.Lock()

def db_baglan():
    """Ağ güvenli DB bağlantısı — busy_timeout + DELETE journal mode"""
    if not DB_PATH: return None
    conn = sqlite3.connect(DB_PATH, timeout=30)
    conn.execute("PRAGMA busy_timeout = 30000")  # 30 saniye kilit bekleme
    conn.execute("PRAGMA journal_mode = DELETE")  # Ağda WAL tehlikeli, DELETE güvenli
    conn.execute("PRAGMA synchronous = FULL")     # Tam disk senkronizasyonu
    conn.execute("PRAGMA locking_mode = NORMAL")  # Her işlem sonrası kilidi bırak
    return conn

def db_log(kul, islem, detay=""):
    if not DB_PATH: return
    try:
        with _db_lock:
            conn = db_baglan()
            if conn:
                with conn:
                    conn.execute("INSERT INTO Loglar(tarih,kul,islem,detay)VALUES(?,?,?,?)",
                        (datetime.now().strftime("%Y-%m-%d %H:%M:%S"),kul,islem,detay))
                conn.close()
    except Exception as e: logging.error(f"db_log:{e}")

class VeriOnbellek:
    """Sık kullanılan DB verilerini cache'ler — performans artışı"""
    def __init__(self):
        self._cache={}; self._zamanlari={}
        self.TTL=120  # saniye — 2dk sonra yenile
    def al(self,anahtar,sorgu_fn):
        """Cache'den al veya sorgu_fn çalıştırıp cache'e ekle"""
        now=time.time()
        if anahtar in self._cache and (now-self._zamanlari.get(anahtar,0))<self.TTL:
            return self._cache[anahtar]
        try:
            veri=sorgu_fn()
            self._cache[anahtar]=veri; self._zamanlari[anahtar]=now
            return veri
        except Exception as e:
            logging.error(f"cache:{anahtar}:{e}")
            return self._cache.get(anahtar,[])
    def temizle(self,anahtar=None):
        if anahtar:
            self._cache.pop(anahtar,None); self._zamanlari.pop(anahtar,None)
        else:
            self._cache.clear(); self._zamanlari.clear()

_onbellek = VeriOnbellek()

# ─── SESSION ─────────────────────────────────────────────────────────────────
class SessionManager:
    TIMEOUT=900
    def __init__(self):
        self._last=time.time(); self._lock=threading.Lock()
        self._cbs=[]; self._on=True
        threading.Thread(target=self._watch,daemon=True).start()
    def ping(self):
        with self._lock: self._last=time.time()
    def register(self,fn): self._cbs.append(fn)
    def stop(self): self._on=False
    def _watch(self):
        while self._on:
            time.sleep(30)
            with self._lock: exp=time.time()-self._last>self.TIMEOUT
            if exp:
                self._on=False
                for fn in self._cbs: fn()

# ─── YEDEKLEME ───────────────────────────────────────────────────────────────
class YedekYoneticisi:
    MAX=30
    def __init__(self,uid):
        self.uid=uid; BACKUP_DIR.mkdir(exist_ok=True)
        self.drive_path=self._load_drive_path()
        threading.Thread(target=self._oto,daemon=True).start()
        threading.Thread(target=self._otosave,daemon=True).start()
    def _load_drive_path(self):
        """Config'den Drive klasör yolunu oku"""
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE,encoding="utf-8") as f: data=json.load(f)
                dp=data.get("drive_path","")
                if dp and os.path.isdir(dp): return dp
            except Exception: pass
        return ""
    def _save_drive_path(self,yol):
        """Config'e Drive klasör yolunu kaydet"""
        self.drive_path=yol
        data={}
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE,encoding="utf-8") as f: data=json.load(f)
            except Exception: pass
        data["drive_path"]=yol
        with open(CONFIG_FILE,"w",encoding="utf-8") as f: json.dump(data,f)
    def set_drive_path(self,yol):
        """Drive klasörünü ayarla ve kaydet"""
        if yol and os.path.isdir(yol):
            self._save_drive_path(yol); return True
        return False
    def _oto(self):
        while True:
            time.sleep(86400)
            try: self.al(otomatik=True)
            except Exception as e: logging.error(f"oto yedek:{e}")
    def _otosave(self):
        """10 dakikada bir otomatik yedek"""
        while True:
            time.sleep(OTOSAVE_SEC)
            try: self.al(otomatik=True)
            except Exception: pass
    def al(self,otomatik=False):
        if not DB_PATH: raise ValueError("DB seçilmemiş")
        ts=datetime.now().strftime("%Y%m%d_%H%M%S")
        e="oto" if otomatik else "manuel"
        h=BACKUP_DIR/f"Miras_{e}_{ts}.db"
        shutil.copy2(DB_PATH,h); self._temizle()
        # Otomatik Drive kopyalama
        drive_ok=False
        if self.drive_path and os.path.isdir(self.drive_path):
            try:
                drive_dosya=os.path.join(self.drive_path,f"Miras_{e}_{ts}.db")
                shutil.copy2(DB_PATH,drive_dosya)
                drive_ok=True
                # Drive'da eski yedekleri temizle (son 10 tane kalsın)
                self._drive_temizle()
            except Exception as ex:
                logging.error(f"Drive kopyalama hatası: {ex}")
        detay=str(h)
        if drive_ok: detay+=f" + Drive"
        db_log(self.uid,"Yedekleme",detay); return str(h)
    def _temizle(self):
        ys=sorted(BACKUP_DIR.glob("*.db"),key=lambda p:p.stat().st_mtime)
        for y in ys[:-self.MAX]: y.unlink(missing_ok=True)
    def _drive_temizle(self):
        """Drive klasöründeki eski yedekleri temizle (son 10 kalsın)"""
        if not self.drive_path or not os.path.isdir(self.drive_path): return
        try:
            drive_files=sorted(
                [f for f in Path(self.drive_path).glob("Miras_*.db")],
                key=lambda p:p.stat().st_mtime)
            for f in drive_files[:-10]: f.unlink(missing_ok=True)
        except Exception as e: logging.error(f"drive_temizle:{e}")
    def listele(self):
        ys=sorted(BACKUP_DIR.glob("*.db"),key=lambda p:p.stat().st_mtime,reverse=True)
        return [{"ad":p.name,"yol":str(p),"boyut":f"{p.stat().st_size/1024:.1f} KB",
                 "tarih":datetime.fromtimestamp(p.stat().st_mtime).strftime("%d.%m.%Y %H:%M")}
                for p in ys]
    def geri(self,yol):
        try: self.al(True); shutil.copy2(yol,DB_PATH); db_log(self.uid,"Geri Yükleme",yol); return True
        except Exception as e: logging.error(f"geri:{e}"); return False

# ─── GEMİNİ AI ───────────────────────────────────────────────────────────────
SISTEM_PROMPT=(
    "Sen T.C. Ardahan İl Tarım ve Orman Müdürlüğü Çayır Mera Şubesi için geliştirilmiş "
    "MERA AI yapay zeka asistanısın. Mera Kanunu(4342), BBHB hesaplamaları, "
    "ihale/tahsis/islah süreçleri ve tarım mevzuatı konularında Türkçe, teknik ve "
    "kurumsal dilde yardımcı olursun. Kendini her zaman MERA AI olarak tanıt."
)

class GeminiAsistan:
    def __init__(self):
        self.model=None; self.chat=None; self.hazir=False
        self.api_key=self._key()
        if self.api_key and GEMINI_OK:
            threading.Thread(target=self._baglan,daemon=True).start()
    def _key(self):
        k=os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
        if k: return k
        if os.path.exists(API_KEY_FILE):
            with open(API_KEY_FILE,encoding="utf-8") as f: k=f.read().strip()
            if k:
                try:
                    import base64
                    return base64.b64decode(k).decode()
                except Exception:
                    return k  # fallback: düz metin
        return None
    def _baglan(self):
        try:
            genai.configure(api_key=self.api_key)
            secilen=None
            # Google'dan mevcut modelleri çek — hangisi varsa onu kullan
            try:
                mevcut=[]
                for m in genai.list_models():
                    n=m.name.replace("models/","")
                    try:
                        if "generateContent" in m.supported_generation_methods:
                            mevcut.append(n)
                    except Exception:
                        mevcut.append(n)
                # Önce flash tercih et (ücretsiz+hızlı), sonra pro, sonra ilk bulunan
                for m in mevcut:
                    if "flash" in m:
                        secilen=m; break
                if not secilen:
                    for m in mevcut:
                        if "pro" in m:
                            secilen=m; break
                if not secilen and mevcut:
                    secilen=mevcut[0]
            except Exception as ex:
                logging.warning(f"list_models: {ex}")
            # list_models çalışmadıysa bilinen modelleri sırayla dene
            if not secilen:
                for dene in ["gemini-2.5-flash","gemini-2.0-flash","gemini-1.5-flash","gemini-1.5-pro"]:
                    try:
                        test=genai.GenerativeModel(dene)
                        test.generate_content("merhaba")
                        secilen=dene; break
                    except Exception:
                        continue
            if not secilen:
                raise ValueError("Kullanılabilir Gemini modeli bulunamadı.")
            gen_config={"max_output_tokens": 2048}
            try:
                self.model=genai.GenerativeModel(secilen,system_instruction=SISTEM_PROMPT,generation_config=gen_config)
            except Exception:
                self.model=genai.GenerativeModel(secilen,generation_config=gen_config)
            self.chat=self.model.start_chat(history=[])
            self.hazir=True; logging.info(f"Gemini bağlandı: {secilen}")
        except Exception as e:
            logging.error(f"Gemini bağlantı hatası: {e}")
            self.hazir=False
    def key_kaydet(self,k):
        self.api_key=k.strip()
        import base64
        encrypted=base64.b64encode(self.api_key.encode()).decode()
        with open(API_KEY_FILE,"w",encoding="utf-8") as f: f.write(encrypted)
        self.hazir=False
        threading.Thread(target=self._baglan,daemon=True).start()
    def yeni_chat(self):
        if self.model: self.chat=self.model.start_chat(history=[])
    def sor(self,mesaj):
        if not self.hazir:
            return ("⚠️ MERA AI aktif değil.\n"
                    "API anahtarı için: aistudio.google.com/app/apikey\n"
                    "MERA AI sekmesi → API Ayarları bölümünden anahtarınızı girin.")
        return self._guvenli_istek(lambda: self.chat.send_message(mesaj).text)
    def tek(self,mesaj):
        if not self.hazir or not self.model: return "⚠️ Gemini bağlı değil."
        return self._guvenli_istek(lambda: self.model.generate_content(mesaj).text)
    def _guvenli_istek(self, fn, max_deneme=3):
        """429 hatasında otomatik bekle ve tekrar dene"""
        import re
        for deneme in range(max_deneme):
            try:
                return fn()
            except Exception as e:
                hata_str=str(e)
                if "429" in hata_str or "quota" in hata_str.lower() or "rate" in hata_str.lower():
                    # Bekleme süresini hatadan çıkar
                    bekleme=45  # varsayılan
                    m=re.search(r'retry in (\d+)',hata_str)
                    if m: bekleme=int(m.group(1))+2
                    m2=re.search(r'retry_delay.*?seconds:\s*(\d+)',hata_str,re.DOTALL)
                    if m2: bekleme=int(m2.group(1))+2
                    if deneme < max_deneme-1:
                        logging.warning(f"Gemini 429 - {bekleme}s bekleniyor (deneme {deneme+1}/{max_deneme})")
                        time.sleep(bekleme)
                        continue
                    else:
                        # Son deneme de başarısız
                        if "per day" in hata_str.lower() or "perday" in hata_str.lower():
                            return ("⚠️ Günlük AI kullanım kotası doldu.\n\n"
                                    "Google ücretsiz plan günlük istek sınırı aşıldı.\n"
                                    "Yarın sıfırlanacak veya farklı bir API key deneyin.\n\n"
                                    "💡 Kota bilgisi: ai.google.dev/gemini-api/docs/rate-limits")
                        else:
                            return (f"⚠️ AI şu an yoğun — {max_deneme} kez denendi.\n\n"
                                    f"Dakikalık istek sınırı aşıldı.\n"
                                    f"Lütfen {bekleme} saniye sonra tekrar deneyin.\n\n"
                                    "💡 Kota bilgisi: ai.google.dev/gemini-api/docs/rate-limits")
                else:
                    logging.error(f"Gemini hata: {e}")
                    return f"❌ AI Hatası: {hata_str[:300]}"
        return "❌ Beklenmeyen hata oluştu."

# ─── PDF MOTORU ───────────────────────────────────────────────────────────────
_FN="Helvetica"; _FNB="Helvetica-Bold"

def _init_fonts():
    global _FN,_FNB
    if not PDF_OK: return
    for r,b in [
        ("C:/Windows/Fonts/arial.ttf","C:/Windows/Fonts/arialbd.ttf"),
        ("C:/Windows/Fonts/calibri.ttf","C:/Windows/Fonts/calibrib.ttf"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ]:
        if os.path.exists(r):
            try:
                pdfmetrics.registerFont(TTFont("TF",r))
                pdfmetrics.registerFont(TTFont("TFB",b if os.path.exists(b) else r))
                _FN="TF"; _FNB="TFB"; return
            except Exception: continue

def uret_pdf(dosya,veri):
    if not PDF_OK: raise ImportError("reportlab kurulu degil")
    _init_fonts()
    fn,fnb=_FN,_FNB
    W,H=A4
    cv=rl_canvas.Canvas(dosya,pagesize=A4)
    def yaz(x,y,t,f=None,s=10,rk=None):
        cv.setFont(f or fn,s)
        cv.setFillColor(rk or rl_colors.black)
        cv.drawString(x*cm,y*cm,str(t))
    def yazm(y,t,f=None,s=10):
        cv.setFont(f or fn,s); cv.setFillColor(rl_colors.black)
        cv.drawCentredString(W/2,y*cm,str(t))
    def cizgi(y,sl=2.0,sr=19.0):
        cv.setStrokeColor(rl_colors.HexColor("#1E5631")); cv.setLineWidth(0.8)
        cv.line(sl*cm,y*cm,sr*cm,y*cm)

    # Başlık bandı
    cv.setFillColor(rl_colors.HexColor("#1E5631"))
    cv.rect(0,H-3.2*cm,W,3.2*cm,fill=1,stroke=0)
    # Alt dekoratif çizgi
    cv.setFillColor(rl_colors.HexColor("#2D8C55"))
    cv.rect(0,H-3.4*cm,W,0.2*cm,fill=1,stroke=0)
    cv.setFillColor(rl_colors.white)
    LOGO="bakanlik_logo.png"
    if os.path.exists(LOGO):
        try: cv.drawImage(LOGO,0.4*cm,H-3.0*cm,width=2.4*cm,height=2.4*cm,preserveAspectRatio=True,mask="auto")
        except Exception: pass
    cv.setFont(fnb,13); cv.drawCentredString(W/2,H-1.2*cm,"T.C. ARDAHAN VALİLİĞİ")
    cv.setFont(fn,10); cv.drawCentredString(W/2,H-1.8*cm,"İl Tarım ve Orman Müdürlüğü")
    cv.setFont(fn,9); cv.drawCentredString(W/2,H-2.4*cm,"Çayır, Mera ve Yem Bitkileri Şube Müdürlüğü")

    y=26.0
    yazm(y,"OTLATMA KAPASİTESİ RAPORU",fnb,12); y-=0.6
    yazm(y,f"ARDAHAN İLİ {veri['ilce'].upper()} İLÇESİ {veri['koy'].upper()} KÖYÜ",fn,10)
    y-=0.9; cizgi(y); y-=0.6
    yaz(2,y,f"Sayı : {veri['rapor_no']}",fn,10)
    yaz(13,y,f"Tarih: {datetime.now().strftime('%d.%m.%Y')}",fn,10)
    y-=0.5; yaz(2,y,f"Konu : {veri['ilce']} İlçesi {veri['koy']} Köyü Mera Otlatma Kapasitesi",fn,10)
    y-=0.5; yaz(2,y,f"Talep Eden: {veri['talep_eden']}   TC: {veri['tc']}",fn,10)
    y-=0.3; cizgi(y)

    ok=veri["ok_bbhb"]
    y-=0.7; yaz(2,y,"HESAPLAMA BİLGİLERİ",fnb,10); y-=0.5
    yaz(2.5,y,f"Mera Alanı: {veri['alan']:.0f} da  |  Yararlanılabilir Yeşil Yem: {veri['yem']:.0f} kg/da  |  Otlatma: {OTLATMA_GUN} gün  |  Günlük İhtiyaç: {GUNLUK_YEM_KG} kg",fn,9)
    y-=0.5
    yaz(2,y,f"OTLATMA KAPASİTESİ: {ok:.2f} BBHB   |   İşletme Sayısı: {veri['aktif']}   |   İşletme Başı: {ok/veri['aktif']:.2f} BBHB" if veri['aktif']>0 else f"OTLATMA KAPASİTESİ: {ok:.2f} BBHB",fnb,11,rl_colors.HexColor("#1E5631"))
    y-=0.3; cizgi(y)

    y-=0.6; yazm(y,"BBHB'YE GÖRE OTLATILABİLECEK HAYVAN SAYILARI",fnb,10); y-=0.4
    CW=[6.0,2.2,4.5,5.3]
    HDRS=["IRK VE TÜR","KATSAYI","TOPLAM HAYVAN SAYISI","İŞLETME BAŞINA"]
    cx=[2.0]
    for w in CW[:-1]: cx.append(cx[-1]+w)
    cv.setFillColor(rl_colors.HexColor("#1E5631"))
    cv.rect(2*cm,(y-0.45)*cm,sum(CW)*cm,0.5*cm,fill=1,stroke=0)
    cv.setFillColor(rl_colors.white); cv.setFont(fnb,8)
    for i,(h,x) in enumerate(zip(HDRS,cx)):
        cv.drawCentredString((x+CW[i]/2)*cm,(y-0.22)*cm,h)
    y-=0.5
    rh=0.42
    for idx,(tur,kat) in enumerate(HAYVAN_TURLERI):
        if y<4.5: cv.showPage(); y=27.0
        toplam=ok/kat if kat>0 else 0
        isb=toplam/veri["aktif"] if veri["aktif"]>0 else 0
        bg=rl_colors.HexColor("#EAF4EE") if idx%2==0 else rl_colors.white
        cv.setFillColor(bg)
        cv.rect(2*cm,(y-rh+0.1)*cm,sum(CW)*cm,rh*cm,fill=1,stroke=0)
        cv.setFillColor(rl_colors.black); cv.setFont(fn,8)
        for i,(txt,x) in enumerate(zip([tur,f"{kat:.2f}",f"{round(toplam)}",f"{round(isb)}"],cx)):
            cv.drawCentredString((x+CW[i]/2)*cm,(y-0.15)*cm,txt)
        cv.setStrokeColor(rl_colors.lightgrey)
        cv.rect(2*cm,(y-rh+0.1)*cm,sum(CW)*cm,rh*cm,fill=0,stroke=1)
        y-=rh
    y-=0.3; cizgi(y); y-=0.5
    if veri.get("aciklama"):
        yaz(2,y,f"Not: {veri['aciklama']}",fn,8)
    y-=0.8
    # Dekoratif kapanış çizgisi (imza yok — rapor bilgi amaçlıdır)
    cizgi(y)
    y-=0.4
    yaz(2,y,f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y')}",fn,9,rl_colors.HexColor("#555555"))
    yaz(12,y,f"Rapor No: {veri['rapor_no']}",fn,9,rl_colors.HexColor("#555555"))
    # Alt bilgi bandı
    cv.setFillColor(rl_colors.HexColor("#1E5631"))
    cv.rect(0,0,W,1.2*cm,fill=1,stroke=0)
    cv.setFillColor(rl_colors.white); cv.setFont(fn,7)
    cv.drawCentredString(W/2,0.5*cm,f"MİRAS Enterprise {VERSIYON} — {HAKLAR}")
    cv.save()

# ─── WORD ŞABLON DOLDURMA ─────────────────────────────────────────────────────
def _yil():
    """Mevcut yılı döndür"""
    return str(datetime.now().year)

def word_katilim_evrak(dosya, veri):
    """İhale Katılım Evrakları — ☐ checkbox, tek sayfa, TESLİM ALAN=personel"""
    if not DOCX_OK: raise ImportError("python-docx kurulu değil")
    doc = DocxDocument()
    yil = _yil()
    style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0); style.paragraph_format.space_after = Pt(1)
    style.paragraph_format.line_spacing = 1.0
    # Sayfa kenar boşlukları 0.6 cm — tek sayfaya sığsın
    for section in doc.sections:
        section.top_margin = Cm(0.6); section.bottom_margin = Cm(0.6)
        section.left_margin = Cm(0.6); section.right_margin = Cm(0.6)
    # Başlık
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("MERA KİRALAMA İHALELERİ KATILIM EVRAKLARI")
    r.bold = True; r.font.size = Pt(13); r.font.name = 'Times New Roman'
    # Kişi bilgileri
    for lbl,val in [("ADI SOYADI:",veri.get("ad_soyad","")),("T.C. NO:",veri.get("tc","")),
        ("İHALEYE GİRİLECEK MERA:",veri.get("mera","")),
        ("ADA/PARSEL:",veri.get("ada_parsel",""))]:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)
        run=p.add_run(f"{lbl}  "); run.bold=True; run.font.size=Pt(10); run.font.name='Times New Roman'
        if val: run2=p.add_run(str(val)); run2.font.size=Pt(10); run2.font.name='Times New Roman'
    # Evrak listesi başlığı
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(2)
    r=p.add_run("ALINAN EVRAKLAR (ZORUNLU OLANLAR)"); r.bold=True; r.font.size=Pt(11); r.font.name='Times New Roman'
    evraklar = [
        "Aile Nüfus Kayıt Örneği","İkametgâh Belgesi","Nüfus Cüzdanı Fotokopisi",
        "İki Adet Vesikalık Fotoğraf","Savcılık İyi Hal Kâğıdı (Sabıka Kaydı)",
        "Geçici Teminatın Yatırıldığını Gösterir Onaylı Dekont veya teminat mektubu",
        "İhaleye ortak olarak katılacakların noter tasdikli ortaklık sözleşmeleri",
        "Onaylı Hayvan Listesi (Büyükbaş Hayvan sayısı)",
        "Kiralama Yapılacak Alana Gelecek olanların (Çoban vb.) Nüfus Cüzdan Fotokopileri",
        "Gerçek Kişiler Adına Katılacakların Noter Tasdikli Vekâletnameleri",
        f"Tüzel kişilerin sicile kayıtlı olduklarını gösterir noter tasdikli vekâletname ve imza sirküsü ({yil})",
        f"Defterdarlığa yatırılacak olan {DOSYA_BEDELI_TL:.2f} ₺'lik Dosya Bedeli Makbuzu",
        "Hayvancılık İşletme Tescil Belgesi",
        f"Hayvanların 15.02.{yil} tarihinden önce işletmesine kayıtlı olduğunu gösteren resmi belge",
        "Hayvanlarını motorlu araçla getirecekleri yönünde noterden taahhütname",
        "SGK ve Maliyeden vadesi geçmiş borcu olmadığına dair belge",
    ]
    for i,ev in enumerate(evraklar):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0)
        run=p.add_run(f"☐ {i+1}. {ev}"); run.font.size=Pt(9); run.font.name='Times New Roman'
    # Teslim alan/eden — personel ve başvuru sahibi
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8)
    hazirlayan=veri.get("hazirlayan","")
    basvuru=veri.get("ad_soyad","")
    r1=p.add_run("TESLİM ALAN:  "); r1.bold=True; r1.font.size=Pt(10); r1.font.name='Times New Roman'
    r2=p.add_run(hazirlayan); r2.font.size=Pt(10); r2.font.name='Times New Roman'
    p.add_run("\t\t\t\t")
    r3=p.add_run("TESLİM EDEN:  "); r3.bold=True; r3.font.size=Pt(10); r3.font.name='Times New Roman'
    r4=p.add_run(basvuru); r4.font.size=Pt(10); r4.font.name='Times New Roman'
    # İmza satırı
    p2=doc.add_paragraph(); p2.paragraph_format.space_before=Pt(14)
    p2.add_run("İmza:\t\t\t\t\t\t\tİmza:").font.size=Pt(10)
    doc.save(dosya)

def word_kiralama_sozlesme(dosya, veri):
    """Kiralama Sözleşmesi — 14 maddelik orijinal şablon birebir eşleşme"""
    if not DOCX_OK: raise ImportError("python-docx kurulu değil")
    doc = DocxDocument()
    yil = _yil()
    style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)
    def _p(txt,bold=False,sz=11,center=False):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
        if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(txt); run.font.size=Pt(sz); run.font.name='Times New Roman'
        if bold: run.bold=True
        return p
    # Başlık
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(6)
    r=p.add_run("4342 SAYILI MERA KANUNU VE 31 TEMMUZ 1998 TARİHLİ MERA YÖNETMELİĞİNİN\n7.VE 13. MADDESİ KAPSAMINDA KİRALANAN MERA-YAYLA ALANLARININ KİRALAMA\nSÖZLEŞMESİDİR.")
    r.bold=True; r.font.size=Pt(11); r.font.name='Times New Roman'
    doc.add_paragraph()
    # Giriş metni
    _p("28 Şubat 1998 tarihli ve 4342 sayılı Mera Kanunu ve bu kanunun usul ve "
       "esaslarını belirleyen ve 31 Temmuz 1998 tarihli ve 23419 sayılı Resmi "
       "Gazete'de yayımlanarak yürürlüğe giren Mera Yönetmeliğinin 7.ve 13. "
       "Maddesine göre otlatma amacıyla kiraya verilen mera, yaylak umuma ait "
       "çayır ve otlak alanı için aşağıda belirtilen sözleşme hükümleri uygulanır.")
    doc.add_paragraph()
    _p("GENEL HÜKÜMLER",True,12)
    doc.add_paragraph()
    _p("Kiralanan Mera-Yayla Bulunduğu Yerin",True)
    _p(f"İli\t    :\tArdahan")
    _p(f"İlçesi        :\t{veri.get('ilce','')}")
    _p(f"Köyü\t:\t{veri.get('koy','')}")
    _p("Kiralanan Mera-Yaylanın Özellikleri")
    _p("Durumu ve sınıfı – İyi, kısmen orta")
    kap=veri.get('kapasite','……………..'); _p(f"Otlatma kapasitesi {kap} adet büyükbaş veya …………………..adet küçükbaş")
    doc.add_paragraph()
    _p("Kiralayanın",True)
    _p(f"Adı Soyadı: {veri.get('ad_soyad','……………………………')}")
    _p(f"Doğum Yeri ve Tarihi: {veri.get('dogum_yeri','………………………….')}")
    _p(f"Nüfusa Kayıtlı Olduğu Yer: ...............İli ……………… İlçesi ………………… Köyü")
    _p(f"İkametgahı: ……………. İli …………..İlçesi ……………… Köyü.")
    _p(f"Tebligat Adresi  ……………. İli ……………..  İlçesi ………………… Köyü.")
    doc.add_paragraph()
    bas=veri.get('baslangic',f'{OTLATMA_BASLANGIC}/{yil}'); bit=veri.get('bitis',f'{OTLATMA_BITIS}/{yil}')
    _p("3- Kiralama Süresi",True)
    _p(f"Kiranın Başlama Tarihi            : {bas}")
    _p(f"Kiralamanın Sona Erme Tarihi: {bit} (4 Ay)")
    doc.add_paragraph()
    _p("4- Otlatılacak Hayvanın Cinsi ve Miktarı :",True)
    doc.add_paragraph()
    _p("5- Otlatma süresi",True)
    _p(f"Otlatma başlangıcı\t   :{bas}")
    _p(f"Otlatmanın sona ermesi: {bit} (4 Ay)")
    # Maddeler — bazıları bold başlık, bazıları normal metin (orijinale uygun)
    _p("6- Kiralanan yer için komisyonca belirlenen ve ilan edilen hayvan sayısından fazla hayvan otlatılmayacaktır. Kiralanan yer başka amaçla kullanılmayacaktır.")
    doc.add_paragraph()
    _p("7- Kiracıya teslim edilen mera, yaylak ve kışlaklara 3. kişiler tarafından herhangi bir suretle yapılacak tecavüzleri, kiracı 7 gün içinde komisyona bildirecektir.",True)
    _p("8-Kiracılar, bölgelerinde huzursuzluk çıkartarak, ekili ve dikili alanlara tespit edildiğinde, Valilik onayı ile sözleşmeleri fesih edilerek kiraladıkları alandan çıkartılacaklar ve bu konuda herhangi bir hak talep edemeyeceklerdir.")
    doc.add_paragraph()
    _p("9- Kiracı bu hakkını devredemez, ortak alamaz, kiraya verilen mera, yaylak veya kışlağın sınırlarını daraltamaz, genişletemez, amacı dışında kullanamaz. Kiralanan yer dışında hayvan otlatamaz ve bundan doğan zararlar kiralayan tarafından ilgililere ödenir.",True)
    _p("10 – Kiracı, ihaleye girdiği sırada dosyasında kulak küpe numaralarını belirttiği hayvanların dışında hayvan meraya götüreceği zaman, söz konusu hayvanları nereden aldığı belgelemek (alım/satımda bedel banka üzerinden transfer edilecektir.) zorundadır.",True)
    doc.add_paragraph()
    _p("11-Komisyonun görevlendirdiği elemanların mera, yaylak ve kışlaklarda yapacakları çalışmaları kiracı tarafından hiçbir suretle engellenmeyecektir. Ayrıca söz konusu merada yapılacak olan hayvan tespitlerinde (Kulak küpe numarası) kiracı (hayvan sahibi) tarafından gerekli tedbirler (hayvanları çite kapatarak kulak küpe numarası tespitine yardımcı olmak) alınacaktır.",True)
    doc.add_paragraph()
    ge_soz=genel_emri_ceza(int(yil))
    _p(f"12-Kiracı (Göçer), {ge_soz['sayi']} sayılı Valilik Genel Emrinde belirtilen esaslara uymak ve yükümlülüklerini yerine getirmek zorundadır.",True)
    doc.add_paragraph()
    _p("13-Göçerler kiraladıkları mera alanına getirecekleri hayvanları ekte yer alan komisyon kararında belirlenen yerlere kadar araçla karayolundan getirmekle yükümlüdür. Hayvanları yaya olarak getiren göçerlerin sözleşmeleri fesih edilecek ve hakkında idari işlem uygulanacaktır.",True)
    doc.add_paragraph()
    _p("14-Yukarıda belirtilen maddelere uymayan kiracıların sözleşmesi tek taraflı olarak İl Mera Komisyonu kararı ile fesih edilecek olup, kiracı herhangi bir hak talep edemeyecektir.",True)
    # İmza bloğu
    doc.add_paragraph(); doc.add_paragraph()
    komisyon=veri.get("komisyon_baskani","Semih CEMBEKLİ")
    p=doc.add_paragraph()
    r1=p.add_run(komisyon); r1.font.name='Times New Roman'; r1.font.size=Pt(11); r1.bold=True
    p.add_run(f"\t\t\t\t\t{veri.get('ad_soyad','')}")
    p2=doc.add_paragraph()
    r2=p2.add_run("Vali Yardımcısı"); r2.font.name='Times New Roman'; r2.font.size=Pt(11)
    p2.add_run(f"\t\t\t\t\t\tTC: {veri.get('tc','')}")
    p3=doc.add_paragraph()
    r3=p3.add_run("İl Mera Komisyonu Başkanı"); r3.font.name='Times New Roman'; r3.font.size=Pt(11)
    p3.add_run("\t\t\t\t\tKiracı")
    doc.save(dosya)

def word_kiralama_sartname(dosya, veri):
    """Kiralama Şartnamesi — 22 madde Genel + 14 Özel + Özel Hükümler birebir orijinal"""
    if not DOCX_OK: raise ImportError("python-docx kurulu değil")
    doc = DocxDocument()
    yil = _yil()
    style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)
    for section in doc.sections:
        section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.2)
    def _p(txt,bold=False,sz=11,center=False):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
        if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(txt); run.font.size=Pt(sz); run.font.name='Times New Roman'
        if bold: run.bold=True
        return p
    ilce=veri.get('ilce','…………'); koy=veri.get('koy','………………….'); ada=veri.get('ada_parsel','……Ada…….Parsel')
    kap=veri.get('kapasite','………'); bas=f"{OTLATMA_BASLANGIC.replace("/",".")}.{yil}"; bit=f"{OTLATMA_BITIS.replace("/",".")}.{yil}"
    bedel=veri.get('bedel','')
    _p("MERA YAYLAK ve KIŞLAK ŞARTNAMESİ",True,14,True)
    doc.add_paragraph()
    _p("GENEL ŞARTLAR",True,12)
    doc.add_paragraph()
    # 22 Genel Madde — orijinal metinle birebir
    _p(f"Madde 1- Ardahan İli {ilce} İlçesi {koy} Köy kadastro alanı içerisinde yer alan "
       f"{ada} numaralı mera, yaylak ve kışlak ve umuma ait çayır, Ardahan İli Mera Komisyonu'nca "
       f"Ardahan İli sınırları içerisinde ikamet edip hayvancılıkla uğraşan üreticilere ihale ile kiraya verilecektir.",True)
    _p(f"Madde 2- Kiraya verilecek mera ,yaylak ve kışlağın otlatma kapasitesi {kap}\nadet Büyükbaş hayvan ile sınırlıdır.",True)
    _p(f"Madde 3- İhale Ardahan İli Mera komisyonunca 2886 Sayılı Devlet İhale Kanununun 51. Maddesi (g) bendi ve 4342 sayılı Mera Kanunu ve bu kanun uyarınca çıkartılan yönetmelik çerçevesinde Pazarlık usulü ile yapılacaktır.",True)
    _p(f"Madde 4- Mera yaylak ve kışlakların kiralama süresi 4 aydır. ({bas}– {bit})",True)
    # Geçici teminat hesaplama (%30)
    tah_bedel=veri.get('tahmini_bedel',0)
    gecici=tah_bedel*0.30 if tah_bedel else 0
    if gecici>0:
        gecici_str=f"{gecici:,.2f}".replace(",","X").replace(".",",").replace("X",".")
        gecici_yazi=sayi_yaziya(gecici)
        _p(f"Madde 5- Geçici teminat, tahmini bedelinin %30 'u olan {gecici_str} TL({gecici_yazi}) Defterdarlık Muhasebesinde İl Müdürlüğümüz (Ardahan İl Tarım ve Orman Müdürlüğü) adına işlem gören {IBAN_EMANET} IBAN numaralı emanet hesabına yatırılacaktır. Ayrıca açıklama kısmına Kiralanacak Alanın Ada/Parsel ve hangi Köyün kadastro sınırları içerisinde olduğu belirtilecektir.",True)
    else:
        _p(f"Madde 5- Geçici teminat, tahmini bedelinin %30 'u olan ………………(……………………) ₺ Defterdarlık Muhasebesinde İl Müdürlüğümüz (Ardahan İl Tarım ve Orman Müdürlüğü) adına işlem gören {IBAN_EMANET} IBAN numaralı emanet hesabına yatırılacaktır. Ayrıca açıklama kısmına Kiralanacak Alanın Ada/Parsel ve hangi Köyün kadastro sınırları içerisinde olduğu belirtilecektir.",True)
    _p(f"Madde 6- Kesin teminat, ihale bedelinin %6'sı olan ......................(……………….........) ₺ sözleşmenin Noterde onaylanmasının ardından Defterdarlık Muhasebesinde İl Müdürlüğümüz (Ardahan İl Tarım ve Orman Müdürlüğü) adına işlem gören {IBAN_EMANET} IBAN numaralı emanet hesabına yatırılacaktır.",True)
    _p("Madde 7- İhale komisyonu gerekçesini kararda belirtmek suretiyle ihaleyi yapıp yapmamakta serbesttir. Komisyonun ihaleyi yapmama kararına itiraz edilemez.",True)
    _p("Madde 8- İhale komisyonu tarafından alınan ihale kararları ita amirlerince karar tarihinden itibaren en geç 15 iş günü içinde onaylanır veya iptal edilir. İta amirlerince karar iptal edilirse ihale hükümsüz sayılır.",True)
    _p("İta amirlerince onaylanan ihale kararı onaylandığı günden itibaren an geç 5 iş günü içinde, kiracıya veya vekiline imzası alınmak suretiyle bildirilir veya iadeli taahhütlü mektupla tebligat adresine postalanır.")
    _p(f"Onaylanan ihale kararlarının yukarda açıklanan şekilde tebliğinden itibaren 15 gün içinde kesin teminatı yatırmak, komisyonca verilecek sözleşmeyi tasdik ettirerek idareye vermek ve bu süre içerisinde kira bedelinin % 25'ni T.C. Ziraat Bankası A.Ş Ankara Kamu Şubesindeki İBAN No: {IBAN_BAKANLIK} No'lu \"Mera Özel Gelir Hesabına\", % 75'ni ilgili köyün özel hesabına (Köy Sandığı) yatırarak dekontların İl Tarım ve Orman Müdürlüğü Çayır, Mera ve Yem Bitkileri Şube Müdürlüğüne vermek zorundadır.",True)
    _p("Komisyon; aynı süre içerisinde Mera, Yaylak veya Kamu Tüzel Kişiliği çayırı kiracıya ihale dokümanlarının arasında bulunan harita üzerinden ve/veya tanzim edilecek tutanakla şartnamede belirtilen sınıf ve evsafa göre teslim eder.")
    _p(f"Kira müddeti {bas} tarihinden itibaren başlar.",True)
    _p("Madde 9- Kiracılık hakkı sona erdiğinde kiraya verilen mera, yaylak, kışlak, kiracı tarafından İl Tarım ve Orman Müdürlüğü'ne (Çayır, Mera ve Yem Bitkileri Şube Müdürlüğü) müracaat ederek ilgili Köy Muhtarınca onaylanmış bir tutanakla veya gerek görülürse mera teknik ekibince incelenerek düzenlenecek rapor sonucuna göre teslim alınır.",True)
    _p("Madde 10- Kira süresi sona erdiği halde kiraya verilen yer kiracı tarafından komisyona müracaat ederek teslim edilmediği takdirde, her geçen gün için bedelin %2'si kadar ceza ödenir. Ayrıca idare tarafından resen tahliye edilir.",True)
    _p("Madde 11- Kiraya verilen mera, yaylak ve umuma ait çayır kiracı tarafından korunacak, değerini düşürmeyecek, özelliğini ve verim gücünü bozmayacak önlemleri almak; tedbirsizlik, dikkatsizlik, ihmal, kusur gibi nedenlerle vuku bulacak zarar ve ziyanları idarenin belirlediği rakam doğrultusunda defaten ödemek zorundadır.",True)
    _p("Madde 12- Kiracı bu hakkını devredemez, ortak alamaz, kiraya verilen mera, yaylak veya kışlağın sınırlarını daraltamaz, genişletemez, amacı dışında kullanamaz. Kiralanan yer dışında hayvan otlatamaz ve bundan doğan zararlar kiralayan tarafından ilgililere ödenir.",True)
    _p("Madde 13-  Sözleşme süresinin bitimi veya süresinden evvel iptali halinde tebligatı müteakip 15 gün içerisinde tahliye edilir.",True)
    _p("Madde 14- Göçerler görevlilerin denetiminde gerekli bilgi ve belgeleri göstermekle yükümlüdür. Ayrıca ihaleyi kazanan şahıslar çobanların kimlik bilgilerini, İl ve İlçe Jandarma Komutanlığına bildirilecek ve Jandarma Komutanlığınca kontrol edilecektir.",True)
    _p("Madde 15- Yukarıda yazılan hususlarla birlikte kiraya verilen mera, yaylak, kışlağın komisyonca belirleyeceği ve bu şartnameye veya kira sözleşmesinin özel şartlar, özel hükümler, genel hükümler kira kontratosu bölümüne ilave edeceği kurallara kiracı tarafından riayet edilmediği takdirde 2886 sayılı yasanın 62. maddesine göre işlem yapılır.",True)
    _p("Madde 16- Bu şartname Mera komisyonunun belirleyeceği diğer hususlarla birlikte sözleşmenin ekine teşkil edilir.",True)
    _p("Madde 17- Geçici teminat kira bedelinin tamamının yatırılmasından sonra, kesin teminat ise kira süresi sonunda ilgiliye ödenecektir.",True)
    _p("Madde 18- Gerektiğinde ihtilafların hal mercii Ardahan İcra Daireleri ve Mahkemeleridir.",True)
    _p("Madde 19- Göçer kiracılar 1774 Sayılı Kimlik Bildirme Kanununa göre 15 gün içerisinde kiraladıkları meranın bağlı bulunduğu Jandarma karakoluna bildirimde bulunacaklardır.",True)
    _p("Madde 20- Kiracılar, bölgelerinde huzursuzluk çıkartarak, ekili ve dikili alanlara zarar verdikleri tespit edildiğinde, Valilik onayı ile bölgeden çıkartılacaklar ve bu konuda herhangi bir hak talep edemeyeceklerdir.",True)
    _p("Madde 21- İhaleye birden fazla katılım olması durumunda en yüksek 1. teklifi veren kişinin ihale yükümlülüklerini komisyonda belirtilen sürede yerine getirmemesi halinde 2. en yüksek teklifi veren kişiye hak tanınır.",True)
    doc.add_paragraph()
    _p("Madde 22- İş bu şartname bu madde dâhil 22(yirmi bir) maddedir.",True)
    doc.add_paragraph()
    _p("Yukarıdaki maddeleri kabul ediyorum.")
    doc.add_paragraph()
    _p(f"Adı Soyadı: {veri.get('ad_soyad','')}",True)
    _p("İmza:",True)
    doc.add_paragraph()
    # ÖZEL ŞARTLAR
    _p("ÖZEL ŞARTLAR",True,12)
    ozel=[
        "1- Kiralanan yer için komisyonca belirlenen ve ilan edilen hayvan sayısından fazla hayvan otlatılmayacaktır. Kiralanan yer başka amaçla kullanılmayacaktır.",
        "2- Kiralamaya ilişkin her türlü vergi ve resmi harçlar kiracıya aittir.",
        "3- Kiracıya teslim edilen mera, yaylak ve kışlaklara 3'üncü kişiler tarafından herhangi bir suretle yapılacak tecavüzleri, kiracı 7 gün içinde komisyona bildirecektir.",
        "4- Mera, yaylak, kışlaktan faydalanacak sürü sahipleri Komisyonun/Jandarmanın vereceği kimlik belgelerini yanında bulunduracaklar. Kimlik belgesi olmayanlar kiraladıkları yere giremeyeceklerdir.",
        "5-Komisyonun görevlendirdiği elemanlar ve İl/İlçe Tarım ve Orman Müdürlüğü personellerince kiralanan mera, yaylak ve kışlakta ve umuma ait çayırlarda yapılacak çalışmalar kiracı tarafından hiçbir surette engellenemez.",
        f"6-Kiraya verilen mera, yaylak, kışlağın 4 aylık kira süresi ({bas}-{bit}) dışında süre uzatımı talebinde bulunulamaz.",
        "7-Kiracı, kiralanan alan, alan üzerinde bulunan bina, koruyucu çit duvar ve buna benzer tesislere zarar veremez, yenilerini tesis edemez.",
        "8-Yapılan ihtara rağmen kiracı tarafından kira bedelinin süresinde ödenmemesi ya da sözleşmede belirtilen diğer hususlara uyulmaması halinde sözleşme fesih edilir.",
        "9- Kiracı tarafından ödenmeyen kira bedeli, 6183 sayılı Amme Alacaklarının Tahsil Usulü Hakkında Kanun hükümlerine göre tahsil edilir.",
        "10- Sözleşmenin hazırlanmasının ardından göçerlerin mera, yaylak ve kışlaklara çıkabilmesi için gerekli belgeler ilgili yerin Mülki Amirliklerince düzenlenecektir.",
        "11- İhaleyi kazanan kimse herhangi bir şekilde üçüncü şahıslara mera, yaylak veya kışlağı kiralayamaz ve amacı dışında kullanılamaz.",
        "12- Hayvanlar yolun gittiği yere kadar motorlu araçlar ile götürülecektir.",
        f"13- Kiracı tespit edilen hayvan sayısı üzerinde hayvan getiremeyecektir. Aksi halde her büyükbaş hayvan için {CEZA_BB_TL:.2f} ₺, küçükbaş hayvan için {CEZA_KB_TL:.2f} ₺ İdari Para Cezası uygulanır.",
        f"14- {int(yil)-1} yılında mera, yaylak, kışlak ve umuma ait çayırları kiralayıp Jandarma tarafından, kiraladıkları yere kira şartnamesinde, sözleşmesinde ve Valilik Genel Emrinde belirtilen esaslara aykırı hareket edenlere yaptırım uygulanır.",
    ]
    for s in ozel: _p(s,True)
    doc.add_paragraph()
    _p("Bu şartnamedeki yazılı hususları olduğu gibi kabul ve taahhüt ederim. Her çeşit tebligat aşağıdaki adresime yapılabilir.")
    doc.add_paragraph()
    _p(f"Adı Soyadı\t:  {veri.get('ad_soyad','')}\t\t\t\tTebligat Adresi\t: {veri.get('adres','')}",True)
    _p("İmza\t  :",True)
    _p(f"İmza Tarihi: {veri.get('tarih',datetime.now().strftime('%d/%m/%Y'))}",True)
    doc.add_paragraph()
    # ÖZEL HÜKÜMLER
    _p("ÖZEL HÜKÜMLER",True,12)
    doc.add_paragraph()
    oh=[
        "Kiracı, kiralanan alan, alan üzerinde bulunan bina, koruyucu çit, duvar ve buna benzer tesislere zarar veremez, yenilerini inşa edemez.",
        "Kiralama ücreti her yıl en geç otlatma mevsimine başlamadan peşin olarak Mera Özel Gelir Hesabına yatırılır.",
        "Bu alanları kiralayanlar, Komisyonlarca belirlenen ıslah, amenajman planlarını ve otlatma planlarına uymakla yükümlüdürler. Komisyonlar, yükümlülüklere uymayanların sözleşmelerini feshederler.",
        "Kiracı tarafından kira bedelinin yapılan ihtara rağmen süresinde ödenmemesi ya da sözleşmede belirtilen diğer hususlara uyulmaması halinde, sözleşme fesih edilir.",
        "Kiracı tarafından ödenmeyen kira bedeli, 6183 sayılı Amme Alacaklarını Tahsil Usulü Hakkında Kanun hükümlerine göre tahsil edilir.",
        "Kiracılar, Otlatma bedelinin tamamını otlatma izni verildiğini belirten sözleşmenin yapılması sırasında Mera Özel Gelir Hesabına yatırır.",
        "Komisyon gerekli gördüğü takdirde bölgelerinin özel şartlarına göre sözleşmeye özel hükümler eklemeye yetkilidirler.",
        "Bu sözleşmeyle Genel Şartlar, Özel Şartlar ve Özel Hükümleri olduğu gibi kabul ve taahhüt ederim. Her çeşit yazışma ve tebligatlar, tebligat adresime yapılabilir.",
        "Göçerlerin yayla ve kışlaklarına hareket edecek olan büyükbaş hayvan sürülerine Şap, Şarbon ve LSD  aşılanmış olması mecburidir. Ayrıca bu hayvanların menşe şahadetnamesi ile tüberküloz ve brusella testi kontrollerinin yapılmış olması gerekmektedir.",
    ]
    for h in oh: _p(h)
    doc.add_paragraph()
    # Başvuru sahibi ad soyad + TC
    _p(f"{veri.get('ad_soyad','')}")
    _p(f"TC: {veri.get('tc','')}")
    doc.add_paragraph()
    komisyon=veri.get("komisyon_baskani","Semih CEMBEKLİ")
    _p(komisyon)
    p=doc.add_paragraph()
    r1=p.add_run("Vali Yardımcısı"); r1.font.name='Times New Roman'; r1.font.size=Pt(11)
    p.add_run(f"\t\t\t\t\t\t{veri.get('ad_soyad','')}")
    _p(f"İl Mera Komisyonu Başkanı \t\t\t\t\t          Kiralayan Şahıs")
    doc.save(dosya)

def word_idari_ceza(dosya, veri):
    """İdari Para Cezası Kararı Oluru — orijinal belgeyle birebir aynı yapı"""
    if not DOCX_OK: raise ImportError("python-docx kurulu değil")
    doc = DocxDocument()
    yil = _yil()
    # Sayfa boyutunu küçült — tek A4'e sığsın
    for section in doc.sections:
        section.top_margin = Cm(1.0); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.5); section.right_margin = Cm(1.5)
    style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(9)
    style.paragraph_format.space_before = Pt(0); style.paragraph_format.space_after = Pt(0)
    # Ana tablo — 26 satır × 8 sütun (orijinal yapı)
    tbl = doc.add_table(rows=26, cols=8, style='Table Grid')
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    def _mc(r,c1,c2,txt,bold=False,sz=9):
        merged=tbl.cell(r,c1).merge(tbl.cell(r,c2))
        merged.text=""
        p=merged.paragraphs[0]; run=p.add_run(str(txt))
        run.font.size=Pt(sz); run.font.name='Times New Roman'
        if bold: run.bold=True
    def _c(r,c,txt,bold=False,sz=9):
        cell=tbl.cell(r,c); cell.text=""
        p=cell.paragraphs[0]; run=p.add_run(str(txt))
        run.font.size=Pt(sz); run.font.name='Times New Roman'
        if bold: run.bold=True
    # Satır 0-1: Başlık
    _mc(0,0,7,"İDARİ PARA CEZASI UYGULANAN ŞAHSIN",True,10)
    _mc(1,0,7,"KİMLİK BİLGİLERİ",True,10)
    tbl.cell(0,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    tbl.cell(1,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    # Satır 2-10: Kimlik
    _mc(2,0,1,"Adı Soyadı",True); _mc(2,2,7,veri.get("ad_soyad",""))
    _mc(3,0,1,"T.C. Kimlik No",True); _mc(3,2,7,veri.get("tc",""))
    _mc(4,0,1,"Baba Adı",True); _mc(4,2,7,veri.get("baba_adi",""))
    _mc(5,0,1,"Doğum Yeri-Tarihi",True); _mc(5,2,7,veri.get("dogum",""))
    _c(6,0,"Nüfusa Kayıtlı Olduğu",True); _c(6,1,"İl",True); _mc(6,2,7,"Ardahan")
    _c(7,0,"Nüfusa Kayıtlı Olduğu",True); _c(7,1,"İlçe",True); _mc(7,2,7,veri.get("ilce","Merkez"))
    _mc(8,0,1,"Belge Türü",True); _c(8,2,"Kimlik"); _c(8,3,"X"); _c(8,4,"Ehliyet"); _c(8,5,""); _c(8,6,"Diğer"); _c(8,7,"")
    _mc(9,0,1,"Plaka No-Ehliyet Ruhsat No",True); _mc(9,2,7,"")
    _mc(10,0,1,"İkametgâh Adresi",True); _mc(10,2,7,veri.get("adres",""))
    # Satır 11-14: Ceza bilgileri
    _mc(11,0,7,"İDARİ PARA CEZASINA İLİŞKİN BİLGİLER",True,10)
    tbl.cell(11,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    tutar_str = veri.get("tutar","")
    if isinstance(tutar_str,(int,float)): tutar_str=para_format(tutar_str).replace(" ₺"," TL")
    _mc(12,0,2,"Miktarı(Rakamla)",True); _mc(12,3,7,tutar_str)
    _mc(13,0,2,"Ödeneceği Yer",True); _mc(13,3,7,ODEME_YERI)
    _mc(14,0,2,"Son Ödeme Tarihi",True); _mc(14,3,7,"Tebliğ tarihinden itibaren en geç 30 (otuz) gün içerisinde")
    # Satır 15-21: Yasal dayanak — yılın Genel Emri'nden dinamik
    _mc(15,0,7,"(*) İDARİ PARA CEZASININ YASAL DAYANAĞI",True,10)
    tbl.cell(15,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    ge=genel_emri_ceza(int(yil))
    dayanak=(f"4342 Sayılı Mera Kanunu 26. ve 27. Maddeleri hükümlerine göre, "
        f"{ge['sayi']} Sayılı Valilik Genel Emrinin {ge['madde']}, "
        f"Çayır Mera ve Yem Bitkileri Şube Müdürlüğü ekiplerinin "
        f"{veri.get('tarih',datetime.now().strftime('%d/%m/%Y'))} tarihinde "
        f"{veri.get('mera','')} mera alanında yaptıkları denetim sonucunda "
        f"{veri.get('ad_soyad','')} isimli şahsın {veri.get('konu','')} "
        f"fiilinden dolayı İdari Para Cezası uygulanmasına karar verilmiştir.")
    _mc(16,0,7,dayanak,False,8)
    for i in range(17,20): _mc(i,0,7,"")
    not_txt=(f"Not: (*) {ge['sayi']} Sayılı Valilik Genel Emrinin {ge['madde']}; "
        "meralara kapasitesinden fazla veya izinsiz hayvan sokan ya da meralardan izinsiz "
        "faydalanan kişilere 4342 sayılı Mera Kanununun 26. maddesinde belirtilen İPC uygulanır.")
    _mc(20,0,7,not_txt,False,7)
    _mc(21,0,7,"")
    # Satır 22-25: İtiraz
    _mc(22,0,7,"KESİLEN İDARİ PARA CEZASINA",True,10)
    tbl.cell(22,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    _mc(23,0,2,"İtiraz Merci",True); _mc(23,3,7,ITIRAZ_MERCI)
    _mc(24,0,2,"Son İtiraz Tarihi",True); _mc(24,3,7,"Tebliğ tarihinden itibaren 15 (onbeş) gün içerisinde")
    _mc(25,0,7,"Kararın tebliğ tarihinden itibaren 15(onbeş) gün içerisinde yetkili Sulh Ceza Hakimliğine başvurulabilir. Süresinde başvurulmaması halinde karar kesinleşir.",False,8)
    # Açıklama paragrafı
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run=p.add_run(f"Yukarıda açık kimliği yazılı {veri.get('ad_soyad','')} hakkında "
        f"4342 Sayılı Mera Kanununun 26. Maddesi gereğince {tutar_str} "
        f"İdari Para Cezası uygulanmasına karar verilmiştir.")
    run.font.size=Pt(10); run.font.name='Times New Roman'
    # İl Müdürü tablosu
    doc.add_paragraph().paragraph_format.space_before=Pt(2)
    tbl2=doc.add_table(rows=2,cols=1); tbl2.alignment=WD_TABLE_ALIGNMENT.CENTER
    il_muduru=veri.get("il_muduru","Muhammet Fatih CİNEVİZ")
    r2=tbl2.cell(0,0).paragraphs[0].add_run(il_muduru); r2.font.size=Pt(10); r2.bold=True; r2.font.name='Times New Roman'
    tbl2.cell(0,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r3=tbl2.cell(1,0).paragraphs[0].add_run("İl Müdürü"); r3.font.size=Pt(10); r3.font.name='Times New Roman'
    tbl2.cell(1,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    # OLUR tablosu
    doc.add_paragraph().paragraph_format.space_before=Pt(2)
    tbl3=doc.add_table(rows=6,cols=1); tbl3.alignment=WD_TABLE_ALIGNMENT.CENTER
    olur_data=[("OLUR",True,11),(f"…../{datetime.now().strftime('%m')}/{yil}",False,10),
        (veri.get("komisyon_baskani","Semih CEMBEKLİ"),True,10),
        (veri.get("evrak_unvan_1","Vali a."),False,10),
        (veri.get("evrak_unvan_2","Vali Yardımcısı"),False,10),
        ("Mera Komisyon Başkanı",False,10)]
    for i,(txt,bld,sz) in enumerate(olur_data):
        cell=tbl3.cell(i,0); cell.text=""
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(txt); run.font.size=Pt(sz); run.bold=bld; run.font.name='Times New Roman'
    doc.save(dosya)

def word_bilgi_notu(dosya, veri):
    """Bilgi Notu Word şablonu"""
    if not DOCX_OK: raise ImportError("python-docx kurulu değil")
    doc = DocxDocument()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BİLGİ NOTU"); r.bold = True; r.font.size = Pt(14)
    doc.add_paragraph()
    tbl = doc.add_table(rows=4, cols=2, style='Table Grid')
    bilgiler = [
        ("Sunulan Makam", veri.get("makam","")),
        ("Konu", veri.get("konu","")),
        ("Hazırlayan Birim", veri.get("birim","Çayır, Mera ve Yem Bitkileri Şube Müdürlüğü")),
        ("Hazırlama Tarihi", veri.get("tarih", datetime.now().strftime("%d.%m.%Y"))),
    ]
    for i,(lbl,val) in enumerate(bilgiler):
        tbl.cell(i,0).text = lbl
        tbl.cell(i,1).text = str(val)
        for cell in tbl.rows[i].cells:
            for p2 in cell.paragraphs:
                for r2 in p2.runs: r2.font.size = Pt(11)
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("İçerik:"); r.bold = True
    doc.add_paragraph(veri.get("icerik",""))
    doc.add_paragraph()
    doc.add_paragraph(f"Hazırlayan: {veri.get('hazirlayan','')}")
    doc.save(dosya)

# ─── İHALE BEDEL HESAPLAMA ───────────────────────────────────────────────────
def para_parse(metin):
    """Türk formatındaki para değerini float'a çevir: 52.507,18 → 52507.18"""
    if not metin: return 0.0
    t=str(metin).strip().replace("₺","").replace("TL","").strip()
    # Türk formatı: 52.507,18 → nokta binlik, virgül kuruş
    if "," in t and "." in t:
        t=t.replace(".","").replace(",",".")
    elif "," in t:
        t=t.replace(",",".")
    # sadece nokta varsa → zaten ondalık (1500.50) veya binlik (1.500)?
    # Eğer noktadan sonra 3+ hane varsa binlik ayracı
    elif "." in t:
        parcalar=t.split(".")
        if len(parcalar)==2 and len(parcalar[1])==3:
            t=t.replace(".","")  # binlik ayracı
    return float(t)

def para_format(sayi):
    """Float'ı Türk para formatına çevir: 52507.18 → 52.507,18 ₺"""
    try:
        s=f"{float(sayi):,.2f}"  # 52,507.18
        # İngilizce → Türkçe: virgül↔nokta swap
        s=s.replace(",","X").replace(".",",").replace("X",".")
        return f"{s} ₺"
    except (ValueError,TypeError):
        return "0,00 ₺"

def ihale_bedel_hesapla(bedel):
    """İhale bedelinden tüm kalemleri hesaplar"""
    try:
        b = para_parse(bedel)
        if b <= 0: return None
        return {
            "toplam": b,
            "bakanlik_25": b * 0.25,
            "koy_75": b * 0.75,
            "kesin_teminat_6": b * 0.06,
            "damga_vergisi": b * 5.69 / 1000,
        }
    except (ValueError, TypeError):
        return None

# ─── ARAYÜZ BİLEŞENLERİ ──────────────────────────────────────────────────────
class ToolTip:
    """Fare ile widget üzerine gelince açıklama gösterir"""
    def __init__(self, widget, text, delay=500):
        self.widget = widget; self.text = text; self.delay = delay
        self.tip_win = None; self._id = None
        widget.bind("<Enter>", self._schedule)
        widget.bind("<Leave>", self._hide)
        widget.bind("<ButtonPress>", self._hide)
    def _schedule(self, event=None):
        self._id = self.widget.after(self.delay, self._show)
    def _show(self):
        if self.tip_win: return
        x = self.widget.winfo_rootx() + 20; y = self.widget.winfo_rooty() + self.widget.winfo_height() + 4
        self.tip_win = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True); tw.wm_geometry(f"+{x}+{y}")
        lbl = tk.Label(tw, text=self.text, bg="#FFF8DC", fg="#333", font=("Segoe UI", 9),
            relief="solid", borderwidth=1, padx=8, pady=4, wraplength=300)
        lbl.pack()
    def _hide(self, event=None):
        if self._id: self.widget.after_cancel(self._id); self._id = None
        if self.tip_win: self.tip_win.destroy(); self.tip_win = None

class MBtn(tk.Button):
    def __init__(self,master,text,command=None,color="#2D8C55",tooltip=None,**kw):
        kw.setdefault("pady",7); kw.setdefault("padx",14)
        super().__init__(master,text=text,command=command,bg=color,fg="white",
            font=("Segoe UI",10,"bold"),relief="flat",activebackground="#5DADE2",cursor="hand2",**kw)
        self._orig=text
        if tooltip: ToolTip(self, tooltip)
    def loading(self,d):
        self.config(state="disabled" if d else "normal",text="⏳ İşleniyor..." if d else self._orig)

class StatKart(tk.Frame):
    def __init__(self,master,baslik,deger,renk="#1E5631",ikon="",**kw):
        super().__init__(master,bg=C_WHITE,relief="flat",highlightbackground="#D0DDD8",highlightthickness=1,**kw)
        tk.Label(self,text=ikon,font=("Segoe UI",22),bg=C_WHITE).pack(pady=(14,2))
        self._v=tk.Label(self,text=str(deger),font=("Segoe UI",26,"bold"),fg=renk,bg=C_WHITE); self._v.pack()
        tk.Label(self,text=baslik,font=("Segoe UI",9),fg="#666",bg=C_WHITE).pack(pady=(2,14))
    def set(self,v): self._v.config(text=str(v))

class StatusBar(tk.Label):
    def __init__(self,master,**kw):
        super().__init__(master,text="  Hazır.",anchor="w",bg="#D5E5DC",fg="#1E5631",font=("Segoe UI",9),padx=12,**kw)
    def set(self,msg): self.config(text=f"  {msg}"); self.update_idletasks()

class AramaFrame(tk.Frame):
    """Canlı arama — yazarken anında filtreler (debounce 200ms)"""
    def __init__(self,master,cb,ph="🔍 Ara...",**kw):
        super().__init__(master,**kw); self._cb=cb; self._timer=None
        self.var=tk.StringVar(); self.var.trace_add("write",self._on_change)
        tk.Label(self,text="🔍",font=("Segoe UI",11),bg=self["bg"]).pack(side="left",padx=(0,4))
        ttk.Entry(self,textvariable=self.var,width=30).pack(side="left")
        tk.Button(self,text="✕",command=lambda:self.var.set(""),relief="flat",cursor="hand2",bg=self["bg"]).pack(side="left",padx=4)
    def _on_change(self,*_):
        if self._timer: self.after_cancel(self._timer)
        self._timer=self.after(200,lambda:self._cb(self.var.get()))

# ─── VERİTABANI ───────────────────────────────────────────────────────────────
def init_db():
    if not DB_PATH: return
    with db_baglan() as c:
        c.execute("PRAGMA foreign_keys = ON")
        c.executescript("""
        CREATE TABLE IF NOT EXISTS Kullanicilar(k_adi TEXT PRIMARY KEY,sifre TEXT NOT NULL,
            yetki TEXT NOT NULL DEFAULT 'Uzman',ad TEXT NOT NULL,unvan TEXT,
            aktif INTEGER NOT NULL DEFAULT 1,fail_count INTEGER NOT NULL DEFAULT 0,lockout_ts TEXT,
            riza_onay INTEGER DEFAULT 0,gorevler TEXT DEFAULT '*');
        CREATE TABLE IF NOT EXISTS Ayarlar(k_adi TEXT PRIMARY KEY,tema TEXT DEFAULT 'Orman Yeşili',
            punto INTEGER DEFAULT 10,sube_mudur TEXT DEFAULT 'Leyla ARSLAN',vali_yardimcisi TEXT DEFAULT 'Semih CEMBEKLİ');
        CREATE TABLE IF NOT EXISTS Loglar(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tarih TEXT NOT NULL,kul TEXT,islem TEXT,detay TEXT);
        CREATE TABLE IF NOT EXISTS Duyurular(id INTEGER PRIMARY KEY AUTOINCREMENT,
            mesaj TEXT NOT NULL,tarih TEXT NOT NULL,gonderen TEXT);
        CREATE TABLE IF NOT EXISTS Mera_Varligi(koy TEXT PRIMARY KEY,ilce TEXT NOT NULL,alan REAL NOT NULL DEFAULT 0,yem REAL NOT NULL DEFAULT 0,turkvet_tarihi TEXT);
        CREATE TABLE IF NOT EXISTS Rapor_Gecmisi(rapor_no TEXT PRIMARY KEY,koy TEXT NOT NULL,
            talep_eden TEXT,tc TEXT,islem_tarihi TEXT NOT NULL,duzenleyen TEXT,aciklama TEXT);
        CREATE TABLE IF NOT EXISTS Ihaleler(id INTEGER PRIMARY KEY AUTOINCREMENT,
            koy TEXT,ilce TEXT,ad_soyad TEXT,tc TEXT,telefon TEXT,adres TEXT,
            bedel REAL,durum TEXT DEFAULT 'Başvuru Alındı',notlar TEXT,tarih TEXT,
            ada_parsel TEXT,kapasite TEXT);
        CREATE TABLE IF NOT EXISTS Ihale_Log(id INTEGER PRIMARY KEY AUTOINCREMENT,
            ihale_id INTEGER,tarih TEXT,personel TEXT,durum TEXT,not_icerik TEXT);
        CREATE TABLE IF NOT EXISTS Islah_Amenajman(id INTEGER PRIMARY KEY AUTOINCREMENT,
            koy TEXT,ilce TEXT,dilekce_tarihi TEXT,talep_eden TEXT,talep_alani REAL,
            talep_aciklama TEXT,verilen_alan REAL DEFAULT 0,verilmeme_neden TEXT,
            is_programi TEXT,durum TEXT DEFAULT 'Bekliyor',kapanma_tarihi TEXT);
        CREATE TABLE IF NOT EXISTS Tahsisler(id INTEGER PRIMARY KEY AUTOINCREMENT,
            koy TEXT,ilce TEXT,ada TEXT,parsel TEXT,kurum TEXT,amac TEXT,alan_ha REAL,
            asama TEXT DEFAULT '1-Dilekçe Alındı',durum TEXT DEFAULT 'Devam Ediyor',
            basvuru_t TEXT,sonuc_t TEXT,notlar TEXT,
            madde14_bent TEXT,ot_geliri REAL,tescil_tarihi TEXT,
            sure_bitis TEXT,sure_tipi TEXT);
        CREATE TABLE IF NOT EXISTS Tahsis_Log(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tahsis_id INTEGER,tarih TEXT,personel TEXT,asama TEXT,aciklama TEXT);
        CREATE TABLE IF NOT EXISTS Muhtarlar(id INTEGER PRIMARY KEY AUTOINCREMENT,
            ilce TEXT,koy TEXT,ad_soyad TEXT,telefon TEXT,email TEXT,notlar TEXT);
        CREATE TABLE IF NOT EXISTS Veri_Kayit(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tarih TEXT,kul TEXT,baslik TEXT,kategori TEXT,detay TEXT);
        CREATE TABLE IF NOT EXISTS Ajanda(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tarih TEXT,sure TEXT,baslik TEXT,tur TEXT DEFAULT 'Hatırlatıcı',
            durum TEXT DEFAULT 'Bekliyor',icerik TEXT,k_adi TEXT);
        CREATE TABLE IF NOT EXISTS Sifre_Talepleri(id INTEGER PRIMARY KEY AUTOINCREMENT,
            k_adi TEXT,tarih TEXT,durum TEXT DEFAULT 'Bekliyor');
        CREATE TABLE IF NOT EXISTS Kayit_Talepleri(id INTEGER PRIMARY KEY AUTOINCREMENT,
            ad TEXT,unvan TEXT,k_adi TEXT,tarih TEXT,durum TEXT DEFAULT 'Bekliyor');
        CREATE TABLE IF NOT EXISTS Sikayetler(id INTEGER PRIMARY KEY AUTOINCREMENT,
            koy TEXT,ilce TEXT,sikayet_eden TEXT,tur TEXT,aciklama TEXT,
            durum TEXT DEFAULT 'Yeni',tarih TEXT,kapanma TEXT,sonuc TEXT);
        CREATE TABLE IF NOT EXISTS Personel_Takvim(id INTEGER PRIMARY KEY AUTOINCREMENT,
            k_adi TEXT,ad TEXT,tarih TEXT,baslangic TEXT,bitis TEXT,
            tur TEXT DEFAULT 'Görev',aciklama TEXT);
        CREATE TABLE IF NOT EXISTS Son_Islemler(id INTEGER PRIMARY KEY AUTOINCREMENT,
            k_adi TEXT,modul TEXT,kayit_adi TEXT,tarih TEXT);
        CREATE TABLE IF NOT EXISTS Idari_Cezalar(id INTEGER PRIMARY KEY AUTOINCREMENT,
            ad_soyad TEXT,tc TEXT,il TEXT DEFAULT 'Ardahan',ilce TEXT,
            mera_koy TEXT,mera_ada_parsel TEXT,yil INTEGER,konu TEXT,
            hayvan_sayisi TEXT,ceza_miktari TEXT,ipc_tutari REAL,
            tarih TEXT,durum TEXT DEFAULT 'Uygulandı',notlar TEXT);
        CREATE TABLE IF NOT EXISTS Ihale_Yerleri(id INTEGER PRIMARY KEY AUTOINCREMENT,
            ilce TEXT,koy TEXT,ada TEXT,parsel TEXT,alan_da REAL,kapasite_bbhb REAL,
            tahmini_bedel REAL,vasif TEXT,durum TEXT DEFAULT 'Aktif',yil TEXT);
        CREATE TABLE IF NOT EXISTS Iletisim_Formu(id INTEGER PRIMARY KEY AUTOINCREMENT,
            gonderen TEXT,konu_tipi TEXT,mesaj TEXT,tarih TEXT,durum TEXT DEFAULT 'Yeni');
        CREATE TABLE IF NOT EXISTS Silme_Talepleri(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tablo TEXT,kayit_id INTEGER,kayit_ozet TEXT,talep_eden TEXT,tarih TEXT,
            durum TEXT DEFAULT 'Bekliyor');
        CREATE TABLE IF NOT EXISTS Dahili_Mesajlar(id INTEGER PRIMARY KEY AUTOINCREMENT,
            gonderen TEXT,alici TEXT,konu TEXT,mesaj TEXT,tarih TEXT,
            okundu INTEGER DEFAULT 0,tur TEXT DEFAULT 'Mesaj');
        CREATE TABLE IF NOT EXISTS Islah_Projeler(id INTEGER PRIMARY KEY AUTOINCREMENT,
            yil TEXT,ilce TEXT,koy TEXT,gubre_da REAL DEFAULT 0,mera_tohum_da REAL DEFAULT 0,
            yem_tohum_da REAL DEFAULT 0,golgelik INTEGER DEFAULT 0,sivat INTEGER DEFAULT 0,
            boru_m REAL DEFAULT 0,coban_evi INTEGER DEFAULT 0,tuzluk INTEGER DEFAULT 0,
            kasinka INTEGER DEFAULT 0,finansman TEXT,notlar TEXT,durum TEXT DEFAULT 'Tamamlandı');
        CREATE TABLE IF NOT EXISTS Iletisim_Bilgileri(anahtar TEXT PRIMARY KEY,deger TEXT);
        CREATE TABLE IF NOT EXISTS Genel_Emri(yil INTEGER PRIMARY KEY,sayi TEXT NOT NULL,bb_ceza REAL NOT NULL,kb_ceza REAL NOT NULL,madde TEXT,guncelleme_tarihi TEXT);
        CREATE TABLE IF NOT EXISTS Kayit_Fotolari(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tablo TEXT NOT NULL,kayit_id INTEGER NOT NULL,dosya_yolu TEXT NOT NULL,
            aciklama TEXT,ekleme_tarihi TEXT,ekleyen TEXT);
        CREATE TABLE IF NOT EXISTS Loglar_Arsiv(id INTEGER PRIMARY KEY AUTOINCREMENT,
            tarih TEXT,kul TEXT,islem TEXT,detay TEXT);
        """)
        # Migration — yeni sütunlar
        for tbl, col, tip in [
            ("Ihaleler","ada_parsel","TEXT"),("Ihaleler","kapasite","TEXT"),
            ("Ihaleler","ihale_yeri_id","INTEGER"),
            ("Ihale_Yerleri","tahmini_bedel","REAL DEFAULT 0"),
            ("Tahsisler","madde14_bent","TEXT"),("Tahsisler","ot_geliri","REAL"),
            ("Tahsisler","tescil_tarihi","TEXT"),
            ("Tahsisler","sure_bitis","TEXT"),
            ("Tahsisler","sure_tipi","TEXT"),
            ("Kullanicilar","riza_onay","INTEGER DEFAULT 0"),
            ("Kullanicilar","sifre_tarih","TEXT"),
            ("Kullanicilar","gorevler","TEXT DEFAULT '*'"),
            ("Ayarlar","vali_yardimcisi","TEXT DEFAULT 'Semih CEMBEKLİ'"),
        ]:
            try: c.execute(f'ALTER TABLE {tbl} ADD COLUMN {col} {tip}')
            except Exception: pass
        # DB indexes — performans
        for idx in [
            "CREATE INDEX IF NOT EXISTS idx_log_tarih ON Loglar(tarih)",
            "CREATE INDEX IF NOT EXISTS idx_ihale_koy ON Ihaleler(koy)",
            "CREATE INDEX IF NOT EXISTS idx_ceza_tc ON Idari_Cezalar(tc)",
            "CREATE INDEX IF NOT EXISTS idx_mera_koy ON Mera_Varligi(koy)",
        ]:
            try: c.execute(idx)
            except Exception: pass
        # WAL mode — çok kullanıcılı erişim
        try: c.execute("PRAGMA journal_mode=DELETE")  # Ağ güvenli mod
        except Exception: pass
        # Log arşivleme — 1 yıldan eski logları arşivle
        try:
            bir_yil_once=(datetime.now()-timedelta(days=365)).strftime("%Y-%m-%d")
            c.execute("INSERT INTO Loglar_Arsiv(tarih,kul,islem,detay) SELECT tarih,kul,islem,detay FROM Loglar WHERE tarih<?",(bir_yil_once,))
            c.execute("DELETE FROM Loglar WHERE tarih<?",(bir_yil_once,))
        except Exception: pass
        # İletişim bilgileri default
        for k,v in [("telefon",DEV_TEL),("email",DEV_MAIL),("whatsapp",DEV_WA)]:
            try: c.execute("INSERT OR IGNORE INTO Iletisim_Bilgileri(anahtar,deger)VALUES(?,?)",(k,v))
            except Exception: pass
        if c.execute("SELECT COUNT(*) FROM Kullanicilar").fetchone()[0]==0:
            c.execute("INSERT INTO Kullanicilar(k_adi,sifre,yetki,ad,unvan,riza_onay)VALUES(?,?,?,?,?,?)",
                ("admin",hash_pw("Admin123!"),"Admin","Sistem Yöneticisi","Ziraat Mühendisi",1))
            c.execute("INSERT INTO Kullanicilar(k_adi,sifre,yetki,ad,unvan,riza_onay)VALUES(?,?,?,?,?,?)",
                ("emre",hash_pw("Emre1234"),"Uzman","Emre ÖZTÜRK","Ziraat Mühendisi",1))
        # Genel Emri varsayılan verileri
        for yil_g,veri_g in GENEL_EMRI_GECMIS.items():
            try: c.execute("INSERT OR IGNORE INTO Genel_Emri(yil,sayi,bb_ceza,kb_ceza,madde,guncelleme_tarihi)VALUES(?,?,?,?,?,?)",
                (yil_g,veri_g["sayi"],veri_g["bb"],veri_g["kb"],veri_g["madde"],datetime.now().strftime("%d.%m.%Y")))
            except Exception: pass

# ─── ANA UYGULAMA ─────────────────────────────────────────────────────────────
class MirasApp:
    MENU_ITEMS_UZMAN = [
        ("🏠","Dashboard","dash"),("🧮","Kapasite","kapasite"),
        ("⚖️","İhale","ihale"),
        ("🌱","Islah/Amenajman","islah"),
        ("📋","Tahsis/Md.14","tahsis"),("🚨","Şikayet","sikayet"),
        ("💰","İdari Ceza","ceza"),("📊","Veri Kayıt","veri"),
        ("👥","Muhtarlar","muhtar"),("📅","Ajanda","ajanda"),
        ("👨‍💼","Personel","personel"),
        ("📈","İstatistik","istatistik"),("📝","Evrak Üretici","evrak"),
        ("🌿","MERA AI","ai"),("📖","Mevzuat","mevzuat"),
        ("📞","İletişim","iletisim"),("⚙️","Ayarlar","ayarlar"),
    ]
    MENU_ITEMS_ADMIN = [("🛡️","Admin","admin")]
    MENU_ITEMS_IZLEYICI_HIDE = ["ihale","islah","tahsis","ceza","evrak"]

    def __init__(self,root):
        self.root=root; self.root.title(f"{PROG_ADI} {VERSIYON}"); self.root.geometry("520x380")
        self.ai=None; self.session=None; self.yedekci=None; self.status=None
        self.u_id=self.u_yetki=self.u_ad=self.u_unvan=None
        self.tema="Orman Yeşili"; self.punto=10; self.sube_mudur="Leyla ARSLAN"
        self._active_menu=None; self._content_frame=None; self._menu_btns={}
        BACKUP_DIR.mkdir(exist_ok=True)
        self._check_db()

    def _get_ai(self):
        """AI'ı ilk kullanımda başlat — açılışı yavaşlatmasın"""
        if self.ai is None:
            self.ai=GeminiAsistan()
        return self.ai

    def gc(self,k): return TEMALAR.get(self.tema,TEMALAR["Orman Yeşili"]).get(k,"#1E5631")

    def _check_db(self):
        global DB_PATH
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE,encoding="utf-8") as f: data=json.load(f)
                yol=data.get("db_path","")
                if yol and os.path.exists(yol):
                    DB_PATH=yol
                    # Ağ DB bütünlük kontrolü
                    try:
                        conn=db_baglan()
                        if conn:
                            r=conn.execute("PRAGMA integrity_check").fetchone()
                            conn.close()
                            if r and r[0]!="ok":
                                messagebox.showwarning("⚠️ DB Uyarı","Veritabanında bütünlük sorunu tespit edildi.\nOtomatik yedekten geri yüklemeyi deneyin.")
                    except Exception as e:
                        logging.error(f"integrity:{e}")
                    # Açılışta otomatik yedek al (ağ güvenliği)
                    try:
                        ts=datetime.now().strftime("%Y%m%d_%H%M")
                        BACKUP_DIR.mkdir(exist_ok=True)
                        import shutil
                        shutil.copy2(DB_PATH,BACKUP_DIR/f"Miras_acilis_{ts}.db")
                    except Exception: pass
                    init_db(); self._build_login(); return
                elif yol:
                    messagebox.showwarning("Veritabanı Bulunamadı",f"Kayıtlı DB bulunamadı:\n{yol}\nLütfen yeniden seçin.")
            except Exception as e: logging.error(f"config:{e}")
        self._build_db_setup()

    def _build_db_setup(self):
        self._clear(); self.root.geometry("560x420")
        f=tk.Frame(self.root,bg=C_WHITE,padx=40,pady=40); f.pack(fill="both",expand=True,padx=20,pady=20)
        tk.Label(f,text="🌿 MİRAS Ağ Kurulumu",font=("Segoe UI",17,"bold"),fg="#1E5631",bg=C_WHITE).pack(pady=(0,8))
        tk.Label(f,text="Ortak bir ağ klasörünü seçerek tüm personelin\naynı veritabanını kullanmasını sağlayın.",
            font=("Segoe UI",10),fg="#555",bg=C_WHITE,justify="center").pack(pady=(0,20))
        MBtn(f,"📂  Mevcut Veritabanını Seç",command=lambda:self._sec_db(False),color=C_INFO,width=36).pack(pady=8)
        MBtn(f,"➕  Yeni Veritabanı Oluştur",command=lambda:self._sec_db(True),color="#2D8C55",width=36).pack(pady=8)
        tk.Label(f,text="İpucu: \\\\sunucu\\tarim\\miras.db gibi ağ yolu kullanın",font=("Segoe UI",8),fg="#999",bg=C_WHITE).pack(pady=(16,0))

    def _sec_db(self,yeni):
        global DB_PATH
        if yeni: yol=filedialog.asksaveasfilename(defaultextension=".db",filetypes=[("SQLite","*.db")])
        else: yol=filedialog.askopenfilename(filetypes=[("SQLite","*.db")])
        if yol:
            DB_PATH=yol
            with open(CONFIG_FILE,"w",encoding="utf-8") as f: json.dump({"db_path":DB_PATH},f)
            init_db(); self._build_login()

    def _build_login(self):
        self._clear(); self.root.geometry("460x720"); self.root.configure(bg="#F0F4F2")
        logo=tk.Frame(self.root,bg="#1E5631",height=180); logo.pack(fill="x"); logo.pack_propagate(False)
        # Logo/amblem alanı
        LOGO_PATH = "miras_logo.png"
        if os.path.exists(LOGO_PATH):
            try:
                self._login_logo = tk.PhotoImage(file=LOGO_PATH)
                # Boyut küçültme (subsample)
                try: self._login_logo = self._login_logo.subsample(max(1, self._login_logo.width()//80))
                except Exception: pass
                tk.Label(logo,image=self._login_logo,bg="#1E5631").pack(pady=(16,4))
            except Exception:
                tk.Label(logo,text="🌿",font=("Segoe UI",42),fg=C_WHITE,bg="#1E5631").pack(pady=(16,0))
        else:
            tk.Label(logo,text="🌿",font=("Segoe UI",42),fg=C_WHITE,bg="#1E5631").pack(pady=(16,0))
        tk.Label(logo,text="MİRAS",font=("Segoe UI",28,"bold"),fg=C_WHITE,bg="#1E5631").pack(pady=(0,0))
        tk.Label(logo,text="Mera İhtiyaç ve Rasyonel Amenajman Sistemi",font=("Segoe UI",8),fg="#A9DFBF",bg="#1E5631").pack()
        tk.Label(logo,text=f"Enterprise {VERSIYON}",font=("Segoe UI",9),fg="#8FCF9F",bg="#1E5631").pack()
        form=tk.Frame(self.root,bg=C_WHITE); form.pack(fill="both",expand=True,padx=50)
        for lbl,attr,show in [("Kullanıcı Adı","e_u",""),("Şifre","e_p","●")]:
            tk.Label(form,text=lbl,bg=C_WHITE,font=("Segoe UI",10,"bold"),fg="#444").pack(anchor="w",pady=(20 if "Ku" in lbl else 14,2))
            e=ttk.Entry(form,font=("Segoe UI",11),show=show); e.pack(fill="x"); setattr(self,attr,e)
        # Şifre göster/gizle
        self._pw_vis=tk.BooleanVar(value=False)
        def _toggle_pw():
            self.e_p.config(show="" if self._pw_vis.get() else "●")
        ttk.Checkbutton(form,text="Şifreyi Göster",variable=self._pw_vis,command=_toggle_pw).pack(anchor="w",pady=(4,0))
        MBtn(form,"SİSTEME GİRİŞ YAP →",command=self._login,width=28).pack(pady=(20,8))
        self.lbl_hata=tk.Label(form,text="",fg=C_DANGER,bg=C_WHITE,font=("Segoe UI",9)); self.lbl_hata.pack()
        lf=tk.Frame(form,bg=C_WHITE); lf.pack(fill="x",pady=10)
        tk.Button(lf,text="🔑 Şifremi Unuttum",command=self._unuttum,fg="#1E5631",bg=C_WHITE,
            relief="flat",cursor="hand2",font=("Segoe UI",9)).pack(side="left")
        tk.Button(lf,text="📝 Kayıt Talebi",command=self._kayit_talep,fg=C_INFO,bg=C_WHITE,
            relief="flat",cursor="hand2",font=("Segoe UI",9)).pack(side="right")
        alt=tk.Frame(self.root,bg="#F0F4F2"); alt.pack(side="bottom",fill="x",pady=8)
        tk.Label(alt,text=f"{HAKLAR} | {PROG_ADI} {VERSIYON}",fg="#333333",bg="#F0F4F2",font=("Segoe UI",9,"bold")).pack()
        self.root.bind("<Return>",lambda e:self._login()); self.e_u.focus()

    def _unuttum(self):
        if not DB_PATH: messagebox.showerror("Hata","Veritabanı seçilmemiş."); return
        u=tkinter.simpledialog.askstring("Şifre Sıfırlama","Kullanıcı adınızı girin:")
        if u:
            with db_baglan() as c:
                ex=c.execute("SELECT COUNT(*) FROM Kullanicilar WHERE k_adi=?",(u,)).fetchone()[0]
            if ex:
                with db_baglan() as c:
                    c.execute("INSERT INTO Sifre_Talepleri(k_adi,tarih)VALUES(?,?)",(u,datetime.now().strftime("%Y-%m-%d")))
                messagebox.showinfo("Tamam","Şifre sıfırlama talebiniz Admin'e iletildi.")
            else: messagebox.showerror("Hata","Bu kullanıcı adı sistemde kayıtlı değil.")

    def _kayit_talep(self):
        if not DB_PATH: return
        win=tk.Toplevel(self.root); win.title("Kayıt Talebi"); win.geometry("340x320")
        win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        tk.Label(win,text="Yeni Personel Kayıt Talebi",font=("Segoe UI",13,"bold"),bg=C_WHITE,fg="#1E5631").pack(pady=16)
        al={}
        for l in ["Ad Soyad:","Ünvan:","Kullanıcı Adı:"]:
            tk.Label(win,text=l,bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=24,pady=(8,0))
            e=ttk.Entry(win,width=32); e.pack(padx=24); al[l]=e
        def _g():
            k=al["Kullanıcı Adı:"].get().strip(); a=al["Ad Soyad:"].get().strip()
            if not k or not a: messagebox.showwarning("Uyarı","Tüm alanları doldurun.",parent=win); return
            with db_baglan() as c:
                c.execute("INSERT INTO Kayit_Talepleri(ad,unvan,k_adi,tarih)VALUES(?,?,?,?)",(a,al["Ünvan:"].get(),k,datetime.now().strftime("%Y-%m-%d")))
            messagebox.showinfo("Tamam","Talebiniz Admin'e iletildi.",parent=win); win.destroy()
        MBtn(win,"Gönder",command=_g,width=20).pack(pady=16)

    def _login(self):
        if not DB_PATH: messagebox.showerror("Hata","Veritabanı seçilmemiş."); return
        u=self.e_u.get().strip(); p=self.e_p.get()
        if not u or not p: self.lbl_hata.config(text="❌ Kullanıcı adı ve şifre gerekli."); return
        try:
            with db_baglan() as conn:
                row=conn.execute("SELECT k_adi,sifre,yetki,ad,unvan,aktif,fail_count,lockout_ts,riza_onay FROM Kullanicilar WHERE k_adi=?",(u,)).fetchone()
                ayr=conn.execute("SELECT tema,punto,sube_mudur FROM Ayarlar WHERE k_adi=?",(u,)).fetchone()
        except Exception as e:
            self.lbl_hata.config(text="❌ Veritabanı bağlantı hatası."); logging.error(f"login:{e}"); return
        if not row: self.lbl_hata.config(text="❌ Kullanıcı bulunamadı."); db_log(u,"Başarısız","Bilinmeyen"); return
        k_adi,hashed,yetki,ad,unvan,aktif,fail_c,lkts=row[:8]
        riza_onay=row[8] if len(row)>8 else 0
        if not aktif: self.lbl_hata.config(text="❌ Hesabınız devre dışı."); return
        if lkts:
            bitis=datetime.fromisoformat(lkts)
            if datetime.now()<bitis:
                kalan=int((bitis-datetime.now()).total_seconds()/60)
                self.lbl_hata.config(text=f"🔒 Hesap kilitli. {kalan+1} dk bekleyin."); return
            else:
                with db_baglan() as c: c.execute("UPDATE Kullanicilar SET fail_count=0,lockout_ts=NULL WHERE k_adi=?",(k_adi,))
                fail_c=0
        if verify_pw(p,hashed):
            with db_baglan() as c: c.execute("UPDATE Kullanicilar SET fail_count=0,lockout_ts=NULL WHERE k_adi=?",(k_adi,))
            db_log(k_adi,"Giriş","Başarılı")
            self.u_id=k_adi; self.u_yetki=yetki; self.u_ad=ad; self.u_unvan=unvan
            if ayr: self.tema,self.punto,self.sube_mudur=ayr[0] or "Orman Yeşili",ayr[1] or 10,ayr[2] or "Leyla ARSLAN"
            self.root.unbind("<Return>"); self.session=SessionManager()
            self.session.register(self._oturum_bitti); self.yedekci=YedekYoneticisi(self.u_id)
            # Rıza kontrolü
            if not riza_onay:
                self._riza_goster()
            else:
                self._sifre_yas_kontrol()
        else:
            self._login_fail(k_adi,fail_c)

    def _sifre_yas_kontrol(self):
        """Madde 13: Şifre ömrü 180 gün kontrolü"""
        try:
            with db_baglan() as c:
                row=c.execute("SELECT sifre_tarih FROM Kullanicilar WHERE k_adi=?",(self.u_id,)).fetchone()
            st=row[0] if row and row[0] else None
            if not st:
                with db_baglan() as c:
                    c.execute("UPDATE Kullanicilar SET sifre_tarih=? WHERE k_adi=?",(datetime.now().strftime("%Y-%m-%d"),self.u_id))
                self._build_app(); return
            sifre_t=datetime.strptime(st,"%Y-%m-%d")
            gecen=(datetime.now()-sifre_t).days
            kalan=SIFRE_OMUR_GUN-gecen
            if kalan<=0:
                messagebox.showwarning("🔒 Şifre Süresi Doldu",
                    f"Şifreniz {SIFRE_OMUR_GUN} günlük ömrünü tamamladı.\n"
                    "Devam etmek için şifrenizi değiştirmeniz gerekiyor.")
                self._zorla_sifre_degistir()
            elif kalan<=SIFRE_UYARI_GUN:
                messagebox.showwarning("⚠️ Şifre Uyarısı",
                    f"Şifrenizin süresi dolmak üzere!\n"
                    f"Kalan gün: {kalan}\n\n"
                    "Lütfen Ayarlar → Şifre sekmesinden şifrenizi değiştirin.")
                self._build_app()
            else:
                self._build_app()
        except Exception as e:
            logging.error(f"sifre_yas:{e}"); self._build_app()
    def _zorla_sifre_degistir(self):
        """Şifre süresi dolmuşsa zorunlu değiştirme ekranı"""
        win=tk.Toplevel(self.root); win.title("Şifre Değiştir"); win.geometry("420x320")
        win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy); win.transient(self.root)
        win.protocol("WM_DELETE_WINDOW",lambda:None)
        tk.Label(win,text="🔒 Zorunlu Şifre Değişikliği",font=("Segoe UI",14,"bold"),fg=C_DANGER,bg=C_WHITE).pack(pady=16)
        tk.Label(win,text="Şifrenizin süresi doldu. Yeni şifre belirleyin.",font=("Segoe UI",10),fg="#555",bg=C_WHITE).pack(pady=(0,12))
        al={}
        for l,k in [("Mevcut Şifre:","eski"),("Yeni Şifre:","yeni"),("Tekrar:","tekrar")]:
            tk.Label(win,text=l,bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=30,pady=(6,0))
            e=ttk.Entry(win,show="●",width=30); e.pack(padx=30); al[k]=e
        def _k():
            eski=al["eski"].get(); yeni=al["yeni"].get(); tekrar=al["tekrar"].get()
            if yeni!=tekrar: messagebox.showerror("Hata","Şifreler eşleşmiyor.",parent=win); return
            ok,msg=strong_pw(yeni)
            if not ok: messagebox.showerror("Hata",msg,parent=win); return
            with db_baglan() as c:
                db_s=c.execute("SELECT sifre FROM Kullanicilar WHERE k_adi=?",(self.u_id,)).fetchone()[0]
            if not verify_pw(eski,db_s): messagebox.showerror("Hata","Mevcut şifre yanlış.",parent=win); return
            with db_baglan() as c:
                c.execute("UPDATE Kullanicilar SET sifre=?,sifre_tarih=? WHERE k_adi=?",(hash_pw(yeni),datetime.now().strftime("%Y-%m-%d"),self.u_id))
            db_log(self.u_id,"Şifre Değiştir","Zorunlu - süre doldu"); win.destroy(); self._build_app()
        MBtn(win,"🔒 Şifreyi Değiştir ve Giriş Yap",command=_k,color=C_DANGER,width=30).pack(pady=16)

    def _login_fail(self,k_adi,fail_c):
        """Login başarısız — ayrı metod olarak"""
        fail_c+=1; lockout=None
        kalan_hak=MAX_LOGIN_FAIL-fail_c
        if fail_c>=MAX_LOGIN_FAIL:
            lockout=(datetime.now()+timedelta(minutes=LOCKOUT_MIN)).isoformat()
            self.lbl_hata.config(text=f"🔒 {MAX_LOGIN_FAIL} hatalı giriş. {LOCKOUT_MIN} dk kilitlendi.")
        else:
            self.lbl_hata.config(text=f"❌ Hatalı şifre. ({fail_c}/{MAX_LOGIN_FAIL}) — {kalan_hak} hakkınız kaldı.")
        with db_baglan() as c: c.execute("UPDATE Kullanicilar SET fail_count=?,lockout_ts=? WHERE k_adi=?",(fail_c,lockout,k_adi))
        db_log(k_adi,"Başarısız",f"{fail_c}/{MAX_LOGIN_FAIL}"); self.e_p.delete(0,tk.END)

    def _riza_goster(self):
        """Açık rıza metni göster ve onay al"""
        win=tk.Toplevel(self.root); win.title("Kullanım Koşulları"); win.geometry("600x520")
        win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy); win.transient(self.root)
        tk.Label(win,text="📋 Kullanım Koşulları ve Açık Rıza",font=("Segoe UI",14,"bold"),fg="#1E5631",bg=C_WHITE).pack(pady=14)
        txt=scrolledtext.ScrolledText(win,font=("Segoe UI",10),wrap="word",bg="#FAFAFA",relief="flat",
            highlightbackground="#D0DDD8",highlightthickness=1,padx=14,pady=12)
        txt.pack(fill="both",expand=True,padx=16,pady=8)
        txt.insert("1.0",RIZA_METNI); txt.config(state="disabled")
        var=tk.BooleanVar()
        ttk.Checkbutton(win,text="Yukarıdaki koşulları okudum ve kabul ediyorum.",variable=var).pack(pady=8)
        def _onayla():
            if not var.get():
                messagebox.showwarning("Uyarı","Koşulları kabul etmeniz gerekiyor.",parent=win); return
            with db_baglan() as c:
                c.execute("UPDATE Kullanicilar SET riza_onay=1 WHERE k_adi=?",(self.u_id,))
            db_log(self.u_id,"Rıza Onay","Kabul edildi"); win.destroy(); self._build_app()
        def _reddet():
            self.u_id=None; win.destroy(); self._build_login()
        bf=tk.Frame(win,bg=C_WHITE); bf.pack(pady=10)
        MBtn(bf,"✅ Kabul Ediyorum",command=_onayla,color="#2D8C55",width=20).pack(side="left",padx=8)
        MBtn(bf,"❌ Reddet",command=_reddet,color=C_DANGER,width=14).pack(side="left",padx=8)

    def _oturum_bitti(self):
        self.root.after(0,lambda:(messagebox.showwarning("Oturum Sona Erdi",f"{SESSION_TIMEOUT_MIN} dakika hareketsizlik nedeniyle oturum kapatıldı."),self._cikis()))

    def _cikis(self):
        # Çıkışta otomatik yedek al
        if self.yedekci:
            try: self.yedekci.al(otomatik=True)
            except Exception: pass
        if self.session: self.session.stop(); self.session=None
        self.u_id=self.u_yetki=self.u_ad=self.u_unvan=None
        self._build_login()

    def _clear(self):
        for w in self.root.winfo_children(): w.destroy()

    # ═══ SOL MENÜ (SIDEBAR) YAPISI ═══════════════════════════════════════════
    def _build_app(self):
        self._clear(); self.root.state("zoomed"); self.root.configure(bg=self.gc("bg"))
        self.root.bind_all("<Motion>",lambda e:self.session and self.session.ping())
        self.root.bind_all("<KeyPress>",lambda e:self.session and self.session.ping())

        # Ana container
        main=tk.Frame(self.root,bg=self.gc("bg")); main.pack(fill="both",expand=True)

        # ── Header ──
        h=tk.Frame(main,bg=self.gc("pri"),height=52); h.pack(fill="x",side="top"); h.pack_propagate(False)
        tk.Label(h,text="🌿 MİRAS",font=("Segoe UI",18,"bold"),fg=C_WHITE,bg=self.gc("pri")).pack(side="left",padx=18,pady=8)
        tk.Label(h,text=VERSIYON,font=("Segoe UI",8),fg="#8FCF9F",bg=self.gc("pri")).pack(side="left")
        right=tk.Frame(h,bg=self.gc("pri")); right.pack(side="right",padx=14)
        # Mini profil fotoğrafı
        foto_dir=Path("miras_profil"); foto_yol=foto_dir/f"{self.u_id}.png"
        if foto_yol.exists():
            try:
                from PIL import Image, ImageTk, ImageDraw
                img=Image.open(str(foto_yol)).resize((36,36))
                mask=Image.new("L",(36,36),0); ImageDraw.Draw(mask).ellipse((0,0,36,36),fill=255)
                img.putalpha(mask)
                self._hdr_foto=ImageTk.PhotoImage(img)
                tk.Label(right,image=self._hdr_foto,bg=self.gc("pri"),cursor="hand2").pack(side="left",padx=(0,6))
            except Exception: pass
        tk.Label(right,text=f"👤 {self.u_ad}  |  {self.u_unvan}  |  {self.u_yetki}",
            font=("Segoe UI",9),fg=C_WHITE,bg=self.gc("pri")).pack(side="left",padx=10)
        MBtn(right,"⏏ Çıkış",command=self._cikis,color=C_DANGER,pady=3,padx=8).pack(side="left")

        # ── Body: Sidebar + Content ──
        body=tk.Frame(main,bg=self.gc("bg")); body.pack(fill="both",expand=True)

        # Sidebar
        self._sidebar=tk.Frame(body,bg=self.gc("side"),width=200); self._sidebar.pack(side="left",fill="y")
        self._sidebar.pack_propagate(False)

        # Sidebar menü butonları
        self._menu_btns={}
        items=list(self.MENU_ITEMS_UZMAN)
        if self.u_yetki=="Admin": items+=self.MENU_ITEMS_ADMIN

        # Scrollable sidebar with visible scrollbar
        s_canvas=tk.Canvas(self._sidebar,bg=self.gc("side"),highlightthickness=0,width=180)
        s_scroll=ttk.Scrollbar(self._sidebar,orient="vertical",command=s_canvas.yview)
        s_canvas.configure(yscrollcommand=s_scroll.set)
        s_scroll.pack(side="right",fill="y")
        s_canvas.pack(side="left",fill="both",expand=True)
        s_inner=tk.Frame(s_canvas,bg=self.gc("side"))
        s_canvas.create_window((0,0),window=s_inner,anchor="nw",width=180)
        s_inner.bind("<Configure>",lambda e:s_canvas.configure(scrollregion=s_canvas.bbox("all")))
        # Mouse wheel scroll - platform uyumlu
        def _sidebar_scroll(event):
            if event.delta: s_canvas.yview_scroll(int(-1*(event.delta/120)),"units")
            elif event.num==4: s_canvas.yview_scroll(-1,"units")
            elif event.num==5: s_canvas.yview_scroll(1,"units")
        s_canvas.bind("<MouseWheel>",_sidebar_scroll)
        s_canvas.bind("<Button-4>",_sidebar_scroll)
        s_canvas.bind("<Button-5>",_sidebar_scroll)
        s_inner.bind("<MouseWheel>",_sidebar_scroll)
        s_inner.bind("<Button-4>",_sidebar_scroll)
        s_inner.bind("<Button-5>",_sidebar_scroll)

        for ikon,baslik,key in items:
            if self.u_yetki=="İzleyici" and key in self.MENU_ITEMS_IZLEYICI_HIDE:
                continue
            btn=tk.Button(s_inner,text=f" {ikon} {baslik}",anchor="w",
                font=("Segoe UI",9),fg=self.gc("side_text"),bg=self.gc("side"),
                relief="flat",cursor="hand2",padx=10,pady=6,
                activebackground=self.gc("side_hover"),activeforeground=C_WHITE,
                command=lambda k=key:self._menu_click(k))
            btn.pack(fill="x",pady=0)
            btn.bind("<Enter>",lambda e,b=btn:b.config(bg=self.gc("side_hover"),fg=C_WHITE))
            btn.bind("<Leave>",lambda e,b=btn,k=key:b.config(
                bg=self.gc("pri") if self._active_menu==k else self.gc("side"),
                fg=C_WHITE if self._active_menu==k else self.gc("side_text")))
            btn.bind("<MouseWheel>",_sidebar_scroll)
            btn.bind("<Button-4>",_sidebar_scroll)
            btn.bind("<Button-5>",_sidebar_scroll)
            self._menu_btns[key]=btn

        # Content area
        content_wrapper=tk.Frame(body,bg=self.gc("bg")); content_wrapper.pack(side="left",fill="both",expand=True)
        self._content_frame=tk.Frame(content_wrapper,bg=self.gc("bg")); self._content_frame.pack(fill="both",expand=True)

        # Footer
        ft=tk.Frame(content_wrapper,bg="#E8E8E8",height=28); ft.pack(fill="x",side="bottom"); ft.pack_propagate(False)
        tk.Label(ft,text=HAKLAR,font=("Segoe UI",8),fg=C_FOOTER,bg="#E8E8E8").pack(side="left",padx=12)
        self.status=StatusBar(ft); self.status.pack(side="right",fill="y")

        # İlk sayfa: Dashboard
        self._menu_click("dash")
        self._duyuru()
        self.root.after(400,self._hosgeldin)
        self.root.after(1600,self._ajanda_uyar)
        self.root.after(2800,self._yedek_uyari)
        self._klavye_kisayollari()
        self.root.after(5000,self._guncelleme_kontrol)
        # Session timeout — hareketsizlik kontrolü
        self._son_aktivite=datetime.now()
        def _aktivite_kaydet(event=None):
            self._son_aktivite=datetime.now()
        # add="+" ile mevcut binding'leri bozmadan ekle
        for evt in ["<Key>","<Button>","<Motion>","<MouseWheel>","<FocusIn>"]:
            try: self.root.bind_all(evt,_aktivite_kaydet,add="+")
            except Exception: pass
        def _session_kontrol():
            try:
                if hasattr(self,"_son_aktivite") and self.u_id:
                    gecen=(datetime.now()-self._son_aktivite).total_seconds()/60
                    if gecen>=SESSION_TIMEOUT_MIN:
                        messagebox.showwarning("⏰ Oturum Süresi Doldu",
                            f"{SESSION_TIMEOUT_MIN} dakikadır işlem yapılmadı.\nGüvenlik için oturum kapatılıyor.")
                        self._cikis()
                        return
                self.root.after(60000,_session_kontrol)
            except Exception: pass
        self.root.after(60000,_session_kontrol)

    def _guncelleme_kontrol(self):
        """GitHub'dan güncelleme kontrolü — arka planda çalışır"""
        def _kontrol():
            try:
                url=f"https://api.github.com/repos/{GITHUB_REPO}/releases/latest"
                req=urllib.request.Request(url,headers={"User-Agent":"MirasEnterprise/1.0","Accept":"application/vnd.github.v3+json"})
                with urllib.request.urlopen(req,timeout=10) as resp:
                    if resp.status!=200: return
                    data=json.loads(resp.read().decode("utf-8"))
                son_v=data.get("tag_name","").strip()
                if not son_v: return
                # Sayısal karşılaştırma — v17.5 → [17,5]
                def _versiyon_ayikla(v):
                    v=v.replace("v","").strip()
                    parcalar=[]
                    for p in v.split("."):
                        try: parcalar.append(int("".join(ch for ch in p if ch.isdigit()) or "0"))
                        except Exception: parcalar.append(0)
                    return parcalar
                mevcut_list=_versiyon_ayikla(VERSIYON)
                yeni_list=_versiyon_ayikla(son_v)
                # Aynı uzunluğa getir
                max_len=max(len(mevcut_list),len(yeni_list))
                mevcut_list+=[0]*(max_len-len(mevcut_list))
                yeni_list+=[0]*(max_len-len(yeni_list))
                if yeni_list>mevcut_list:
                    dl_url=data.get("html_url","") or f"https://github.com/{GITHUB_REPO}/releases/latest"
                    self.root.after(0,lambda:self._guncelleme_bildir(son_v,dl_url))
            except Exception as e:
                logging.debug(f"Güncelleme kontrol: {e}")
        threading.Thread(target=_kontrol,daemon=True).start()

    def _guncelleme_bildir(self,yeni_v,url):
        if not url or not url.startswith("http"):
            url=f"https://github.com/{GITHUB_REPO}/releases/latest"
        if messagebox.askyesno("🔄 Güncelleme Mevcut",
            f"Yeni sürüm: {yeni_v}\nMevcut: {VERSIYON}\n\n"
            f"GitHub: {GITHUB_REPO}\nİndirme sayfasını açmak ister misiniz?"):
            # webbrowser.open başarısız olabilir — birden fazla yöntem dene
            acildi=False
            try:
                if webbrowser.open(url,new=2): acildi=True
            except Exception as e:
                logging.warning(f"webbrowser: {e}")
            if not acildi:
                try:
                    if sys.platform=="win32":
                        os.startfile(url); acildi=True
                    elif sys.platform=="darwin":
                        os.system(f'open "{url}"'); acildi=True
                    else:
                        os.system(f'xdg-open "{url}"'); acildi=True
                except Exception as e:
                    logging.warning(f"startfile: {e}")
            if not acildi:
                # Son çare — URL'yi panoya kopyala + göster
                try:
                    self.root.clipboard_clear()
                    self.root.clipboard_append(url)
                    self.root.update()
                    messagebox.showinfo("📋 Link Kopyalandı",
                        f"Tarayıcı otomatik açılamadı.\n\n"
                        f"İndirme linki panoya kopyalandı:\n{url}\n\n"
                        f"Tarayıcınızı açıp Ctrl+V ile yapıştırın.")
                except Exception:
                    messagebox.showinfo("İndirme Linki",
                        f"Tarayıcı açılamadı.\nLütfen manuel olarak gidin:\n\n{url}")

    def _menu_click(self,key):
        """Sidebar menü tıklama — ilgili sayfayı yükle"""
        # Aktif menü stilini güncelle
        for k,btn in self._menu_btns.items():
            if k==key:
                btn.config(bg=self.gc("pri"),fg=C_WHITE)
            else:
                btn.config(bg=self.gc("side"),fg=self.gc("side_text"))
        self._active_menu=key

        # İçerik alanını temizle
        for w in self._content_frame.winfo_children(): w.destroy()

        # Madde 12: Modül bazlı görev kontrolü
        # Kapasite raporu herkese açık, Admin her yere erişir
        MODUL_GOREV_MAP={"ihale":"İhale","ihale_yer":"İhale","sikayet":"Şikayet",
            "tahsis":"Tahsis","ceza":"Ceza","islah":"Islah"}
        if key in MODUL_GOREV_MAP and self.u_yetki!="Admin":
            try:
                with db_baglan() as c:
                    gorevler=c.execute("SELECT gorevler FROM Kullanicilar WHERE k_adi=?",(self.u_id,)).fetchone()
                gorev_str=gorevler[0] if gorevler and gorevler[0] else "*"
                if gorev_str!="*" and MODUL_GOREV_MAP[key] not in gorev_str:
                    tk.Label(self._content_frame,text=f"🔒 Bu modüle erişim yetkiniz bulunmuyor.\n\nGörevleriniz: {gorev_str}\n\nAdmin'den görev ataması isteyin.",
                        font=("Segoe UI",13),fg=C_DANGER,bg=self.gc("bg"),justify="center").pack(expand=True)
                    return
            except Exception: pass

        # İlgili sayfayı yükle
        builder=getattr(self,f"_t_{key}",None)
        if builder:
            try: builder(self._content_frame)
            except Exception as e:
                logging.error(f"Sayfa yükleme hatası ({key}): {e}")
                tk.Label(self._content_frame,text=f"❌ Sayfa yüklenirken hata oluştu:\n{e}",
                    font=("Segoe UI",12),fg=C_DANGER,bg=self.gc("bg")).pack(expand=True)

    def _hosgeldin(self):
        win=tk.Toplevel(self.root); win.overrideredirect(True); win.configure(bg=self.gc("pri"))
        w,h=440,110; rx=self.root.winfo_x()+(self.root.winfo_width()-w)//2; ry=self.root.winfo_y()+52
        win.geometry(f"{w}x{h}+{rx}+{ry}"); win.attributes("-alpha",0.93)
        tk.Label(win,text=f"Hoş Geldiniz, {self.u_ad}! 🌿",font=("Segoe UI",16,"bold"),fg=C_WHITE,bg=self.gc("pri")).pack(pady=(18,4))
        tk.Label(win,text=self.u_unvan,font=("Segoe UI",11),fg="#A9DFBF",bg=self.gc("pri")).pack()
        self.root.after(3200,win.destroy)

    def _duyuru(self):
        try:
            with db_baglan() as c:
                r=c.execute("SELECT mesaj,tarih,gonderen,id FROM Duyurular ORDER BY id DESC LIMIT 1").fetchone()
            if r:
                # 7 gün kontrolü
                try:
                    dt=datetime.strptime(r[1],"%d.%m.%Y")
                    if (datetime.now()-dt).days>7: return
                except Exception: pass
                messagebox.showinfo("📢 MİRAS Duyurusu",f"{r[0]}\n\n— {r[2]} ({r[1]})")
        except Exception: pass

    def _tv(self,parent,cols,h=14):
        frame=tk.Frame(parent,bg=self.gc("bg")); frame.pack(fill="both",expand=True)
        tv=ttk.Treeview(frame,columns=[c[0] for c in cols],show="headings",height=h)
        for col,w,bas in cols:
            tv.heading(col,text=bas,command=lambda _c=col,_t=tv:self._sort(_t,_c))
            tv.column(col,width=w,anchor="center")
        sb=ttk.Scrollbar(frame,orient="vertical",command=tv.yview)
        tv.configure(yscrollcommand=sb.set)
        tv.pack(side="left",fill="both",expand=True); sb.pack(side="left",fill="y")
        return tv

    @staticmethod
    def _sort(tree,col):
        """Sütun başlığına tıklayınca sırala — tekrar tıklayınca ters çevir"""
        data=[(tree.set(k,col),k) for k in tree.get_children("")]
        # Mevcut sıralama yönünü kontrol et
        reverse = getattr(tree, f"_sort_reverse_{col}", False)
        try: data.sort(key=lambda t:float(t[0].replace(",","").replace(".","",1).replace("₺","").strip() or 0), reverse=reverse)
        except (ValueError,AttributeError): data.sort(key=lambda t:t[0].lower(), reverse=reverse)
        for i,(_,k) in enumerate(data): tree.move(k,"",i)
        # Toggle yön
        setattr(tree, f"_sort_reverse_{col}", not reverse)
        # Başlıktaki ok işareti
        for c in tree["columns"]:
            text = tree.heading(c)["text"].replace(" ▲","").replace(" ▼","")
            tree.heading(c, text=text)
        arrow = " ▼" if reverse else " ▲"
        cur_text = tree.heading(col)["text"].replace(" ▲","").replace(" ▼","")
        tree.heading(col, text=cur_text + arrow)

    def _son_islem_kaydet(self,modul,kayit_adi):
        try:
            with db_baglan() as c:
                c.execute("INSERT INTO Son_Islemler(k_adi,modul,kayit_adi,tarih)VALUES(?,?,?,?)",
                    (self.u_id,modul,kayit_adi,datetime.now().strftime("%d.%m.%Y %H:%M")))
                c.execute("DELETE FROM Son_Islemler WHERE id NOT IN (SELECT id FROM Son_Islemler WHERE k_adi=? ORDER BY id DESC LIMIT 20)",(self.u_id,))
        except Exception: pass

    def _guvenli_sil(self, tablo, kayit_id, kayit_ozet, sil_sql, sil_params, yenile_fn):
        """Madde 2: Admin direkt siler, diğerleri talep oluşturur"""
        if self.u_yetki=="Admin":
            if messagebox.askyesno("Onay",f"'{kayit_ozet}' silinsin mi?"):
                with db_baglan() as c: c.execute(sil_sql, sil_params)
                db_log(self.u_id,"Silme",f"{tablo}:{kayit_ozet}"); yenile_fn()
        else:
            if messagebox.askyesno("Silme Talebi",f"'{kayit_ozet}' için Admin'e silme talebi gönderilsin mi?"):
                with db_baglan() as c:
                    c.execute("INSERT INTO Silme_Talepleri(tablo,kayit_id,kayit_ozet,talep_eden,tarih)VALUES(?,?,?,?,?)",
                        (tablo,kayit_id,kayit_ozet,self.u_ad,datetime.now().strftime("%d.%m.%Y %H:%M")))
                db_log(self.u_id,"Silme Talebi",f"{tablo}:{kayit_ozet}")
                messagebox.showinfo("Tamam","Silme talebiniz Admin'e iletildi.")

    def _yedek_uyari(self):
        try:
            yedekler=sorted(BACKUP_DIR.glob("*.db"),key=lambda p:p.stat().st_mtime)
            if not yedekler:
                messagebox.showwarning("💾 Yedek Uyarısı","Hiç yedek alınmamış!\nAyarlar → Yedekleme sekmesinden yedek alın."); return
            son=datetime.fromtimestamp(yedekler[-1].stat().st_mtime)
            fark=(datetime.now()-son).days
            if fark>=7: messagebox.showwarning("💾 Yedek Uyarısı",f"Son yedekten {fark} gün geçti!\nAyarlar → Yedekleme sekmesinden yedek alın.")
        except Exception as e: logging.error(f"yedek_uyari:{e}")

    def _ajanda_uyar(self):
        try:
            bugun=datetime.now().strftime("%d.%m.%Y")
            with db_baglan() as c:
                rows=c.execute("SELECT baslik,tarih,sure FROM Ajanda WHERE k_adi=? AND durum='Bekliyor' AND tarih<=? ORDER BY tarih,sure",(self.u_id,bugun)).fetchall()
            if rows:
                liste="\n".join(f"• {r[1]} {r[2]} — {r[0]}" for r in rows[:5])
                mesaj=f"{len(rows)} bekleyen etkinlik var:\n\n{liste}"
                if len(rows)>5: mesaj+=f"\n... ve {len(rows)-5} daha"
                messagebox.showinfo("📅 Ajanda Hatırlatıcı",mesaj)
        except Exception as e: logging.error(f"ajanda_uyar:{e}")
        # Tahsis süre uyarısı
        self._tahsis_sure_uyar()

    def _tahsis_sure_uyar(self):
        """Madde 11: Tahsis süre takibi — son 7 gün uyarı"""
        try:
            with db_baglan() as c:
                rows=c.execute("SELECT id,koy,sure_bitis,sure_tipi FROM Tahsisler WHERE durum='Devam Ediyor' AND sure_bitis IS NOT NULL AND sure_bitis!=''").fetchall()
            uyarilar=[]
            for r in rows:
                try:
                    bitis=datetime.strptime(r[2],"%d.%m.%Y")
                    kalan=(bitis-datetime.now()).days
                    if kalan<=7 and kalan>=0:
                        uyarilar.append(f"⚠️ {r[1]} — {r[3]}: {kalan} gün kaldı! ({r[2]})")
                    elif kalan<0:
                        uyarilar.append(f"🔴 {r[1]} — {r[3]}: Süre {abs(kalan)} gün önce doldu!")
                except Exception: pass
            if uyarilar:
                messagebox.showwarning("📋 Tahsis Süre Uyarısı","\n".join(uyarilar))
        except Exception as e: logging.error(f"tahsis_sure:{e}")

    def _klavye_kisayollari(self):
        self.root.bind_all("<F5>",lambda e:self._menu_click("dash"))
        self.root.bind_all("<Control-n>",lambda e:self._menu_click("kapasite"))
        self.root.bind_all("<Control-h>",lambda e:self._koy_gecmis())
        self.root.bind_all("<Control-f>",lambda e:self._hizli_kisi_ara())
        self.root.bind_all("<F1>",lambda e:messagebox.showinfo("⌨️ Klavye Kısayolları",
            "F5        → Dashboard Yenile\nCtrl+N    → Kapasite\nCtrl+H    → Köy Geçmişi\nCtrl+F    → Hızlı Kişi Arama\nF1        → Bu pencere"))

    def _hizli_kisi_ara(self):
        """Ctrl+F — TC veya isimle tüm modüllerde arama"""
        win=tk.Toplevel(self.root); win.title("🔍 Hızlı Kişi Arama"); win.geometry("720x520")
        win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        win.transient(self.root)
        tk.Label(win,text="🔍 Hızlı Kişi Arama",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=(14,4))
        tk.Label(win,text="TC numarası veya isim yazın — tüm modüllerdeki kayıtlar listelenir",
            font=("Segoe UI",9),fg="#666",bg=C_WHITE).pack(pady=(0,8))
        sf=tk.Frame(win,bg=C_WHITE); sf.pack(fill="x",padx=14)
        e=ttk.Entry(sf,width=30,font=("Segoe UI",12)); e.pack(side="left",padx=(0,8))
        tv=self._tv(win,[("modul",100,"Modül"),("ad",150,"Ad Soyad"),("tc",110,"TC"),
            ("detay",340,"Detay")],14)
        def _ara(event=None):
            aranan=e.get().strip().lower()
            if len(aranan)<2: return
            tv.delete(*tv.get_children())
            if not DB_PATH: return
            try:
                with db_baglan() as c:
                    # İhaleler
                    for r in c.execute("SELECT ad_soyad,tc,koy,ilce,durum FROM Ihaleler WHERE LOWER(ad_soyad) LIKE ? OR tc LIKE ?",
                        (f"%{aranan}%",f"%{aranan}%")).fetchall():
                        tv.insert("","end",values=("⚖️ İhale",r[0],r[1],f"{r[2]}/{r[3]} — {r[4]}"))
                    # Cezalar
                    for r in c.execute("SELECT ad_soyad,tc,mera_koy,ipc_tutari,tarih FROM Idari_Cezalar WHERE LOWER(ad_soyad) LIKE ? OR tc LIKE ?",
                        (f"%{aranan}%",f"%{aranan}%")).fetchall():
                        tv.insert("","end",values=("💰 Ceza",r[0],r[1],f"{r[2]} — {r[3]}₺ ({r[4]})"))
                    # Tahsisler
                    for r in c.execute("SELECT kurum,koy,ilce,amac,durum FROM Tahsisler WHERE LOWER(kurum) LIKE ? OR LOWER(koy) LIKE ?",
                        (f"%{aranan}%",f"%{aranan}%")).fetchall():
                        tv.insert("","end",values=("📋 Tahsis",r[0],"",f"{r[1]}/{r[2]} — {r[3]} ({r[4]})"))
                    # Şikayetler
                    for r in c.execute("SELECT sikayet_eden,koy,konu,durum FROM Sikayetler WHERE LOWER(sikayet_eden) LIKE ? OR LOWER(koy) LIKE ?",
                        (f"%{aranan}%",f"%{aranan}%")).fetchall():
                        tv.insert("","end",values=("🚨 Şikayet",r[0],"",f"{r[1]} — {r[2]} ({r[3]})"))
                    # Raporlar
                    for r in c.execute("SELECT talep_eden,tc,koy,rapor_no FROM Rapor_Gecmisi WHERE LOWER(talep_eden) LIKE ? OR tc LIKE ?",
                        (f"%{aranan}%",f"%{aranan}%")).fetchall():
                        tv.insert("","end",values=("📄 Rapor",r[0],r[1],f"{r[2]} — {r[3]}"))
            except Exception as ex: logging.error(f"hizli_ara:{ex}")
        e.bind("<KeyRelease>",lambda ev:win.after(300,_ara))
        MBtn(sf,"🔍 Ara",command=_ara,color=self.gc("acc"),width=10).pack(side="left")
        e.focus_set()

    # ═══ 1. DASHBOARD ════════════════════════════════════════════════════════
    def _t_dash(self,p):
        bg=self.gc("bg")
        # Scrollable dashboard
        dcv=tk.Canvas(p,bg=bg,highlightthickness=0); dsb=ttk.Scrollbar(p,orient="vertical",command=dcv.yview)
        dcv.configure(yscrollcommand=dsb.set); dsb.pack(side="right",fill="y"); dcv.pack(side="left",fill="both",expand=True)
        f=tk.Frame(dcv,bg=bg); dwid=dcv.create_window((0,0),window=f,anchor="nw")
        dcv.bind("<Configure>",lambda e:dcv.itemconfig(dwid,width=e.width))
        f.bind("<Configure>",lambda e:dcv.configure(scrollregion=dcv.bbox("all")))
        def _dscroll(e):
            try: dcv.yview_scroll(-1*(e.delta//120),"units")
            except Exception: pass
        dcv.bind("<MouseWheel>",_dscroll); f.bind("<MouseWheel>",_dscroll)
        tk.Label(f,text="📊 Genel Durum Paneli",font=("Segoe UI",16,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")

        # Bildirim çubuğu
        self._bildirim_frame=tk.Frame(f,bg="#FEF9E7",highlightbackground="#F0C040",highlightthickness=1)
        self._bildirim_frame.pack(fill="x",pady=(8,4))
        self.lbl_bildirim=tk.Label(self._bildirim_frame,text="📢 Bildirimler yükleniyor...",bg="#FEF9E7",
            font=("Segoe UI",9),fg="#7D6608",padx=10,pady=6)
        self.lbl_bildirim.pack(anchor="w")

        # Hızlı erişim butonları
        hf=tk.Frame(f,bg=bg); hf.pack(fill="x",pady=(4,4))
        tk.Label(hf,text="⚡ Hızlı Erişim:",bg=bg,font=("Segoe UI",9,"bold"),fg="#666").pack(side="left",padx=(0,8))
        for txt,key,renk in [("🧮 Yeni Rapor","kapasite",self.gc("acc")),("⚖️ Yeni İhale","ihale",C_INFO),
                              ("🚨 Yeni Şikayet","sikayet",C_WARN),("💰 Yeni Ceza","ceza",C_DANGER),
                              ("🔍 Kişi Ara","_ara",self.gc("pri"))]:
            cmd=lambda k=key:self._hizli_kisi_ara() if k=="_ara" else self._menu_click(k)
            if self.u_yetki!="İzleyici" or key=="_ara":
                tk.Button(hf,text=txt,command=cmd,bg=renk,fg=C_WHITE,
                    font=("Segoe UI",8,"bold"),relief="flat",cursor="hand2",padx=10,pady=4).pack(side="left",padx=4)

        # 📌 Bugün Ne Yapmalıyım?
        bugun_f=tk.LabelFrame(f,text="  📌 Bugün Ne Yapmalıyım?  ",bg="#FFF8E1",
            font=("Segoe UI",10,"bold"),fg="#E65100",padx=10,pady=6)
        bugun_f.pack(fill="x",pady=(0,6))
        self.lbl_bugun=tk.Label(bugun_f,text="Yükleniyor...",bg="#FFF8E1",font=("Segoe UI",9),fg="#333",justify="left",wraplength=900)
        self.lbl_bugun.pack(anchor="w")

        kf=tk.Frame(f,bg=bg); kf.pack(fill="x",pady=8)
        self.dk1=StatKart(kf,"Sistemdeki Köy","—",self.gc("pri"),"🏘️")
        self.dk2=StatKart(kf,"Düzenlenen Rapor","—",C_INFO,"📄")
        self.dk3=StatKart(kf,"Bekleyen Islah","—",C_DANGER,"🌱")
        self.dk4=StatKart(kf,"Aktif İhale","—",C_WARN,"⚖️")
        self.dk5=StatKart(kf,"Bekleyen Tahsis","—","#8E44AD","📋")
        for k in (self.dk1,self.dk2,self.dk3,self.dk4,self.dk5): k.pack(side="left",fill="both",expand=True,padx=4)
        lf=tk.LabelFrame(f,text="  Son İşlem Kayıtları  ",bg=C_WHITE,font=("Segoe UI",10,"bold"),fg=self.gc("pri"))
        lf.pack(fill="both",expand=True,pady=(0,8))
        self.tv_log=self._tv(lf,[("t",150,"Zaman"),("k",120,"Personel"),("i",140,"İşlem"),("d",500,"Detay")],10)
        sf=tk.LabelFrame(f,text="  ⏱ Son İşlemlerim  ",bg=C_WHITE,font=("Segoe UI",10,"bold"),fg=self.gc("pri"))
        sf.pack(fill="x",pady=(0,8))
        self.tv_son=self._tv(sf,[("modul",120,"Modül"),("kayit",280,"Kayıt"),("tarih",160,"Tarih")],4)
        MBtn(f,"🔄  Paneli Yenile",command=self._dash_yenile,color=self.gc("acc"),width=18).pack(pady=6)
        self._dash_yenile()

    def _dash_yenile(self):
        if not DB_PATH: return
        bildirimler=[]
        bugun=datetime.now().strftime("%d.%m.%Y")
        try:
            with db_baglan() as c:
                try: self.dk1.set(c.execute("SELECT COUNT(*) FROM Mera_Varligi").fetchone()[0])
                except Exception: self.dk1.set("—")
                try: self.dk2.set(c.execute("SELECT COUNT(*) FROM Rapor_Gecmisi").fetchone()[0])
                except Exception: self.dk2.set("—")
                try: self.dk3.set(c.execute("SELECT COUNT(*) FROM Islah_Amenajman WHERE durum='Bekliyor'").fetchone()[0])
                except Exception: self.dk3.set("—")
                try:
                    ihale_c=c.execute("SELECT COUNT(*) FROM Ihaleler WHERE durum NOT IN('Tamamlandı','İptal Edildi')").fetchone()[0]
                    self.dk4.set(ihale_c)
                    if ihale_c>0: bildirimler.append(f"⚖️ {ihale_c} aktif ihale")
                except Exception: self.dk4.set("—")
                try:
                    tahsis_c=c.execute("SELECT COUNT(*) FROM Tahsisler WHERE durum='Devam Ediyor'").fetchone()[0]
                    self.dk5.set(tahsis_c)
                    if tahsis_c>0: bildirimler.append(f"📋 {tahsis_c} devam eden tahsis")
                except Exception: self.dk5.set("—")
                try:
                    sik_c=c.execute("SELECT COUNT(*) FROM Sikayetler WHERE durum='Yeni'").fetchone()[0]
                    if sik_c>0: bildirimler.append(f"🚨 {sik_c} yeni şikayet")
                except Exception: pass
                try:
                    aj_c=c.execute("SELECT COUNT(*) FROM Ajanda WHERE k_adi=? AND durum='Bekliyor' AND tarih<=?",(self.u_id,bugun)).fetchone()[0]
                    if aj_c>0: bildirimler.append(f"📅 {aj_c} bekleyen etkinlik")
                except Exception: pass
                if self.u_yetki=="Admin":
                    try:
                        silt_c=c.execute("SELECT COUNT(*) FROM Silme_Talepleri WHERE durum='Bekliyor'").fetchone()[0]
                        if silt_c>0: bildirimler.append(f"🗑 {silt_c} silme talebi")
                    except Exception: pass
                    try:
                        ilet_c=c.execute("SELECT COUNT(*) FROM Iletisim_Formu WHERE durum='Yeni'").fetchone()[0]
                        if ilet_c>0: bildirimler.append(f"📝 {ilet_c} iletişim formu")
                    except Exception: pass
                    try:
                        kayit_c=c.execute("SELECT COUNT(*) FROM Kayit_Talepleri WHERE durum='Bekliyor'").fetchone()[0]
                        if kayit_c>0: bildirimler.append(f"📝 {kayit_c} kayıt talebi")
                    except Exception: pass
                    try:
                        sifre_c=c.execute("SELECT COUNT(*) FROM Sifre_Talepleri WHERE durum='Bekliyor'").fetchone()[0]
                        if sifre_c>0: bildirimler.append(f"🔑 {sifre_c} şifre sıfırlama talebi")
                    except Exception: pass
                try:
                    msj_c=c.execute("SELECT COUNT(*) FROM Dahili_Mesajlar WHERE alici=? AND okundu=0",(self.u_id,)).fetchone()[0]
                    if msj_c>0: bildirimler.append(f"💬 {msj_c} okunmamış mesaj")
                except Exception: pass
                try:
                    sure_rows=c.execute("SELECT COUNT(*) FROM Tahsisler WHERE durum='Devam Ediyor' AND sure_bitis IS NOT NULL AND sure_bitis!=''").fetchone()[0]
                    if sure_rows>0: bildirimler.append(f"⏰ {sure_rows} süreli tahsis")
                except Exception: pass
                # Bildirim göster
                if bildirimler:
                    self.lbl_bildirim.config(text="📢 " + "  |  ".join(bildirimler))
                else:
                    self.lbl_bildirim.config(text="✅ Bekleyen bildirim yok.")
                # Log tablosu
                try:
                    self.tv_log.delete(*self.tv_log.get_children())
                    if self.u_yetki!="Admin":
                        log_rows=c.execute("SELECT tarih,kul,islem,detay FROM Loglar WHERE kul=? ORDER BY id DESC LIMIT 30",(self.u_id,)).fetchall()
                    else:
                        log_rows=c.execute("SELECT tarih,kul,islem,detay FROM Loglar ORDER BY id DESC LIMIT 30").fetchall()
                    for r in log_rows: self.tv_log.insert("","end",values=r)
                except Exception: pass
                try:
                    self.tv_son.delete(*self.tv_son.get_children())
                    for r in c.execute("SELECT modul,kayit_adi,tarih FROM Son_Islemler WHERE k_adi=? ORDER BY id DESC LIMIT 20",(self.u_id,)).fetchall():
                        self.tv_son.insert("","end",values=r)
                except Exception: pass
        except Exception as e: logging.error(f"dash:{e}")
        if self.status: self.status.set(f"✔  Panel güncellendi — {datetime.now().strftime('%H:%M:%S')}")
        # Bugün Ne Yapmalıyım?
        if hasattr(self,"lbl_bugun"):
            bugun_items=[]
            try:
                with db_baglan() as c2:
                    try:
                        for a in c2.execute("SELECT baslik FROM Ajanda WHERE k_adi=? AND tarih=? AND durum='Bekliyor'",(self.u_id,bugun)).fetchall():
                            bugun_items.append(f"📅 {a[0]}")
                    except Exception: pass
                    try:
                        for s in c2.execute("SELECT koy,sure_bitis,sure_tipi FROM Tahsisler WHERE durum='Devam Ediyor' AND sure_bitis IS NOT NULL AND sure_bitis!=''").fetchall():
                            try:
                                bitis=datetime.strptime(s[1],"%d.%m.%Y"); kalan=(bitis-datetime.now()).days
                                if 0<=kalan<=7: bugun_items.append(f"⏰ {s[0]} — {s[2]}: {kalan} gün kaldı")
                                elif kalan<0: bugun_items.append(f"🔴 {s[0]} — {s[2]}: {abs(kalan)} gün gecikmiş!")
                            except Exception: pass
                    except Exception: pass
                    try:
                        for s in c2.execute("SELECT koy,konu FROM Sikayetler WHERE durum='Yeni'").fetchall():
                            bugun_items.append(f"🚨 Şikayet: {s[0]} — {s[1]}")
                    except Exception: pass
                    if self.u_yetki=="Admin":
                        try:
                            bst=c2.execute("SELECT COUNT(*) FROM Silme_Talepleri WHERE durum='Bekliyor'").fetchone()[0]
                            if bst>0: bugun_items.append(f"🗑 {bst} silme talebi bekliyor")
                        except Exception: pass
                        try:
                            bkt=c2.execute("SELECT COUNT(*) FROM Kayit_Talepleri WHERE durum='Bekliyor'").fetchone()[0]
                            if bkt>0: bugun_items.append(f"📝 {bkt} kayıt talebi bekliyor")
                        except Exception: pass
                        try:
                            bst2=c2.execute("SELECT COUNT(*) FROM Sifre_Talepleri WHERE durum='Bekliyor'").fetchone()[0]
                            if bst2>0: bugun_items.append(f"🔑 {bst2} şifre sıfırlama talebi")
                        except Exception: pass
            except Exception: pass
            if bugun_items:
                self.lbl_bugun.config(text="\n".join(bugun_items[:8]),fg="#333")
            else:
                self.lbl_bugun.config(text="✅ Bugün bekleyen görev yok — günün kutlu olsun!",fg="#2D8C55")

    # ═══ 2. KAPASİTE ════════════════════════════════════════════════════════
    def _t_kapasite(self,p):
        bg=self.gc("bg")
        cw=tk.Canvas(p,bg=bg,highlightthickness=0); vsb=ttk.Scrollbar(p,orient="vertical",command=cw.yview)
        cw.configure(yscrollcommand=vsb.set); vsb.pack(side="right",fill="y"); cw.pack(side="left",fill="both",expand=True)
        ana=tk.Frame(cw,bg=bg); wid=cw.create_window((0,0),window=ana,anchor="nw")
        cw.bind("<Configure>",lambda e:cw.itemconfig(wid,width=e.width))
        ana.bind("<Configure>",lambda e:cw.configure(scrollregion=cw.bbox("all")))
        def _kap_scroll(e):
            try: cw.yview_scroll(-1*(e.delta//120),"units")
            except Exception: pass
        cw.bind("<MouseWheel>",_kap_scroll); ana.bind("<MouseWheel>",_kap_scroll)
        px=20
        top=tk.LabelFrame(ana,text="  Mera Verisi & Rapor Hesaplama  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=14,pady=10)
        top.pack(fill="x",padx=px,pady=(14,6))
        r1=tk.Frame(top,bg=bg); r1.pack(fill="x",pady=4)
        if self.u_yetki!="İzleyici":
            MBtn(r1,"📥  Excel'den Mera Verisi Yükle",command=self._excel_mera,color=C_INFO,width=26).pack(side="left",padx=(0,14))
            MBtn(r1,"➕  Manuel Köy Ekle/Güncelle",command=self._koy_manuel,color="#8E44AD",width=26).pack(side="left")
        r2=tk.Frame(top,bg=bg); r2.pack(fill="x",pady=8)
        for lbl,attr,tip,w in [("Köy Seç:","kap_koy","combo",22),("Aktif İşletme:","kap_aktif","entry",8),
               ("Talep Eden:","kap_talep","entry",20),("TC No:","kap_tc","entry",14)]:
            tk.Label(r2,text=lbl,bg=bg,font=("Segoe UI",10)).pack(side="left",padx=(0,4))
            w_obj=ttk.Combobox(r2,width=w,state="readonly") if tip=="combo" else ttk.Entry(r2,width=w)
            w_obj.pack(side="left",padx=(0,14)); setattr(self,attr,w_obj)
        # Köy seçince otomatik bilgi
        self.kap_koy.bind("<<ComboboxSelected>>",self._kap_koy_sec)
        self.lbl_kap_oto=tk.Label(r2,text="",bg=bg,font=("Segoe UI",9),fg=self.gc("pri"))
        self.lbl_kap_oto.pack(side="left")
        r3=tk.Frame(top,bg=bg); r3.pack(fill="x",pady=4)
        tk.Label(r3,text="Ek Açıklama:",bg=bg,font=("Segoe UI",10)).pack(side="left",padx=(0,6))
        self.kap_acik=ttk.Entry(r3,width=50); self.kap_acik.pack(side="left")
        if self.u_yetki!="İzleyici":
            MBtn(top,"🧮  Hesapla ve Kurumsal PDF Üret",command=self._hesapla_pdf,color=self.gc("acc"),
                tooltip="Köy seçin, bilgileri doldurun → otomatik BBHB hesaplanır ve kurumsal PDF rapor oluşturulur").pack(anchor="w",pady=10)
        kv=tk.LabelFrame(ana,text="  Kayıtlı Mera Verileri  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"))
        kv.pack(fill="x",padx=px,pady=(0,8))
        self.tv_mera=self._tv(kv,[("koy",130,"Köy"),("ilce",90,"İlçe"),("alan",90,"Alan (da)"),("yem",110,"Yeşil Yem (kg/da)"),("ok",110,"OK (BBHB)")],6)
        rf=tk.LabelFrame(ana,text="  Düzenlenen Raporlar  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"))
        rf.pack(fill="both",expand=True,padx=px,pady=(0,14))
        AramaFrame(rf,self._filtre_rap,bg=bg).pack(anchor="w",padx=10,pady=6)
        self.tv_rap=self._tv(rf,[("no",80,"Rapor No"),("koy",110,"Köy"),("talep",150,"Talep Eden"),
            ("tc",110,"TC No"),("tarih",100,"Tarih"),("per",120,"Düzenleyen")],8)
        bf=tk.Frame(rf,bg=bg); bf.pack(pady=6)
        MBtn(bf,"🔄 Yenile",command=self._yenile_kap,color=C_INFO,width=14).pack(side="left",padx=6)
        if self.u_yetki!="İzleyici":
            MBtn(bf,"🗑 Seçili Raporu Sil",command=self._sil_rap,color=C_DANGER,width=18).pack(side="left",padx=6)
        self._yenile_kap()

    def _kap_koy_sec(self,event=None):
        """Köy seçilince alan/yem/ilçe/aktif işletme otomatik doldur"""
        koy=self.kap_koy.get()
        if not koy or not DB_PATH: return
        try:
            with db_baglan() as c:
                mera=c.execute("SELECT ilce,alan,yem FROM Mera_Varligi WHERE koy=?",(koy,)).fetchone()
                muhtar=c.execute("SELECT ad_soyad,telefon FROM Muhtarlar WHERE koy=? LIMIT 1",(koy,)).fetchone()
            if mera:
                ok=(mera[1]*mera[2])/PAYDA
                info=f"✅ {mera[0]} | {mera[1]:.0f} da | Yeşil Yem: {mera[2]:.0f} kg/da | {ok:.1f} BBHB"
                if muhtar: info+=f" | Muhtar: {muhtar[0]}"
                self.lbl_kap_oto.config(text=info)
            else:
                self.lbl_kap_oto.config(text="⚠️ Bu köyün mera verisi yok — önce Excel'den yükleyin",fg=C_DANGER)
        except Exception: pass

    def _yenile_kap(self):
        if not DB_PATH: return
        try:
            with db_baglan() as c:
                koyler=[r[0] for r in c.execute("SELECT koy FROM Mera_Varligi ORDER BY koy")]
                self.kap_koy["values"]=koyler
                self.tv_mera.delete(*self.tv_mera.get_children())
                for r in c.execute("SELECT koy,ilce,alan,yem FROM Mera_Varligi ORDER BY koy"):
                    ok=(r[2]*r[3])/PAYDA
                    self.tv_mera.insert("","end",values=(r[0],r[1],f"{r[2]:.0f}",f"{r[3]:.0f}",f"{ok:.2f}"))
                self._all_rap=c.execute("SELECT rapor_no,koy,talep_eden,tc,islem_tarihi,duzenleyen FROM Rapor_Gecmisi ORDER BY rowid DESC").fetchall()
            self._filtre_rap("")
        except Exception as e: logging.error(f"yenile_kap:{e}")

    def _filtre_rap(self,a):
        self.tv_rap.delete(*self.tv_rap.get_children()); a=a.lower()
        for r in getattr(self,"_all_rap",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            self.tv_rap.insert("","end",values=r)

    def _excel_mera(self):
        yol=filedialog.askopenfilename(filetypes=[("Excel","*.xlsx *.xls")])
        if not yol: return
        try:
            df=pd.read_excel(yol)
            if len(df.columns)<4: messagebox.showerror("Hata","4 sütun gerekli: koy|ilce|alan_da|yesil_yem_kg_da"); return
            with db_baglan() as c:
                kayit=0
                for _,r in df.iterrows():
                    v=r.tolist()
                    c.execute("INSERT OR REPLACE INTO Mera_Varligi(koy,ilce,alan,yem)VALUES(?,?,?,?)",(str(v[0]),str(v[1]),float(v[2]),float(v[3])))
                    kayit+=1
            db_log(self.u_id,"Excel Yükleme",f"{kayit} köy"); self._yenile_kap()
            messagebox.showinfo("Tamam",f"{kayit} köy verisi sisteme yüklendi.")
        except Exception as e: messagebox.showerror("Excel Hatası",str(e))

    def _koy_manuel(self):
        win=tk.Toplevel(self.root); win.title("Köy Mera Verisi"); win.geometry("420x320"); win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        tk.Label(win,text="Köy Mera Verisi Ekle/Güncelle",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=14)
        tk.Label(win,text="⚠️ Kapasite raporu için: Yararlanılabilir Yeşil Yem (kg/da) girilir.\n    Muhammen bedel için: Kuru Ot Verimi ayrı hesaplanır.",
            font=("Segoe UI",8),fg="#888",bg=C_WHITE).pack(anchor="w",padx=24)
        al={}
        for lbl in ["Köy Adı:","İlçe:","Mera Alanı (da):","Yararlanılabilir Yeşil Yem (kg/da):"]:
            tk.Label(win,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=24,pady=(6,0))
            e=ttk.Entry(win,width=32); e.pack(padx=24); al[lbl]=e
        def _k():
            try:
                koy=al["Köy Adı:"].get().strip(); ilce=al["İlçe:"].get().strip()
                alan=float(al["Mera Alanı (da):"].get()); yem=float(al["Yararlanılabilir Yeşil Yem (kg/da):"].get())
                if not koy or not ilce: raise ValueError("Köy ve İlçe zorunlu.")
                with db_baglan() as c: c.execute("INSERT OR REPLACE INTO Mera_Varligi(koy,ilce,alan,yem)VALUES(?,?,?,?)",(koy,ilce,alan,yem))
                ok=(alan*yem)/PAYDA
                db_log(self.u_id,"Mera Verisi",f"{koy} {alan}da {yem}kg/da OK:{ok:.1f}"); self._yenile_kap()
                messagebox.showinfo("Tamam",f"'{koy}' kaydedildi.\nAlan: {alan:.0f} da | Yeşil Yem: {yem:.0f} kg/da\nOtlatma Kapasitesi: {ok:.1f} BBHB",parent=win); win.destroy()
            except ValueError as e: messagebox.showerror("Hata",str(e),parent=win)
        MBtn(win,"Kaydet",command=_k,width=20).pack(pady=16)

    def _hesapla_pdf(self):
        koy=self.kap_koy.get(); talep=self.kap_talep.get().strip(); tc=self.kap_tc.get().strip()
        try: aktif=int(self.kap_aktif.get())
        except ValueError: messagebox.showerror("Hata","Aktif işletme sayısı sayı olmalı."); return
        # Veri doğrulama
        if not veri_dogrula([("Köy",koy,"str"),("Talep Eden",talep,"str"),("TC No",tc,"str"),("Aktif İşletme",aktif,"int+")]): return
        if aktif<=0: messagebox.showerror("Hata","Aktif işletme > 0 olmalı."); return
        if not tc_kontrol_ve_devam(tc): return
        try:
            with db_baglan() as c:
                mera=c.execute("SELECT alan,yem,ilce FROM Mera_Varligi WHERE koy=?",(koy,)).fetchone()
        except Exception as e: messagebox.showerror("DB Hatası",str(e)); return
        if not mera: messagebox.showwarning("Eksik",f"'{koy}' mera verisi yok."); return
        alan,yem,ilce=mera; ok=(alan*yem)/PAYDA
        try:
            with db_baglan() as c:
                yil=datetime.now().year
                son=c.execute("SELECT rapor_no FROM Rapor_Gecmisi WHERE rapor_no LIKE ? ORDER BY rowid DESC LIMIT 1",(f"{yil}-%",)).fetchone()
                sira=(int(son[0].split("-")[1])+1) if son else 1
                rap_no=f"{yil}-{sira:03d}"
                c.execute("INSERT INTO Rapor_Gecmisi(rapor_no,koy,talep_eden,tc,islem_tarihi,duzenleyen,aciklama)VALUES(?,?,?,?,?,?,?)",
                    (rap_no,koy,talep,tc,datetime.now().strftime("%d.%m.%Y"),self.u_ad,self.kap_acik.get()))
        except Exception as e: messagebox.showerror("DB Hatası",str(e)); return
        db_log(self.u_id,"Kapasite Raporu",f"{koy} {rap_no} OK:{ok:.2f}BBHB")
        self._son_islem_kaydet("Rapor",f"{koy} — {rap_no}")
        if not PDF_OK:
            messagebox.showwarning("PDF Yok",f"reportlab kurulu değil.\n\n{koy}: {ok:.2f} BBHB")
            self._yenile_kap(); return
        dosya=filedialog.asksaveasfilename(defaultextension=".pdf",initialfile=f"Rapor_{koy}_{rap_no}.pdf",filetypes=[("PDF","*.pdf")])
        if not dosya: self._yenile_kap(); return
        try:
            uret_pdf(dosya,{"ilce":ilce,"koy":koy,"alan":alan,"yem":yem,"aktif":aktif,"ok_bbhb":ok,"rapor_no":rap_no,"talep_eden":talep,"tc":tc,"aciklama":self.kap_acik.get()})
            self._yenile_kap()
            messagebox.showinfo("Rapor Hazır",f"✅ Rapor No: {rap_no}\n{koy}\nOtlatma Kapasitesi: {ok:.2f} BBHB\nİşletme Başı: {ok/aktif:.2f} BBHB\n\nPDF: {dosya}")
        except Exception as e: messagebox.showerror("PDF Hatası",str(e)); logging.error(f"pdf:{e}")

    def _sil_rap(self):
        sel=self.tv_rap.selection()
        if not sel: messagebox.showwarning("Seçim","Raporu seçin."); return
        no=self.tv_rap.item(sel[0])["values"][0]
        self._guvenli_sil("Rapor_Gecmisi",no,f"Rapor {no}","DELETE FROM Rapor_Gecmisi WHERE rapor_no=?",(no,),self._yenile_kap)

    # ═══ 3. İHALE TAKİP ═════════════════════════════════════════════════════
    def _t_ihale(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2); yf=ttk.Frame(nb2); bf2=ttk.Frame(nb2); mf=ttk.Frame(nb2)
        nb2.add(lf,text="  📋 İhale Listesi  "); nb2.add(nf,text="  ➕ Yeni Başvuru  ")
        nb2.add(yf,text="  📍 İhale Yerleri  "); nb2.add(bf2,text="  💰 Bedel Hesaplama  ")
        nb2.add(mf,text="  💰 Muhammen Bedel  ")
        # Liste
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,self._filtre_ihale,bg=bg).pack(anchor="w",pady=(0,6))
        df2=tk.Frame(ll,bg=bg); df2.pack(fill="x",pady=(0,6))
        tk.Label(df2,text="Durum:",bg=bg,font=("Segoe UI",10)).pack(side="left",padx=(0,6))
        self.ih_df=ttk.Combobox(df2,values=["Tümü"]+IHALE_DURUMLARI,state="readonly",width=22)
        self.ih_df.set("Tümü"); self.ih_df.bind("<<ComboboxSelected>>",lambda e:self._filtre_ihale(getattr(self,"_ih_ara","")))
        self.ih_df.pack(side="left")
        self.tv_ihale=self._tv(ll,[("id",40,"#"),("koy",100,"Köy"),("ilce",70,"İlçe"),("ad",120,"Başvuran"),
            ("tc",100,"TC"),("tel",100,"Telefon"),("bedel",80,"Bedel"),("durum",120,"Durum"),("tarih",90,"Tarih")],14)
        self.tv_ihale.tag_configure("aktif",background="#EAF4EE")
        self.tv_ihale.tag_configure("tamamlandi",background="#EBF5FB")
        self.tv_ihale.tag_configure("iptal",background="#FDEDEC")
        btnf=tk.Frame(ll,bg=bg); btnf.pack(pady=8)
        MBtn(btnf,"🔄 Yenile",command=self._yenile_ihale,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(btnf,"✏️ Durum Güncelle",command=self._ihale_durum,color=C_WARN,width=18).pack(side="left",padx=4)
        MBtn(btnf,"📋 Süreç",command=self._ihale_surec,color="#8E44AD",width=12).pack(side="left",padx=4)
        MBtn(btnf,"📄 Evrak Üret",command=self._ihale_evrak_uret,color=self.gc("acc"),width=14).pack(side="left",padx=4)
        MBtn(btnf,"✏️ Güncelle",command=self._ihale_guncelle,color=C_INFO,width=12).pack(side="left",padx=4)
        MBtn(btnf,"📷 Fotoğraflar",command=self._ihale_foto,color="#16A085",width=14).pack(side="left",padx=4)
        MBtn(btnf,"📊 Excel Export",command=self._ihale_excel,color="#8E44AD",width=14).pack(side="left",padx=4)
        MBtn(btnf,"🗑 Sil",command=self._ihale_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        # Yeni başvuru — İhale Yerinden Seçerek
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(nn,text="➕ Yeni İhale Başvurusu",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,8))
        tk.Label(nn,text="1) Önce İhale Yerini seçin — yer bilgileri otomatik dolar. 2) Kişi bilgilerini doldurun.",
            font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,8))
        # Yer seçim
        ysf=tk.LabelFrame(nn,text="  📍 İhale Yeri Seç  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=10,pady=6)
        ysf.pack(fill="x",pady=(0,8))
        self.cb_ihale_yer=ttk.Combobox(ysf,state="readonly",width=60); self.cb_ihale_yer.pack(side="left",padx=(0,10))
        self.cb_ihale_yer.bind("<<ComboboxSelected>>",self._ihale_yer_sec)
        MBtn(ysf,"🔄",command=self._ihale_yer_yukle,color=C_INFO,width=4).pack(side="left")
        # Otomatik dolan alanlar
        self.ih_oto=tk.Label(nn,text="↑ Yukarıdan yer seçin",bg=bg,font=("Segoe UI",10),fg="#888")
        self.ih_oto.pack(anchor="w",pady=(0,6))
        # Kişi bilgileri
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1,padx=16,pady=12)
        card.pack(fill="x",pady=(0,8))
        self.ih_g={}
        for lbl,key,tip,r,col in [("Ad Soyad:","ad","entry",0,0),("TC No:","tc","entry",0,2),
            ("Telefon:","tel","entry",1,0),("Adres:","adres","entry",1,2),
            ("Teklif Bedeli (₺):","bedel","entry",2,0)]:
            tk.Label(card,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).grid(row=r,column=col,padx=(4,4),pady=6,sticky="e")
            w_obj=ttk.Entry(card,width=20)
            w_obj.grid(row=r,column=col+1,padx=(0,14),pady=6,sticky="w"); self.ih_g[key]=w_obj
        tk.Label(card,text="Notlar:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=3,column=0,padx=(4,4),pady=6,sticky="e")
        self.ih_g["notlar"]=tk.Text(card,width=56,height=2,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=6,pady=4)
        self.ih_g["notlar"].grid(row=3,column=1,columnspan=3,padx=(0,14),pady=6,sticky="w")
        MBtn(nn,"✅  Başvuruyu Kaydet",command=self._ihale_kaydet,color=self.gc("acc"),width=34).pack(anchor="w",pady=8)
        # İhale Yerleri tab (eskiden ayrı sekmeydi)
        self._build_ihale_yerleri(yf,bg)
        # Bedel hesaplama
        bh=tk.Frame(bf2,bg=bg); bh.pack(fill="both",expand=True,padx=40,pady=30)
        tk.Label(bh,text="💰 İhale Bedel Hesaplayıcı",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")
        tk.Label(bh,text="İhale bedelini girin, tüm kalemler otomatik hesaplansın.",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,14))
        ef=tk.Frame(bh,bg=bg); ef.pack(anchor="w")
        tk.Label(ef,text="İhale Bedeli (₺):",bg=bg,font=("Segoe UI",11,"bold")).pack(side="left",padx=(0,8))
        self.e_bedel_hesap=ttk.Entry(ef,width=18,font=("Segoe UI",12)); self.e_bedel_hesap.pack(side="left",padx=(0,12))
        MBtn(ef,"Hesapla",command=self._bedel_hesapla,color=self.gc("acc"),width=12).pack(side="left")
        self.lbl_bedel_sonuc=tk.Label(bh,text="",bg=bg,font=("Segoe UI",11),fg="#333",justify="left")
        self.lbl_bedel_sonuc.pack(anchor="w",pady=16)
        # Muhammen Bedel tab
        self._build_muhammen_bedel(mf,bg)
        self._yenile_ihale(); self._ihale_yer_yukle()

    def _build_ihale_yerleri(self,parent,bg):
        """İhale Yerleri tab içeriği"""
        ll=tk.Frame(parent,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(ll,text="📍 İhaleye Çıkan Mera Yerleri Havuzu",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,6))
        AramaFrame(ll,self._filtre_iyer,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_iyer=self._tv(ll,[("id",40,"#"),("ilce",80,"İlçe"),("koy",120,"Köy"),("ada",60,"Ada"),
            ("parsel",60,"Parsel"),("alan",80,"Alan(da)"),("kap",90,"Kapasite"),("bedel",110,"Tah.Bedel"),("vasif",80,"Vasıf"),("yil",50,"Yıl")],10)
        bff=tk.Frame(ll,bg=bg); bff.pack(pady=6)
        MBtn(bff,"➕ Yeni Yer",command=self._iyer_ekle_popup,color=self.gc("acc"),width=14).pack(side="left",padx=4)
        MBtn(bff,"🔄 Yenile",command=self._yenile_iyer,color=C_INFO,width=12).pack(side="left",padx=4)
        MBtn(bff,"📥 Import",command=self._iyer_import,color="#8E44AD",width=12).pack(side="left",padx=4)
        MBtn(bff,"📊 Export",command=self._iyer_export,color=self.gc("pri"),width=12).pack(side="left",padx=4)
        MBtn(bff,"🗑 Sil",command=self._iyer_sil,color=C_DANGER,width=8).pack(side="left",padx=4)
        self._yenile_iyer()

    def _build_muhammen_bedel(self,parent,bg):
        """Muhammen Bedel tab içeriği"""
        bh=tk.Frame(parent,bg=bg); bh.pack(fill="both",expand=True,padx=30,pady=20)
        tk.Label(bh,text="💰 Muhammen Bedel Hesaplayıcı (Tahdit Raporu)",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")
        tk.Label(bh,text="Formül: Kapasite(BBHB) = Alan × Kuru Ot Verimi / (Otlatma Gün × 12,5 kg)\nMuhammen Bedel = Kapasite × Otlatma Gün × 12,5 × Kuru Ot Fiyatı",
            font=("Segoe UI",9),fg="#666",bg=bg,justify="left").pack(anchor="w",pady=(4,14))
        mf2=tk.Frame(bh,bg=bg); mf2.pack(fill="x")
        self.mb_g={}
        # İlçe/Köy/Ada/Parsel satırı
        for i,(l,k,w) in enumerate([("İlçe:","ilce",12),("Köy:","koy",12),("Ada:","ada",6),("Parsel:","parsel",6)]):
            tk.Label(mf2,text=l,bg=bg,font=("Segoe UI",10)).grid(row=0,column=i*2,padx=(0,4),sticky="w")
            if k=="ilce":
                w_obj=ttk.Combobox(mf2,values=ILCELER,state="readonly",width=w)
            else:
                w_obj=ttk.Entry(mf2,width=w)
            w_obj.grid(row=0,column=i*2+1,padx=(0,10)); self.mb_g[k]=w_obj
        # Hesaplama alanları
        for i,(l,k,d,w) in enumerate([("Mera Alanı (da):","alan","",14),("Vasıf:","vasif","",14),
            ("Otlatma Gün:","gun","",8),("Kuru Ot Fiyatı (₺/kg):","fiyat","",10)]):
            tk.Label(mf2,text=l,bg=bg,font=("Segoe UI",10)).grid(row=1,column=i*2,padx=(0,4),sticky="w",pady=(6,0))
            if k=="vasif":
                w_obj=ttk.Combobox(mf2,values=MERA_VASIF,state="readonly",width=w)
                w_obj.bind("<<ComboboxSelected>>",self._mb_vasif_sec)
            else: w_obj=ttk.Entry(mf2,width=w)
            w_obj.grid(row=1,column=i*2+1,padx=(0,10),pady=(6,0)); self.mb_g[k]=w_obj
        tk.Label(bh,text="Kuru Ot Verimi (kg/da):",bg=bg,font=("Segoe UI",10,"bold")).pack(anchor="w",pady=(10,2))
        self.mb_kov=tk.Label(bh,text="— vasıf seçin —",bg=bg,font=("Segoe UI",12),fg=self.gc("pri"))
        self.mb_kov.pack(anchor="w")
        bf_mb=tk.Frame(bh,bg=bg); bf_mb.pack(anchor="w",pady=10)
        MBtn(bf_mb,"💰 Hesapla",command=self._mb_hesapla,color=self.gc("acc"),width=16).pack(side="left",padx=(0,8))
        MBtn(bf_mb,"📄 PDF Çıktı Al",command=self._mb_pdf_uret,color=self.gc("pri"),width=16).pack(side="left")
        self.lbl_mb_sonuc=tk.Label(bh,text="",bg=bg,font=("Segoe UI",11),fg="#333",justify="left")
        self.lbl_mb_sonuc.pack(anchor="w",pady=8)

    def _ihale_yer_yukle(self):
        """İhale yerleri combobox'ını doldur"""
        try:
            with db_baglan() as c:
                rows=c.execute("SELECT id,ilce,koy,ada,parsel,alan_da,kapasite_bbhb,COALESCE(tahmini_bedel,0) FROM Ihale_Yerleri WHERE durum='Aktif' ORDER BY ilce,koy").fetchall()
            self._ihale_yer_data={f"{r[1]}/{r[2]} — Ada:{r[3]} Parsel:{r[4]} ({r[5]:.0f} da)":r for r in rows}
            self.cb_ihale_yer["values"]=list(self._ihale_yer_data.keys())
        except Exception: self._ihale_yer_data={}

    def _ihale_yer_sec(self,event=None):
        """İhale yeri seçilince bilgileri otomatik doldur"""
        sec=self.cb_ihale_yer.get()
        r=self._ihale_yer_data.get(sec)
        if not r: return
        bedel_t=r[7] if len(r)>7 else 0
        self.ih_oto.config(text=f"✅ {r[1]}/{r[2]} — Ada: {r[3]}, Parsel: {r[4]}, Alan: {r[5]:.2f} da, Kap: {r[6]:.2f} BBHB, Tah.Bedel: {para_format(bedel_t)}",fg=self.gc("pri"))
        # Seçilen yeri gizli tutuyoruz kaydederken kullanmak için
        self._sec_ihale_yer=r

    def _iyer_ekle_popup(self):
        """İhale yeri ekleme popup"""
        win=tk.Toplevel(self.root); win.title("Yeni İhale Yeri"); win.geometry("500x420"); win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        tk.Label(win,text="📍 Yeni İhale Yeri Ekle",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=14)
        iy_g={}
        for lbl,key,tip in [("İlçe:","ilce","combo"),("Köy:","koy","entry"),("Ada:","ada","entry"),("Parsel:","parsel","entry"),
            ("Alan (da):","alan","entry"),("Vasıf:","vasif","vasif"),("Tahmini Bedel (₺):","bedel","entry"),("Yıl:","yil","entry")]:
            f=tk.Frame(win,bg=C_WHITE); f.pack(fill="x",padx=20,pady=3)
            tk.Label(f,text=lbl,bg=C_WHITE,font=("Segoe UI",10),width=16,anchor="w").pack(side="left")
            if tip=="combo": w=ttk.Combobox(f,values=ILCELER,state="readonly",width=20)
            elif tip=="vasif": w=ttk.Combobox(f,values=MERA_VASIF,state="readonly",width=20)
            else: w=ttk.Entry(f,width=22)
            if key=="yil" and isinstance(w,ttk.Entry): w.insert(0,_yil())
            w.pack(side="left"); iy_g[key]=w
        def _kaydet():
            koy=iy_g["koy"].get().strip()
            if not koy: messagebox.showerror("Hata","Köy zorunlu.",parent=win); return
            alan=float(iy_g["alan"].get() or 0); vasif=iy_g["vasif"].get()
            kov=OT_VERIM_KURU.get(vasif,90); kap=(alan*kov)/(135*GUNLUK_OT_BBHB) if alan>0 else 0
            bedel=para_parse(iy_g["bedel"].get())
            with db_baglan() as c:
                c.execute("INSERT INTO Ihale_Yerleri(ilce,koy,ada,parsel,alan_da,kapasite_bbhb,tahmini_bedel,vasif,yil)VALUES(?,?,?,?,?,?,?,?,?)",
                    (iy_g["ilce"].get(),koy,iy_g["ada"].get(),iy_g["parsel"].get(),alan,kap,bedel,vasif,iy_g["yil"].get()))
            db_log(self.u_id,"İhale Yeri Ekle",f"{koy} {bedel}₺"); win.destroy(); self._yenile_iyer(); self._ihale_yer_yukle()
        MBtn(win,"✅ Kaydet",command=_kaydet,color=self.gc("acc"),width=24).pack(pady=14)

    def _bedel_hesapla(self):
        sonuc=ihale_bedel_hesapla(self.e_bedel_hesap.get())
        if not sonuc: messagebox.showerror("Hata","Geçerli bir bedel girin.\nÖrnek: 52.507,18 veya 52507.18"); return
        self.lbl_bedel_sonuc.config(text=
            f"📊 İhale Bedeli: {para_format(sonuc['toplam'])}\n\n"
            f"💰 Bakanlık Hesabına (%25): {para_format(sonuc['bakanlik_25'])}\n"
            f"🏘️ Köy Hesabına (%75): {para_format(sonuc['koy_75'])}\n"
            f"🔐 Kesin Teminat (%6): {para_format(sonuc['kesin_teminat_6'])}\n"
            f"📋 Damga Vergisi (‰5.69): {para_format(sonuc['damga_vergisi'])}\n"
            f"{'─'*40}\n"
            f"IBAN (Bakanlık): {IBAN_BAKANLIK}\n"
            f"IBAN (Emanet): {IBAN_EMANET}")

    def _yenile_ihale(self):
        if not DB_PATH: return
        try:
            with db_baglan() as c:
                self._all_ihale=c.execute("SELECT id,koy,ilce,ad_soyad,tc,telefon,bedel,durum,tarih FROM Ihaleler ORDER BY id DESC").fetchall()
            self._ih_ara=getattr(self,"_ih_ara",""); self._filtre_ihale(self._ih_ara)
        except Exception as e: logging.error(f"yenile_ihale:{e}")

    def _filtre_ihale(self,a):
        self._ih_ara=a.lower(); filtre=self.ih_df.get()
        self.tv_ihale.delete(*self.tv_ihale.get_children())
        for r in getattr(self,"_all_ihale",[]):
            if filtre!="Tümü" and r[7]!=filtre: continue
            if self._ih_ara and self._ih_ara not in " ".join(str(x) for x in r).lower(): continue
            tag="iptal" if "İptal" in str(r[7]) else "tamamlandi" if "Tamamlandı" in str(r[7]) else "aktif"
            # Bedel sütununu Türk formatında göster
            row=list(r)
            try: row[6]=para_format(row[6])
            except Exception: pass
            self.tv_ihale.insert("","end",values=row,tags=(tag,))

    def _ihale_kaydet(self):
        try:
            yer=getattr(self,"_sec_ihale_yer",None)
            if not yer: raise ValueError("Önce İhale Yerini seçin!")
            ad=self.ih_g["ad"].get().strip(); tc=self.ih_g["tc"].get().strip()
            bedel_str=self.ih_g["bedel"].get()
            if not veri_dogrula([("Ad Soyad",ad,"str"),("Teklif Bedeli",bedel_str,"float+")]): return
            bedel=para_parse(bedel_str)
            if not tc_kontrol_ve_devam(tc): return
            koy=yer[2]; ilce=yer[1]; ada_parsel=f"{yer[3]}/{yer[4]}"; kapasite=f"{yer[6]:.1f} BBHB"
            notlar=self.ih_g["notlar"].get("1.0",tk.END).strip()
            with db_baglan() as conn:
                cur=conn.cursor()
                cur.execute("INSERT INTO Ihaleler(koy,ilce,ad_soyad,tc,telefon,adres,bedel,durum,notlar,tarih,ada_parsel,kapasite,ihale_yeri_id)VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (koy,ilce,ad,tc,self.ih_g["tel"].get(),self.ih_g["adres"].get(),bedel,"Başvuru Alındı",notlar,
                     datetime.now().strftime("%d.%m.%Y"),ada_parsel,kapasite,yer[0]))
                ih_id=cur.lastrowid
                cur.execute("INSERT INTO Ihale_Log(ihale_id,tarih,personel,durum,not_icerik)VALUES(?,?,?,?,?)",
                    (ih_id,datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad,"Başvuru Alındı","Yeni başvuru."))
            db_log(self.u_id,"Yeni İhale",f"{koy} {bedel}₺")
            self._son_islem_kaydet("İhale",koy)
            for key,w in self.ih_g.items():
                if isinstance(w,(ttk.Entry,ttk.Combobox)): w.set("") if isinstance(w,ttk.Combobox) else w.delete(0,tk.END)
                elif isinstance(w,tk.Text): w.delete("1.0",tk.END)
            self._sec_ihale_yer=None; self.ih_oto.config(text="↑ Yukarıdan yer seçin",fg="#888")
            self.cb_ihale_yer.set("")
            self._yenile_ihale(); messagebox.showinfo("Tamam",f"'{koy}' ihale başvurusu kaydedildi.\n{ada_parsel} — {kapasite}")
        except ValueError as e: messagebox.showerror("Hata",str(e))
        except Exception as e: messagebox.showerror("DB Hatası",str(e))

    def _ihale_durum(self):
        sel=self.tv_ihale.selection()
        if not sel: messagebox.showwarning("Seçim","İhale seçin."); return
        ih_id=self.tv_ihale.item(sel[0])["values"][0]
        win=tk.Toplevel(self.root); win.title("Durum Güncelle"); win.geometry("400x280"); win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        tk.Label(win,text="Yeni Durum:",bg=C_WHITE,font=("Segoe UI",11,"bold")).pack(pady=(16,4))
        cb=ttk.Combobox(win,values=IHALE_DURUMLARI,state="readonly",width=30); cb.pack(padx=20)
        tk.Label(win,text="Not:",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=20,pady=(10,2))
        not_e=tk.Text(win,height=4,width=42,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        not_e.pack(padx=20)
        def _k():
            yeni=cb.get(); acik=not_e.get("1.0",tk.END).strip()
            if not yeni: messagebox.showwarning("Uyarı","Durum seçin.",parent=win); return
            with db_baglan() as c:
                c.execute("UPDATE Ihaleler SET durum=? WHERE id=?",(yeni,ih_id))
                c.execute("INSERT INTO Ihale_Log(ihale_id,tarih,personel,durum,not_icerik)VALUES(?,?,?,?,?)",
                    (ih_id,datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad,yeni,acik))
            db_log(self.u_id,"İhale Durum",f"ID:{ih_id}→{yeni}"); self._yenile_ihale(); win.destroy()
        MBtn(win,"Güncelle",command=_k,width=18).pack(pady=12)

    def _ihale_surec(self):
        sel=self.tv_ihale.selection()
        if not sel: messagebox.showwarning("Seçim","İhale seçin."); return
        vals=self.tv_ihale.item(sel[0])["values"]; ih_id=vals[0]; koy=vals[1]
        win=tk.Toplevel(self.root); win.title(f"Süreç — {koy}"); win.geometry("700x400"); win.configure(bg=C_WHITE)
        tk.Label(win,text=f"İhale Süreç Geçmişi: {koy}",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=12)
        tv=ttk.Treeview(win,columns=("t","p","d","n"),show="headings",height=14)
        for col,w,bas in [("t",140,"Tarih"),("p",130,"Personel"),("d",130,"Durum"),("n",280,"Not")]:
            tv.heading(col,text=bas); tv.column(col,width=w)
        tv.pack(fill="both",expand=True,padx=12,pady=8)
        try:
            with db_baglan() as c:
                for r in c.execute("SELECT tarih,personel,durum,not_icerik FROM Ihale_Log WHERE ihale_id=? ORDER BY id",(ih_id,)).fetchall():
                    tv.insert("","end",values=r)
        except Exception as e: logging.error(f"surec:{e}")

    def _ihale_evrak_uret(self):
        """Seçili ihale için Word evrak üret"""
        sel=self.tv_ihale.selection()
        if not sel: messagebox.showwarning("Seçim","İhale seçin."); return
        vals=self.tv_ihale.item(sel[0])["values"]
        ih_id=vals[0]
        if not DOCX_OK: messagebox.showerror("Hata","python-docx kurulu değil.\npip install python-docx"); return
        try:
            with db_baglan() as c:
                r=c.execute("SELECT koy,ilce,ad_soyad,tc,telefon,adres,bedel,ada_parsel,kapasite FROM Ihaleler WHERE id=?",(ih_id,)).fetchone()
        except Exception: return
        if not r: return
        yil=_yil()
        veri={"koy":r[0],"ilce":r[1],"ad_soyad":r[2],"tc":r[3],"telefon":r[4],"adres":r[5],
              "bedel":r[6],"ada_parsel":r[7],"kapasite":r[8],"mera":f"{r[1]}/{r[0]}",
              "tarih":datetime.now().strftime("%d/%m/%Y"),
              "komisyon_baskani":getattr(self,"vali_yardimcisi","Semih CEMBEKLİ"),
              "baslangic":f"{OTLATMA_BASLANGIC}/{yil}","bitis":f"{OTLATMA_BITIS}/{yil}",
              "dogum_yeri":"","hayvan_bilgi":"","hazirlayan":self.u_ad}
        # Tahmini bedel — ihale yerine bağlı kayıttan çek
        try:
            with db_baglan() as c:
                iy=c.execute("SELECT ilce,koy,ada,parsel,tahmini_bedel FROM Ihale_Yerleri WHERE koy=? AND ilce=? LIMIT 1",(r[0],r[1])).fetchone()
                if iy and iy[4]:
                    veri["tahmini_bedel"]=iy[4]
                    veri["gecici_teminat"]=iy[4]*0.30
                else:
                    veri["tahmini_bedel"]=0; veri["gecici_teminat"]=0
        except Exception: veri["tahmini_bedel"]=0; veri["gecici_teminat"]=0
        # İl Müdürü ve unvan bilgileri
        try:
            with db_baglan() as c:
                for row in c.execute("SELECT anahtar,deger FROM Iletisim_Bilgileri WHERE anahtar IN('il_muduru','evrak_unvan')"):
                    if row[0]=="il_muduru" and row[1]: veri["il_muduru"]=row[1]
                    elif row[0]=="evrak_unvan" and row[1]:
                        parts=row[1].split("\n") if "\n" in row[1] else [row[1]]
                        if len(parts)>=1: veri["evrak_unvan_1"]=parts[0]
                        if len(parts)>=2: veri["evrak_unvan_2"]=parts[1]
        except Exception: pass
        win=tk.Toplevel(self.root); win.title("Evrak Seç"); win.geometry("380x320"); win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        tk.Label(win,text="Hangi evrakı üretmek istiyorsunuz?",font=("Segoe UI",12,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=14)
        def _uret(tip):
            dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile=f"{tip}_{r[0]}_{r[2]}.docx",filetypes=[("Word","*.docx")])
            if not dosya: return
            try:
                if tip=="Katılım Evrakları": word_katilim_evrak(dosya,veri)
                elif tip=="Kiralama Sözleşmesi": word_kiralama_sozlesme(dosya,veri)
                elif tip=="Kiralama Şartnamesi": word_kiralama_sartname(dosya,veri)
                messagebox.showinfo("Tamam",f"Evrak üretildi:\n{dosya}"); win.destroy()
            except Exception as e: messagebox.showerror("Hata",str(e))
        MBtn(win,"📋 Katılım Evrakları",command=lambda:_uret("Katılım Evrakları"),color=self.gc("acc"),width=30).pack(pady=6)
        MBtn(win,"📝 Kiralama Sözleşmesi (14 madde)",command=lambda:_uret("Kiralama Sözleşmesi"),color=C_INFO,width=30).pack(pady=6)
        MBtn(win,"📄 Kiralama Şartnamesi (22 madde)",command=lambda:_uret("Kiralama Şartnamesi"),color="#8E44AD",width=30).pack(pady=6)

    def _ihale_guncelle(self):
        """Seçili ihale kaydını güncelle"""
        sel=self.tv_ihale.selection()
        if not sel: messagebox.showwarning("Seçim","Güncellenecek ihaleyi seçin."); return
        ih_id=self.tv_ihale.item(sel[0])["values"][0]
        try:
            with db_baglan() as c:
                r=c.execute("SELECT koy,ilce,ad_soyad,tc,telefon,adres,bedel,ada_parsel,kapasite,durum FROM Ihaleler WHERE id=?",(ih_id,)).fetchone()
        except Exception: return
        if not r: return
        win=tk.Toplevel(self.root); win.title("İhale Güncelle"); win.geometry("480x450"); win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        tk.Label(win,text="✏️ İhale Kaydını Güncelle",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=10)
        gg={}
        for i,(lbl,key,val) in enumerate([("Ad Soyad:","ad",r[2]),("TC:","tc",r[3]),("Telefon:","tel",r[4]),
            ("Adres:","adres",r[5]),("Teklif Bedeli (₺):","bedel",r[6] or ""),("Durum:","durum",r[9] or "Başvuru Alındı")]):
            f=tk.Frame(win,bg=C_WHITE); f.pack(fill="x",padx=20,pady=2)
            tk.Label(f,text=lbl,bg=C_WHITE,font=("Segoe UI",10),width=16,anchor="w").pack(side="left")
            if key=="durum":
                w=ttk.Combobox(f,values=IHALE_DURUMLARI,state="readonly",width=22); w.set(val or "Başvuru Alındı")
            else:
                w=ttk.Entry(f,width=26); w.insert(0,str(val) if val else "")
            w.pack(side="left"); gg[key]=w
        def _kaydet():
            with db_baglan() as c:
                c.execute("UPDATE Ihaleler SET ad_soyad=?,tc=?,telefon=?,adres=?,bedel=?,durum=? WHERE id=?",
                    (gg["ad"].get(),gg["tc"].get(),gg["tel"].get(),gg["adres"].get(),gg["bedel"].get(),gg["durum"].get(),ih_id))
            db_log(self.u_id,"İhale Güncelle",f"#{ih_id} {gg['ad'].get()}")
            win.destroy(); self._yenile_ihale()
            messagebox.showinfo("Tamam","İhale kaydı güncellendi.")
        MBtn(win,"💾 Güncelle",command=_kaydet,color=self.gc("acc"),width=22).pack(pady=14)

    def _ihale_excel(self):
        """İhaleleri Excel'e aktar"""
        try:
            with db_baglan() as c:
                rows=c.execute("SELECT id,koy,ilce,ad_soyad,tc,telefon,adres,bedel,ada_parsel,kapasite,durum,tarih,notlar FROM Ihaleler ORDER BY id DESC").fetchall()
            if not rows: messagebox.showwarning("Veri Yok","Kayıt yok."); return
            yol=filedialog.asksaveasfilename(defaultextension=".xlsx",initialfile="Ihaleler.xlsx",filetypes=[("Excel","*.xlsx")])
            if yol:
                pd.DataFrame(rows,columns=["ID","Köy","İlçe","Ad Soyad","TC","Telefon","Adres","Bedel","Ada/Parsel","Kapasite","Durum","Tarih","Notlar"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"İhaleler dışa aktarıldı:\n{yol}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _ihale_foto(self):
        """Seçili ihaleye foto yönetimi"""
        sel=self.tv_ihale.selection()
        if not sel: messagebox.showwarning("Seçim","İhale seçin."); return
        v=self.tv_ihale.item(sel[0])["values"]
        self._kayit_foto_yonet("Ihaleler",v[0],f"İhale: {v[3]} ({v[1]})")

    def _ceza_foto(self):
        """Seçili cezaya foto yönetimi"""
        sel=self.tv_ceza.selection()
        if not sel: messagebox.showwarning("Seçim","Ceza seçin."); return
        v=self.tv_ceza.item(sel[0])["values"]
        self._kayit_foto_yonet("Idari_Cezalar",v[0],f"Ceza: {v[1]} — {v[4]}")

    def _sik_foto(self):
        """Seçili şikayete foto yönetimi"""
        sel=self.tv_sik.selection()
        if not sel: messagebox.showwarning("Seçim","Şikayet seçin."); return
        v=self.tv_sik.item(sel[0])["values"]
        self._kayit_foto_yonet("Sikayetler",v[0],f"Şikayet: {v[1]}")

    def _sik_excel(self):
        """Şikayetleri Excel'e aktar"""
        try:
            with db_baglan() as c:
                rows=c.execute("SELECT id,koy,ilce,sikayet_eden,tur,aciklama,durum,tarih FROM Sikayetler ORDER BY id DESC").fetchall()
            if not rows: messagebox.showwarning("Veri Yok","Kayıt yok."); return
            yol=filedialog.asksaveasfilename(defaultextension=".xlsx",initialfile="Sikayetler.xlsx",filetypes=[("Excel","*.xlsx")])
            if yol:
                pd.DataFrame(rows,columns=["ID","Köy","İlçe","Şikayet Eden","Tür","Açıklama","Durum","Tarih"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"Şikayetler dışa aktarıldı:\n{yol}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _ihale_sil(self):
        sel=self.tv_ihale.selection()
        if not sel: return
        ih_id=self.tv_ihale.item(sel[0])["values"][0]; koy=self.tv_ihale.item(sel[0])["values"][1]
        self._guvenli_sil("Ihaleler",ih_id,f"İhale: {koy}","DELETE FROM Ihaleler WHERE id=?",(ih_id,),self._yenile_ihale)

    # ═══ 4. ISLAH ════════════════════════════════════════════════════════════
    def _t_islah(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2); pf=ttk.Frame(nb2)
        nb2.add(lf,text="  📋 Islah Listesi  "); nb2.add(nf,text="  ➕ Yeni Dilekçe  "); nb2.add(pf,text="  📊 Proje Takip Cetveli  ")
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,self._filtre_islah,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_islah=self._tv(ll,[("id",40,"#"),("koy",110,"Köy"),("ilce",70,"İlçe"),("talep",120,"Talep Eden"),
            ("dt",90,"Dilekçe T."),("talan",70,"Talep (da)"),("valan",70,"Verilen (da)"),("durum",100,"Durum")],12)
        self.tv_islah.tag_configure("bekliyor",background="#FEF9E7")
        self.tv_islah.tag_configure("verildi",background="#EAFAF1")
        self.tv_islah.tag_configure("verilmedi",background="#FDEDEC")
        bf=tk.Frame(ll,bg=bg); bf.pack(pady=8)
        MBtn(bf,"🔄 Yenile",command=self._yenile_islah,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(bf,"✏️ Sonuç Gir",command=self._islah_sonuc,color=C_WARN,width=16).pack(side="left",padx=4)
        MBtn(bf,"🗑 Sil",command=self._islah_sil,color=C_DANGER,width=12).pack(side="left",padx=4)
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=16)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.48,anchor="center",width=640,height=480)
        tk.Label(card,text="Yeni Islah/Amenajman Dilekçesi",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=4,pady=(18,14))
        self.isl_g={}
        for lbl,key,tip,r,col in [("Köy:","koy","entry",1,0),("İlçe:","ilce","combo",1,2),
            ("Talep Eden:","talep","entry",2,0),("Dilekçe Tarihi:","dt","entry",2,2),
            ("Talep Alanı (da):","alan","entry",3,0),("İş Programı:","ip","entry",3,2)]:
            tk.Label(card,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).grid(row=r,column=col,padx=(20,6),pady=8,sticky="e")
            w_obj=ttk.Combobox(card,values=ILCELER,state="readonly",width=18) if tip=="combo" else ttk.Entry(card,width=20)
            if key=="dt": w_obj.insert(0,datetime.now().strftime("%d.%m.%Y"))
            w_obj.grid(row=r,column=col+1,padx=(0,14),pady=8,sticky="w"); self.isl_g[key]=w_obj
        tk.Label(card,text="Açıklama:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=4,column=0,padx=(20,6),pady=8,sticky="ne")
        self.isl_g["acik"]=tk.Text(card,width=46,height=4,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        self.isl_g["acik"].grid(row=4,column=1,columnspan=3,padx=(0,20),pady=8,sticky="w")
        MBtn(card,"✅  Dilekçeyi Kaydet",command=self._islah_kaydet,color=self.gc("acc"),width=34).grid(row=5,column=0,columnspan=4,pady=18)
        # Proje Takip Cetveli
        pp=tk.Frame(pf,bg=bg); pp.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(pp,text="📊 Çayır Mera Islah ve Amenajman Projeleri — Proje Takip Cetveli",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,8))
        # Yıl seçimi
        yf2=tk.Frame(pp,bg=bg); yf2.pack(fill="x",pady=(0,6))
        tk.Label(yf2,text="Yıl:",bg=bg,font=("Segoe UI",10,"bold")).pack(side="left")
        self.prj_yil=ttk.Combobox(yf2,values=[str(y) for y in range(2020,2031)],state="readonly",width=8)
        self.prj_yil.set(_yil()); self.prj_yil.pack(side="left",padx=6)
        self.prj_yil.bind("<<ComboboxSelected>>",lambda e:self._yenile_proje())
        AramaFrame(yf2,self._filtre_proje,bg=bg).pack(side="left",padx=10)
        self.tv_prj=self._tv(pp,[("id",35,"#"),("ilce",70,"İlçe"),("koy",100,"Köy"),
            ("gub",65,"Gübre"),("mt",65,"M.Tohum"),("yt",65,"Y.Tohum"),
            ("gol",50,"Gölg."),("siv",50,"Sıvat"),("bor",55,"Boru"),
            ("cob",50,"Ç.Evi"),("tuz",50,"Tuzlk"),("kas",50,"K.Kzğ"),
            ("fin",100,"Finans"),("dur",80,"Durum")],10)
        bf2=tk.Frame(pp,bg=bg); bf2.pack(pady=6)
        MBtn(bf2,"➕ Yeni Proje",command=self._proje_ekle,color=self.gc("acc"),width=14).pack(side="left",padx=4)
        MBtn(bf2,"🔄 Yenile",command=self._yenile_proje,color=C_INFO,width=12).pack(side="left",padx=4)
        MBtn(bf2,"📥 Excel Import",command=self._proje_import,color="#8E44AD",width=14).pack(side="left",padx=4)
        MBtn(bf2,"📊 Excel Export",command=self._proje_export,color=self.gc("pri"),width=14).pack(side="left",padx=4)
        MBtn(bf2,"🗑 Sil",command=self._proje_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        self._yenile_islah(); self._yenile_proje()

    def _yenile_islah(self):
        if not DB_PATH: return
        try:
            with db_baglan() as c:
                self._all_islah=c.execute("SELECT id,koy,ilce,talep_eden,dilekce_tarihi,talep_alani,verilen_alan,durum FROM Islah_Amenajman ORDER BY id DESC").fetchall()
            self._filtre_islah("")
        except Exception as e: logging.error(f"yenile_islah:{e}")

    def _filtre_islah(self,a):
        self.tv_islah.delete(*self.tv_islah.get_children()); a=a.lower()
        for r in getattr(self,"_all_islah",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            d=str(r[7]); tag="bekliyor" if d=="Bekliyor" else "verildi" if "Verildi" in d else "verilmedi"
            self.tv_islah.insert("","end",values=r,tags=(tag,))

    def _islah_kaydet(self):
        try:
            koy=self.isl_g["koy"].get().strip(); talep=self.isl_g["talep"].get().strip()
            if not koy or not talep: raise ValueError("Köy ve talep eden zorunlu.")
            alan=float(self.isl_g["alan"].get() or 0)
            acik=self.isl_g["acik"].get("1.0",tk.END).strip()
            with db_baglan() as c:
                c.execute("INSERT INTO Islah_Amenajman(koy,ilce,dilekce_tarihi,talep_eden,talep_alani,talep_aciklama,is_programi,durum)VALUES(?,?,?,?,?,?,?,?)",
                    (koy,self.isl_g["ilce"].get(),self.isl_g["dt"].get(),talep,alan,acik,self.isl_g["ip"].get(),"Bekliyor"))
            db_log(self.u_id,"Islah Dilekçe",koy); self._son_islem_kaydet("Islah",koy)
            for key,w in self.isl_g.items():
                if isinstance(w,ttk.Entry): w.delete(0,tk.END)
                elif isinstance(w,ttk.Combobox): w.set("")
                elif isinstance(w,tk.Text): w.delete("1.0",tk.END)
            self.isl_g["dt"].insert(0,datetime.now().strftime("%d.%m.%Y"))
            self._yenile_islah(); messagebox.showinfo("Tamam",f"'{koy}' dilekçesi kaydedildi.")
        except ValueError as e: messagebox.showerror("Hata",str(e))

    def _islah_sonuc(self):
        sel=self.tv_islah.selection()
        if not sel: messagebox.showwarning("Seçim","Kayıt seçin."); return
        is_id=self.tv_islah.item(sel[0])["values"][0]; koy=self.tv_islah.item(sel[0])["values"][1]
        win=tk.Toplevel(self.root); win.title(f"Sonuç — {koy}"); win.geometry("460x340"); win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        tk.Label(win,text=f"Islah Sonucu: {koy}",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=14)
        tk.Label(win,text="Durum:",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=20)
        cb=ttk.Combobox(win,values=["Bekliyor","Islah Yapılıyor","Islah Verildi","Islah Verilmedi","Kısmi Verildi"],state="readonly",width=30)
        cb.pack(padx=20,pady=4)
        tk.Label(win,text="Verilen Alan (da):",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=20)
        e_v=ttk.Entry(win,width=20); e_v.insert(0,"0"); e_v.pack(padx=20,pady=4,anchor="w")
        tk.Label(win,text="Açıklama:",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=20)
        txt=tk.Text(win,height=3,width=44,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        txt.pack(padx=20,pady=4)
        def _k():
            try:
                v=float(e_v.get() or 0); n=txt.get("1.0",tk.END).strip(); d=cb.get() or "Bekliyor"
                with db_baglan() as c:
                    c.execute("UPDATE Islah_Amenajman SET verilen_alan=?,verilmeme_neden=?,durum=?,kapanma_tarihi=? WHERE id=?",
                        (v,n,d,datetime.now().strftime("%d.%m.%Y"),is_id))
                db_log(self.u_id,"Islah Sonuç",f"ID:{is_id} {d}"); self._yenile_islah(); win.destroy()
            except ValueError as e: messagebox.showerror("Hata",str(e),parent=win)
        MBtn(win,"Sonucu Kaydet",command=_k,width=20).pack(pady=10)

    def _islah_sil(self):
        sel=self.tv_islah.selection()
        if not sel: return
        is_id=self.tv_islah.item(sel[0])["values"][0]; koy=self.tv_islah.item(sel[0])["values"][1]
        self._guvenli_sil("Islah_Amenajman",is_id,f"Islah: {koy}","DELETE FROM Islah_Amenajman WHERE id=?",(is_id,),self._yenile_islah)

    # ─── PROJE TAKİP CETVELİ ─────────────────────────────────────────────────
    ISLAH_FAALIYETLER = ["Gübreleme (da)","Mera Tohum (da)","Yem Bitkisi Tohum (da)",
        "Gölgelik (adet)","Sıvat (adet)","Kangal Boru (m)","Çoban Evi (adet)","Tuzluk (adet)","Kaşınma Kazığı (adet)"]
    FINANSMAN = ["DAP","AKAK","Bakanlık","İl Özel İdaresi","Köylere Hizmet Götürme","Diğer"]

    def _yenile_proje(self):
        if not DB_PATH: return
        yil=self.prj_yil.get() if hasattr(self,"prj_yil") else _yil()
        try:
            with db_baglan() as c:
                self._all_prj=c.execute("SELECT id,ilce,koy,gubre_da,mera_tohum_da,yem_tohum_da,golgelik,sivat,boru_m,coban_evi,tuzluk,kasinka,finansman,durum FROM Islah_Projeler WHERE yil=? ORDER BY ilce,koy",(yil,)).fetchall()
            self._filtre_proje("")
        except Exception as e: logging.error(f"proje:{e}")

    def _filtre_proje(self,a):
        if not hasattr(self,"tv_prj"): return
        self.tv_prj.delete(*self.tv_prj.get_children()); a=a.lower()
        for r in getattr(self,"_all_prj",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            self.tv_prj.insert("","end",values=r)

    def _proje_ekle(self):
        win=tk.Toplevel(self.root); win.title("Yeni Proje Kaydı"); win.geometry("560x520"); win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        tk.Label(win,text="📊 Yeni Islah Projesi",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=14)
        pg={}
        for lbl,key in [("İlçe:","ilce"),("Köy:","koy")]:
            f=tk.Frame(win,bg=C_WHITE); f.pack(fill="x",padx=20,pady=4)
            tk.Label(f,text=lbl,bg=C_WHITE,font=("Segoe UI",10),width=18,anchor="w").pack(side="left")
            if key=="ilce":
                w=ttk.Combobox(f,values=ILCELER,state="readonly",width=20)
            else:
                w=ttk.Entry(f,width=22)
            w.pack(side="left"); pg[key]=w
        tk.Label(win,text="── Harcama Faaliyetleri ──",font=("Segoe UI",10,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=(10,4))
        faaliyet_keys=["gubre","mt","yt","gol","siv","bor","cob","tuz","kas"]
        faaliyet_labels=["Gübreleme (da)","Mera Tohum (da)","Yem Bitkisi Tohum (da)","Gölgelik (adet)","Sıvat (adet)","Kangal Boru (m)","Çoban Evi (adet)","Tuzluk (adet)","Kaşınma Kazığı (adet)"]
        for lbl,key in zip(faaliyet_labels,faaliyet_keys):
            f=tk.Frame(win,bg=C_WHITE); f.pack(fill="x",padx=20,pady=2)
            tk.Label(f,text=lbl+":",bg=C_WHITE,font=("Segoe UI",9),width=22,anchor="w").pack(side="left")
            e=ttk.Entry(f,width=12); e.insert(0,"0"); e.pack(side="left"); pg[key]=e
        f2=tk.Frame(win,bg=C_WHITE); f2.pack(fill="x",padx=20,pady=6)
        tk.Label(f2,text="Finansman:",bg=C_WHITE,font=("Segoe UI",10),width=18,anchor="w").pack(side="left")
        pg["fin"]=ttk.Combobox(f2,values=self.FINANSMAN,state="readonly",width=20); pg["fin"].pack(side="left")
        def _kaydet():
            koy=pg["koy"].get().strip()
            if not koy: messagebox.showerror("Hata","Köy zorunlu.",parent=win); return
            try:
                with db_baglan() as c:
                    c.execute("INSERT INTO Islah_Projeler(yil,ilce,koy,gubre_da,mera_tohum_da,yem_tohum_da,golgelik,sivat,boru_m,coban_evi,tuzluk,kasinka,finansman)VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",
                        (self.prj_yil.get(),pg["ilce"].get(),koy,
                         float(pg["gubre"].get() or 0),float(pg["mt"].get() or 0),float(pg["yt"].get() or 0),
                         int(pg["gol"].get() or 0),int(pg["siv"].get() or 0),float(pg["bor"].get() or 0),
                         int(pg["cob"].get() or 0),int(pg["tuz"].get() or 0),int(pg["kas"].get() or 0),pg["fin"].get()))
                db_log(self.u_id,"Proje Ekle",koy); win.destroy(); self._yenile_proje()
            except Exception as e: messagebox.showerror("Hata",str(e),parent=win)
        MBtn(win,"✅ Kaydet",command=_kaydet,color=self.gc("acc"),width=24).pack(pady=12)

    def _proje_import(self):
        yol=filedialog.askopenfilename(filetypes=[("Excel","*.xlsx")])
        if not yol: return
        try:
            df=pd.read_excel(yol)
            n=0
            with db_baglan() as c:
                for _,r in df.iterrows():
                    v=r.tolist()
                    if len(v)>=5:
                        c.execute("INSERT INTO Islah_Projeler(yil,ilce,koy,gubre_da,mera_tohum_da,yem_tohum_da,golgelik,sivat,boru_m,coban_evi,tuzluk,kasinka,finansman)VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",
                            (self.prj_yil.get(),str(v[0] if len(v)>0 else ""),str(v[1] if len(v)>1 else ""),
                             float(v[2] or 0) if len(v)>2 else 0,float(v[3] or 0) if len(v)>3 else 0,
                             float(v[4] or 0) if len(v)>4 else 0,int(v[5] or 0) if len(v)>5 else 0,
                             int(v[6] or 0) if len(v)>6 else 0,float(v[7] or 0) if len(v)>7 else 0,
                             int(v[8] or 0) if len(v)>8 else 0,int(v[9] or 0) if len(v)>9 else 0,
                             int(v[10] or 0) if len(v)>10 else 0,str(v[11]) if len(v)>11 else ""))
                        n+=1
            self._yenile_proje(); messagebox.showinfo("Tamam",f"{n} proje kaydı yüklendi.")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _proje_export(self):
        if not getattr(self,"_all_prj",None): return
        yol=filedialog.asksaveasfilename(defaultextension=".xlsx",filetypes=[("Excel","*.xlsx")])
        if yol:
            try:
                pd.DataFrame(self._all_prj,columns=["ID","İlçe","Köy","Gübreleme","M.Tohum","Y.Tohum","Gölgelik","Sıvat","Boru(m)","Ç.Evi","Tuzluk","K.Kazığı","Finansman","Durum"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"Dışa aktarıldı:\n{yol}")
            except Exception as e: messagebox.showerror("Hata",str(e))

    def _proje_sil(self):
        sel=self.tv_prj.selection()
        if not sel: return
        pid=self.tv_prj.item(sel[0])["values"][0]; koy=self.tv_prj.item(sel[0])["values"][2]
        self._guvenli_sil("Islah_Projeler",pid,f"Proje: {koy}","DELETE FROM Islah_Projeler WHERE id=?",(pid,),self._yenile_proje)

    # ═══ 5. TAHSİS / MADDE 14 ═══════════════════════════════════════════════
    def _t_tahsis(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2); tf=ttk.Frame(nb2); of=ttk.Frame(nb2)
        nb2.add(lf,text="  📋 Tahsis Listesi  "); nb2.add(nf,text="  ➕ Yeni Tahsis/Md.14  ")
        nb2.add(tf,text="  📝 TER (Teknik Ekip Raporu)  "); nb2.add(of,text="  💰 TAD Ot Bedeli  ")
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,lambda a:None,bg=bg).pack(anchor="w",pady=(0,6))
        # Aşama göstergesi
        af=tk.Frame(ll,bg=bg); af.pack(fill="x",pady=(0,8))
        tk.Label(af,text="Aşama Durumu:",bg=bg,font=("Segoe UI",9,"bold")).pack(side="left",padx=(0,8))
        self._al=[]
        for a in TAHSIS_ASAMALARI:
            l=tk.Label(af,text=f" {a} ",bg="#E0E0E0",fg="#555",font=("Segoe UI",8),padx=6,pady=3)
            l.pack(side="left",padx=2); self._al.append(l)
        self.tv_tahsis=self._tv(ll,[("id",40,"#"),("koy",100,"Köy"),("ilce",70,"İlçe"),("ada",60,"Ada"),
            ("parsel",60,"Parsel"),("kurum",120,"Kurum"),("amac",120,"Amaç"),("alan",70,"Alan (ha)"),
            ("asama",140,"Aşama"),("durum",90,"Durum"),("bt",90,"Başvuru T.")],12)
        self.tv_tahsis.bind("<<TreeviewSelect>>",self._tahsis_sec)
        self.tv_tahsis.tag_configure("devam",background="#EBF5FB")
        self.tv_tahsis.tag_configure("tamam",background="#EAFAF1")
        self.tv_tahsis.tag_configure("red",background="#FDEDEC")
        bf=tk.Frame(ll,bg=bg); bf.pack(pady=8)
        MBtn(bf,"🔄 Yenile",command=self._yenile_tahsis,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(bf,"▶️ Sonraki Aşama",command=self._tahsis_ileri,color=self.gc("acc"),width=16).pack(side="left",padx=4)
        MBtn(bf,"❌ Reddet/Kapat",command=self._tahsis_red,color=C_DANGER,width=14).pack(side="left",padx=4)
        MBtn(bf,"📋 Süreç",command=self._tahsis_gecmis,color="#8E44AD",width=10).pack(side="left",padx=4)
        MBtn(bf,"📊 Excel Export",command=self._tahsis_excel,color="#8E44AD",width=14).pack(side="left",padx=4)
        # Yeni
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=16)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.45,anchor="center",width=680,height=540)
        tk.Label(card,text="Yeni Tahsis / Madde 14 Başvurusu",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=4,pady=(18,14))
        self.tah_g={}
        for lbl,key,tip,r,col in [("Köy:","koy","entry",1,0),("İlçe:","ilce","combo",1,2),("Ada No:","ada","entry",2,0),
            ("Parsel:","parsel","entry",2,2),("Kurum:","kurum","entry",3,0),("Alan (ha):","alan","entry",3,2),
            ("Tahsis Amacı:","amac","entry",4,0),("Md.14 Bent:","madde14","md14combo",4,2)]:
            tk.Label(card,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).grid(row=r,column=col,padx=(20,6),pady=8,sticky="e")
            if tip=="combo": w_obj=ttk.Combobox(card,values=ILCELER,state="readonly",width=18)
            elif tip=="md14combo": w_obj=ttk.Combobox(card,values=["—","(a) Maden/Petrol","(b) Turizm","(c) Kamu Yatırımı","(d) Yerleşim/Muhafaza","(e) Köy Kanunu","(f) Güvenlik","(g) Doğal Afet"],state="readonly",width=22)
            else: w_obj=ttk.Entry(card,width=20)
            w_obj.grid(row=r,column=col+1,padx=(0,14),pady=8,sticky="w"); self.tah_g[key]=w_obj
        tk.Label(card,text="20 Yıllık Ot Geliri (₺):",bg=C_WHITE,font=("Segoe UI",10)).grid(row=5,column=0,padx=(20,6),pady=8,sticky="e")
        self.tah_g["ot_geliri"]=ttk.Entry(card,width=20); self.tah_g["ot_geliri"].grid(row=5,column=1,padx=(0,14),pady=8,sticky="w")
        tk.Label(card,text="Süre Bitiş Tarihi:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=5,column=2,padx=(20,6),pady=8,sticky="e")
        self.tah_g["sure_bitis"]=ttk.Entry(card,width=14); self.tah_g["sure_bitis"].grid(row=5,column=3,padx=(0,14),pady=8,sticky="w")
        tk.Label(card,text="Süre Tipi:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=6,column=0,padx=(20,6),pady=8,sticky="e")
        self.tah_g["sure_tipi"]=ttk.Combobox(card,values=["Askı İlanı (30 gün)","İtiraz Süresi (60 gün)","Tescil Süresi","Yatırım Başlama (2 yıl)","Diğer"],state="readonly",width=24)
        self.tah_g["sure_tipi"].grid(row=6,column=1,padx=(0,14),pady=8,sticky="w")
        tk.Label(card,text="Notlar:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=7,column=0,padx=(20,6),pady=8,sticky="ne")
        self.tah_g["notlar"]=tk.Text(card,width=46,height=3,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        self.tah_g["notlar"].grid(row=7,column=1,columnspan=3,padx=(0,20),pady=8,sticky="w")
        MBtn(card,"✅  Başvuruyu Kaydet",command=self._tahsis_kaydet,color=self.gc("acc"),width=34).grid(row=8,column=0,columnspan=4,pady=18)
        # ═══ TER — Teknik Ekip Raporu ═══
        self._build_ter(tf,bg)
        # ═══ TAD Ot Bedeli ═══
        self._build_tad_ot(of,bg)
        self._yenile_tahsis()

    def _build_tad_ot(self,parent,bg):
        """TAD 20 Yıllık Ot Bedeli Hesaplama — Bakanlık formülü"""
        f=tk.Frame(parent,bg=bg); f.pack(fill="both",expand=True,padx=20,pady=14)
        tk.Label(f,text="💰 TAD — 20 Yıllık Ot Bedeli Hesaplama",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")
        tk.Label(f,text="Formül: 20 Yıllık Ot Bedeli = TAD Alanı(da) × Üretilen Kuru Ot Verimi(kg/da) × Birim Fiyat(₺/kg) × 20 × Katsayı",
            font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(2,6))
        # Bakanlık notu
        nf=tk.Frame(f,bg="#FEF9E7",highlightbackground="#F0C040",highlightthickness=1,padx=10,pady=6); nf.pack(fill="x",pady=(0,10))
        tk.Label(nf,text="📌 Tarım ve Orman Bakanlığı Bitkisel Üretim Genel Müdürlüğü'nün 31.07.2025 tarih ve\n"
            "E-37234586-115.02-20335113 sayılı yazısı gereğince, tahsis amacı değişikliği işlemlerinde\n"
            "mevcut vasıf yerine en az bir üst vasıf esas alınmaktadır.",
            font=("Segoe UI",8),fg="#7D6608",bg="#FEF9E7",justify="left").pack(anchor="w")
        gf=tk.Frame(f,bg=bg); gf.pack(fill="x",pady=8)
        self.tad_g={}
        for i,(l,k,w) in enumerate([("TAD Alanı (da):","alan",12),("Mevcut Vasıf:","vasif",12),
            ("Kuru Ot Fiyatı (₺/kg):","fiyat",10)]):
            tk.Label(gf,text=l,bg=bg,font=("Segoe UI",10)).grid(row=0,column=i*2,padx=(0,4),sticky="w")
            if k=="vasif":
                w_obj=ttk.Combobox(gf,values=MERA_VASIF,state="readonly",width=w)
            else: w_obj=ttk.Entry(gf,width=w)
            w_obj.grid(row=0,column=i*2+1,padx=(0,12)); self.tad_g[k]=w_obj
        tk.Label(gf,text="Faaliyet Konusu:",bg=bg,font=("Segoe UI",10)).grid(row=1,column=0,padx=(0,4),pady=8,sticky="w")
        self.tad_g["faaliyet"]=ttk.Combobox(gf,values=list(TAD_KATSAYILAR.keys()),state="readonly",width=40)
        self.tad_g["faaliyet"].grid(row=1,column=1,columnspan=5,padx=(0,12),pady=8,sticky="w")
        MBtn(f,"💰 Hesapla",command=self._tad_hesapla,color=self.gc("acc"),width=16).pack(anchor="w",pady=6)
        self.lbl_tad=tk.Label(f,text="",bg=bg,font=("Segoe UI",11),fg="#333",justify="left"); self.lbl_tad.pack(anchor="w",pady=8)

    def _tad_hesapla(self):
        vasif=self.tad_g["vasif"].get(); faaliyet=self.tad_g["faaliyet"].get()
        if not vasif: messagebox.showerror("Hata","Vasıf seçin."); return
        if not faaliyet: messagebox.showerror("Hata","Faaliyet konusu seçin."); return
        sonuc=tad_ot_bedeli_hesapla(self.tad_g["alan"].get(),vasif,self.tad_g["fiyat"].get(),faaliyet)
        if not sonuc: messagebox.showerror("Hata","Tüm alanları doğru doldurun."); return
        self.lbl_tad.config(text=
            f"📊 TAD Alanı: {sonuc['alan']:,.1f} da\n"
            f"🌱 Mevcut Vasıf: {sonuc['vasif']} ({sonuc['kov']} kg/da) → Üst Vasıf: {sonuc['ust_vasif']} ({sonuc['kov_ust']} kg/da)\n"
            f"💲 Kuru Ot Fiyatı: {sonuc['fiyat']:.2f} ₺/kg\n"
            f"📋 Faaliyet: {sonuc['faaliyet']} (Katsayı: {sonuc['katsayi']})\n\n"
            f"{'─'*50}\n"
            f"📦 Yıllık Üretim: {sonuc['yillik_uretim']:,.0f} kg ({sonuc['alan']:,.1f} da × {sonuc['kov_ust']} kg/da)\n"
            f"💰 Yıllık Gelir: {para_format(sonuc['yillik_gelir'])}\n"
            f"📅 20 Yıllık Ot Bedeli: {para_format(sonuc['ot_bedeli_20'])}\n\n"
            f"🏷️ TAD ÜCRETİ (Katsayı×Ot Bedeli): {para_format(sonuc['tad_ucreti'])}")

    def _build_ter(self,parent,bg):
        """TER — Teknik Ekip Raporu formu"""
        canvas=tk.Canvas(parent,bg=bg,highlightthickness=0); canvas.pack(side="left",fill="both",expand=True)
        sb=ttk.Scrollbar(parent,orient="vertical",command=canvas.yview); sb.pack(side="right",fill="y")
        canvas.configure(yscrollcommand=sb.set)
        f=tk.Frame(canvas,bg=bg); canvas.create_window((0,0),window=f,anchor="nw")
        f.bind("<Configure>",lambda e:canvas.configure(scrollregion=canvas.bbox("all")))
        def _ter_scroll(e):
            try: canvas.yview_scroll(-1*(e.delta//120),"units")
            except Exception: pass
        canvas.bind("<MouseWheel>",_ter_scroll)
        f.bind("<MouseWheel>",_ter_scroll)
        # Alt widget'lara da bind et (canvas içi scroll)
        def _bind_scroll_recursive(widget):
            widget.bind("<MouseWheel>",_ter_scroll)
            for child in widget.winfo_children():
                try: _bind_scroll_recursive(child)
                except Exception: pass
        f.bind("<Map>",lambda e:_bind_scroll_recursive(f))
        tk.Label(f,text="📝 TAHSİS AMACI DEĞİŞİKLİĞİ — TEKNİK EKİP RAPORU (TER)",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",padx=14,pady=(10,4))
        tk.Label(f,text="Tüm alanları doldurun. Hesaplamalar otomatik yapılır. Sonunda Word raporu oluşturulur.",
            font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",padx=14,pady=(0,10))
        self.ter={}
        # A) Köy bilgileri
        sec_a=tk.LabelFrame(f,text="  (I-A) Köy Genel Bilgileri  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=10,pady=6)
        sec_a.pack(fill="x",padx=14,pady=4)
        for i,(l,k,w) in enumerate([("Başvuru Açıklaması:","basvuru_aciklama",60),("İnceleme Tarihi:","inceleme_tarihi",14),
            ("İl:","il",14),("İlçe:","ilce",14),("Köy:","koy",18),("Mevkii:","mevkii",18),
            ("Nüfus:","nufus",8),("Hane Sayısı:","hane",8),("Aktif İşletme:","aktif_isletme",8)]):
            r=i//3; c=(i%3)*2
            tk.Label(sec_a,text=l,bg=bg,font=("Segoe UI",9)).grid(row=r,column=c,padx=(0,4),pady=4,sticky="w")
            if k=="ilce": e=ttk.Combobox(sec_a,values=ILCELER,state="readonly",width=w)
            elif k=="il": e=ttk.Entry(sec_a,width=w); e.insert(0,"Ardahan")
            elif k=="inceleme_tarihi": e=ttk.Entry(sec_a,width=w); e.insert(0,datetime.now().strftime("%d.%m.%Y"))
            elif k=="basvuru_aciklama": e=ttk.Entry(sec_a,width=w)
            else: e=ttk.Entry(sec_a,width=w)
            e.grid(row=r,column=c+1,padx=(0,12),pady=4,sticky="w"); self.ter[k]=e
        # B) Hayvan Varlığı
        sec_b=tk.LabelFrame(f,text="  (I-C) Mevcut Hayvan Varlığı (BAŞ)  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=10,pady=6)
        sec_b.pack(fill="x",padx=14,pady=4)
        self.ter_hayvan={}
        hayvan_satirlari=[("Kültür İnek","k_inek",1.0),("Melez İnek","m_inek",0.75),("Yerli İnek","y_inek",0.5),
            ("Kültür Dana-Düve","k_dana",0.6),("Melez Dana-Düve","m_dana",0.45),("Yerli Dana-Düve","y_dana",0.3),
            ("Boğa","boga",1.5),("Öküz","okuz",0.6),("Koyun","koyun",0.1),("Keçi","keci",0.08),
            ("Kuzu-Oğlak","kuzu",0.04),("At","at",0.5),("Katır","katir",0.4),("Eşek","esek",0.3),
            ("Manda Dişi","manda_d",0.75),("Manda Erkek","manda_e",0.9)]
        for i,(ad,key,kat) in enumerate(hayvan_satirlari):
            r=i//4; c=(i%4)*3
            tk.Label(sec_b,text=f"{ad} ({kat}):",bg=bg,font=("Segoe UI",8)).grid(row=r,column=c,padx=(0,2),pady=2,sticky="w")
            e=ttk.Entry(sec_b,width=6); e.insert(0,"0"); e.grid(row=r,column=c+1,padx=(0,8),pady=2); self.ter_hayvan[key]=e
        self.lbl_ter_bbhb=tk.Label(sec_b,text="TOPLAM BBHB: —",font=("Segoe UI",11,"bold"),fg=self.gc("pri"),bg=bg)
        self.lbl_ter_bbhb.grid(row=5,column=0,columnspan=8,pady=6)
        MBtn(sec_b,"🔄 BBHB Hesapla",command=self._ter_bbhb_hesapla,color=C_INFO,width=16).grid(row=5,column=8,columnspan=4,pady=6)
        # C-D) Kaba Yem Kaynakları
        sec_c=tk.LabelFrame(f,text="  (I-D) Kaba Yem Kaynakları — Yem Bitkileri  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=10,pady=6)
        sec_c.pack(fill="x",padx=14,pady=4)
        self.ter_yem={}
        yem_turleri=[("Yonca","yonca",850),("Korunga","korunga",867),("Fiğ","fig",753),
            ("Buğday","bugday",345),("Arpa","arpa",320),("Yulaf","yulaf",792),("Çavdar","cavdar",600)]
        tk.Label(sec_c,text="Yem Bitkisi",bg=bg,font=("Segoe UI",8,"bold")).grid(row=0,column=0,padx=4)
        tk.Label(sec_c,text="Ekiliş(da)",bg=bg,font=("Segoe UI",8,"bold")).grid(row=0,column=1,padx=4)
        tk.Label(sec_c,text="Verim(kg/da)",bg=bg,font=("Segoe UI",8,"bold")).grid(row=0,column=2,padx=4)
        tk.Label(sec_c,text="Üretim(ton)",bg=bg,font=("Segoe UI",8,"bold")).grid(row=0,column=3,padx=4)
        for i,(ad,key,verim) in enumerate(yem_turleri):
            tk.Label(sec_c,text=ad+":",bg=bg,font=("Segoe UI",8)).grid(row=i+1,column=0,padx=4,pady=1,sticky="w")
            ek=ttk.Entry(sec_c,width=8); ek.insert(0,"0"); ek.grid(row=i+1,column=1,padx=2,pady=1); self.ter_yem[key+"_ek"]=ek
            vr=ttk.Entry(sec_c,width=8); vr.insert(0,str(verim)); vr.grid(row=i+1,column=2,padx=2,pady=1); self.ter_yem[key+"_vr"]=vr
            ur=tk.Label(sec_c,text="—",bg=bg,font=("Segoe UI",8),width=8); ur.grid(row=i+1,column=3,padx=2,pady=1); self.ter_yem[key+"_ur"]=ur
        # Çayır-Mera Alanları
        tk.Label(sec_c,text="── Çayır-Mera Alanları ──",bg=bg,font=("Segoe UI",8,"bold")).grid(row=len(yem_turleri)+1,column=0,columnspan=4,pady=(6,2))
        for i,(l,k) in enumerate([("Toplam Mera (da):","mera_toplam"),("İlin Yağış Kuşağı:","yagis")]):
            tk.Label(sec_c,text=l,bg=bg,font=("Segoe UI",8)).grid(row=len(yem_turleri)+2,column=i*2,padx=4,pady=2,sticky="w")
            e=ttk.Entry(sec_c,width=10)
            if k=="yagis": e.insert(0,"500-650 mm")
            e.grid(row=len(yem_turleri)+2,column=i*2+1,padx=2,pady=2); self.ter[k]=e
        self.lbl_ter_yem=tk.Label(sec_c,text="",bg=bg,font=("Segoe UI",9,"bold"),fg=self.gc("pri"))
        self.lbl_ter_yem.grid(row=len(yem_turleri)+3,column=0,columnspan=4,pady=4)
        MBtn(sec_c,"🔄 Yem Üretimi Hesapla",command=self._ter_yem_hesapla,color=C_INFO,width=20).grid(row=len(yem_turleri)+3,column=4,pady=4)
        # E-L) Kaba Yem İhtiyacı
        sec_el=tk.LabelFrame(f,text="  (I-E/L) Kaba Yem İhtiyacı & Karşılama  ",bg=bg,font=("Segoe UI",10,"bold"),fg=C_WARN,padx=10,pady=6)
        sec_el.pack(fill="x",padx=14,pady=4)
        self.lbl_ter_kaba=tk.Label(sec_el,text="↑ Önce BBHB ve Yem Üretimi hesaplayın",bg=bg,font=("Segoe UI",10),fg="#888")
        self.lbl_ter_kaba.pack(anchor="w")
        MBtn(sec_el,"📊 Kaba Yem Analizi",command=self._ter_kaba_hesapla,color=C_WARN,width=20).pack(anchor="w",pady=4)
        # M-O) Metin alanları
        sec_m=tk.LabelFrame(f,text="  (I-M/N/O) Değerlendirme  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=10,pady=6)
        sec_m.pack(fill="x",padx=14,pady=4)
        for l,k in [("M) Kaba yem açığı nasıl karşılanacak:","yem_acigi_cozum"),
            ("N) Yatırımın çevreye ve meraya etkisi:","cevre_etki"),("O) Alternatif alan:","alternatif")]:
            tk.Label(sec_m,text=l,bg=bg,font=("Segoe UI",8,"bold")).pack(anchor="w")
            e=ttk.Entry(sec_m,width=80); e.pack(fill="x",pady=2); self.ter[k]=e
        # P) Çiftçi görüşleri
        sec_p=tk.LabelFrame(f,text="  (I-P) Çiftçi Hane Reisi Görüşleri  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=10,pady=6)
        sec_p.pack(fill="x",padx=14,pady=4)
        pf=tk.Frame(sec_p,bg=bg); pf.pack(fill="x")
        for i,(l,k) in enumerate([("Olumlu:","olumlu"),("Olumsuz:","olumsuz"),("Bildirmeyen:","bildirmeyen"),("Ulaşılamayan:","ulasilamayan")]):
            tk.Label(pf,text=l,bg=bg,font=("Segoe UI",8)).grid(row=0,column=i*2,padx=4,pady=2,sticky="w")
            e=ttk.Entry(pf,width=6); e.insert(0,"0"); e.grid(row=0,column=i*2+1,padx=2,pady=2); self.ter[k]=e
        # D) Mera Parselleri
        sec_d=tk.LabelFrame(f,text="  (II) TAD İstenen Mera Parselleri  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=10,pady=6)
        sec_d.pack(fill="x",padx=14,pady=4)
        for i,(l,k,w) in enumerate([("Ada/Parsel:","ada_parsel",14),("Toplam Yüzölçümü (da):","yuzolcumu",12),
            ("TAD İstenen Alan (da):","tad_alan",12),("Mera Sınıfı:","mera_sinifi",12),
            ("Arazi Sınıfı:","arazi_sinifi",12)]):
            tk.Label(sec_d,text=l,bg=bg,font=("Segoe UI",9)).grid(row=0,column=i*2,padx=(0,4),pady=4,sticky="w")
            if k=="mera_sinifi": e=ttk.Combobox(sec_d,values=MERA_VASIF,state="readonly",width=w)
            elif k=="arazi_sinifi": e=ttk.Combobox(sec_d,values=["I","II","III","IV","V","VI","VII"],state="readonly",width=w)
            else: e=ttk.Entry(sec_d,width=w)
            e.grid(row=0,column=i*2+1,padx=(0,8),pady=4,sticky="w"); self.ter[k]=e
        # Ot Bedeli
        sec_e=tk.LabelFrame(f,text="  Ot Bedeli Hesaplama  ",bg=bg,font=("Segoe UI",10,"bold"),fg=C_DANGER,padx=10,pady=6)
        sec_e.pack(fill="x",padx=14,pady=4)
        for i,(l,k,w) in enumerate([("Kuru Ot Fiyatı (₺/kg):","ot_fiyat",10),("Faaliyet:","faaliyet",36)]):
            tk.Label(sec_e,text=l,bg=bg,font=("Segoe UI",9)).grid(row=0,column=i*2,padx=(0,4),pady=4,sticky="w")
            if k=="faaliyet": e=ttk.Combobox(sec_e,values=list(TAD_KATSAYILAR.keys()),state="readonly",width=w)
            else: e=ttk.Entry(sec_e,width=w)
            e.grid(row=0,column=i*2+1,padx=(0,8),pady=4,sticky="w"); self.ter[k]=e
        self.lbl_ter_ot=tk.Label(sec_e,text="",font=("Segoe UI",10,"bold"),fg=C_DANGER,bg=bg)
        self.lbl_ter_ot.grid(row=1,column=0,columnspan=6,pady=4)
        MBtn(sec_e,"💰 Ot Bedeli Hesapla",command=self._ter_ot_hesapla,color=C_DANGER,width=20).grid(row=1,column=6,pady=4)
        # Metin alanları
        sec_f=tk.LabelFrame(f,text="  (III) Sonuç & İmzacılar  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=10,pady=6)
        sec_f.pack(fill="x",padx=14,pady=4)
        tk.Label(sec_f,text="Sonuç/Kanaat Metni:",bg=bg,font=("Segoe UI",9)).pack(anchor="w")
        self.ter["sonuc"]=tk.Text(sec_f,height=4,width=80,font=("Segoe UI",9),wrap="word",relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=6,pady=4)
        self.ter["sonuc"].pack(fill="x",pady=4)
        tk.Label(sec_f,text="İmzacılar (her raporda farklı ekip olabilir):",bg=bg,font=("Segoe UI",9,"bold")).pack(anchor="w",pady=(8,4))
        self.ter_imza={}
        imza_fr=tk.Frame(sec_f,bg=bg); imza_fr.pack(fill="x")
        for i,(l,k) in enumerate([("Teknik Ekip Başkanı","baskan"),("Üye 1","uye1"),("Üye 2","uye2"),("Üye 3","uye3")]):
            tk.Label(imza_fr,text=l+":",bg=bg,font=("Segoe UI",8)).grid(row=0,column=i*2,padx=(0,2),pady=2,sticky="w")
            e=ttk.Entry(imza_fr,width=18); e.grid(row=0,column=i*2+1,padx=(0,8),pady=2); self.ter_imza[k+"_ad"]=e
            tk.Label(imza_fr,text="Unvan:",bg=bg,font=("Segoe UI",8)).grid(row=1,column=i*2,padx=(0,2),pady=2,sticky="w")
            e2=ttk.Entry(imza_fr,width=18); e2.grid(row=1,column=i*2+1,padx=(0,8),pady=2); self.ter_imza[k+"_unvan"]=e2
            tk.Label(imza_fr,text="Kurum:",bg=bg,font=("Segoe UI",8)).grid(row=2,column=i*2,padx=(0,2),pady=2,sticky="w")
            e3=ttk.Entry(imza_fr,width=18); e3.grid(row=2,column=i*2+1,padx=(0,8),pady=2); self.ter_imza[k+"_kurum"]=e3
        # Butonlar
        bf3=tk.Frame(f,bg=bg); bf3.pack(fill="x",padx=14,pady=10)
        MBtn(bf3,"📄 Word TER Raporu Oluştur",command=self._ter_word_uret,color=self.gc("pri"),width=28,
            tooltip="Tüm alanları doldurup BBHB ve yem hesapladıktan sonra Word raporu oluşturur").pack(side="left",padx=6)
        MBtn(bf3,"📊 Excel Export",command=self._ter_excel_uret,color="#8E44AD",width=14,
            tooltip="TER verilerini 3 sayfalık Excel dosyası olarak dışa aktarır").pack(side="left",padx=6)
        MBtn(bf3,"🔄 BBHB Hesapla",command=self._ter_bbhb_hesapla,color=C_INFO,width=16).pack(side="left",padx=6)

    def _ter_bbhb_hesapla(self):
        """TER — BBHB hesapla"""
        hayvan_kat={"k_inek":1.0,"m_inek":0.75,"y_inek":0.5,"k_dana":0.6,"m_dana":0.45,"y_dana":0.3,
            "boga":1.5,"okuz":0.6,"koyun":0.1,"keci":0.08,"kuzu":0.04,"at":0.5,"katir":0.4,"esek":0.3,
            "manda_d":0.75,"manda_e":0.9}
        toplam=0
        for k,kat in hayvan_kat.items():
            try: toplam+=int(self.ter_hayvan[k].get() or 0)*kat
            except Exception: pass
        self.lbl_ter_bbhb.config(text=f"TOPLAM BBHB: {toplam:.2f}")
        self._ter_bbhb=toplam

    def _ter_yem_hesapla(self):
        """TER — Yem bitkileri üretim hesapla"""
        yem_turleri=["yonca","korunga","fig","bugday","arpa","yulaf","cavdar"]
        toplam_uretim=0
        for key in yem_turleri:
            try:
                ek=float(self.ter_yem[key+"_ek"].get() or 0)
                vr=float(self.ter_yem[key+"_vr"].get() or 0)
                ur=(ek*vr)/1000  # ton
                self.ter_yem[key+"_ur"].config(text=f"{ur:.1f}")
                toplam_uretim+=ur
            except Exception: self.ter_yem[key+"_ur"].config(text="—")
        self._ter_yem_toplam=toplam_uretim
        self.lbl_ter_yem.config(text=f"TOPLAM YEM BİTKİSİ ÜRETİMİ: {toplam_uretim:,.1f} Ton/Yıl")

    def _ter_kaba_hesapla(self):
        """TER — Kaba yem ihtiyacı ve karşılama oranı"""
        bbhb=getattr(self,"_ter_bbhb",0)
        yem_toplam=getattr(self,"_ter_yem_toplam",0)
        if bbhb<=0:
            messagebox.showwarning("Uyarı","Önce BBHB hesaplayın."); return
        # F = BBHB × 500 × 0.1 × 365 / 1000 (ton/yıl)
        ihtiyac = bbhb * 500 * 0.1 * 365 / 1000
        # Mera üretimi (basit tahmin — mera alanı × verim)
        try: mera_alan=float(self.ter.get("mera_toplam",ttk.Entry()).get() or 0)
        except Exception: mera_alan=0
        vasif=self.ter.get("mera_sinifi",ttk.Combobox())
        v=vasif.get() if hasattr(vasif,'get') else ""
        mera_verim={"Zayıf":90,"Orta":180,"İyi":270,"Çok İyi":360}.get(v,180)
        mera_uretim=(mera_alan*mera_verim)/1000
        toplam_uretim=yem_toplam+mera_uretim
        # TAD istenen parsel üretimi
        try: tad_alan=float(self.ter.get("tad_alan",ttk.Entry()).get() or 0)
        except Exception: tad_alan=0
        tad_uretim=(tad_alan*mera_verim)/1000
        tad_sonrasi=toplam_uretim-tad_uretim
        karsilama_once=(toplam_uretim/ihtiyac*100) if ihtiyac>0 else 0
        karsilama_sonra=(tad_sonrasi/ihtiyac*100) if ihtiyac>0 else 0
        acik_once=toplam_uretim-ihtiyac
        acik_sonra=tad_sonrasi-ihtiyac
        self.lbl_ter_kaba.config(text=
            f"E) Toplam Kaba Yem: {toplam_uretim:,.1f} ton (Yem Bitkileri: {yem_toplam:,.1f} + Çayır-Mera: {mera_uretim:,.1f})\n"
            f"F) Kaba Yem İhtiyacı: {ihtiyac:,.1f} ton (BBHB: {bbhb:.1f} × 500 × 0.1 × 365 / 1000)\n"
            f"G) TAD Öncesi Karşılama: %{karsilama_once:.1f} | Açık: {acik_once:,.1f} ton\n"
            f"I) TAD İstenen Parsellerin Üretimi: {tad_uretim:,.1f} ton\n"
            f"K) TAD Sonrası Karşılama: %{karsilama_sonra:.1f} | Açık: {acik_sonra:,.1f} ton")
        self._ter_kaba_veri={"ihtiyac":ihtiyac,"toplam":toplam_uretim,"mera":mera_uretim,
            "yem":yem_toplam,"tad_uretim":tad_uretim,"karsilama_once":karsilama_once,
            "karsilama_sonra":karsilama_sonra,"acik_once":acik_once,"acik_sonra":acik_sonra}

    def _ter_ot_hesapla(self):
        """TER — Ot bedeli hesapla"""
        try:
            vasif=self.ter["mera_sinifi"].get()
            alan=float(self.ter["tad_alan"].get() or 0)
            fiyat=para_parse(self.ter["ot_fiyat"].get() or "0")
            faaliyet=self.ter["faaliyet"].get()
            sonuc=tad_ot_bedeli_hesapla(alan,vasif,fiyat,faaliyet)
            if sonuc:
                self.lbl_ter_ot.config(text=f"Üst Vasıf: {sonuc['ust_vasif']} ({sonuc['kov_ust']} kg/da) | "
                    f"20 Yıl: {para_format(sonuc['ot_bedeli_20'])} | Katsayı: {sonuc['katsayi']} | "
                    f"TAD Ücreti: {para_format(sonuc['tad_ucreti'])}")
                self._ter_ot_sonuc=sonuc
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _ter_excel_uret(self):
        """TER — Excel export"""
        dosya=filedialog.asksaveasfilename(defaultextension=".xlsx",initialfile=f"TER_{self.ter['koy'].get()}.xlsx",filetypes=[("Excel","*.xlsx")])
        if not dosya: return
        try:
            # Hayvan varlığı
            hayvan_kat={"k_inek":("Kültür İnek",1.0),"m_inek":("Melez İnek",0.75),"y_inek":("Yerli İnek",0.5),
                "k_dana":("Kültür Dana-Düve",0.6),"m_dana":("Melez Dana-Düve",0.45),"y_dana":("Yerli Dana-Düve",0.3),
                "boga":("Boğa",1.5),"okuz":("Öküz",0.6),"koyun":("Koyun",0.1),"keci":("Keçi",0.08),
                "kuzu":("Kuzu-Oğlak",0.04),"at":("At",0.5),"katir":("Katır",0.4),"esek":("Eşek",0.3),
                "manda_d":("Manda Dişi",0.75),"manda_e":("Manda Erkek",0.9)}
            hayvan_rows=[]
            for key,(ad,kat) in hayvan_kat.items():
                bas=int(self.ter_hayvan.get(key,ttk.Entry()).get() or 0)
                hayvan_rows.append({"Tür":ad,"Katsayı":kat,"Baş":bas,"BBHB":bas*kat})
            df_hay=pd.DataFrame(hayvan_rows)
            # Yem kaynakları
            yem_keys=["yonca","korunga","fig","bugday","arpa","yulaf","cavdar"]
            yem_adlar=["Yonca","Korunga","Fiğ","Buğday","Arpa","Yulaf","Çavdar"]
            yem_rows=[]
            for ad,key in zip(yem_adlar,yem_keys):
                ek=float(self.ter_yem.get(key+"_ek",ttk.Entry()).get() or 0)
                vr=float(self.ter_yem.get(key+"_vr",ttk.Entry()).get() or 0)
                yem_rows.append({"Yem Bitkisi":ad,"Ekiliş(da)":ek,"Verim(kg/da)":vr,"Üretim(ton)":(ek*vr)/1000})
            df_yem=pd.DataFrame(yem_rows)
            # Genel bilgi
            genel=[{"Bilgi":"İl","Değer":self.ter.get("il",ttk.Entry()).get()},
                {"Bilgi":"İlçe","Değer":self.ter.get("ilce",ttk.Combobox()).get()},
                {"Bilgi":"Köy","Değer":self.ter.get("koy",ttk.Entry()).get()},
                {"Bilgi":"Mevkii","Değer":self.ter.get("mevkii",ttk.Entry()).get()},
                {"Bilgi":"Aktif İşletme","Değer":self.ter.get("aktif_isletme",ttk.Entry()).get()},
                {"Bilgi":"Ada/Parsel","Değer":self.ter.get("ada_parsel",ttk.Entry()).get()},
                {"Bilgi":"TAD Alan(da)","Değer":self.ter.get("tad_alan",ttk.Entry()).get()},
                {"Bilgi":"Mera Sınıfı","Değer":self.ter.get("mera_sinifi",ttk.Combobox()).get()},
                {"Bilgi":"Toplam BBHB","Değer":f"{getattr(self,'_ter_bbhb',0):.2f}"}]
            kaba=getattr(self,"_ter_kaba_veri",None)
            if kaba:
                genel.extend([
                    {"Bilgi":"Kaba Yem İhtiyacı(ton)","Değer":f"{kaba['ihtiyac']:,.1f}"},
                    {"Bilgi":"Toplam Üretim(ton)","Değer":f"{kaba['toplam']:,.1f}"},
                    {"Bilgi":"TAD Öncesi Karşılama %","Değer":f"{kaba['karsilama_once']:.1f}"},
                    {"Bilgi":"TAD Sonrası Karşılama %","Değer":f"{kaba['karsilama_sonra']:.1f}"}])
            df_gen=pd.DataFrame(genel)
            with pd.ExcelWriter(dosya) as w:
                df_gen.to_excel(w,sheet_name="Genel Bilgi",index=False)
                df_hay.to_excel(w,sheet_name="Hayvan Varlığı",index=False)
                df_yem.to_excel(w,sheet_name="Yem Kaynakları",index=False)
            messagebox.showinfo("Tamam",f"TER Excel oluşturuldu:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _ter_word_uret(self):
        """TER — Word raporu oluştur"""
        if not DOCX_OK: messagebox.showerror("Eksik","pip install python-docx"); return
        dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile=f"TER_{self.ter['koy'].get()}.docx",filetypes=[("Word","*.docx")])
        if not dosya: return
        try:
            ter_no = oto_dosya_no("TER", "Tahsisler")
            doc=DocxDocument()
            style=doc.styles['Normal']; style.font.name='Times New Roman'; style.font.size=Pt(12)
            # Başlık
            h=doc.add_paragraph(); h.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=h.add_run("TAHSİS AMACI DEĞİŞİKLİĞİ (T.A.D.) İLE İLGİLİ\nMERA İNCELEME RAPORU"); r.bold=True; r.font.size=Pt(14)
            doc.add_paragraph(f"Rapor No: {ter_no}")
            # Başvuru açıklaması
            doc.add_paragraph(self.ter["basvuru_aciklama"].get())
            doc.add_paragraph(f"İnceleme Yerine Gidiş Tarihi: {self.ter['inceleme_tarihi'].get()}")
            # Köy bilgileri tablosu
            t=doc.add_table(rows=5,cols=2,style='Table Grid')
            for i,(l,k) in enumerate([("İLİ","il"),("İLÇESİ","ilce"),("MAHALLE-KÖY","koy"),("MEVKİİ","mevkii"),("AKTİF İŞLETME","aktif_isletme")]):
                t.rows[i].cells[0].text=l
                w=self.ter[k]
                t.rows[i].cells[1].text=w.get() if isinstance(w,(ttk.Entry,ttk.Combobox)) else ""
            # Hayvan varlığı
            doc.add_paragraph("C) MEVCUT HAYVAN VARLIĞI (BAŞ)").runs[0].bold=True
            ht=doc.add_table(rows=1,cols=3,style='Table Grid')
            ht.rows[0].cells[0].text="Tür"; ht.rows[0].cells[1].text="Baş"; ht.rows[0].cells[2].text="BBHB"
            hayvan_kat={"k_inek":("Kültür İnek",1.0),"m_inek":("Melez İnek",0.75),"y_inek":("Yerli İnek",0.5),
                "k_dana":("Kültür Dana-Düve",0.6),"m_dana":("Melez Dana-Düve",0.45),"y_dana":("Yerli Dana-Düve",0.3),
                "boga":("Boğa",1.5),"okuz":("Öküz",0.6),"koyun":("Koyun",0.1),"keci":("Keçi",0.08),
                "kuzu":("Kuzu-Oğlak",0.04),"at":("At",0.5),"katir":("Katır",0.4),"esek":("Eşek",0.3),
                "manda_d":("Manda Dişi",0.75),"manda_e":("Manda Erkek",0.9)}
            toplam_bbhb=0
            for key,(ad,kat) in hayvan_kat.items():
                bas=int(self.ter_hayvan[key].get() or 0)
                if bas>0:
                    bbhb=bas*kat; toplam_bbhb+=bbhb
                    row=ht.add_row(); row.cells[0].text=ad; row.cells[1].text=str(bas); row.cells[2].text=f"{bbhb:.2f}"
            row=ht.add_row(); row.cells[0].text="TOPLAM"; row.cells[1].text=""; row.cells[2].text=f"{toplam_bbhb:.2f}"
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs: r.bold=True
            # Parsel bilgileri
            doc.add_paragraph("")
            # D) Yem Kaynakları
            doc.add_paragraph("D) KABA YEM KAYNAKLARI").runs[0].bold=True
            doc.add_paragraph("1- MEVCUT YEM BİTKİLERİ EKİLİŞ VERİM VE ÜRETİM DURUMLARI").runs[0].bold=True
            yt=doc.add_table(rows=1,cols=4,style='Table Grid')
            yt.rows[0].cells[0].text="YEM BİTKİSİ"; yt.rows[0].cells[1].text="Ekiliş(da)"
            yt.rows[0].cells[2].text="Verim(Kg/da)"; yt.rows[0].cells[3].text="Üretim(Ton)"
            for c in yt.rows[0].cells:
                for p in c.paragraphs:
                    for r in p.runs: r.bold=True; r.font.size=Pt(9)
            yem_toplam=0
            yem_keys=["yonca","korunga","fig","bugday","arpa","yulaf","cavdar"]
            yem_adlar=["Yonca(*)","Korunga(*)","Fiğ(*)","Buğday(*)","Arpa(*)","Yulaf(*)","Çavdar(*)"]
            for ad,key in zip(yem_adlar,yem_keys):
                try:
                    ek=float(self.ter_yem.get(key+"_ek",ttk.Entry()).get() or 0)
                    vr=float(self.ter_yem.get(key+"_vr",ttk.Entry()).get() or 0)
                    ur=(ek*vr)/1000
                except Exception: ek=vr=ur=0
                if ek>0 or ur>0:
                    row=yt.add_row(); row.cells[0].text=ad; row.cells[1].text=f"{ek:.0f}"
                    row.cells[2].text=f"{vr:.0f}"; row.cells[3].text=f"{ur:.3f}"
                    yem_toplam+=ur
            row=yt.add_row(); row.cells[0].text="TOPLAM"; row.cells[3].text=f"{yem_toplam:.3f}"
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs: r.bold=True
            doc.add_paragraph("*   :Verim değeri; yeşil ot olarak alınmıştır.").runs[0].font.size=Pt(8)
            # E-L Kaba Yem Analizi
            kaba=getattr(self,"_ter_kaba_veri",None)
            if kaba:
                doc.add_paragraph("")
                doc.add_paragraph(f"E) TOPLAM KABA YEM ÜRETİMİ: {kaba['toplam']:,.1f} Ton/Yıl").runs[0].bold=True
                doc.add_paragraph(f"   1-Yem Bitkileri: {kaba['yem']:,.1f} ton  |  2-Çayır-Mera: {kaba['mera']:,.1f} ton")
                doc.add_paragraph(f"F) KABA YEM İHTİYACI: {kaba['ihtiyac']:,.1f} Ton").runs[0].bold=True
                doc.add_paragraph(f"G) TAD ÖNCESİ KARŞILAMA ORANI: %{kaba['karsilama_once']:.1f}")
                doc.add_paragraph(f"H) TAD ÖNCESİ KABA YEM AÇIĞI: {kaba['acik_once']:,.1f} ton")
                doc.add_paragraph(f"I) TAD İSTENEN PARSELLERİN ÜRETİMİ: {kaba['tad_uretim']:,.1f} ton")
                doc.add_paragraph(f"K) TAD SONRASI KARŞILAMA ORANI: %{kaba['karsilama_sonra']:.1f}")
                doc.add_paragraph(f"L) TAD SONRASI KABA YEM AÇIĞI: {kaba['acik_sonra']:,.1f} ton")
            # M-O Değerlendirme
            for l,k in [("M) KABA YEM AÇIĞININ KARŞILANMASI:","yem_acigi_cozum"),
                ("N) YATIRIMIN ÇEVREYE VE MERAYA ETKİSİ:","cevre_etki"),("O) ALTERNATİF ALAN:","alternatif")]:
                val=self.ter.get(k,ttk.Entry())
                txt=val.get() if hasattr(val,'get') else ""
                if txt: doc.add_paragraph(f"{l} {txt}")
            # P) Çiftçi görüşleri
            doc.add_paragraph("")
            doc.add_paragraph("P) ÇİFTÇİ HANE REİSLERİNİN GÖRÜŞLERİ").runs[0].bold=True
            gt=doc.add_table(rows=2,cols=5,style='Table Grid')
            gt.rows[0].cells[0].text="Görüş"; gt.rows[0].cells[1].text="Olumlu"
            gt.rows[0].cells[2].text="Olumsuz"; gt.rows[0].cells[3].text="Bildirmeyen"; gt.rows[0].cells[4].text="Ulaşılamayan"
            for k,i in [("olumlu",1),("olumsuz",2),("bildirmeyen",3),("ulasilamayan",4)]:
                val=self.ter.get(k,ttk.Entry())
                gt.rows[1].cells[i].text=val.get() if hasattr(val,'get') else "0"
            gt.rows[1].cells[0].text="Çiftçi Hane Reisi Sayısı"
            doc.add_paragraph("")
            # (II) Parsel bilgileri
            doc.add_paragraph("(II) TAD İSTENEN MERA PARSELLERİ").runs[0].bold=True
            pt=doc.add_table(rows=2,cols=4,style='Table Grid')
            pt.rows[0].cells[0].text="Ada/Parsel"; pt.rows[0].cells[1].text="Yüzölçümü (da)"
            pt.rows[0].cells[2].text="TAD Alan (da)"; pt.rows[0].cells[3].text="Mera Sınıfı"
            pt.rows[1].cells[0].text=self.ter["ada_parsel"].get()
            pt.rows[1].cells[1].text=self.ter["yuzolcumu"].get()
            pt.rows[1].cells[2].text=self.ter["tad_alan"].get()
            pt.rows[1].cells[3].text=self.ter["mera_sinifi"].get()
            # Ot bedeli
            ot=getattr(self,"_ter_ot_sonuc",None)
            if ot:
                doc.add_paragraph(f"KURU OT VERİMİ: 500-650 yağış kuşağında yer alan Ardahan ilinin dekara kuru ot verimi {ot['kov_ust']} Kg/da dır.").runs[0].bold=True
                doc.add_paragraph(f"Faaliyet: {ot['faaliyet']} — Katsayı: {ot['katsayi']}")
                doc.add_paragraph(f"Kuru ot birim fiyatı: {ot['fiyat']:.2f} TL/kg")
                doc.add_paragraph(f"Yıllık Verim: {ot['yillik_uretim']:,.0f} kg")
                doc.add_paragraph(f"20 Yıllık Ot Bedeli: {para_format(ot['ot_bedeli_20'])}")
                doc.add_paragraph(f"TAD Ücreti (Katsayı ile): {para_format(ot['tad_ucreti'])}").runs[0].bold=True
            # Sonuç
            doc.add_paragraph("(III) SONUÇ, NETİCE VE KANAAT").runs[0].bold=True
            doc.add_paragraph(self.ter["sonuc"].get("1.0",tk.END).strip())
            doc.add_paragraph(f"Rapor Düzenleme Tarihi: {datetime.now().strftime('%d.%m.%Y')}")
            # İmzacılar
            doc.add_paragraph("")
            it=doc.add_table(rows=3,cols=4,style='Table Grid')
            for i,rol in enumerate(["baskan","uye1","uye2","uye3"]):
                it.rows[0].cells[i].text=self.ter_imza.get(f"{rol}_ad",ttk.Entry()).get()
                it.rows[1].cells[i].text=self.ter_imza.get(f"{rol}_unvan",ttk.Entry()).get()
                it.rows[2].cells[i].text=self.ter_imza.get(f"{rol}_kurum",ttk.Entry()).get()
            doc.save(dosya)
            db_log(self.u_id,"TER Rapor",self.ter["koy"].get())
            messagebox.showinfo("Tamam",f"TER Raporu oluşturuldu:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e)); logging.error(f"ter_word:{e}")

    def _yenile_tahsis(self):
        if not DB_PATH: return
        try:
            with db_baglan() as c:
                rows=c.execute("SELECT id,koy,ilce,ada,parsel,kurum,amac,alan_ha,asama,durum,basvuru_t FROM Tahsisler ORDER BY id DESC").fetchall()
            self.tv_tahsis.delete(*self.tv_tahsis.get_children())
            for r in rows:
                d=str(r[9]); tag="tamam" if d=="Tamamlandı" else "red" if d=="Reddedildi" else "devam"
                self.tv_tahsis.insert("","end",values=r,tags=(tag,))
        except Exception as e: logging.error(f"yenile_tahsis:{e}")

    def _tahsis_sec(self,event=None):
        sel=self.tv_tahsis.selection()
        if not sel: return
        asama=str(self.tv_tahsis.item(sel[0])["values"][8])
        mi=next((i for i,a in enumerate(TAHSIS_ASAMALARI) if a==asama),0)
        for i,l in enumerate(self._al):
            if i<mi: l.config(bg="#A9DFBF",fg="#1E5631")
            elif i==mi: l.config(bg=self.gc("pri"),fg=C_WHITE)
            else: l.config(bg="#E0E0E0",fg="#555")

    def _tahsis_kaydet(self):
        try:
            koy=self.tah_g["koy"].get().strip()
            if not koy: raise ValueError("Köy adı zorunlu.")
            alan=float(self.tah_g["alan"].get() or 0)
            ot_geliri=float(self.tah_g["ot_geliri"].get() or 0)
            notlar=self.tah_g["notlar"].get("1.0",tk.END).strip()
            md14=self.tah_g["madde14"].get()
            with db_baglan() as conn:
                cur=conn.cursor()
                sure_bitis=self.tah_g["sure_bitis"].get().strip() if "sure_bitis" in self.tah_g else ""
                sure_tipi=self.tah_g["sure_tipi"].get() if "sure_tipi" in self.tah_g else ""
                cur.execute("INSERT INTO Tahsisler(koy,ilce,ada,parsel,kurum,amac,alan_ha,asama,durum,basvuru_t,notlar,madde14_bent,ot_geliri,sure_bitis,sure_tipi)VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (koy,self.tah_g["ilce"].get(),self.tah_g["ada"].get(),self.tah_g["parsel"].get(),
                     self.tah_g["kurum"].get(),self.tah_g["amac"].get(),alan,TAHSIS_ASAMALARI[0],"Devam Ediyor",
                     datetime.now().strftime("%d.%m.%Y"),notlar,md14 if md14!="—" else "",ot_geliri,sure_bitis,sure_tipi))
                tid=cur.lastrowid
                cur.execute("INSERT INTO Tahsis_Log(tahsis_id,tarih,personel,asama,aciklama)VALUES(?,?,?,?,?)",
                    (tid,datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad,TAHSIS_ASAMALARI[0],"Başvuru alındı."))
            db_log(self.u_id,"Tahsis Başvuru",koy); self._son_islem_kaydet("Tahsis",koy)
            for k,w in self.tah_g.items():
                if isinstance(w,ttk.Entry): w.delete(0,tk.END)
                elif isinstance(w,ttk.Combobox): w.set("")
                elif isinstance(w,tk.Text): w.delete("1.0",tk.END)
            self._yenile_tahsis(); messagebox.showinfo("Tamam",f"'{koy}' tahsis başvurusu kaydedildi.")
        except ValueError as e: messagebox.showerror("Hata",str(e))

    def _tahsis_ileri(self):
        sel=self.tv_tahsis.selection()
        if not sel: messagebox.showwarning("Seçim","Tahsis seçin."); return
        vals=self.tv_tahsis.item(sel[0])["values"]; tid=vals[0]; koy=vals[1]; asama=str(vals[8])
        if asama==TAHSIS_ASAMALARI[-1]: messagebox.showinfo("Bilgi","Son aşamada."); return
        mi=next((i for i,a in enumerate(TAHSIS_ASAMALARI) if a==asama),0)
        sonraki=TAHSIS_ASAMALARI[mi+1]
        acik=tkinter.simpledialog.askstring("Not",f"'{koy}' → '{sonraki}'\nAçıklama:") or ""
        with db_baglan() as c:
            c.execute("UPDATE Tahsisler SET asama=? WHERE id=?",(sonraki,tid))
            if sonraki==TAHSIS_ASAMALARI[-1]:
                c.execute("UPDATE Tahsisler SET durum='Tamamlandı',sonuc_t=? WHERE id=?",(datetime.now().strftime("%d.%m.%Y"),tid))
            c.execute("INSERT INTO Tahsis_Log(tahsis_id,tarih,personel,asama,aciklama)VALUES(?,?,?,?,?)",
                (tid,datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad,sonraki,acik or "—"))
        db_log(self.u_id,"Tahsis Aşama",f"ID:{tid}→{sonraki}"); self._yenile_tahsis()

    def _tahsis_red(self):
        sel=self.tv_tahsis.selection()
        if not sel: return
        tid=self.tv_tahsis.item(sel[0])["values"][0]; koy=self.tv_tahsis.item(sel[0])["values"][1]
        n=tkinter.simpledialog.askstring("Red",f"'{koy}' red/kapatma nedeni:")
        if n is not None:
            with db_baglan() as c:
                c.execute("UPDATE Tahsisler SET durum='Reddedildi',sonuc_t=? WHERE id=?",(datetime.now().strftime("%d.%m.%Y"),tid))
                c.execute("INSERT INTO Tahsis_Log(tahsis_id,tarih,personel,asama,aciklama)VALUES(?,?,?,?,?)",
                    (tid,datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad,"REDDEDİLDİ",n))
            db_log(self.u_id,"Tahsis Red",f"ID:{tid}"); self._yenile_tahsis()

    def _tahsis_excel(self):
        """Tahsisleri Excel'e aktar"""
        try:
            with db_baglan() as c:
                rows=c.execute("SELECT id,koy,ilce,tahsis_yapilan,madde14_bent,ot_geliri,sure_bitis,sure_tipi,asama,tarih FROM Tahsisler ORDER BY id DESC").fetchall()
            if not rows: messagebox.showwarning("Veri Yok","Kayıt yok."); return
            yol=filedialog.asksaveasfilename(defaultextension=".xlsx",initialfile="Tahsisler.xlsx",filetypes=[("Excel","*.xlsx")])
            if yol:
                pd.DataFrame(rows,columns=["ID","Köy","İlçe","Tahsis Yapılan","Madde 14 Bendi","Ot Geliri","Süre Bitiş","Süre Tipi","Aşama","Tarih"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"Tahsisler dışa aktarıldı:\n{yol}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _tahsis_gecmis(self):
        sel=self.tv_tahsis.selection()
        if not sel: return
        tid=self.tv_tahsis.item(sel[0])["values"][0]; koy=self.tv_tahsis.item(sel[0])["values"][1]
        win=tk.Toplevel(self.root); win.title(f"Süreç — {koy}"); win.geometry("640x360"); win.configure(bg=C_WHITE)
        tk.Label(win,text=f"Tahsis Süreç Geçmişi: {koy}",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=12)
        tv=ttk.Treeview(win,columns=("t","p","a","n"),show="headings",height=14)
        for col,w,bas in [("t",140,"Tarih"),("p",130,"Personel"),("a",150,"Aşama"),("n",280,"Açıklama")]:
            tv.heading(col,text=bas); tv.column(col,width=w)
        tv.pack(fill="both",expand=True,padx=12,pady=8)
        try:
            with db_baglan() as c:
                for r in c.execute("SELECT tarih,personel,asama,aciklama FROM Tahsis_Log WHERE tahsis_id=? ORDER BY id",(tid,)).fetchall():
                    tv.insert("","end",values=r)
        except Exception as e: logging.error(f"tahsis_gecmis:{e}")

    # ═══ 6. ŞİKAYET TAKİP ═══════════════════════════════════════════════════
    def _t_sikayet(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2)
        nb2.add(lf,text="  📋 Şikayet Listesi  "); nb2.add(nf,text="  ➕ Yeni Şikayet  ")
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,self._filtre_sik,bg=bg).pack(anchor="w",pady=(0,6))
        df2=tk.Frame(ll,bg=bg); df2.pack(fill="x",pady=(0,6))
        tk.Label(df2,text="Durum:",bg=bg,font=("Segoe UI",10)).pack(side="left",padx=(0,6))
        self.sik_df=ttk.Combobox(df2,values=["Tümü","Yeni","İncelemede","Sonuçlandı","Kapatıldı"],state="readonly",width=16)
        self.sik_df.set("Tümü"); self.sik_df.bind("<<ComboboxSelected>>",lambda e:self._filtre_sik(""))
        self.sik_df.pack(side="left")
        self.tv_sik=self._tv(ll,[("id",40,"#"),("koy",110,"Köy"),("ilce",70,"İlçe"),
            ("sikayet_eden",130,"Şikayet Eden"),("tur",120,"Tür"),("tarih",90,"Tarih"),("durum",110,"Durum")],14)
        self.tv_sik.tag_configure("yeni",background="#FEF9E7")
        self.tv_sik.tag_configure("sonuclandi",background="#EAFAF1")
        bf=tk.Frame(ll,bg=bg); bf.pack(pady=8)
        MBtn(bf,"🔄 Yenile",command=self._yenile_sik,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(bf,"✏️ Güncelle",command=self._sik_guncelle,color=C_WARN,width=14).pack(side="left",padx=4)
        MBtn(bf,"📷 Fotoğraflar",command=self._sik_foto,color="#16A085",width=14).pack(side="left",padx=4)
        MBtn(bf,"📊 Excel Export",command=self._sik_excel,color="#8E44AD",width=14).pack(side="left",padx=4)
        MBtn(bf,"🗑 Sil",command=self._sik_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        # Yeni
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=16)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.46,anchor="center",width=620,height=480)
        tk.Label(card,text="Yeni Şikayet / İhbar",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=4,pady=(18,14))
        self.sik_g={}
        for lbl,key,tip,r,col in [("Köy:","koy","entry",1,0),("İlçe:","ilce","combo",1,2),
            ("Şikayet Eden:","sikayet_eden","entry",2,0),("Tarih:","tarih","entry",2,2),
            ("Tür:","tur","siktip",3,0)]:
            tk.Label(card,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).grid(row=r,column=col,padx=(20,6),pady=8,sticky="e")
            if tip=="combo": w_obj=ttk.Combobox(card,values=ILCELER,state="readonly",width=18)
            elif tip=="siktip": w_obj=ttk.Combobox(card,state="readonly",width=28,values=["Kaçak Otlatma","Mera İşgali","İzinsiz Yapılaşma","Mera Tahrip","Hayvan Zararı","Diğer"])
            else: w_obj=ttk.Entry(card,width=20)
            if key=="tarih" and isinstance(w_obj,ttk.Entry): w_obj.insert(0,datetime.now().strftime("%d.%m.%Y"))
            w_obj.grid(row=r,column=col+1,padx=(0,14),pady=8,sticky="w"); self.sik_g[key]=w_obj
        tk.Label(card,text="Açıklama:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=4,column=0,padx=(20,6),pady=8,sticky="ne")
        self.sik_g["aciklama"]=tk.Text(card,width=46,height=5,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        self.sik_g["aciklama"].grid(row=4,column=1,columnspan=3,padx=(0,20),pady=8,sticky="w")
        MBtn(card,"✅  Kaydet",command=self._sik_kaydet,color=self.gc("acc"),width=34).grid(row=5,column=0,columnspan=4,pady=18)
        self._yenile_sik()

    def _yenile_sik(self):
        if not DB_PATH: return
        try:
            with db_baglan() as c:
                self._all_sik=c.execute("SELECT id,koy,ilce,sikayet_eden,tur,tarih,durum FROM Sikayetler ORDER BY id DESC").fetchall()
            self._filtre_sik("")
        except Exception as e: logging.error(f"yenile_sik:{e}")

    def _filtre_sik(self,a):
        self.tv_sik.delete(*self.tv_sik.get_children()); a=a.lower()
        filtre=self.sik_df.get()
        for r in getattr(self,"_all_sik",[]):
            if filtre!="Tümü" and r[6]!=filtre: continue
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            tag="sonuclandi" if r[6] in ("Sonuçlandı","Kapatıldı") else "yeni"
            self.tv_sik.insert("","end",values=r,tags=(tag,))

    def _sik_kaydet(self):
        try:
            koy=self.sik_g["koy"].get().strip()
            if not koy: raise ValueError("Köy adı zorunlu.")
            acik=self.sik_g["aciklama"].get("1.0",tk.END).strip()
            with db_baglan() as c:
                c.execute("INSERT INTO Sikayetler(koy,ilce,sikayet_eden,tur,aciklama,durum,tarih)VALUES(?,?,?,?,?,?,?)",
                    (koy,self.sik_g["ilce"].get(),self.sik_g["sikayet_eden"].get(),self.sik_g["tur"].get(),acik,"Yeni",self.sik_g["tarih"].get()))
            db_log(self.u_id,"Şikayet Kayıt",koy); self._son_islem_kaydet("Şikayet",koy)
            for k,w in self.sik_g.items():
                if isinstance(w,ttk.Entry): w.delete(0,tk.END)
                elif isinstance(w,ttk.Combobox): w.set("")
                elif isinstance(w,tk.Text): w.delete("1.0",tk.END)
            self.sik_g["tarih"].insert(0,datetime.now().strftime("%d.%m.%Y"))
            self._yenile_sik(); messagebox.showinfo("Tamam",f"'{koy}' şikayet kaydedildi.")
        except ValueError as e: messagebox.showerror("Hata",str(e))

    def _sik_guncelle(self):
        sel=self.tv_sik.selection()
        if not sel: messagebox.showwarning("Seçim","Şikayet seçin."); return
        sid=self.tv_sik.item(sel[0])["values"][0]
        win=tk.Toplevel(self.root); win.title("Güncelle"); win.geometry("420x260"); win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        tk.Label(win,text="Yeni Durum:",bg=C_WHITE,font=("Segoe UI",11,"bold")).pack(pady=(16,4))
        cb=ttk.Combobox(win,values=["Yeni","İncelemede","Sonuçlandı","Kapatıldı"],state="readonly",width=28); cb.pack(padx=20)
        tk.Label(win,text="Sonuç:",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=20,pady=(10,2))
        txt=tk.Text(win,height=3,width=44,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6); txt.pack(padx=20)
        def _k():
            yeni=cb.get(); sonuc=txt.get("1.0",tk.END).strip()
            if not yeni: messagebox.showwarning("Uyarı","Durum seçin.",parent=win); return
            kapanma=datetime.now().strftime("%d.%m.%Y") if yeni in ["Sonuçlandı","Kapatıldı"] else None
            with db_baglan() as c:
                c.execute("UPDATE Sikayetler SET durum=?,sonuc=?,kapanma=? WHERE id=?",(yeni,sonuc,kapanma,sid))
            db_log(self.u_id,"Şikayet Güncelle",f"ID:{sid}→{yeni}"); self._yenile_sik(); win.destroy()
        MBtn(win,"Güncelle",command=_k,width=18).pack(pady=12)

    def _sik_sil(self):
        sel=self.tv_sik.selection()
        if not sel: return
        sid=self.tv_sik.item(sel[0])["values"][0]
        if messagebox.askyesno("Onay","Silinsin mi?"):
            with db_baglan() as c: c.execute("DELETE FROM Sikayetler WHERE id=?",(sid,))
            self._yenile_sik()

    # ═══ 7. İDARİ PARA CEZASI ════════════════════════════════════════════════
    def _t_ceza(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2); mf=ttk.Frame(nb2)
        nb2.add(lf,text="  📋 Ceza Listesi  "); nb2.add(nf,text="  ➕ Yeni Ceza  "); nb2.add(mf,text="  🔁 Mükerrer Cezalar  ")
        # Liste
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,self._filtre_ceza,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_ceza=self._tv(ll,[("id",40,"#"),("ad",130,"Ad Soyad"),("tc",100,"TC No"),
            ("ilce",70,"İlçe"),("mera",120,"Mera"),("yil",50,"Yıl"),("konu",180,"Konu"),
            ("tutar",90,"Tutar (₺)"),("tarih",90,"Tarih")],14)
        self.tv_ceza.tag_configure("mukerrer",background="#FDEDEC")
        bf=tk.Frame(ll,bg=bg); bf.pack(pady=8)
        MBtn(bf,"🔄 Yenile",command=self._yenile_ceza,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(bf,"📄 Evrak Üret",command=self._ceza_evrak,color=self.gc("acc"),width=14).pack(side="left",padx=4)
        MBtn(bf,"📷 Fotoğraflar",command=self._ceza_foto,color="#16A085",width=14).pack(side="left",padx=4)
        MBtn(bf,"📊 Excel Export",command=self._ceza_excel,color="#8E44AD",width=14).pack(side="left",padx=4)
        MBtn(bf,"🗑 Sil",command=self._ceza_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        # Yeni ceza
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=16)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.45,anchor="center",width=660,height=540)
        tk.Label(card,text="Yeni İdari Para Cezası Kaydı",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=4,pady=(18,14))
        self.cz_g={}
        for lbl,key,tip,r,col in [("Ad Soyad:","ad","entry",1,0),("TC No:","tc","entry",1,2),
            ("İlçe:","ilce","combo",2,0),("Mera Köy:","mera_koy","entry",2,2),
            ("Ada/Parsel:","ada_parsel","entry",3,0),("Yıl:","yil","entry",3,2),
            ("Konu:","konu","ceza_tip",4,0),("Hayvan Sayısı:","hayvan","entry",4,2),
            ("Ceza Miktarı (₺/baş):","birim","entry",5,0),("İPC Tutarı (₺):","tutar","entry",5,2)]:
            tk.Label(card,text=lbl,bg=C_WHITE,font=("Segoe UI",10)).grid(row=r,column=col,padx=(20,6),pady=8,sticky="e")
            if tip=="combo": w_obj=ttk.Combobox(card,values=ILCELER,state="readonly",width=18)
            elif tip=="ceza_tip": w_obj=ttk.Combobox(card,values=CEZA_TURLERI,state="readonly",width=30)
            else: w_obj=ttk.Entry(card,width=20)
            if key=="yil": w_obj.insert(0,str(datetime.now().year))
            w_obj.grid(row=r,column=col+1,padx=(0,14),pady=8,sticky="w"); self.cz_g[key]=w_obj
        tk.Label(card,text="Notlar:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=6,column=0,padx=(20,6),pady=8,sticky="ne")
        self.cz_g["notlar"]=tk.Text(card,width=46,height=2,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        self.cz_g["notlar"].grid(row=6,column=1,columnspan=3,padx=(0,20),pady=8,sticky="w")
        MBtn(card,"✅  Cezayı Kaydet",command=self._ceza_kaydet,color=C_DANGER,width=34).grid(row=7,column=0,columnspan=4,pady=18)
        # Mükerrer
        mm=tk.Frame(mf,bg=bg); mm.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(mm,text="🔁 Birden Fazla Ceza Uygulanan Şahıslar",font=("Segoe UI",13,"bold"),fg=C_DANGER,bg=bg).pack(anchor="w",pady=(0,10))
        self.tv_mukerrer=self._tv(mm,[("ad",160,"Ad Soyad"),("tc",110,"TC No"),("sayi",60,"Adet"),("konular",250,"Konular"),("tarihler",250,"Ceza Tarihleri")],14)
        MBtn(mm,"🔄 Yenile",command=self._yenile_mukerrer,color=C_INFO,width=14).pack(pady=8)
        self._yenile_ceza()
        self._yenile_mukerrer()

    def _yenile_ceza(self):
        if not DB_PATH: return
        try:
            with db_baglan() as c:
                self._all_ceza=c.execute("SELECT id,ad_soyad,tc,ilce,mera_koy,yil,konu,ipc_tutari,tarih FROM Idari_Cezalar ORDER BY id DESC").fetchall()
                # Mükerrer TC'leri bul
                mukerrer_tc=set()
                for r in c.execute("SELECT tc FROM Idari_Cezalar GROUP BY tc HAVING COUNT(*)>1").fetchall():
                    mukerrer_tc.add(str(r[0]))
                self._mukerrer_tc=mukerrer_tc
            self._filtre_ceza("")
        except Exception as e: logging.error(f"yenile_ceza:{e}")

    def _filtre_ceza(self,a):
        self.tv_ceza.delete(*self.tv_ceza.get_children()); a=a.lower()
        for r in getattr(self,"_all_ceza",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            tag="mukerrer" if str(r[2]) in getattr(self,"_mukerrer_tc",set()) else ""
            # Tutar sütununu Türk formatında göster
            row=list(r)
            try: row[7]=para_format(row[7])
            except Exception: pass
            self.tv_ceza.insert("","end",values=row,tags=(tag,) if tag else ())

    def _yenile_mukerrer(self):
        if not DB_PATH: return
        try:
            self.tv_mukerrer.delete(*self.tv_mukerrer.get_children())
            with db_baglan() as c:
                for r in c.execute("""SELECT ad_soyad,tc,COUNT(*) as cnt,
                    GROUP_CONCAT(DISTINCT konu) as konular,
                    GROUP_CONCAT(tarih,' / ') as tarihler
                    FROM Idari_Cezalar GROUP BY tc HAVING cnt>1 ORDER BY cnt DESC""").fetchall():
                    self.tv_mukerrer.insert("","end",values=r)
        except Exception as e: logging.error(f"mukerrer:{e}")

    def _ceza_kaydet(self):
        try:
            ad=self.cz_g["ad"].get().strip(); tc=self.cz_g["tc"].get().strip()
            tutar_str=self.cz_g["tutar"].get()
            if not veri_dogrula([("Ad Soyad",ad,"str"),("Ceza Tutarı",tutar_str,"float+")]): return
            if not tc_kontrol_ve_devam(tc): return
            tutar=para_parse(tutar_str)
            notlar=self.cz_g["notlar"].get("1.0",tk.END).strip()
            # Mükerrer kontrolü
            dosya_no = oto_dosya_no("IPC", "Idari_Cezalar")
            with db_baglan() as c:
                onceki=c.execute("SELECT COUNT(*) FROM Idari_Cezalar WHERE tc=?",(tc,)).fetchone()[0]
                if onceki>0:
                    messagebox.showwarning("⚠️ Mükerrer Ceza",f"Bu kişiye daha önce {onceki} adet ceza uygulanmış!\nTC: {tc}")
                c.execute("INSERT INTO Idari_Cezalar(ad_soyad,tc,ilce,mera_koy,mera_ada_parsel,yil,konu,hayvan_sayisi,ceza_miktari,ipc_tutari,tarih,notlar)VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                    (ad,tc,self.cz_g["ilce"].get(),self.cz_g["mera_koy"].get(),self.cz_g["ada_parsel"].get(),
                     int(self.cz_g["yil"].get() or datetime.now().year),self.cz_g["konu"].get(),
                     self.cz_g["hayvan"].get(),self.cz_g["birim"].get(),tutar,
                     datetime.now().strftime("%d.%m.%Y"),notlar))
            db_log(self.u_id,"İdari Ceza",f"{dosya_no} {ad} {tutar}₺"); self._son_islem_kaydet("Ceza",ad)
            for k,w in self.cz_g.items():
                if isinstance(w,ttk.Entry): w.delete(0,tk.END)
                elif isinstance(w,ttk.Combobox): w.set("")
                elif isinstance(w,tk.Text): w.delete("1.0",tk.END)
            self.cz_g["yil"].insert(0,str(datetime.now().year))
            self._yenile_ceza(); self._yenile_mukerrer()
            messagebox.showinfo("Tamam",f"'{ad}' ceza kaydı oluşturuldu.\nDosya No: {dosya_no}")
        except ValueError as e: messagebox.showerror("Hata",str(e))

    def _ceza_evrak(self):
        sel=self.tv_ceza.selection()
        if not sel: messagebox.showwarning("Seçim","Ceza seçin."); return
        if not DOCX_OK: messagebox.showerror("Hata","python-docx kurulu değil."); return
        cid=self.tv_ceza.item(sel[0])["values"][0]
        try:
            with db_baglan() as c:
                r=c.execute("SELECT ad_soyad,tc,ilce,mera_koy,konu,ipc_tutari,tarih FROM Idari_Cezalar WHERE id=?",(cid,)).fetchone()
        except Exception: return
        dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile=f"Ceza_{r[0]}.docx",filetypes=[("Word","*.docx")])
        if not dosya: return
        try:
            # İl Müdürü ve unvan bilgilerini DB'den çek
            il_muduru="Muhammet Fatih CİNEVİZ"; unvan1="Vali a."; unvan2="Vali Yardımcısı"
            try:
                with db_baglan() as c:
                    for row in c.execute("SELECT anahtar,deger FROM Iletisim_Bilgileri WHERE anahtar IN('il_muduru','evrak_unvan')"):
                        if row[0]=="il_muduru" and row[1]: il_muduru=row[1]
                        elif row[0]=="evrak_unvan" and row[1]:
                            parts=row[1].split("\n") if "\n" in row[1] else [row[1]]
                            if len(parts)>=1: unvan1=parts[0]
                            if len(parts)>=2: unvan2=parts[1]
            except Exception: pass
            komisyon=getattr(self,"vali_yardimcisi","Semih CEMBEKLİ")
            word_idari_ceza(dosya,{"ad_soyad":r[0],"tc":r[1],"ilce":r[2],"mera":f"{r[2]}/{r[3]}","konu":r[4],"tutar":r[5],"tarih":r[6],
                "baba_adi":"","dogum":"","adres":"","il_muduru":il_muduru,"komisyon_baskani":komisyon,
                "evrak_unvan_1":unvan1,"evrak_unvan_2":unvan2})
            messagebox.showinfo("Tamam",f"Evrak üretildi:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _ceza_excel(self):
        if not getattr(self,"_all_ceza",None): messagebox.showwarning("Veri Yok","Kayıt yok."); return
        yol=filedialog.asksaveasfilename(defaultextension=".xlsx",filetypes=[("Excel","*.xlsx")])
        if yol:
            try:
                pd.DataFrame(self._all_ceza,columns=["ID","Ad Soyad","TC","İlçe","Mera","Yıl","Konu","Tutar","Tarih"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"Dışa aktarıldı:\n{yol}")
            except Exception as e: messagebox.showerror("Hata",str(e))

    def _ceza_sil(self):
        sel=self.tv_ceza.selection()
        if not sel: return
        cid=self.tv_ceza.item(sel[0])["values"][0]; ad=self.tv_ceza.item(sel[0])["values"][1]
        self._guvenli_sil("Idari_Cezalar",cid,f"Ceza: {ad}","DELETE FROM Idari_Cezalar WHERE id=?",(cid,),lambda:(self._yenile_ceza(),self._yenile_mukerrer()))

    # ═══ 8. VERİ KAYIT ══════════════════════════════════════════════════════
    def _t_veri(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2)
        nb2.add(lf,text="  📋 Kayıtlar  "); nb2.add(nf,text="  ➕ Yeni Kayıt  ")
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,self._filtre_veri,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_veri=self._tv(ll,[("id",40,"#"),("t",120,"Tarih"),("k",110,"Personel"),("kat",110,"Kategori"),("bas",200,"Başlık"),("det",350,"Detay")],16)
        bf=tk.Frame(ll,bg=bg); bf.pack(pady=6)
        MBtn(bf,"🔄 Yenile",command=self._yenile_veri,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(bf,"📊 Excel Export",command=self._veri_excel,color=self.gc("acc"),width=16).pack(side="left",padx=4)
        if self.u_yetki!="İzleyici":
            MBtn(bf,"🗑 Sil",command=self._veri_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=20)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.45,anchor="center",width=560,height=380)
        tk.Label(card,text="Yeni Veri Kaydı",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=2,pady=(18,14))
        tk.Label(card,text="Başlık:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=1,column=0,padx=(28,6),pady=8,sticky="e")
        self.veri_bas=ttk.Entry(card,width=38); self.veri_bas.grid(row=1,column=1,padx=(0,28),sticky="w")
        tk.Label(card,text="Kategori:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=2,column=0,padx=(28,6),pady=8,sticky="e")
        self.veri_kat=ttk.Combobox(card,width=28,state="readonly",values=["Genel","Mera","Hayvancılık","İdari","Teknik","Diğer"])
        self.veri_kat.set("Genel"); self.veri_kat.grid(row=2,column=1,padx=(0,28),sticky="w")
        tk.Label(card,text="Detay:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=3,column=0,padx=(28,6),pady=8,sticky="ne")
        self.veri_det=tk.Text(card,width=36,height=6,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        self.veri_det.grid(row=3,column=1,padx=(0,28),pady=8,sticky="w")
        MBtn(card,"💾  Kaydet",command=self._veri_kaydet,color=self.gc("acc"),width=28).grid(row=4,column=0,columnspan=2,pady=14)
        self._yenile_veri()

    def _yenile_veri(self):
        if not DB_PATH: return
        try:
            with db_baglan() as c: self._all_veri=c.execute("SELECT id,tarih,kul,kategori,baslik,detay FROM Veri_Kayit ORDER BY id DESC").fetchall()
            self._filtre_veri("")
        except Exception as e: logging.error(f"yenile_veri:{e}")

    def _filtre_veri(self,a):
        self.tv_veri.delete(*self.tv_veri.get_children()); a=a.lower()
        for r in getattr(self,"_all_veri",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            self.tv_veri.insert("","end",values=r)

    def _veri_kaydet(self):
        bas=self.veri_bas.get().strip()
        if not bas: messagebox.showwarning("Uyarı","Başlık zorunlu."); return
        det=self.veri_det.get("1.0",tk.END).strip()
        with db_baglan() as c:
            c.execute("INSERT INTO Veri_Kayit(tarih,kul,baslik,kategori,detay)VALUES(?,?,?,?,?)",
                (datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad,bas,self.veri_kat.get(),det))
        db_log(self.u_id,"Veri Kayıt",bas); self.veri_bas.delete(0,tk.END); self.veri_det.delete("1.0",tk.END); self._yenile_veri()

    def _veri_excel(self):
        if not getattr(self,"_all_veri",None): messagebox.showwarning("Veri Yok","Kayıt yok."); return
        yol=filedialog.asksaveasfilename(defaultextension=".xlsx",filetypes=[("Excel","*.xlsx")])
        if yol:
            try:
                pd.DataFrame(self._all_veri,columns=["ID","Tarih","Personel","Kategori","Başlık","Detay"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"Dışa aktarıldı:\n{yol}")
            except Exception as e: messagebox.showerror("Hata",str(e))

    def _veri_sil(self):
        sel=self.tv_veri.selection()
        if not sel: return
        vid=self.tv_veri.item(sel[0])["values"][0]
        if messagebox.askyesno("Onay","Silinsin mi?"):
            with db_baglan() as c: c.execute("DELETE FROM Veri_Kayit WHERE id=?",(vid,))
            self._yenile_veri()

    # ═══ 9. MUHTARLAR ════════════════════════════════════════════════════════
    def _t_muhtar(self,p):
        bg=self.gc("bg"); f=tk.Frame(p,bg=bg); f.pack(fill="both",expand=True,padx=16,pady=14)
        tk.Label(f,text="👥 Muhtarlar Rehberi",font=("Segoe UI",15,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        if self.u_yetki!="İzleyici":
            form=tk.LabelFrame(f,text="  Yeni Muhtar  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=10,pady=8)
            form.pack(fill="x",pady=(0,10))
            self.muh_g={}
            for i,(lbl,key,w) in enumerate([("İlçe:","ilce",12),("Köy:","koy",16),("Ad Soyad:","ad",20),("Telefon:","tel",14),("E-posta:","email",20)]):
                tk.Label(form,text=lbl,bg=bg,font=("Segoe UI",10)).grid(row=0,column=i*2,padx=(0,4),sticky="w")
                w_obj=ttk.Combobox(form,values=ILCELER,state="readonly",width=w) if lbl=="İlçe:" else ttk.Entry(form,width=w)
                w_obj.grid(row=0,column=i*2+1,padx=(0,10)); self.muh_g[key]=w_obj
            MBtn(form,"➕ Ekle",command=self._muh_ekle).grid(row=0,column=10,padx=8)
        AramaFrame(f,self._filtre_muh,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_muh=self._tv(f,[("id",40,"#"),("ilce",90,"İlçe"),("koy",140,"Köy"),("ad",160,"Ad Soyad"),("tel",130,"Telefon"),("email",180,"E-posta")],18)
        bf=tk.Frame(f,bg=bg); bf.pack(pady=6)
        MBtn(bf,"🔄 Yenile",command=self._yenile_muh,color=C_INFO,width=14).pack(side="left",padx=4)
        if self.u_yetki!="İzleyici":
            MBtn(bf,"📥 Excel Import",command=self._muh_import,color="#8E44AD",width=14).pack(side="left",padx=4)
            MBtn(bf,"📊 Excel Export",command=self._muh_export,color=self.gc("acc"),width=14).pack(side="left",padx=4)
            MBtn(bf,"🗑 Sil",command=self._muh_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        self._tum_muh=[]; self._yenile_muh()

    def _yenile_muh(self):
        if not DB_PATH: return
        try:
            with db_baglan() as c: self._tum_muh=c.execute("SELECT id,ilce,koy,ad_soyad,telefon,email FROM Muhtarlar ORDER BY ilce,koy").fetchall()
            self._filtre_muh("")
        except Exception as e: logging.error(f"yenile_muh:{e}")

    def _filtre_muh(self,a):
        self.tv_muh.delete(*self.tv_muh.get_children()); a=a.lower()
        for r in self._tum_muh:
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            self.tv_muh.insert("","end",values=r)

    def _muh_ekle(self):
        koy=self.muh_g["koy"].get().strip(); ad=self.muh_g["ad"].get().strip()
        if not koy or not ad: messagebox.showwarning("Uyarı","Köy ve ad zorunlu."); return
        with db_baglan() as c: c.execute("INSERT INTO Muhtarlar(ilce,koy,ad_soyad,telefon,email)VALUES(?,?,?,?,?)",(self.muh_g["ilce"].get(),koy,ad,self.muh_g["tel"].get(),self.muh_g["email"].get()))
        db_log(self.u_id,"Muhtar Ekle",koy)
        for k,w in self.muh_g.items():
            if isinstance(w,ttk.Entry): w.delete(0,tk.END)
            elif isinstance(w,ttk.Combobox): w.set("")
        self._yenile_muh()

    def _muh_sil(self):
        sel=self.tv_muh.selection()
        if not sel: return
        mid=self.tv_muh.item(sel[0])["values"][0]
        if messagebox.askyesno("Onay","Silinsin mi?"):
            with db_baglan() as c: c.execute("DELETE FROM Muhtarlar WHERE id=?",(mid,))
            self._yenile_muh()

    def _muh_import(self):
        """Excel'den muhtar bilgilerini toplu yükle"""
        yol=filedialog.askopenfilename(filetypes=[("Excel","*.xlsx *.xls")])
        if not yol: return
        try:
            df=pd.read_excel(yol)
            if len(df.columns)<3:
                messagebox.showerror("Hata","En az 3 sütun gerekli: İlçe | Köy | Ad Soyad\nOpsiyonel: Telefon | E-posta"); return
            with db_baglan() as c:
                kayit=0
                for _,r in df.iterrows():
                    v=r.tolist()
                    ilce=str(v[0]) if len(v)>0 else ""
                    koy=str(v[1]) if len(v)>1 else ""
                    ad=str(v[2]) if len(v)>2 else ""
                    tel=str(v[3]) if len(v)>3 else ""
                    email=str(v[4]) if len(v)>4 else ""
                    if koy and ad:
                        c.execute("INSERT INTO Muhtarlar(ilce,koy,ad_soyad,telefon,email)VALUES(?,?,?,?,?)",(ilce,koy,ad,tel,email))
                        kayit+=1
            db_log(self.u_id,"Muhtar Import",f"{kayit} kayıt"); self._yenile_muh()
            messagebox.showinfo("Tamam",f"{kayit} muhtar bilgisi yüklendi.")
        except Exception as e: messagebox.showerror("Excel Hatası",str(e))

    def _muh_export(self):
        """Muhtar bilgilerini Excel'e aktar"""
        if not self._tum_muh: messagebox.showwarning("Veri Yok","Kayıt yok."); return
        yol=filedialog.asksaveasfilename(defaultextension=".xlsx",initialfile="Muhtarlar.xlsx",filetypes=[("Excel","*.xlsx")])
        if yol:
            try:
                pd.DataFrame(self._tum_muh,columns=["ID","İlçe","Köy","Ad Soyad","Telefon","E-posta"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"Dışa aktarıldı:\n{yol}")
            except Exception as e: messagebox.showerror("Hata",str(e))

    # ═══ 10. AJANDA ══════════════════════════════════════════════════════════
    def _t_ajanda(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        lf=ttk.Frame(nb2); nf=ttk.Frame(nb2)
        nb2.add(lf,text="  📅 Tüm Notlar  "); nb2.add(nf,text="  ➕ Yeni Hatırlatıcı  ")
        ll=tk.Frame(lf,bg=bg); ll.pack(fill="both",expand=True,padx=14,pady=12)
        AramaFrame(ll,self._filtre_aj,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_aj=self._tv(ll,[("id",40,"#"),("t",100,"Tarih"),("s",70,"Saat"),("bas",190,"Başlık"),("tur",90,"Tür"),("dur",90,"Durum"),("ic",280,"İçerik")],14)
        self.tv_aj.tag_configure("bekliyor",background="#FEF9E7")
        self.tv_aj.tag_configure("tamamlandi",background="#EAFAF1")
        self.tv_aj.tag_configure("gecmis",background="#FDEDEC")
        bf=tk.Frame(ll,bg=bg); bf.pack(pady=8)
        MBtn(bf,"🔄 Yenile",command=self._yenile_aj,color=C_INFO,width=14).pack(side="left",padx=4)
        MBtn(bf,"✅ Tamamlandı",command=lambda:self._aj_dur("Tamamlandı"),color=self.gc("acc"),width=16).pack(side="left",padx=4)
        MBtn(bf,"🗑 Sil",command=self._aj_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        nn=tk.Frame(nf,bg=bg); nn.pack(fill="both",expand=True,padx=20,pady=20)
        card=tk.Frame(nn,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
        card.place(relx=0.5,rely=0.45,anchor="center",width=560,height=460)
        tk.Label(card,text="Yeni Ajanda Notu",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).grid(row=0,column=0,columnspan=2,pady=(18,14))
        self.aj_g={}
        for i,(l,k,d) in enumerate([("Başlangıç Tarihi:","t",datetime.now().strftime("%d.%m.%Y")),
                                     ("Bitiş Tarihi:","t_bitis",""),
                                     ("Saat:","s","09:00"),("Başlık:","bas","")],1):
            tk.Label(card,text=l,bg=C_WHITE,font=("Segoe UI",10)).grid(row=i,column=0,padx=(28,6),pady=8,sticky="e")
            e=ttk.Entry(card,width=28); e.insert(0,d); e.grid(row=i,column=1,padx=(0,28),sticky="w"); self.aj_g[k]=e
        tk.Label(card,text="",bg=C_WHITE,fg="#888",font=("Segoe UI",8)).grid(row=2,column=1,sticky="w",padx=(0,28))
        lbl_hint=tk.Label(card,text="💡 Boş bırakırsanız tek gün kaydedilir.\nÖrn: 5 günlük izin için bitiş tarihi girin.",bg=C_WHITE,fg="#888",font=("Segoe UI",8),justify="left")
        lbl_hint.grid(row=2,column=1,sticky="se",padx=(0,28))
        tk.Label(card,text="Tür:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=5,column=0,padx=(28,6),pady=8,sticky="e")
        self.aj_g["tur"]=ttk.Combobox(card,values=["Hatırlatıcı","Toplantı","Denetim","Randevu","İzin","Saha Ziyareti","Diğer"],state="readonly",width=26)
        self.aj_g["tur"].set("Hatırlatıcı"); self.aj_g["tur"].grid(row=5,column=1,padx=(0,28),sticky="w")
        tk.Label(card,text="İçerik:",bg=C_WHITE,font=("Segoe UI",10)).grid(row=6,column=0,padx=(28,6),pady=8,sticky="ne")
        self.aj_g["ic"]=tk.Text(card,width=28,height=4,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=8,pady=6)
        self.aj_g["ic"].grid(row=6,column=1,padx=(0,28),pady=8,sticky="w")
        MBtn(card,"💾  Kaydet",command=self._aj_kaydet,color=self.gc("acc"),width=28).grid(row=7,column=0,columnspan=2,pady=16)
        self._yenile_aj()

    def _yenile_aj(self):
        if not DB_PATH: return
        try:
            with db_baglan() as c: self._all_aj=c.execute("SELECT id,tarih,sure,baslik,tur,durum,icerik FROM Ajanda WHERE k_adi=? ORDER BY tarih DESC,sure",(self.u_id,)).fetchall()
            self._filtre_aj("")
        except Exception as e: logging.error(f"yenile_aj:{e}")

    def _filtre_aj(self,a):
        self.tv_aj.delete(*self.tv_aj.get_children()); a=a.lower(); bugun=datetime.now().strftime("%d.%m.%Y")
        for r in getattr(self,"_all_aj",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            d=str(r[5]); tag="tamamlandi" if d=="Tamamlandı" else "gecmis" if r[1]<bugun else "bekliyor"
            self.tv_aj.insert("","end",values=r,tags=(tag,))

    def _aj_kaydet(self):
        bas=self.aj_g["bas"].get().strip()
        if not bas: messagebox.showwarning("Uyarı","Başlık zorunlu."); return
        ic=self.aj_g["ic"].get("1.0",tk.END).strip()
        tarih_bas=self.aj_g["t"].get().strip()
        tarih_bit=self.aj_g["t_bitis"].get().strip()
        saat=self.aj_g["s"].get()
        tur=self.aj_g["tur"].get()
        if tarih_bit:
            # Tarih aralığı — her gün için ayrı kayıt
            try:
                d_bas=datetime.strptime(tarih_bas,"%d.%m.%Y")
                d_bit=datetime.strptime(tarih_bit,"%d.%m.%Y")
                if d_bit<d_bas:
                    messagebox.showerror("Hata","Bitiş tarihi başlangıçtan önce olamaz."); return
                gun_sayisi=(d_bit-d_bas).days+1
                if gun_sayisi>60:
                    messagebox.showerror("Hata","En fazla 60 günlük aralık girilebilir."); return
                with db_baglan() as c:
                    for i in range(gun_sayisi):
                        gun=(d_bas+timedelta(days=i)).strftime("%d.%m.%Y")
                        etiket=f"{bas} ({i+1}/{gun_sayisi})" if gun_sayisi>1 else bas
                        c.execute("INSERT INTO Ajanda(tarih,sure,baslik,tur,durum,icerik,k_adi)VALUES(?,?,?,?,?,?,?)",
                            (gun,saat,etiket,tur,"Bekliyor",ic,self.u_id))
                self._yenile_aj()
                messagebox.showinfo("Tamam",f"{gun_sayisi} günlük kayıt oluşturuldu.\n{tarih_bas} — {tarih_bit}")
            except ValueError:
                messagebox.showerror("Hata","Tarih formatı hatalı.\nDoğru format: 01.04.2026"); return
        else:
            # Tek gün
            with db_baglan() as c:
                c.execute("INSERT INTO Ajanda(tarih,sure,baslik,tur,durum,icerik,k_adi)VALUES(?,?,?,?,?,?,?)",
                    (tarih_bas,saat,bas,tur,"Bekliyor",ic,self.u_id))
            self._yenile_aj()

    def _aj_dur(self,d):
        sel=self.tv_aj.selection()
        if not sel: return
        aid=self.tv_aj.item(sel[0])["values"][0]
        with db_baglan() as c: c.execute("UPDATE Ajanda SET durum=? WHERE id=?",(d,aid))
        self._yenile_aj()

    def _aj_sil(self):
        sel=self.tv_aj.selection()
        if not sel: return
        aid=self.tv_aj.item(sel[0])["values"][0]
        if messagebox.askyesno("Onay","Silinsin mi?"):
            with db_baglan() as c: c.execute("DELETE FROM Ajanda WHERE id=?",(aid,))
            self._yenile_aj()

    # ═══ 11. PERSONEL TAKVİMİ ════════════════════════════════════════════════
    def _t_personel(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        tf=ttk.Frame(nb2); ef=ttk.Frame(nb2)
        nb2.add(ef,text="  👥 Şube Ekibi  "); nb2.add(tf,text="  📅 Görev Takvimi  ")
        # Şube Ekibi — profil kartları
        ekip=tk.Frame(ef,bg=bg); ekip.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(ekip,text="👥 Şube Personeli",font=("Segoe UI",15,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        self._ekip_frame=tk.Frame(ekip,bg=bg); self._ekip_frame.pack(fill="both",expand=True)
        bf0=tk.Frame(ekip,bg=bg); bf0.pack(pady=8)
        MBtn(bf0,"📸 Profil Fotoğrafımı Değiştir",command=lambda:self._profil_foto_goster(self.u_id),color=self.gc("acc"),width=26).pack(side="left",padx=4)
        MBtn(bf0,"🔄 Yenile",command=self._ekip_yenile,color=C_INFO,width=12).pack(side="left",padx=4)
        self._ekip_yenile()
        # Görev Takvimi
        f=tk.Frame(tf,bg=bg); f.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(f,text="📅 Personel Takvimi",font=("Segoe UI",15,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        self.tv_personel=self._tv(f,[("id",40,"#"),("ad",150,"Personel"),("tarih",90,"Tarih"),("bas",70,"Başlangıç"),("bit",70,"Bitiş"),("tur",100,"Tür"),("acik",260,"Açıklama")],16)
        bf=tk.Frame(f,bg=bg); bf.pack(pady=8)
        MBtn(bf,"🔄 Yenile",command=self._yenile_personel,color=C_INFO,width=14).pack(side="left",padx=4)
        if self.u_yetki!="İzleyici":
            MBtn(bf,"➕ Yeni Görev",command=self._personel_ekle_popup,color=self.gc("acc"),width=14).pack(side="left",padx=4)
            MBtn(bf,"🗑 Sil",command=self._personel_sil,color=C_DANGER,width=10).pack(side="left",padx=4)
        self._yenile_personel()

    def _ekip_yenile(self):
        """Şube ekibi profil kartları — WhatsApp tarzı"""
        for w in self._ekip_frame.winfo_children(): w.destroy()
        if not DB_PATH: return
        try:
            with db_baglan() as c:
                personeller=c.execute("SELECT k_adi,ad,unvan,yetki,aktif FROM Kullanicilar WHERE aktif=1 ORDER BY ad").fetchall()
        except Exception: return
        foto_dir=Path("miras_profil")
        row_frame=None
        for i,(ka,ad,unvan,yetki,aktif) in enumerate(personeller):
            if i%4==0:
                row_frame=tk.Frame(self._ekip_frame,bg=self.gc("bg")); row_frame.pack(fill="x",pady=4)
            card=tk.Frame(row_frame,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1,padx=12,pady=10,cursor="hand2")
            card.pack(side="left",fill="both",expand=True,padx=4)
            # Fotoğraf
            foto_yol=foto_dir/f"{ka}.png"
            img_label=None
            if foto_yol.exists():
                try:
                    from PIL import Image, ImageTk, ImageDraw
                    img=Image.open(str(foto_yol)).resize((64,64))
                    mask=Image.new("L",(64,64),0); ImageDraw.Draw(mask).ellipse((0,0,64,64),fill=255)
                    img.putalpha(mask)
                    photo=ImageTk.PhotoImage(img)
                    img_label=tk.Label(card,image=photo,bg=C_WHITE); img_label.image=photo
                    img_label.pack(pady=(4,2))
                except Exception:
                    tk.Label(card,text="👤",font=("Segoe UI",24),bg="#E8E8E8",fg="#888",width=3,height=1).pack(pady=(4,2))
            else:
                tk.Label(card,text="👤",font=("Segoe UI",24),bg="#E8E8E8",fg="#888",width=3,height=1).pack(pady=(4,2))
            tk.Label(card,text=ad or ka,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack()
            tk.Label(card,text=unvan or "",font=("Segoe UI",8),fg="#666",bg=C_WHITE).pack()
            yetki_renk={"Admin":C_DANGER,"Uzman":C_INFO,"İzleyici":"#888"}.get(yetki,"#888")
            tk.Label(card,text=yetki,font=("Segoe UI",7,"bold"),fg=yetki_renk,bg=C_WHITE).pack()
            # Tıklayınca profil göster
            card.bind("<Button-1>",lambda e,k=ka:self._profil_foto_goster(k))

    def _yenile_personel(self):
        if not DB_PATH: return
        try:
            with db_baglan() as c:
                rows=c.execute("SELECT id,ad,tarih,baslangic,bitis,tur,aciklama FROM Personel_Takvim ORDER BY tarih DESC").fetchall()
            self.tv_personel.delete(*self.tv_personel.get_children())
            for r in rows: self.tv_personel.insert("","end",values=r)
        except Exception as e: logging.error(f"yenile_personel:{e}")

    def _personel_ekle_popup(self):
        win=tk.Toplevel(self.root); win.title("Yeni Görev/İzin"); win.geometry("440x380"); win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        tk.Label(win,text="Yeni Görev/İzin/Toplantı",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=14)
        al={}
        for l,k,d in [("Personel Adı:","ad",""),("Tarih:","tarih",datetime.now().strftime("%d.%m.%Y")),
                       ("Başlangıç:","bas","08:00"),("Bitiş:","bit","17:00")]:
            tk.Label(win,text=l,bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=24,pady=(6,0))
            e=ttk.Entry(win,width=28); e.insert(0,d); e.pack(padx=24); al[k]=e
        tk.Label(win,text="Tür:",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=24,pady=(6,0))
        al["tur"]=ttk.Combobox(win,values=["Görev","İzin","Toplantı","Denetim","Saha Ziyareti","Eğitim"],state="readonly",width=26)
        al["tur"].set("Görev"); al["tur"].pack(padx=24)
        tk.Label(win,text="Açıklama:",bg=C_WHITE,font=("Segoe UI",10)).pack(anchor="w",padx=24,pady=(6,0))
        al["acik"]=tk.Text(win,width=30,height=3,font=("Segoe UI",10)); al["acik"].pack(padx=24)
        def _k():
            ad=al["ad"].get().strip()
            if not ad: messagebox.showwarning("Uyarı","Personel adı zorunlu.",parent=win); return
            acik=al["acik"].get("1.0",tk.END).strip()
            with db_baglan() as c:
                c.execute("INSERT INTO Personel_Takvim(k_adi,ad,tarih,baslangic,bitis,tur,aciklama)VALUES(?,?,?,?,?,?,?)",
                    (self.u_id,ad,al["tarih"].get(),al["bas"].get(),al["bit"].get(),al["tur"].get(),acik))
            self._yenile_personel(); win.destroy()
        MBtn(win,"💾 Kaydet",command=_k,width=20).pack(pady=12)

    def _personel_sil(self):
        sel=self.tv_personel.selection()
        if not sel: return
        pid=self.tv_personel.item(sel[0])["values"][0]
        if messagebox.askyesno("Onay","Silinsin mi?"):
            with db_baglan() as c: c.execute("DELETE FROM Personel_Takvim WHERE id=?",(pid,))
            self._yenile_personel()

    # ═══ 12. HARİTA ══════════════════════════════════════════════════════════
    # ═══ İHALE YERLERİ HAVUZU + MUHAMMEN BEDEL ════════════════════════════════
    # İhale Yerleri + Muhammen artık _t_ihale içinde

    def _mb_vasif_sec(self,event=None):
        v=self.mb_g["vasif"].get()
        kov=OT_VERIM_KURU.get(v,"—")
        self.mb_kov.config(text=f"{kov} kg/da (Ardahan 500-650mm yağış kuşağı)")

    def _mb_hesapla(self):
        v=self.mb_g["vasif"].get()
        if not v: messagebox.showerror("Hata","Vasıf seçin."); return
        kov=OT_VERIM_KURU.get(v)
        sonuc=muhammen_bedel_hesapla(self.mb_g["alan"].get(),kov,self.mb_g["gun"].get(),self.mb_g["fiyat"].get())
        if not sonuc: messagebox.showerror("Hata","Tüm alanları doğru doldurun."); return
        self.lbl_mb_sonuc.config(text=
            f"📊 Mera Alanı: {sonuc['alan']:,.2f} da | Kuru Ot: {sonuc['kuru_ot']:.0f} kg/da\n"
            f"⏱ Otlatma: {sonuc['gun']:.0f} gün | Fiyat: {sonuc['fiyat']:.2f} ₺/kg\n\n"
            f"🐄 Otlatma Kapasitesi: {sonuc['bbhb']:,.2f} BBHB ({sonuc['kbhb']:,.2f} KBHB)\n\n"
            f"💰 MUHAMMEN BEDEL: {para_format(sonuc['bedel'])}\n"
            f"{'─'*40}\n"
            f"IBAN (Emanet): {IBAN_EMANET}")
        self._mb_sonuc=sonuc

    def _mb_pdf_uret(self):
        """Tahdit Raporu — kurumsal PDF çıktısı"""
        sonuc=getattr(self,"_mb_sonuc",None)
        if not sonuc: messagebox.showwarning("Uyarı","Önce Muhammen Bedel hesaplayın."); return
        if not PDF_OK: messagebox.showerror("Eksik","reportlab kurulu değil."); return
        dosya=filedialog.asksaveasfilename(defaultextension=".pdf",initialfile="Tahdit_Raporu.pdf",filetypes=[("PDF","*.pdf")])
        if not dosya: return
        try:
            _init_fonts(); fn,fnb=_FN,_FNB; W,H=A4
            cv=rl_canvas.Canvas(dosya,pagesize=A4)
            # Başlık bandı
            cv.setFillColor(rl_colors.HexColor("#1E5631"))
            cv.rect(0,H-3.2*cm,W,3.2*cm,fill=1,stroke=0)
            cv.setFillColor(rl_colors.HexColor("#2D8C55"))
            cv.rect(0,H-3.4*cm,W,0.2*cm,fill=1,stroke=0)
            cv.setFillColor(rl_colors.white)
            LOGO="bakanlik_logo.png"
            if os.path.exists(LOGO):
                try: cv.drawImage(LOGO,0.4*cm,H-3.0*cm,width=2.4*cm,height=2.4*cm,preserveAspectRatio=True,mask="auto")
                except Exception: pass
            cv.setFont(fnb,13); cv.drawCentredString(W/2,H-1.2*cm,"T.C. ARDAHAN VALİLİĞİ")
            cv.setFont(fn,10); cv.drawCentredString(W/2,H-1.8*cm,"İl Tarım ve Orman Müdürlüğü")
            cv.setFont(fn,9); cv.drawCentredString(W/2,H-2.4*cm,"MERA KOMİSYON BAŞKANLIĞINA")
            y=25.5
            cv.setFillColor(rl_colors.black)
            cv.setFont(fnb,14); cv.drawCentredString(W/2,y*cm,"TAHDİT RAPORU"); y-=1.0
            def _y(x,yt,t,f=fn,s=10,c=rl_colors.black):
                cv.setFont(f,s); cv.setFillColor(c); cv.drawString(x*cm,yt*cm,str(t))
            # İlçe/Köy/Ada/Parsel bilgileri
            ilce_t=self.mb_g.get("ilce",""); koy_t=self.mb_g.get("koy",""); ada_t=self.mb_g.get("ada",""); parsel_t=self.mb_g.get("parsel","")
            if hasattr(ilce_t,"get"): ilce_t=ilce_t.get()
            if hasattr(koy_t,"get"): koy_t=koy_t.get()
            if hasattr(ada_t,"get"): ada_t=ada_t.get()
            if hasattr(parsel_t,"get"): parsel_t=parsel_t.get()
            if ilce_t or koy_t:
                _y(2,y,f"İlçe: {ilce_t}   Köy: {koy_t}   Ada: {ada_t}   Parsel: {parsel_t}",fnb,11); y-=0.6
            _y(2,y,f"Alan: {sonuc['alan']:,.2f} da",fnb,11); _y(10,y,f"Vasıf: {self.mb_g['vasif'].get()}",fn,11); y-=0.6
            _y(2,y,f"Otlatma Gün: {sonuc['gun']:.0f}",fn,11); _y(10,y,f"Kuru Ot Verimi: {sonuc['kuru_ot']:.0f} kg/da",fn,11); y-=0.6
            _y(2,y,f"Kuru Ot Fiyatı: {sonuc['fiyat']:.2f} ₺/kg",fn,11); y-=0.8
            cv.setStrokeColor(rl_colors.HexColor("#1E5631")); cv.setLineWidth(1)
            cv.line(2*cm,y*cm,19*cm,y*cm); y-=0.8
            _y(2,y,"OTLATMA KAPASİTESİ",fnb,12,rl_colors.HexColor("#1E5631")); y-=0.7
            _y(3,y,f"BBHB: {sonuc['bbhb']:,.1f}",fnb,14,rl_colors.HexColor("#1E5631"))
            _y(10,y,f"KBHB: {sonuc['kbhb']:,.2f}",fn,12); y-=1.0
            cv.line(2*cm,y*cm,19*cm,y*cm); y-=0.8
            _y(2,y,"MUHAMMEN BEDEL",fnb,12,rl_colors.HexColor("#C0392B")); y-=0.8
            _y(3,y,para_format(sonuc['bedel']),fnb,18,rl_colors.HexColor("#C0392B")); y-=1.0
            _y(2,y,f"Formül: {sonuc['bbhb']:,.2f} × {sonuc['gun']:.0f} × 12,5 × {sonuc['fiyat']:.2f} = {para_format(sonuc['bedel'])}",fn,9); y-=1.0
            cv.line(2*cm,y*cm,19*cm,y*cm); y-=0.6
            _y(2,y,f"Tarih: {datetime.now().strftime('%d.%m.%Y')}",fn,9,rl_colors.HexColor("#555555"))
            # Alt bilgi
            cv.setFillColor(rl_colors.HexColor("#1E5631"))
            cv.rect(0,0,W,1.2*cm,fill=1,stroke=0)
            cv.setFillColor(rl_colors.white); cv.setFont(fn,7)
            cv.drawCentredString(W/2,0.5*cm,f"MİRAS Enterprise {VERSIYON} — {HAKLAR}")
            cv.save()
            messagebox.showinfo("Tamam",f"Tahdit Raporu PDF oluşturuldu:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _yenile_iyer(self):
        if not DB_PATH: return
        try:
            with db_baglan() as c:
                self._all_iyer=c.execute("SELECT id,ilce,koy,ada,parsel,alan_da,kapasite_bbhb,tahmini_bedel,vasif,yil FROM Ihale_Yerleri ORDER BY yil DESC,ilce,koy").fetchall()
            self._filtre_iyer("")
        except Exception as e: logging.error(f"yenile_iyer:{e}")

    def _filtre_iyer(self,a):
        self.tv_iyer.delete(*self.tv_iyer.get_children()); a=a.lower()
        for r in getattr(self,"_all_iyer",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            row=list(r)
            try: row[7]=para_format(row[7])
            except Exception: pass
            self.tv_iyer.insert("","end",values=row)

    # _iyer_kaydet kaldırıldı — _iyer_ekle_popup kullanılıyor

    def _iyer_import(self):
        yol=filedialog.askopenfilename(filetypes=[("Excel","*.xlsx")])
        if not yol: return
        try:
            df=pd.read_excel(yol)
            with db_baglan() as c:
                n=0
                for _,r in df.iterrows():
                    v=r.tolist()
                    if len(v)>=5:
                        c.execute("INSERT INTO Ihale_Yerleri(ilce,koy,ada,parsel,alan_da,vasif,yil)VALUES(?,?,?,?,?,?,?)",
                            (str(v[0]),str(v[1]),str(v[2]),str(v[3]),float(v[4] or 0),str(v[5]) if len(v)>5 else "",_yil()))
                        n+=1
            self._yenile_iyer(); messagebox.showinfo("Tamam",f"{n} ihale yeri yüklendi.")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _iyer_export(self):
        if not getattr(self,"_all_iyer",None): return
        yol=filedialog.asksaveasfilename(defaultextension=".xlsx",filetypes=[("Excel","*.xlsx")])
        if yol:
            try:
                pd.DataFrame(self._all_iyer,columns=["ID","İlçe","Köy","Ada","Parsel","Alan(da)","Kapasite","Tah.Bedel","Vasıf","Yıl"]).to_excel(yol,index=False)
                messagebox.showinfo("Tamam",f"Dışa aktarıldı:\n{yol}")
            except Exception as e: messagebox.showerror("Hata",str(e))

    def _iyer_sil(self):
        sel=self.tv_iyer.selection()
        if not sel: return
        if messagebox.askyesno("Onay","Silinsin mi?"):
            with db_baglan() as c: c.execute("DELETE FROM Ihale_Yerleri WHERE id=?",(self.tv_iyer.item(sel[0])["values"][0],))
            self._yenile_iyer()

    # ═══ 13. İSTATİSTİK + İHALE/CEZA ÖZETİ ═════════════════════════════════
    def _t_istatistik(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        gf=ttk.Frame(nb2); ihf=ttk.Frame(nb2); czf=ttk.Frame(nb2)
        nb2.add(gf,text="  📈 Genel  "); nb2.add(ihf,text="  ⚖️ İhale Özeti  "); nb2.add(czf,text="  💰 Ceza Özeti  ")
        # Genel
        f=tk.Frame(gf,bg=bg); f.pack(fill="both",expand=True,padx=14,pady=12)
        kf=tk.Frame(f,bg=bg); kf.pack(fill="x",pady=(0,10))
        self.ist_k={}
        for bas,key,renk,ikon in [("Toplam Köy","koy",self.gc("pri"),"🏘️"),("Toplam Rapor","rapor",C_INFO,"📄"),
            ("Aktif İhale","ihale",C_WARN,"⚖️"),("İdari Ceza","ceza",C_DANGER,"💰")]:
            k=StatKart(kf,bas,"—",renk,ikon); k.pack(side="left",fill="both",expand=True,padx=4)
            self.ist_k[key]=k
        if MPL_OK:
            self.ist_canvas_frame=tk.Frame(f,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1)
            self.ist_canvas_frame.pack(fill="both",expand=True)
        MBtn(f,"📊 Yenile",command=self._ist_ozet,color=self.gc("acc"),width=14).pack(pady=6)
        # İhale Özeti (yıl bazlı)
        ihh=tk.Frame(ihf,bg=bg); ihh.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(ihh,text="⚖️ İhale Özeti (Yıl Bazlı)",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        self.tv_ih_ozet=self._tv(ihh,[("yil",60,"Yıl"),("toplam",70,"Toplam"),("aktif",70,"Aktif"),
            ("tamam",80,"Tamamlanan"),("gelir",120,"Toplam Gelir"),("bakanlik",120,"Bakanlık(%25)"),
            ("koy",120,"Köy(%75)"),("damga",100,"Damga V.")],10)
        MBtn(ihh,"🔄 Yenile",command=self._ihale_ozet_yenile,color=C_INFO,width=14).pack(pady=6)
        # Ceza Özeti (yıl bazlı)
        czz=tk.Frame(czf,bg=bg); czz.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(czz,text="💰 Ceza Özeti (Yıl Bazlı)",font=("Segoe UI",14,"bold"),fg=C_DANGER,bg=bg).pack(anchor="w",pady=(0,10))
        self.tv_cz_ozet=self._tv(czz,[("yil",60,"Yıl"),("sayi",70,"Kişi"),("toplam",120,"Toplam Tutar"),
            ("ilce",140,"İlçe Dağılımı"),("konu",200,"Konu Dağılımı"),("mukerrer",80,"Mükerrer")],10)
        MBtn(czz,"🔄 Yenile",command=self._ceza_ozet_yenile,color=C_INFO,width=14).pack(pady=6)
        self._ist_ozet(); self._ihale_ozet_yenile(); self._ceza_ozet_yenile()

    def _ihale_ozet_yenile(self):
        if not DB_PATH: return
        self.tv_ih_ozet.delete(*self.tv_ih_ozet.get_children())
        try:
            with db_baglan() as c:
                yillar=c.execute("SELECT DISTINCT substr(tarih,-4) as y FROM Ihaleler ORDER BY y DESC").fetchall()
                for (yil,) in yillar:
                    if not yil: continue
                    toplam=c.execute("SELECT COUNT(*) FROM Ihaleler WHERE tarih LIKE ?",(f"%{yil}",)).fetchone()[0]
                    aktif=c.execute("SELECT COUNT(*) FROM Ihaleler WHERE tarih LIKE ? AND durum NOT IN('Tamamlandı','İptal Edildi')",(f"%{yil}",)).fetchone()[0]
                    tamam=c.execute("SELECT COUNT(*) FROM Ihaleler WHERE tarih LIKE ? AND durum='Tamamlandı'",(f"%{yil}",)).fetchone()[0]
                    gelir=c.execute("SELECT COALESCE(SUM(bedel),0) FROM Ihaleler WHERE tarih LIKE ? AND durum='Tamamlandı'",(f"%{yil}",)).fetchone()[0]
                    self.tv_ih_ozet.insert("","end",values=(yil,toplam,aktif,tamam,
                        para_format(gelir),para_format(gelir*0.25),para_format(gelir*0.75),para_format(gelir*5.69/1000)))
        except Exception as e: logging.error(f"ihale_ozet:{e}")

    def _ceza_ozet_yenile(self):
        if not DB_PATH: return
        self.tv_cz_ozet.delete(*self.tv_cz_ozet.get_children())
        try:
            with db_baglan() as c:
                yillar=c.execute("SELECT DISTINCT yil FROM Idari_Cezalar ORDER BY yil DESC").fetchall()
                for (yil,) in yillar:
                    if not yil: continue
                    sayi=c.execute("SELECT COUNT(DISTINCT tc) FROM Idari_Cezalar WHERE yil=?",(yil,)).fetchone()[0]
                    toplam=c.execute("SELECT COALESCE(SUM(ipc_tutari),0) FROM Idari_Cezalar WHERE yil=?",(yil,)).fetchone()[0]
                    ilceler=c.execute("SELECT ilce||'('||COUNT(*)||')' FROM Idari_Cezalar WHERE yil=? GROUP BY ilce",(yil,)).fetchall()
                    konular=c.execute("SELECT konu||'('||COUNT(*)||')' FROM Idari_Cezalar WHERE yil=? GROUP BY konu",(yil,)).fetchall()
                    mukerrer=c.execute("SELECT COUNT(*) FROM (SELECT tc FROM Idari_Cezalar WHERE yil=? GROUP BY tc HAVING COUNT(*)>1)",(yil,)).fetchone()[0]
                    self.tv_cz_ozet.insert("","end",values=(yil,sayi,para_format(toplam),
                        ", ".join(r[0] for r in ilceler),", ".join(r[0] for r in konular),mukerrer))
        except Exception as e: logging.error(f"ceza_ozet:{e}")

    def _ist_ozet(self):
        if not DB_PATH or not MPL_OK: return
        try:
            with db_baglan() as c:
                veri={
                    "koy":c.execute("SELECT COUNT(*) FROM Mera_Varligi").fetchone()[0],
                    "rapor":c.execute("SELECT COUNT(*) FROM Rapor_Gecmisi").fetchone()[0],
                    "ihale":c.execute("SELECT COUNT(*) FROM Ihaleler WHERE durum NOT IN('Tamamlandı','İptal Edildi')").fetchone()[0],
                    "ceza":c.execute("SELECT COUNT(*) FROM Idari_Cezalar").fetchone()[0],
                }
        except Exception: return
        for key,k in self.ist_k.items(): k.set(veri.get(key,"—"))
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        for w in self.ist_canvas_frame.winfo_children(): w.destroy()
        fig,axes=plt.subplots(1,2,figsize=(12,4),facecolor=self.gc("bg"))
        fig.suptitle(f"MİRAS Genel Özet — {datetime.now().strftime('%d.%m.%Y')}",fontsize=12,fontweight="bold",color=self.gc("pri"))
        ax1=axes[0]
        labels=["Köy","Rapor","İhale","Ceza"]
        vals=[veri.get("koy",0),veri.get("rapor",0),veri.get("ihale",0),veri.get("ceza",0)]
        renkler=["#1E5631","#2980B9","#D68910","#C0392B"]
        ax1.bar(labels,[max(1,v) for v in vals],color=renkler)
        ax1.set_title("Genel Durum",color=self.gc("pri"),fontsize=10)
        ax2=axes[1]
        try:
            with db_baglan() as c:
                rows=c.execute("SELECT ilce,COUNT(*) FROM Mera_Varligi GROUP BY ilce ORDER BY COUNT(*) DESC").fetchall()
            if rows:
                ilceler=[r[0] for r in rows]; sayilar=[r[1] for r in rows]
                ax2.barh(ilceler,sayilar,color=self.gc("acc"))
                ax2.set_title("İlçe Bazlı Köy Dağılımı",color=self.gc("pri"),fontsize=10)
        except Exception: pass
        fig.tight_layout()
        canvas=FigureCanvasTkAgg(fig,master=self.ist_canvas_frame)
        canvas.draw(); canvas.get_tk_widget().pack(fill="both",expand=True)
        plt.close(fig)

    # ═══ 14. EVRAK ÜRETİCİ ══════════════════════════════════════════════════
    def _t_evrak(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        bf=ttk.Frame(nb2); kf=ttk.Frame(nb2); cf=ttk.Frame(nb2)
        nb2.add(bf,text="  📋 Bilgi Notu  "); nb2.add(kf,text="  📑 Katılım Evrakları  "); nb2.add(cf,text="  💰 Ceza Oluru  ")
        # Bilgi Notu
        bn=tk.Frame(bf,bg=bg); bn.pack(fill="both",expand=True,padx=20,pady=16)
        tk.Label(bn,text="📋 Bilgi Notu Oluşturucu",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        tk.Label(bn,text="Bilgi notunun bilgilerini girin, Word dosyası oluşturulsun.",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,8))
        self.bn_g={}
        for l,k,d in [("Sunulan Makam:","makam",""),("Konu:","konu",""),("İçerik:","icerik","")]:
            tk.Label(bn,text=l,bg=bg,font=("Segoe UI",10)).pack(anchor="w",pady=(6,0))
            if k=="icerik":
                w_obj=tk.Text(bn,height=10,font=("Segoe UI",10),wrap="word",relief="flat",bd=1,highlightbackground="#D0DDD8",highlightthickness=1,padx=10,pady=8)
            else:
                w_obj=ttk.Entry(bn,width=60); w_obj.insert(0,d)
            w_obj.pack(fill="x",pady=2); self.bn_g[k]=w_obj
        MBtn(bn,"📄 Word Olarak Kaydet",command=self._bilgi_notu_kaydet,color=self.gc("acc"),width=24).pack(anchor="w",pady=12)
        # Katılım Evrakları
        ke=tk.Frame(kf,bg=bg); ke.pack(fill="both",expand=True,padx=20,pady=16)
        tk.Label(ke,text="📑 İhale Katılım Evrakları Şablonu",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        tk.Label(ke,text="Başvuran bilgilerini girin, katılım evrakları Word dosyası oluşturulsun.",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,8))
        self.ke_g={}
        for l,k in [("Ad Soyad:","ad_soyad"),("TC No:","tc"),("Mera Adı:","mera"),("Ada/Parsel:","ada_parsel")]:
            tk.Label(ke,text=l,bg=bg,font=("Segoe UI",10)).pack(anchor="w",pady=(6,0))
            e=ttk.Entry(ke,width=40); e.pack(anchor="w"); self.ke_g[k]=e
        MBtn(ke,"📄 Word Olarak Kaydet",command=self._katilim_evrak_kaydet,color=self.gc("acc"),width=24).pack(anchor="w",pady=12)
        # Ceza Oluru
        co=tk.Frame(cf,bg=bg); co.pack(fill="both",expand=True,padx=20,pady=16)
        tk.Label(co,text="💰 İdari Para Cezası Oluru Şablonu",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        tk.Label(co,text="Ceza bilgilerini girin, olur Word dosyası oluşturulsun.",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,8))
        self.co_g={}
        for l,k in [("Ad Soyad:","ad_soyad"),("TC No:","tc"),("Mera/Köy:","mera"),("İhlal Konusu:","konu"),("Ceza Tutarı (₺):","tutar")]:
            tk.Label(co,text=l,bg=bg,font=("Segoe UI",10)).pack(anchor="w",pady=(6,0))
            e=ttk.Entry(co,width=40); e.pack(anchor="w"); self.co_g[k]=e
        MBtn(co,"📄 Word Olarak Kaydet",command=self._ceza_oluru_kaydet,color=C_DANGER,width=24).pack(anchor="w",pady=12)
        tk.Label(p,text="💡 AI ile otomatik evrak üretmek için: Sidebar → 🌿 MERA AI → AI Evrak Üret sekmesi",
            font=("Segoe UI",9),fg="#888",bg=bg).pack(anchor="w",padx=14,pady=(0,4))

    def _katilim_evrak_kaydet(self):
        if not DOCX_OK: messagebox.showerror("Hata","python-docx kurulu değil."); return
        dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile="Katilim_Evraklari.docx",filetypes=[("Word","*.docx")])
        if not dosya: return
        try:
            veri={k:e.get() for k,e in self.ke_g.items()}
            word_katilim_evrak(dosya,veri)
            messagebox.showinfo("Tamam",f"Evrak üretildi:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _ceza_oluru_kaydet(self):
        if not DOCX_OK: messagebox.showerror("Hata","python-docx kurulu değil."); return
        dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile="Ceza_Oluru.docx",filetypes=[("Word","*.docx")])
        if not dosya: return
        try:
            veri={k:e.get() for k,e in self.co_g.items()}
            veri["tarih"]=datetime.now().strftime("%d.%m.%Y"); veri["baba_adi"]=""; veri["dogum"]=""; veri["adres"]=""
            word_idari_ceza(dosya,veri)
            messagebox.showinfo("Tamam",f"Evrak üretildi:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _evrak_uret_ai(self):
        self.btn_ev.loading(True)
        prompt=(f"T.C. Ardahan Valiliği İl Tarım ve Orman Müdürlüğü adına resmi {self.ev_tur.get()} hazırla.\n"
               f"KONU: {self.ev_kon.get().strip()}\nTARİH: {datetime.now().strftime('%d.%m.%Y')}\nHAZIRLAYAN: {self.u_ad}\n"
               f"{'TALİMAT: '+self.ev_tal.get().strip() if self.ev_tal.get().strip() else ''}\n"
               "Türk resmi yazışma kurallarına uy.")
        def _g(t): self.ev_yan.config(state="normal"); self.ev_yan.delete("1.0",tk.END); self.ev_yan.insert(tk.END,t); self.ev_yan.config(state="disabled")
        _g("⏳ MERA AI hazırlıyor...\nKota aşımında otomatik yeniden denenecek, lütfen bekleyin.")
        def _bg():
            r=self._get_ai().tek(prompt)
            self.root.after(0,lambda:(_g(r),self.btn_ev.loading(False)))
        threading.Thread(target=_bg,daemon=True).start()

    def _bilgi_notu_kaydet(self):
        if not DOCX_OK: messagebox.showerror("Hata","python-docx kurulu değil."); return
        dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile="Bilgi_Notu.docx",filetypes=[("Word","*.docx")])
        if not dosya: return
        icerik=self.bn_g["icerik"].get("1.0",tk.END).strip() if isinstance(self.bn_g["icerik"],tk.Text) else self.bn_g["icerik"].get()
        try:
            word_bilgi_notu(dosya,{"makam":self.bn_g["makam"].get(),"konu":self.bn_g["konu"].get(),"icerik":icerik,"hazirlayan":self.u_ad})
            messagebox.showinfo("Tamam",f"Bilgi notu üretildi:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    # ═══ 15. MERA AI ═════════════════════════════════════════════════════════
    def _t_ai(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        chat_f=ttk.Frame(nb2); evrak_f=ttk.Frame(nb2); key_f=ttk.Frame(nb2)
        nb2.add(chat_f,text="  💬 Sohbet  "); nb2.add(evrak_f,text="  📝 AI Evrak Üret  ")
        if self.u_yetki=="Admin":
            nb2.add(key_f,text="  🔑 API Ayarları  ")
        # ── Durum bandı ──
        durum_renk="#2D8C55" if self._get_ai().hazir else "#C0392B"
        durum_txt="✅ MERA AI Aktif — Gemini bağlı" if self._get_ai().hazir else "❌ MERA AI Pasif — API Key gerekli"
        # ── 1) SOHBET ──
        cc=tk.Frame(chat_f,bg=bg); cc.pack(fill="both",expand=True,padx=14,pady=12)
        df=tk.Frame(cc,bg=durum_renk,height=30); df.pack(fill="x",pady=(0,8)); df.pack_propagate(False)
        tk.Label(df,text=durum_txt,fg=C_WHITE,bg=durum_renk,font=("Segoe UI",9,"bold")).pack(side="left",padx=10)
        if not self._get_ai().hazir and self.u_yetki=="Admin":
            tk.Button(df,text="🔑 API Key Gir →",command=lambda:nb2.select(key_f),fg=C_WHITE,bg=durum_renk,relief="flat",font=("Segoe UI",9,"bold"),cursor="hand2").pack(side="right",padx=10)
        self.chat_alan=scrolledtext.ScrolledText(cc,font=("Segoe UI",10),wrap="word",bg=C_WHITE,relief="flat",bd=1,
            highlightbackground="#D0DDD8",highlightthickness=1,padx=14,pady=12,state="disabled")
        self.chat_alan.pack(fill="both",expand=True)
        self.chat_alan.tag_configure("siz",foreground="#1A5276",font=("Segoe UI",10,"bold"))
        self.chat_alan.tag_configure("ai",foreground=self.gc("pri"))
        self.chat_alan.tag_configure("sis",foreground="#999",font=("Segoe UI",8,"italic"))
        hf=tk.Frame(cc,bg=bg); hf.pack(fill="x",pady=6)
        for s in ["Otlatma kapasitesi?","Tespit tahdit nedir?","Madde 14 süreci?","İhale şartları?","BBHB hesaplama?"]:
            tk.Button(hf,text=s,command=lambda q=s:self._chat_hizli(q),bg="#E8F5E9",fg=self.gc("pri"),font=("Segoe UI",8),relief="flat",cursor="hand2",padx=8,pady=3).pack(side="left",padx=4)
        gf=tk.Frame(cc,bg=bg); gf.pack(fill="x",pady=(6,0))
        self.chat_txt=tk.Text(gf,height=3,font=("Segoe UI",10),relief="flat",bd=1,highlightbackground="#A9DFBF",highlightthickness=1,padx=10,pady=8)
        self.chat_txt.pack(side="left",fill="x",expand=True,padx=(0,10))
        self.chat_txt.bind("<Return>",lambda e:(self._chat_gonder(),"break")[1])
        sg=tk.Frame(gf,bg=bg); sg.pack(side="left")
        self.btn_cg=MBtn(sg,"➤ Gönder",command=self._chat_gonder,color=self.gc("pri"),width=12); self.btn_cg.pack(pady=2)
        MBtn(sg,"🗑 Temizle",command=self._chat_temizle,color="#888",width=12).pack(pady=2)
        self._chat_yaz("sis","🌿 MERA AI'ye Hoş Geldiniz!\nMera mevzuatı, BBHB hesaplama, ihale ve tahsis süreçleri hakkında sorular sorabilirsiniz.\n"+"─"*50+"\n")
        if not self._get_ai().hazir:
            self._chat_yaz("sis","⚠️ AI bağlı değil. Üstteki '🔑 API Key Gir →' butonuna tıklayın.\n"
                "veya MERA AI → API Ayarları sekmesinden key girin.\n"
                "Ücretsiz key almak için: aistudio.google.com/app/apikey\n"+"─"*50+"\n")
        # ── 2) AI EVRAK ÜRET ──
        ee=tk.Frame(evrak_f,bg=bg); ee.pack(fill="both",expand=True,padx=14,pady=12)
        tk.Label(ee,text="✨ Yapay Zeka ile Otomatik Evrak Üretici",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,8))
        tk.Label(ee,text="Evrak türünü ve konuyu seçin, MERA AI resmi formatta hazırlasın.",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,10))
        ctrl=tk.Frame(ee,bg=bg); ctrl.pack(fill="x",pady=(0,10))
        for i,(l,attr,vals,d,w) in enumerate([
            ("Evrak Türü:","ev_tur",["Bilgi Notu","Makam Oluru","Mera Tespit Tutanağı","Otlatma Yasağı Kararı","Islah Proje Talebi","Şikayet Değerlendirme","İhale Duyurusu"],"Bilgi Notu",24),
            ("Konu/Yer:","ev_kon",None,"",28),("Ek Talimat:","ev_tal",None,"",20)]):
            tk.Label(ctrl,text=l,bg=bg,font=("Segoe UI",10)).grid(row=0,column=i*2,padx=(0,4),sticky="w")
            w_obj=(ttk.Combobox(ctrl,values=vals,state="readonly",width=w) if vals else ttk.Entry(ctrl,width=w))
            if vals: w_obj.set(d)
            w_obj.grid(row=0,column=i*2+1,padx=(0,10)); setattr(self,attr,w_obj)
        self.btn_ev=MBtn(ctrl,"✨ MERA AI Üret",command=self._evrak_uret_ai,color=self.gc("pri")); self.btn_ev.grid(row=0,column=6,padx=10)
        self.ev_yan=scrolledtext.ScrolledText(ee,font=("Segoe UI",10),wrap="word",bg="#F9F9FF",relief="flat",bd=1,highlightbackground="#C8B4E8",highlightthickness=1,padx=14,pady=12,state="disabled")
        self.ev_yan.pack(fill="both",expand=True)
        bf_ev=tk.Frame(ee,bg=bg); bf_ev.pack(pady=6)
        MBtn(bf_ev,"📋 Kopyala",command=lambda:(self.root.clipboard_clear(),self.root.clipboard_append(self.ev_yan.get("1.0",tk.END))),color="#555",width=14).pack(side="left",padx=4)
        MBtn(bf_ev,"📄 Word Olarak Kaydet",command=self._ai_evrak_word,color=self.gc("acc"),width=20).pack(side="left",padx=4)
        # ── 3) API AYARLARI (sadece Admin) ──
        if self.u_yetki=="Admin":
            aa=tk.Frame(key_f,bg=bg); aa.pack(fill="both",expand=True,padx=40,pady=30)
            tk.Label(aa,text="🔑 MERA AI — Gemini API Anahtarı",font=("Segoe UI",16,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")
            tk.Label(aa,text="Google AI Studio'dan ücretsiz API key alın ve aşağıya yapıştırın.",font=("Segoe UI",10),fg="#555",bg=bg).pack(anchor="w",pady=(6,4))
            # Adım adım rehber
            rehber_f=tk.Frame(aa,bg="#F0F4F2",highlightbackground="#D0DDD8",highlightthickness=1,padx=16,pady=12)
            rehber_f.pack(fill="x",pady=(8,16))
            tk.Label(rehber_f,text="📌 API Key Nasıl Alınır?",font=("Segoe UI",11,"bold"),fg=self.gc("pri"),bg="#F0F4F2").pack(anchor="w")
            for adim in ["1️⃣  aistudio.google.com/app/apikey adresine gidin",
                         "2️⃣  Google hesabınızla giriş yapın",
                         "3️⃣  'Create API Key' butonuna tıklayın",
                         "4️⃣  Oluşan anahtarı kopyalayıp aşağıya yapıştırın"]:
                tk.Label(rehber_f,text=adim,font=("Segoe UI",10),fg="#333",bg="#F0F4F2").pack(anchor="w",pady=2)
            tk.Label(aa,text="API Anahtarı:",bg=bg,font=("Segoe UI",11,"bold")).pack(anchor="w",pady=(0,4))
            self.ai_key_e=ttk.Entry(aa,width=60,show="•",font=("Segoe UI",11))
            if self._get_ai().api_key: self.ai_key_e.insert(0,self._get_ai().api_key)
            self.ai_key_e.pack(anchor="w",pady=6,fill="x")
            # Durum göstergesi
            dur_f=tk.Frame(aa,bg=bg); dur_f.pack(anchor="w",pady=(4,12))
            self.lbl_ai_durum=tk.Label(dur_f,text=durum_txt,font=("Segoe UI",10,"bold"),fg=durum_renk,bg=bg)
            self.lbl_ai_durum.pack(side="left")
            def _kk():
                k=self.ai_key_e.get().strip()
                if not k: messagebox.showwarning("Uyarı","API anahtarı boş."); return
                self._get_ai().key_kaydet(k)
                self.lbl_ai_durum.config(text="⏳ Bağlanıyor...",fg="#D68910")
                def _kontrol():
                    if self._get_ai().hazir:
                        self.lbl_ai_durum.config(text="✅ MERA AI Aktif — Gemini bağlı!",fg="#2D8C55")
                        messagebox.showinfo("MERA AI","✅ Başarıyla aktifleştirildi!\n\nArtık Sohbet ve AI Evrak Üret sekmelerini kullanabilirsiniz.")
                    else:
                        self.lbl_ai_durum.config(text="❌ Bağlantı kurulamadı — anahtarı kontrol edin",fg="#C0392B")
                        messagebox.showwarning("MERA AI","Bağlantı kurulamadı.\nAPI anahtarını kontrol edin.")
                self.root.after(3000,_kontrol)
            MBtn(aa,"✅ Aktifleştir",command=_kk,color=self.gc("pri"),width=30).pack(anchor="w",pady=4)
            tk.Label(aa,text="💡 API anahtarı miras_gemini_key.txt dosyasına kaydedilir.\nTüm kullanıcılar aynı anahtarı kullanır.",font=("Segoe UI",9),fg="#888",bg=bg).pack(anchor="w",pady=(12,0))

    def _ai_evrak_word(self):
        """AI tarafından üretilen evrakı Word olarak kaydet"""
        icerik=self.ev_yan.get("1.0",tk.END).strip()
        if not icerik or "⏳" in icerik: messagebox.showwarning("Uyarı","Önce evrak üretin."); return
        if not DOCX_OK: messagebox.showerror("Hata","python-docx kurulu değil."); return
        dosya=filedialog.asksaveasfilename(defaultextension=".docx",initialfile=f"AI_Evrak_{datetime.now().strftime('%Y%m%d')}.docx",filetypes=[("Word","*.docx")])
        if not dosya: return
        try:
            doc=DocxDocument()
            for satir in icerik.split("\n"):
                if satir.strip(): doc.add_paragraph(satir)
            doc.save(dosya)
            messagebox.showinfo("Tamam",f"Evrak kaydedildi:\n{dosya}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _chat_yaz(self,tag,metin):
        self.chat_alan.config(state="normal"); self.chat_alan.insert(tk.END,metin+"\n",tag)
        self.chat_alan.see(tk.END); self.chat_alan.config(state="disabled")

    def _chat_hizli(self,s): self.chat_txt.delete("1.0",tk.END); self.chat_txt.insert("1.0",s); self._chat_gonder()

    def _chat_gonder(self):
        m=self.chat_txt.get("1.0",tk.END).strip()
        if not m: return
        self.chat_txt.delete("1.0",tk.END); self._chat_yaz("siz",f"👤 Siz:\n{m}"); self.btn_cg.loading(True)
        self._chat_yaz("sis","⏳ Yanıt bekleniyor... (kota aşımında otomatik yeniden denenecek)")
        def _bg():
            r=self._get_ai().sor(m)
            self.root.after(0,lambda:(self._chat_yaz("ai",f"🌿 MERA AI:\n{r}\n{'─'*40}"),self.btn_cg.loading(False)))
        threading.Thread(target=_bg,daemon=True).start()

    def _chat_temizle(self):
        self._get_ai().yeni_chat()
        self.chat_alan.config(state="normal"); self.chat_alan.delete("1.0",tk.END); self.chat_alan.config(state="disabled")

    # ═══ 16. MEVZUAT REHBERİ ═════════════════════════════════════════════════
    def _t_mevzuat(self,p):
        bg=self.gc("bg"); f=tk.Frame(p,bg=bg); f.pack(fill="both",expand=True,padx=20,pady=16)
        tk.Label(f,text="📖 Mevzuat Rehberi & Bilgi Kartları",font=("Segoe UI",15,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,4))
        tk.Label(f,text="4342 sayılı Mera Kanunu ve Mera Yönetmeliği kapsamında temel kavramlar",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,12))
        # Scrollable kartlar
        cw=tk.Canvas(f,bg=bg,highlightthickness=0); vsb=ttk.Scrollbar(f,orient="vertical",command=cw.yview)
        cw.configure(yscrollcommand=vsb.set); vsb.pack(side="right",fill="y"); cw.pack(side="left",fill="both",expand=True)
        inner=tk.Frame(cw,bg=bg); wid=cw.create_window((0,0),window=inner,anchor="nw")
        cw.bind("<Configure>",lambda e:cw.itemconfig(wid,width=e.width))
        inner.bind("<Configure>",lambda e:cw.configure(scrollregion=cw.bbox("all")))
        for key,kart in MEVZUAT_KARTLARI.items():
            card=tk.Frame(inner,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1,padx=16,pady=12)
            card.pack(fill="x",pady=6,padx=4)
            tk.Label(card,text=kart["baslik"],font=("Segoe UI",12,"bold"),fg=self.gc("pri"),bg=C_WHITE,anchor="w").pack(fill="x")
            tk.Label(card,text=kart["tanim"],font=("Segoe UI",10),fg="#333",bg=C_WHITE,wraplength=800,justify="left",anchor="w").pack(fill="x",pady=(4,2))
            tk.Label(card,text=f"📌 {kart['kanun']}",font=("Segoe UI",9,"italic"),fg=C_INFO,bg=C_WHITE,anchor="w").pack(fill="x")
            tk.Label(card,text=kart["detay"],font=("Segoe UI",9),fg="#555",bg=C_WHITE,wraplength=800,justify="left",anchor="w").pack(fill="x",pady=(4,0))

    # ═══ 17. İLETİŞİM ════════════════════════════════════════════════════════
    def _t_iletisim(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        bf=ttk.Frame(nb2); ff=ttk.Frame(nb2); wf=ttk.Frame(nb2)
        nb2.add(bf,text="  📞 Bilgiler  "); nb2.add(ff,text="  📝 İletişim Formu  "); nb2.add(wf,text="  💬 WhatsApp Şablon  ")
        if self.u_yetki=="Admin":
            af=ttk.Frame(nb2); nb2.add(af,text="  ⚙️ Düzenle  ")
        # WhatsApp Mesaj Şablonları
        ww=tk.Frame(wf,bg=bg); ww.pack(fill="both",expand=True,padx=20,pady=14)
        tk.Label(ww,text="💬 WhatsApp Mesaj Şablonları",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")
        tk.Label(ww,text="Şablon seçin, kişi bilgilerini doldurun → mesaj otomatik oluşsun → kopyala veya WhatsApp'ta aç",
            font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(2,10))
        sablonlar={
            "İhale Kazanan Bildirim":f"Sayın {{AD}},\n\n{KURUM_IL} İli {{İLÇE}} İlçesi {{KÖY}} köyü mera kiralama ihalesini kazandığınızı bildiririz.\n\nİhale Bedeli: {{BEDEL}}\nKesin Teminat (%6): İhale bedelinin %6'sını {IBAN_EMANET} IBAN numaralı hesaba yatırmanız gerekmektedir.\n\nSözleşme için 15 gün içinde müdürlüğümüze başvurunuz.\n\n{KURUM_SUBE}",
            "Ceza Tebligatı":f"Sayın {{AD}},\n\n{GENEL_EMRI_GECMIS.get(int(_yil()),{}).get('sayi','')} sayılı Valilik Genel Emri gereğince {{KÖY}} mera alanında yapılan denetimde {{KONU}} fiilinden dolayı {{TUTAR}} İdari Para Cezası uygulanmıştır.\n\nÖdeme: {ODEME_YERI}\nSon ödeme: Tebliğden itibaren 30 gün\nİtiraz: {ITIRAZ_MERCI} (15 gün)\n\n{KURUM_SUBE}",
            "Süre Dolum Uyarısı":f"Sayın {{AD}},\n\n{KURUM_IL} İli {{KÖY}} köyü mera kiralama süreniz {{TARİH}} tarihinde dolmaktadır.\n\nMerayı İl Tarım ve Orman Müdürlüğü'ne teslim etmeniz gerekmektedir. Gecikmeler halinde günlük kira bedelinin %2'si kadar ceza uygulanır.\n\n{KURUM_SUBE}",
            "Genel Duyuru":"Sayın {{AD}},\n\n{{MESAJ}}\n\nBilgilerinize sunarız.\n\nÇayır, Mera ve Yem Bitkileri Şube Müdürlüğü",
        }
        tk.Label(ww,text="Şablon:",bg=bg,font=("Segoe UI",10,"bold")).pack(anchor="w")
        self.wa_sablon=ttk.Combobox(ww,values=list(sablonlar.keys()),state="readonly",width=40)
        self.wa_sablon.pack(anchor="w",pady=(0,8))
        tk.Label(ww,text="Telefon No (başında 90):",bg=bg,font=("Segoe UI",10)).pack(anchor="w")
        self.wa_tel=ttk.Entry(ww,width=20); self.wa_tel.pack(anchor="w",pady=(0,8))
        tk.Label(ww,text="Mesaj Önizleme (değişkenleri doldurun):",bg=bg,font=("Segoe UI",10,"bold")).pack(anchor="w")
        self.wa_mesaj=tk.Text(ww,height=10,font=("Segoe UI",10),wrap="word",relief="flat",bd=1,
            highlightbackground="#D0DDD8",highlightthickness=1,padx=10,pady=8)
        self.wa_mesaj.pack(fill="both",expand=True,pady=(0,8))
        def _sablon_sec(event=None):
            s=self.wa_sablon.get()
            if s in sablonlar:
                self.wa_mesaj.delete("1.0",tk.END); self.wa_mesaj.insert("1.0",sablonlar[s])
        self.wa_sablon.bind("<<ComboboxSelected>>",_sablon_sec)
        bff=tk.Frame(ww,bg=bg); bff.pack(anchor="w",pady=6)
        def _wa_gonder():
            tel=self.wa_tel.get().strip(); mesaj=self.wa_mesaj.get("1.0",tk.END).strip()
            if not tel or not mesaj: messagebox.showwarning("Uyarı","Telefon ve mesaj zorunlu."); return
            import urllib.parse
            url=f"https://wa.me/{tel}?text={urllib.parse.quote(mesaj)}"
            import webbrowser; webbrowser.open(url)
        def _kopyala():
            mesaj=self.wa_mesaj.get("1.0",tk.END).strip()
            if mesaj:
                self.root.clipboard_clear(); self.root.clipboard_append(mesaj)
                messagebox.showinfo("Tamam","Mesaj panoya kopyalandı.")
        MBtn(bff,"💬 WhatsApp'ta Aç",command=_wa_gonder,color="#25D366",width=18).pack(side="left",padx=4)
        MBtn(bff,"📋 Kopyala",command=_kopyala,color=C_INFO,width=12).pack(side="left",padx=4)
        # Bilgiler
        bi=tk.Frame(bf,bg=bg); bi.pack(fill="both",expand=True,padx=40,pady=30)
        tk.Label(bi,text="📞 Geliştirici İletişim",font=("Segoe UI",18,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,20))
        card=tk.Frame(bi,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1,padx=30,pady=24); card.pack(fill="x")
        # DB'den iletişim bilgilerini çek
        try:
            with db_baglan() as c:
                ilet={r[0]:r[1] for r in c.execute("SELECT anahtar,deger FROM Iletisim_Bilgileri")}
        except Exception: ilet={"telefon":DEV_TEL,"email":DEV_MAIL,"whatsapp":DEV_WA}
        for lbl,val,ikon in [("Ad Soyad",DEV_ADI,"👤"),("Ünvan",DEV_UNVAN,"🎓"),
                              ("E-posta",ilet.get("email",DEV_MAIL),"📧"),
                              ("Telefon",ilet.get("telefon",DEV_TEL),"📱"),
                              ("Program",f"{PROG_ADI} {VERSIYON}","💻")]:
            rf=tk.Frame(card,bg=C_WHITE); rf.pack(fill="x",pady=6)
            tk.Label(rf,text=f"{ikon}  {lbl}:",font=("Segoe UI",11,"bold"),fg="#333",bg=C_WHITE,width=14,anchor="w").pack(side="left")
            tk.Label(rf,text=val,font=("Segoe UI",11),fg=self.gc("pri"),bg=C_WHITE).pack(side="left",padx=8)
        # WhatsApp butonu
        wa=ilet.get("whatsapp",DEV_WA)
        def _wa():
            if messagebox.askyesno("WhatsApp","Sayfadan ayrılıp WhatsApp Web açılacak.\nDevam?"):
                import webbrowser; webbrowser.open(f"https://wa.me/{wa}")
        MBtn(bi,"💬 WhatsApp ile Mesaj Gönder",command=_wa,color="#25D366",width=28).pack(anchor="w",pady=16)
        tk.Label(bi,text="⚠️ Hata bildirimi için ekran görüntüsünü ve miras_debug.log dosyasını gönderin.",
            font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w")
        # İletişim formu
        fm=tk.Frame(ff,bg=bg); fm.pack(fill="both",expand=True,padx=30,pady=20)
        tk.Label(fm,text="📝 İstek / Şikayet / Bilgi Edinme Formu",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        tk.Label(fm,text=f"Gönderen: {self.u_ad}",font=("Segoe UI",10),fg="#555",bg=bg).pack(anchor="w",pady=(0,8))
        tk.Label(fm,text="Konu:",bg=bg,font=("Segoe UI",10,"bold")).pack(anchor="w")
        self.il_konu=ttk.Combobox(fm,values=["İstek","Şikayet","Bilgi Edinme","Öneri","Hata Bildirimi"],state="readonly",width=28)
        self.il_konu.set("İstek"); self.il_konu.pack(anchor="w",pady=(0,10))
        tk.Label(fm,text="Mesaj:",bg=bg,font=("Segoe UI",10,"bold")).pack(anchor="w")
        self.il_mesaj=tk.Text(fm,height=8,font=("Segoe UI",10),wrap="word",relief="flat",bd=1,
            highlightbackground="#D0DDD8",highlightthickness=1,padx=10,pady=8)
        self.il_mesaj.pack(fill="x",pady=(0,10))
        MBtn(fm,"📨 Gönder",command=self._iletisim_gonder,color=self.gc("acc"),width=20).pack(anchor="w")
        # Admin düzenle
        if self.u_yetki=="Admin":
            ae=tk.Frame(af,bg=bg); ae.pack(fill="both",expand=True,padx=40,pady=30)
            tk.Label(ae,text="⚙️ İletişim Bilgilerini Düzenle",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,14))
            self.il_ed={}
            for l,k,d in [("Telefon:","telefon",ilet.get("telefon","")),("E-posta:","email",ilet.get("email","")),("WhatsApp No:","whatsapp",ilet.get("whatsapp",""))]:
                tk.Label(ae,text=l,bg=bg,font=("Segoe UI",10,"bold")).pack(anchor="w",pady=(8,0))
                e=ttk.Entry(ae,width=40); e.insert(0,d); e.pack(anchor="w"); self.il_ed[k]=e
            MBtn(ae,"💾 Kaydet",command=self._iletisim_kaydet,color=C_INFO,width=20).pack(anchor="w",pady=14)

    def _iletisim_gonder(self):
        mesaj=self.il_mesaj.get("1.0",tk.END).strip()
        if not mesaj: messagebox.showwarning("Uyarı","Mesaj boş."); return
        with db_baglan() as c:
            c.execute("INSERT INTO Iletisim_Formu(gonderen,konu_tipi,mesaj,tarih)VALUES(?,?,?,?)",
                (self.u_ad,self.il_konu.get(),mesaj,datetime.now().strftime("%d.%m.%Y %H:%M")))
        db_log(self.u_id,"İletişim Formu",self.il_konu.get())
        self.il_mesaj.delete("1.0",tk.END)
        messagebox.showinfo("Tamam","Mesajınız Admin'e iletildi.")

    def _iletisim_kaydet(self):
        with db_baglan() as c:
            for k,e in self.il_ed.items():
                c.execute("INSERT OR REPLACE INTO Iletisim_Bilgileri(anahtar,deger)VALUES(?,?)",(k,e.get()))
        messagebox.showinfo("Tamam","İletişim bilgileri güncellendi.")

    def _db_degistir(self,yeni=False):
        """Çoklu DB — farklı veritabanına geç"""
        global DB_PATH
        if yeni:
            yol=filedialog.asksaveasfilename(defaultextension=".db",initialfile=f"miras_{_yil()}.db",filetypes=[("SQLite","*.db")])
        else:
            yol=filedialog.askopenfilename(filetypes=[("SQLite","*.db")])
        if not yol: return
        if messagebox.askyesno("Veritabanı Değiştir",f"Veritabanı değiştirilecek:\n{yol}\n\nProgram yeniden başlatılacak. Devam?"):
            try:
                if self.yedekci: self.yedekci.al(otomatik=True)
            except Exception: pass
            DB_PATH=yol
            with open(CONFIG_FILE,"w",encoding="utf-8") as f:
                data={}
                try:
                    with open(CONFIG_FILE) as ff: data=json.load(ff)
                except Exception: pass
                data["db_path"]=DB_PATH
                json.dump(data,f)
            init_db()
            db_log(self.u_id,"DB Değiştir",yol)
            self._cikis()
            messagebox.showinfo("Tamam",f"Veritabanı değiştirildi:\n{yol}\nLütfen tekrar giriş yapın.")

    def _db_arsivle(self):
        """Mevcut DB'yi kopyalayarak arşivle"""
        if not DB_PATH: return
        ts=datetime.now().strftime("%Y%m%d_%H%M")
        varsayilan=DB_PATH.replace(".db",f"_arsiv_{ts}.db")
        yol=filedialog.asksaveasfilename(defaultextension=".db",initialfile=os.path.basename(varsayilan),filetypes=[("SQLite","*.db")])
        if yol:
            try:
                shutil.copy2(DB_PATH,yol)
                db_log(self.u_id,"DB Arşiv",yol)
                messagebox.showinfo("Tamam",f"Veritabanı arşivlendi:\n{yol}\nBoyut: {os.path.getsize(yol)/1024:.0f} KB")
            except Exception as e: messagebox.showerror("Hata",str(e))

    # ═══ KAYIT FOTOĞRAFLARI (saha denetim, ceza, ihale) ═══════════════════════
    def _kayit_foto_yonet(self,tablo,kayit_id,baslik=""):
        """İhale/Ceza/Şikayet kaydına fotoğraf ekle, listele, sil"""
        win=tk.Toplevel(self.root); win.title(f"📷 Fotoğraflar — {baslik}"); win.geometry("720x540")
        win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        tk.Label(win,text=f"📷 {baslik} — Fotoğraf Yönetimi",font=("Segoe UI",13,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=(12,4))
        tk.Label(win,text="Saha fotoğraflarını bu kayda bağlayın. Fotoğraflar 'miras_fotolar' klasöründe saklanır.",
            font=("Segoe UI",9),fg="#666",bg=C_WHITE).pack(pady=(0,8))
        foto_dir=Path("miras_fotolar"); foto_dir.mkdir(exist_ok=True)
        # Liste
        tv=ttk.Treeview(win,columns=("id","dosya","aciklama","tarih","ekleyen"),show="headings",height=10)
        for c,w,b in [("id",40,"#"),("dosya",180,"Dosya"),("aciklama",220,"Açıklama"),("tarih",100,"Tarih"),("ekleyen",120,"Ekleyen")]:
            tv.heading(c,text=b); tv.column(c,width=w)
        tv.pack(fill="both",expand=True,padx=14,pady=6)
        def _yenile():
            tv.delete(*tv.get_children())
            try:
                with db_baglan() as c:
                    for r in c.execute("SELECT id,dosya_yolu,aciklama,ekleme_tarihi,ekleyen FROM Kayit_Fotolari WHERE tablo=? AND kayit_id=? ORDER BY id DESC",(tablo,kayit_id)).fetchall():
                        dosya_adi=os.path.basename(r[1]) if r[1] else ""
                        tv.insert("","end",values=(r[0],dosya_adi,r[2] or "",r[3] or "",r[4] or ""))
            except Exception as e: logging.error(f"foto_yenile:{e}")
        _yenile()
        def _ekle():
            yol=filedialog.askopenfilename(parent=win,filetypes=[("Resim","*.jpg *.jpeg *.png *.bmp"),("Tümü","*.*")])
            if not yol: return
            aciklama=tkinter.simpledialog.askstring("Açıklama","Bu fotoğraf için kısa açıklama (opsiyonel):",parent=win) or ""
            try:
                ts=datetime.now().strftime("%Y%m%d_%H%M%S")
                uzanti=os.path.splitext(yol)[1]
                yeni_ad=f"{tablo}_{kayit_id}_{ts}{uzanti}"
                hedef=foto_dir/yeni_ad
                shutil.copy2(yol,hedef)
                with db_baglan() as c:
                    c.execute("INSERT INTO Kayit_Fotolari(tablo,kayit_id,dosya_yolu,aciklama,ekleme_tarihi,ekleyen)VALUES(?,?,?,?,?,?)",
                        (tablo,kayit_id,str(hedef),aciklama,datetime.now().strftime("%d.%m.%Y %H:%M"),self.u_ad))
                db_log(self.u_id,"Foto Ekle",f"{tablo}#{kayit_id}")
                _yenile()
                messagebox.showinfo("Tamam","Fotoğraf eklendi.",parent=win)
            except Exception as e: messagebox.showerror("Hata",str(e),parent=win)
        def _goster():
            sel=tv.selection()
            if not sel: return
            fid=tv.item(sel[0])["values"][0]
            try:
                with db_baglan() as c:
                    r=c.execute("SELECT dosya_yolu FROM Kayit_Fotolari WHERE id=?",(fid,)).fetchone()
                if r and os.path.exists(r[0]):
                    os.startfile(r[0]) if sys.platform=="win32" else os.system(f'xdg-open "{r[0]}"')
                else:
                    messagebox.showwarning("Bulunamadı","Dosya bulunamadı.",parent=win)
            except Exception as e: messagebox.showerror("Hata",str(e),parent=win)
        def _sil():
            sel=tv.selection()
            if not sel: return
            if not messagebox.askyesno("Onay","Fotoğraf silinsin mi?",parent=win): return
            fid=tv.item(sel[0])["values"][0]
            try:
                with db_baglan() as c:
                    r=c.execute("SELECT dosya_yolu FROM Kayit_Fotolari WHERE id=?",(fid,)).fetchone()
                    c.execute("DELETE FROM Kayit_Fotolari WHERE id=?",(fid,))
                if r and os.path.exists(r[0]):
                    try: os.remove(r[0])
                    except Exception: pass
                db_log(self.u_id,"Foto Sil",f"{tablo}#{kayit_id}")
                _yenile()
            except Exception as e: messagebox.showerror("Hata",str(e),parent=win)
        bf=tk.Frame(win,bg=C_WHITE); bf.pack(pady=10)
        MBtn(bf,"📷 Fotoğraf Ekle",command=_ekle,color=self.gc("acc"),width=18).pack(side="left",padx=4)
        MBtn(bf,"👁️ Göster",command=_goster,color=C_INFO,width=12).pack(side="left",padx=4)
        MBtn(bf,"🗑 Sil",command=_sil,color=C_DANGER,width=10).pack(side="left",padx=4)

    # ═══ PROFİL FOTOĞRAFI ═════════════════════════════════════════════════════
    def _profil_foto_goster(self,k_adi=None,parent=None):
        """Profil fotoğrafı göster — WhatsApp tarzı yuvarlak avatar"""
        hedef=k_adi or self.u_id
        foto_dir=Path("miras_profil"); foto_dir.mkdir(exist_ok=True)
        foto_yol=foto_dir/f"{hedef}.png"
        win=tk.Toplevel(parent or self.root); win.title(f"Profil — {hedef}"); win.geometry("340x420")
        win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        # Fotoğraf alanı
        foto_frame=tk.Frame(win,bg=C_WHITE); foto_frame.pack(pady=20)
        self._profil_img=None
        if foto_yol.exists():
            try:
                from PIL import Image, ImageTk, ImageDraw
                img=Image.open(str(foto_yol)).resize((160,160))
                # Yuvarlak mask
                mask=Image.new("L",(160,160),0)
                ImageDraw.Draw(mask).ellipse((0,0,160,160),fill=255)
                img.putalpha(mask)
                self._profil_img=ImageTk.PhotoImage(img)
                tk.Label(foto_frame,image=self._profil_img,bg=C_WHITE).pack()
            except ImportError:
                # PIL yoksa düz göster
                try:
                    self._profil_img=tk.PhotoImage(file=str(foto_yol))
                    tk.Label(foto_frame,image=self._profil_img,bg=C_WHITE).pack()
                except Exception:
                    tk.Label(foto_frame,text="👤",font=("Segoe UI",60),bg="#E0E0E0",fg="#888",width=6,height=2).pack()
        else:
            tk.Label(foto_frame,text="👤",font=("Segoe UI",60),bg="#E0E0E0",fg="#888",width=6,height=2).pack()
        # İsim ve ünvan
        try:
            with db_baglan() as c:
                row=c.execute("SELECT ad,unvan,yetki FROM Kullanicilar WHERE k_adi=?",(hedef,)).fetchone()
            if row:
                tk.Label(win,text=row[0],font=("Segoe UI",16,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=(4,0))
                tk.Label(win,text=f"{row[1] or ''} — {row[2]}",font=("Segoe UI",10),fg="#666",bg=C_WHITE).pack()
        except Exception: pass
        # Butonlar
        if hedef==self.u_id:
            def _foto_sec():
                yol=filedialog.askopenfilename(filetypes=[("Resim","*.png *.jpg *.jpeg *.gif *.bmp")])
                if not yol: return
                try:
                    # Kopyala ve yeniden boyutlandır
                    try:
                        from PIL import Image
                        img=Image.open(yol).resize((200,200)); img.save(str(foto_yol))
                    except ImportError:
                        shutil.copy2(yol,str(foto_yol))
                    db_log(self.u_id,"Profil Foto","Güncellendi")
                    messagebox.showinfo("Tamam","Profil fotoğrafı güncellendi.",parent=win)
                    win.destroy(); self._profil_foto_goster(hedef)
                except Exception as e: messagebox.showerror("Hata",str(e),parent=win)
            def _foto_sil():
                if foto_yol.exists():
                    foto_yol.unlink()
                    messagebox.showinfo("Tamam","Fotoğraf kaldırıldı.",parent=win)
                    win.destroy(); self._profil_foto_goster(hedef)
            MBtn(win,"📸 Fotoğraf Seç",command=_foto_sec,color=self.gc("acc"),width=20).pack(pady=6)
            MBtn(win,"🗑 Fotoğrafı Kaldır",command=_foto_sil,color="#888",width=20).pack()

    # ═══ 18. AYARLAR ═════════════════════════════════════════════════════════
    def _t_ayarlar(self,p):
        bg=self.gc("bg"); nb2=ttk.Notebook(p); nb2.pack(fill="both",expand=True,padx=10,pady=10)
        gf=ttk.Frame(nb2); sf=ttk.Frame(nb2); yf=ttk.Frame(nb2); df=ttk.Frame(nb2); gef=ttk.Frame(nb2)
        nb2.add(gf,text="  🎨 Görünüm  "); nb2.add(sf,text="  🔑 Şifre  "); nb2.add(yf,text="  💾 Yedekleme  ")
        nb2.add(df,text="  🗄️ Veritabanı  "); nb2.add(gef,text="  📜 Genel Emir  ")
        # Her sekmeyi kendi builder metodu ile doldur
        self._build_gorunum(gf,bg)
        self._build_sifre(sf,bg)
        self._build_yedekleme(yf,bg)
        self._build_db_yonetimi(df,bg)
        # Genel Emir — yıl bazlı ceza güncelleme
        ge=tk.Frame(gef,bg=bg); ge.pack(fill="both",expand=True,padx=30,pady=20)
        tk.Label(ge,text="📜 Valilik Genel Emri — İdari Para Cezası Miktarları",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")
        tk.Label(ge,text="Her yıl Valilik Genel Emri ile belirlenen ceza miktarlarını buradan güncelleyin.\n"
            "Güncelleme yapıldığında tüm Word evrakları ve ceza hesaplamaları otomatik güncellenir.",
            font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(4,14))
        # Mevcut geçmiş tablosu
        tk.Label(ge,text="Yıl Bazlı Ceza Geçmişi:",font=("Segoe UI",11,"bold"),bg=bg).pack(anchor="w",pady=(0,4))
        self.tv_ge=self._tv(ge,[("yil",60,"Yıl"),("sayi",120,"Genel Emir Sayı"),("bb",100,"BB Ceza (₺)"),("kb",100,"KB Ceza (₺)"),("madde",200,"Madde/Fıkra")],5)
        self._ge_yenile()
        # Güncelleme formu
        if self.u_yetki=="Admin":
            uf=tk.LabelFrame(ge,text="  Yeni Yıl / Güncelleme  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=14,pady=10)
            uf.pack(fill="x",pady=10)
            self.ge_g={}
            for i,(l,k,d,w) in enumerate([("Yıl:","yil",str(datetime.now().year),6),
                ("Genel Emir Sayı:","sayi",f"{datetime.now().year}/01",12),
                ("BB Ceza (₺):","bb",f"{CEZA_BB_TL:.2f}",10),("KB Ceza (₺):","kb",f"{CEZA_KB_TL:.2f}",10),
                ("Madde/Fıkra:","madde","18. Madde 9. Fıkra",24)]):
                tk.Label(uf,text=l,bg=bg,font=("Segoe UI",10)).grid(row=0,column=i*2,padx=(0,4),sticky="w")
                e=ttk.Entry(uf,width=w); e.insert(0,d); e.grid(row=0,column=i*2+1,padx=(0,10)); self.ge_g[k]=e
            def _ge_kaydet():
                try:
                    yil=int(self.ge_g["yil"].get())
                    sayi=self.ge_g["sayi"].get().strip()
                    bb=para_parse(self.ge_g["bb"].get()); kb=para_parse(self.ge_g["kb"].get())
                    madde=self.ge_g["madde"].get().strip()
                    if bb<=0 or kb<=0: messagebox.showerror("Hata","Ceza miktarları 0'dan büyük olmalı."); return
                    if not sayi: messagebox.showerror("Hata","Genel Emir sayısı zorunlu."); return
                    with db_baglan() as c:
                        c.execute("INSERT OR REPLACE INTO Genel_Emri(yil,sayi,bb_ceza,kb_ceza,madde,guncelleme_tarihi)VALUES(?,?,?,?,?,?)",
                            (yil,sayi,bb,kb,madde,datetime.now().strftime("%d.%m.%Y")))
                    # Global sabitleri güncelle (aktif yıl için)
                    global CEZA_BB_TL, CEZA_KB_TL
                    if yil==datetime.now().year:
                        CEZA_BB_TL=bb; CEZA_KB_TL=kb
                    db_log(self.u_id,"Genel Emir Güncelle",f"{sayi}: BB={bb:.2f}₺ KB={kb:.2f}₺ ({madde})")
                    self._ge_yenile()
                    messagebox.showinfo("Tamam",f"{yil} yılı Genel Emir ceza miktarları güncellendi.\n\n"
                        f"Genel Emir: {sayi}\n"
                        f"BB Ceza: {bb:.2f} ₺\nKB Ceza: {kb:.2f} ₺\n"
                        f"Madde: {madde}")
                except Exception as e: messagebox.showerror("Hata",str(e))
            MBtn(uf,"💾 Kaydet / Güncelle",command=_ge_kaydet,color=self.gc("acc"),width=20).grid(row=1,column=0,columnspan=10,pady=10)
            # Bilgi notu
            nf=tk.Frame(ge,bg="#FEF9E7",highlightbackground="#F0C040",highlightthickness=1,padx=10,pady=6); nf.pack(fill="x",pady=6)
            tk.Label(nf,text="📌 Ceza miktarları her yıl Ardahan Valiliği Genel Emri ile belirlenir.\n"
                "2025/01 sayılı: BB=360,00₺ KB=36,00₺ (18. Madde h bendi)\n"
                "2026/01 sayılı: BB=451,50₺ KB=45,15₺ (18. Madde 9. fıkrası)\n"
                "Ceza oluru oluşturulurken ilgili yılın Genel Emri otomatik kullanılır.",
                font=("Segoe UI",8),fg="#7D6608",bg="#FEF9E7",justify="left").pack(anchor="w")

    def _ge_yenile(self):
        """Genel Emir tablosunu yenile — DB'den oku"""
        if not hasattr(self,"tv_ge"): return
        self.tv_ge.delete(*self.tv_ge.get_children())
        gosterilen_yillar=set()
        # Önce DB'den oku (güncel veriler)
        if DB_PATH:
            try:
                with db_baglan() as c:
                    for r in c.execute("SELECT yil,sayi,bb_ceza,kb_ceza,madde FROM Genel_Emri ORDER BY yil").fetchall():
                        self.tv_ge.insert("","end",values=(r[0],r[1],f"{r[2]:.2f}",f"{r[3]:.2f}",r[4] or ""))
                        gosterilen_yillar.add(r[0])
            except Exception: pass
        # DB'de olmayan yılları hardcoded geçmişten ekle
        for yil,data in sorted(GENEL_EMRI_GECMIS.items()):
            if yil not in gosterilen_yillar:
                self.tv_ge.insert("","end",values=(yil,data["sayi"],f"{data['bb']:.2f}",f"{data['kb']:.2f}",data["madde"]))

    def _build_db_yonetimi(self,parent,bg):
        """Veritabanı yönetim sekmesi içeriği"""
        dd=tk.Frame(parent,bg=bg); dd.pack(fill="both",expand=True,padx=40,pady=30)
        tk.Label(dd,text="🗄️ Veritabanı Yönetimi",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,10))
        tk.Label(dd,text=f"Aktif DB: {DB_PATH or 'Seçilmemiş'}",font=("Segoe UI",10,"bold"),fg=C_INFO,bg=bg).pack(anchor="w",pady=(0,6))
        try:
            db_boyut=os.path.getsize(DB_PATH)/1024/1024 if DB_PATH and os.path.exists(DB_PATH) else 0
            tk.Label(dd,text=f"Boyut: {db_boyut:.2f} MB",font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,14))
        except Exception: pass
        tk.Label(dd,text="Farklı yıllar için farklı veritabanı kullanabilirsiniz.\nÖrnek: miras_2025.db, miras_2026.db",
            font=("Segoe UI",9),fg="#666",bg=bg).pack(anchor="w",pady=(0,14))
        MBtn(dd,"📂 Farklı Veritabanı Aç",command=self._db_degistir,color=C_INFO,width=28).pack(anchor="w",pady=4)
        MBtn(dd,"➕ Yeni Veritabanı Oluştur",command=lambda:self._db_degistir(yeni=True),color="#2D8C55",width=28).pack(anchor="w",pady=4)
        MBtn(dd,"📋 Mevcut DB'yi Kopyala (Arşiv)",command=self._db_arsivle,color="#8E44AD",width=28).pack(anchor="w",pady=4)
        tk.Label(dd,text="\n⚠️ Veritabanı değiştirince program yeniden başlatılır.",
            font=("Segoe UI",9),fg=C_DANGER,bg=bg).pack(anchor="w")

    def _build_gorunum(self,parent,bg):
        """Görünüm ayarları sekmesi"""
        gg=tk.Frame(parent,bg=bg); gg.pack(fill="both",expand=True,padx=40,pady=30)
        tk.Label(gg,text="Tema:",bg=bg,font=("Segoe UI",11,"bold")).pack(anchor="w")
        self.cb_tema=ttk.Combobox(gg,values=list(TEMALAR.keys()),state="readonly"); self.cb_tema.set(self.tema); self.cb_tema.pack(anchor="w",pady=(4,16))
        tk.Label(gg,text="Yazı Boyutu:",bg=bg,font=("Segoe UI",11,"bold")).pack(anchor="w")
        self.cb_pt=ttk.Combobox(gg,values=[8,9,10,11,12,14],state="readonly",width=6); self.cb_pt.set(self.punto); self.cb_pt.pack(anchor="w",pady=(4,16))
        if self.u_yetki=="Admin":
            tk.Label(gg,text="Şube Müdürü V. (PDF):",bg=bg,font=("Segoe UI",11,"bold")).pack(anchor="w")
            self.e_sm=ttk.Entry(gg,width=28); self.e_sm.insert(0,self.sube_mudur); self.e_sm.pack(anchor="w",pady=(4,10))
            tk.Label(gg,text="Vali Yardımcısı:",bg=bg,font=("Segoe UI",11,"bold")).pack(anchor="w")
            self.e_vy=ttk.Entry(gg,width=28); self.e_vy.insert(0,getattr(self,"vali_yardimcisi","Semih CEMBEKLİ")); self.e_vy.pack(anchor="w",pady=(4,10))
            tk.Label(gg,text="İl Müdürü:",bg=bg,font=("Segoe UI",11,"bold")).pack(anchor="w")
            self.e_im=ttk.Entry(gg,width=28)
            try:
                with db_baglan() as c:
                    r=c.execute("SELECT deger FROM Iletisim_Bilgileri WHERE anahtar='il_muduru'").fetchone()
                if r: self.e_im.insert(0,r[0])
            except Exception: pass
            self.e_im.pack(anchor="w",pady=(4,10))
            tk.Label(gg,text="Evrak Unvanı (imza bloğu):",bg=bg,font=("Segoe UI",11,"bold")).pack(anchor="w")
            tk.Label(gg,text="Örnek: 'Vali a. Vali Yardımcısı İl Mera Komisyonu Başkanı' veya 'İl Müdür V.'",bg=bg,font=("Segoe UI",8),fg="#888").pack(anchor="w")
            self.e_unvan=ttk.Entry(gg,width=40)
            try:
                with db_baglan() as c:
                    r=c.execute("SELECT deger FROM Iletisim_Bilgileri WHERE anahtar='evrak_unvan'").fetchone()
                if r: self.e_unvan.insert(0,r[0])
                else: self.e_unvan.insert(0,"Vali a. Vali Yardımcısı\nİl Mera Komisyonu Başkanı")
            except Exception: self.e_unvan.insert(0,"Vali a. Vali Yardımcısı\nİl Mera Komisyonu Başkanı")
            self.e_unvan.pack(anchor="w",pady=(4,16))
        def _kg():
            sm=self.e_sm.get().strip() if self.u_yetki=="Admin" and hasattr(self,"e_sm") else self.sube_mudur
            vy=self.e_vy.get().strip() if self.u_yetki=="Admin" and hasattr(self,"e_vy") else getattr(self,"vali_yardimcisi","Semih CEMBEKLİ")
            with db_baglan() as c:
                c.execute("INSERT OR REPLACE INTO Ayarlar(k_adi,tema,punto,sube_mudur,vali_yardimcisi)VALUES(?,?,?,?,?)",(self.u_id,self.cb_tema.get(),int(self.cb_pt.get()),sm,vy))
                if self.u_yetki=="Admin":
                    if hasattr(self,"e_im"):
                        c.execute("INSERT OR REPLACE INTO Iletisim_Bilgileri(anahtar,deger)VALUES('il_muduru',?)",(self.e_im.get().strip(),))
                    if hasattr(self,"e_unvan"):
                        c.execute("INSERT OR REPLACE INTO Iletisim_Bilgileri(anahtar,deger)VALUES('evrak_unvan',?)",(self.e_unvan.get().strip(),))
            messagebox.showinfo("Kaydedildi","Bir sonraki girişte aktif olur.")
        MBtn(gg,"💾 Kaydet",command=_kg,color=C_INFO,width=22).pack(anchor="w")

    def _build_sifre(self,parent,bg):
        """Şifre değiştirme sekmesi"""
        # Şifre
        ss=tk.Frame(parent,bg=bg); ss.pack(fill="both",expand=True,padx=40,pady=30)
        tk.Label(ss,text="Şifre Değiştir",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(0,16))
        self.sw={}
        for l,k in [("Mevcut Şifre:","eski"),("Yeni Şifre:","yeni"),("Tekrar:","tekrar")]:
            tk.Label(ss,text=l,bg=bg,font=("Segoe UI",10)).pack(anchor="w",pady=(8,2))
            e=ttk.Entry(ss,show="●",width=30); e.pack(anchor="w"); self.sw[k]=e
        def _sk():
            eski=self.sw["eski"].get(); yeni=self.sw["yeni"].get(); tekrar=self.sw["tekrar"].get()
            if yeni!=tekrar: messagebox.showerror("Hata","Şifreler eşleşmiyor."); return
            ok,msg=strong_pw(yeni)
            if not ok: messagebox.showerror("Hata",msg); return
            with db_baglan() as c:
                db_s=c.execute("SELECT sifre FROM Kullanicilar WHERE k_adi=?",(self.u_id,)).fetchone()[0]
            if not verify_pw(eski,db_s): messagebox.showerror("Hata","Mevcut şifre yanlış."); return
            with db_baglan() as c: c.execute("UPDATE Kullanicilar SET sifre=?,sifre_tarih=? WHERE k_adi=?",(hash_pw(yeni),datetime.now().strftime("%Y-%m-%d"),self.u_id))
            db_log(self.u_id,"Şifre Değiştir","Başarılı")
            for e in self.sw.values(): e.delete(0,tk.END)
            messagebox.showinfo("Tamam","Şifreniz değiştirildi.")
        MBtn(ss,"🔒 Güncelle",command=_sk,color=C_DANGER,width=22).pack(anchor="w",pady=16)
        # Şifre ömrü bilgisi
        try:
            with db_baglan() as c:
                st=c.execute("SELECT sifre_tarih FROM Kullanicilar WHERE k_adi=?",(self.u_id,)).fetchone()
            if st and st[0]:
                gecen=(datetime.now()-datetime.strptime(st[0],"%Y-%m-%d")).days
                kalan=max(0,SIFRE_OMUR_GUN-gecen)
                renk="#2D8C55" if kalan>SIFRE_UYARI_GUN else "#D68910" if kalan>0 else "#C0392B"
                tk.Label(ss,text=f"🔑 Şifre ömrü: {kalan} gün kaldı ({SIFRE_OMUR_GUN} günden {gecen}. gün)",
                    font=("Segoe UI",10,"bold"),fg=renk,bg=bg).pack(anchor="w",pady=(0,8))
        except Exception: pass
        # Giriş Geçmişi (son 10)
        tk.Label(ss,text="📋 Son 10 Giriş Denemesi",font=("Segoe UI",11,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w",pady=(16,6))
        tv_giris=ttk.Treeview(ss,columns=("t","i","d"),show="headings",height=6)
        for col,w,bas in [("t",150,"Tarih"),("i",100,"Sonuç"),("d",200,"Detay")]:
            tv_giris.heading(col,text=bas); tv_giris.column(col,width=w)
        tv_giris.tag_configure("basarili",background="#EAFAF1")
        tv_giris.tag_configure("basarisiz",background="#FDEDEC")
        tv_giris.tag_configure("kilitli",background="#FEF9E7")
        tv_giris.pack(fill="x",pady=4)
        try:
            with db_baglan() as c:
                for r in c.execute("SELECT tarih,islem,detay FROM Loglar WHERE kul=? AND islem IN('Giriş','Başarısız') ORDER BY id DESC LIMIT 10",(self.u_id,)).fetchall():
                    tag="basarili" if "Giriş" in r[1] else "kilitli" if "Kilit" in str(r[2]) else "basarisiz"
                    tv_giris.insert("","end",values=r,tags=(tag,))
        except Exception: pass

    def _build_yedekleme(self,parent,bg):
        """Yedekleme sekmesi"""
        # Yedekleme
        yy=tk.Frame(parent,bg=bg); yy.pack(fill="both",expand=True,padx=20,pady=16)
        tk.Label(yy,text="💾 Yedekleme & Geri Yükleme",font=("Segoe UI",15,"bold"),fg=self.gc("pri"),bg=bg).pack(anchor="w")
        tk.Label(yy,text=f"Yerel Yedek: {BACKUP_DIR.absolute()}",font=("Segoe UI",9),fg="#888",bg=bg).pack(anchor="w",pady=(0,4))
        # Drive klasör bilgisi
        drive_f=tk.Frame(yy,bg=bg); drive_f.pack(fill="x",pady=(0,10))
        dp=self.yedekci.drive_path
        drive_durum="✅ "+dp if dp else "❌ Ayarlanmamış — aşağıdan klasör seçin"
        self.lbl_drive_yol=tk.Label(drive_f,text=f"☁️ Google Drive: {drive_durum}",font=("Segoe UI",9,"bold"),fg="#1A5276" if dp else "#C0392B",bg=bg)
        self.lbl_drive_yol.pack(side="left")
        if dp:
            tk.Label(drive_f,text="  (Her yedeklemede otomatik kopyalanır)",font=("Segoe UI",8),fg="#2D8C55",bg=bg).pack(side="left")
        bff=tk.Frame(yy,bg=bg); bff.pack(fill="x",pady=(0,12))
        MBtn(bff,"💾 Şimdi Yedekle",command=self._yedek_al,width=18).pack(side="left",padx=6)
        MBtn(bff,"🔄 Listeyi Yenile",command=self._yedek_yenile,color=C_INFO,width=16).pack(side="left",padx=6)
        MBtn(bff,"⏮ Yedeğe Dön",command=self._yedek_geri,color=C_DANGER,width=16).pack(side="left",padx=6)
        MBtn(bff,"☁️ Drive Klasörü Seç",command=self._drive_klasor_sec,color="#F39C12",width=20).pack(side="left",padx=6)
        MBtn(bff,"☁️ Şimdi Drive'a Kopyala",command=self._drive_kopyala,color="#1A5276",width=22).pack(side="left",padx=6)
        self.tv_yedek=self._tv(yy,[("ad",280,"Dosya"),("tarih",160,"Tarih"),("boyut",100,"Boyut"),("tur",100,"Tür")],12)
        self._yedek_yenile()

    def _yedek_al(self):
        try:
            yol=self.yedekci.al(False); self._yedek_yenile()
            drive_msg=""
            if self.yedekci.drive_path:
                drive_msg=f"\n☁️ Drive'a da kopyalandı:\n{self.yedekci.drive_path}"
            messagebox.showinfo("Tamam",f"💾 Yerel yedek:\n{yol}{drive_msg}")
        except Exception as e: messagebox.showerror("Hata",str(e))

    def _yedek_yenile(self):
        self.tv_yedek.delete(*self.tv_yedek.get_children())
        for y in self.yedekci.listele():
            tur="🤖 Oto" if "oto" in y["ad"] else "✋ Manuel"
            self.tv_yedek.insert("","end",values=(y["ad"],y["tarih"],y["boyut"],tur),tags=(y["yol"],))

    def _yedek_geri(self):
        sel=self.tv_yedek.selection()
        if not sel: messagebox.showwarning("Seçim","Yedek seçin."); return
        yol=self.tv_yedek.item(sel[0])["tags"][0]
        if messagebox.askyesno("Onay",f"Mevcut veri ÜZERİNE yazılacak!\nDevam?"):
            if self.yedekci.geri(yol): messagebox.showinfo("Tamam","Geri yükleme başarılı.")
            else: messagebox.showerror("Hata","Geri yükleme başarısız.")

    def _drive_klasor_sec(self):
        """Google Drive sync klasörünü seç ve kaydet"""
        mevcut=self.yedekci.drive_path or ""
        hedef=filedialog.askdirectory(title="Google Drive Klasörünü Seçin (örn: G:\\My Drive\\Miras_Yedek)",initialdir=mevcut or None)
        if hedef:
            if self.yedekci.set_drive_path(hedef):
                db_log(self.u_id,"Drive Klasör Ayarla",hedef)
                messagebox.showinfo("Tamam",
                    f"☁️ Drive klasörü kaydedildi:\n{hedef}\n\n"
                    "Artık her yedeklemede (manuel + otomatik günlük)\n"
                    "veritabanı bu klasöre de otomatik kopyalanacak.\n\n"
                    "Google Drive Desktop açık olduğu sürece\n"
                    "dosyalar otomatik buluta sync olur.")
                # UI'ı yenile
                self._menu_click("ayarlar")
            else:
                messagebox.showerror("Hata","Geçersiz klasör yolu.")

    def _drive_kopyala(self):
        """Manuel olarak şu anda Drive'a kopyala"""
        if not self.yedekci.drive_path:
            messagebox.showwarning("Drive Ayarlanmamış",
                "Önce '☁️ Drive Klasörü Seç' butonuyla\n"
                "Google Drive sync klasörünüzü seçin.\n\n"
                "Örnek: G:\\My Drive\\Miras_Yedek")
            return
        if not os.path.isdir(self.yedekci.drive_path):
            messagebox.showerror("Hata",f"Drive klasörü bulunamadı:\n{self.yedekci.drive_path}\n\nKlasörü yeniden seçin.")
            return
        try:
            ts=datetime.now().strftime("%Y%m%d_%H%M%S")
            hedef_dosya=os.path.join(self.yedekci.drive_path,f"Miras_manuel_{ts}.db")
            shutil.copy2(DB_PATH,hedef_dosya)
            self.yedekci._drive_temizle()
            db_log(self.u_id,"Drive Kopyala",hedef_dosya)
            messagebox.showinfo("Tamam",f"☁️ Drive'a kopyalandı!\n{hedef_dosya}\n\nGoogle Drive Desktop sync edecektir.")
        except Exception as e: messagebox.showerror("Hata",str(e))

    # ═══ 19. ADMİN ═══════════════════════════════════════════════════════════
    def _t_admin(self,p):
        bg=self.gc("bg")
        # Scrollable admin panel
        acv=tk.Canvas(p,bg=bg,highlightthickness=0); asb=ttk.Scrollbar(p,orient="vertical",command=acv.yview)
        acv.configure(yscrollcommand=asb.set); asb.pack(side="right",fill="y"); acv.pack(side="left",fill="both",expand=True)
        f=tk.Frame(acv,bg=bg); awid=acv.create_window((0,0),window=f,anchor="nw")
        acv.bind("<Configure>",lambda e:acv.itemconfig(awid,width=e.width))
        f.bind("<Configure>",lambda e:acv.configure(scrollregion=acv.bbox("all")))
        def _ascroll(e):
            try: acv.yview_scroll(-1*(e.delta//120),"units")
            except Exception: pass
        acv.bind("<MouseWheel>",_ascroll); f.bind("<MouseWheel>",_ascroll)
        # Duyuru
        df=tk.LabelFrame(f,text="  📢 Duyuru  ",bg=C_WHITE,padx=10,pady=8); df.pack(fill="x",pady=5)
        self.e_dy=ttk.Entry(df,width=60); self.e_dy.pack(side="left",padx=10)
        def _dy():
            if self.e_dy.get():
                with db_baglan() as c: c.execute("INSERT INTO Duyurular(mesaj,tarih,gonderen)VALUES(?,?,?)",(self.e_dy.get(),datetime.now().strftime("%d.%m.%Y"),self.u_ad))
                messagebox.showinfo("Tamam","Duyuru eklendi."); self.e_dy.delete(0,tk.END)
        MBtn(df,"📢 Gönder",command=_dy,color=C_DANGER).pack(side="left")
        def _dy_sil():
            with db_baglan() as c: c.execute("DELETE FROM Duyurular")
            messagebox.showinfo("Tamam","Tüm duyurular silindi.")
        MBtn(df,"🗑 Tüm Duyuruları Sil",command=_dy_sil,color="#888",width=18).pack(side="left",padx=8)
        # Kullanıcılar
        tk.Label(f,text="Kullanıcı Yönetimi",font=("Segoe UI",12,"bold"),bg=bg,fg=self.gc("pri")).pack(anchor="w",pady=(14,6))
        self.tv_kul=self._tv(f,[("k",120,"Kullanıcı"),("ad",160,"Ad"),("unvan",160,"Ünvan"),("y",90,"Yetki"),("aktif",70,"Aktif"),("hata",60,"Hata")],8)
        bff=tk.Frame(f,bg=bg); bff.pack(fill="x",pady=(0,10))
        for txt,fn,renk in [("🔄 Yenile",self._kul_yenile,"#2D8C55"),("🔒 Kilit Kaldır",self._kul_kilit,C_WARN),
            ("❌ Pasif",self._kul_pasif,C_DANGER),("✅ Aktif",self._kul_aktif,"#2D8C55"),
            ("🔑 Şifre Sıfırla",self._kul_sifre,C_INFO),("✏️ K.Adı Değiştir",self._kul_adi_degistir,"#8E44AD"),
            ("📋 Görev Ata",self._kul_gorev_ata,"#1A5276")]:
            MBtn(bff,txt,command=fn,color=renk,width=16).pack(side="left",padx=3)
        # Yeni kullanıcı
        nf=tk.LabelFrame(f,text="  ➕ Yeni Kullanıcı  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=14,pady=10); nf.pack(fill="x",pady=6)
        self.kul_g={}
        for i,(l,k) in enumerate([("K.Adı","ka"),("Ad Soyad","ad"),("Ünvan","un"),("Şifre","si")]):
            tk.Label(nf,text=l,bg=bg,font=("Segoe UI",10)).grid(row=0,column=i*2,padx=(0,4))
            e=ttk.Entry(nf,width=14,show="*" if k=="si" else ""); e.grid(row=0,column=i*2+1,padx=(0,10)); self.kul_g[k]=e
        tk.Label(nf,text="Yetki",bg=bg,font=("Segoe UI",10)).grid(row=0,column=8)
        self.kul_g["y"]=ttk.Combobox(nf,values=["Uzman","Admin","İzleyici"],state="readonly",width=10); self.kul_g["y"].set("Uzman"); self.kul_g["y"].grid(row=0,column=9,padx=(0,10))
        MBtn(nf,"➕ Ekle",command=self._kul_ekle).grid(row=0,column=10)
        # Log filtresi (sadece admin)
        lff=tk.LabelFrame(f,text="  📋 Log Filtreleme  ",bg=bg,font=("Segoe UI",10,"bold"),fg=self.gc("pri"),padx=10,pady=8); lff.pack(fill="both",expand=True,pady=6)
        AramaFrame(lff,self._filtre_admin_log,bg=bg).pack(anchor="w",pady=(0,6))
        self.tv_admin_log=self._tv(lff,[("t",150,"Zaman"),("k",100,"Personel"),("i",120,"İşlem"),("d",400,"Detay")],10)
        # Talepler
        tk.Label(f,text="Şifre / Kayıt Talepleri",font=("Segoe UI",11,"bold"),bg=bg).pack(anchor="w",pady=(8,4))
        self.tv_st=self._tv(f,[("id",40,"#"),("ka",120,"Kullanıcı"),("t",100,"Tarih"),("d",100,"Durum")],3)
        self.tv_kt=self._tv(f,[("id",40,"#"),("ad",140,"Ad"),("un",140,"Ünvan"),("ka",110,"K.Adı"),("t",100,"Tarih")],3)
        bf2=tk.Frame(f,bg=bg); bf2.pack(pady=4)
        MBtn(bf2,"✅ Şifre Onayla",command=self._sifre_onayla,color=C_INFO,width=18).pack(side="left",padx=6)
        MBtn(bf2,"✅ Kayıt Onayla",command=self._kayit_onayla,color="#2D8C55",width=18).pack(side="left",padx=6)
        # Silme Talepleri
        tk.Label(f,text="🗑 Silme Talepleri",font=("Segoe UI",11,"bold"),bg=bg,fg=C_DANGER).pack(anchor="w",pady=(8,4))
        self.tv_silt=self._tv(f,[("id",40,"#"),("tablo",100,"Modül"),("ozet",200,"Kayıt"),("talep",120,"Talep Eden"),("t",120,"Tarih")],3)
        bf3=tk.Frame(f,bg=bg); bf3.pack(pady=4)
        MBtn(bf3,"✅ Onayla (Sil)",command=self._silme_onayla,color=C_DANGER,width=16).pack(side="left",padx=6)
        MBtn(bf3,"❌ Reddet",command=self._silme_reddet,color="#888",width=14).pack(side="left",padx=6)
        # İletişim Talepleri
        tk.Label(f,text="📝 İletişim Formları",font=("Segoe UI",11,"bold"),bg=bg,fg=C_INFO).pack(anchor="w",pady=(8,4))
        self.tv_ilet=self._tv(f,[("id",40,"#"),("gon",120,"Gönderen"),("konu",100,"Konu"),("mesaj",280,"Mesaj"),("t",120,"Tarih")],3)
        MBtn(f,"✅ Okundu",command=self._ilet_okundu,color=C_INFO,width=14).pack(anchor="w",padx=6,pady=4)
        self._kul_yenile()

    def _kul_yenile(self):
        self.tv_kul.delete(*self.tv_kul.get_children())
        self.tv_st.delete(*self.tv_st.get_children()); self.tv_kt.delete(*self.tv_kt.get_children())
        try:
            with db_baglan() as c:
                for r in c.execute("SELECT k_adi,ad,unvan,yetki,aktif,fail_count FROM Kullanicilar"):
                    self.tv_kul.insert("","end",values=(r[0],r[1],r[2],r[3],"✔" if r[4] else "✘",r[5]))
                for r in c.execute("SELECT id,k_adi,tarih,durum FROM Sifre_Talepleri WHERE durum='Bekliyor'"):
                    self.tv_st.insert("","end",values=r)
                for r in c.execute("SELECT id,ad,unvan,k_adi,tarih FROM Kayit_Talepleri WHERE durum='Bekliyor'"):
                    self.tv_kt.insert("","end",values=r)
                # Loglar
                self._all_admin_log=c.execute("SELECT tarih,kul,islem,detay FROM Loglar ORDER BY id DESC LIMIT 200").fetchall()
                # Silme talepleri
                if hasattr(self,'tv_silt'):
                    self.tv_silt.delete(*self.tv_silt.get_children())
                    for r in c.execute("SELECT id,tablo,kayit_ozet,talep_eden,tarih FROM Silme_Talepleri WHERE durum='Bekliyor'"):
                        self.tv_silt.insert("","end",values=r)
                # İletişim formları
                if hasattr(self,'tv_ilet'):
                    self.tv_ilet.delete(*self.tv_ilet.get_children())
                    for r in c.execute("SELECT id,gonderen,konu_tipi,mesaj,tarih FROM Iletisim_Formu WHERE durum='Yeni' ORDER BY id DESC"):
                        self.tv_ilet.insert("","end",values=r)
            self._filtre_admin_log("")
        except Exception as e: logging.error(f"kul_yenile:{e}")

    def _filtre_admin_log(self,a):
        self.tv_admin_log.delete(*self.tv_admin_log.get_children()); a=a.lower()
        for r in getattr(self,"_all_admin_log",[]):
            if a and a not in " ".join(str(x) for x in r).lower(): continue
            self.tv_admin_log.insert("","end",values=r)

    def _sec_kul(self):
        sel=self.tv_kul.selection()
        if not sel: messagebox.showwarning("Seçim","Kullanıcı seçin."); return None
        return self.tv_kul.item(sel[0])["values"][0]

    def _kul_kilit(self):
        k=self._sec_kul()
        if k:
            with db_baglan() as c: c.execute("UPDATE Kullanicilar SET fail_count=0,lockout_ts=NULL WHERE k_adi=?",(k,))
            db_log(self.u_id,"Kilit Kaldır",k); self._kul_yenile(); messagebox.showinfo("Tamam",f"'{k}' kilidi kaldırıldı.")

    def _kul_pasif(self):
        k=self._sec_kul()
        if not k or k==self.u_id: messagebox.showwarning("Uyarı","Kendinizi pasifleştiremezsiniz."); return
        if k:
            with db_baglan() as c: c.execute("UPDATE Kullanicilar SET aktif=0 WHERE k_adi=?",(k,))
            db_log(self.u_id,"Pasifleştir",k); self._kul_yenile()

    def _kul_aktif(self):
        k=self._sec_kul()
        if k:
            with db_baglan() as c: c.execute("UPDATE Kullanicilar SET aktif=1 WHERE k_adi=?",(k,))
            self._kul_yenile()

    def _kul_sifre(self):
        k=self._sec_kul()
        if not k: return
        y=tkinter.simpledialog.askstring("Şifre",f"'{k}' için yeni şifre:",show="*")
        if y:
            ok,msg=strong_pw(y)
            if not ok: messagebox.showerror("Hata",msg); return
            with db_baglan() as c: c.execute("UPDATE Kullanicilar SET sifre=? WHERE k_adi=?",(hash_pw(y),k))
            db_log(self.u_id,"Şifre Sıfırla",k); messagebox.showinfo("Tamam","Şifre güncellendi.")

    def _kul_adi_degistir(self):
        k=self._sec_kul()
        if not k: return
        yeni=tkinter.simpledialog.askstring("K.Adı Değiştir",f"'{k}' için yeni kullanıcı adı:")
        if yeni and yeni.strip():
            yeni=yeni.strip()
            try:
                with db_baglan() as c:
                    c.execute("UPDATE Kullanicilar SET k_adi=? WHERE k_adi=?",(yeni,k))
                    c.execute("UPDATE Ayarlar SET k_adi=? WHERE k_adi=?",(yeni,k))
                db_log(self.u_id,"K.Adı Değiştir",f"{k}→{yeni}"); self._kul_yenile()
                messagebox.showinfo("Tamam",f"'{k}' → '{yeni}' olarak değiştirildi.")
            except sqlite3.IntegrityError: messagebox.showerror("Hata","Bu kullanıcı adı zaten alınmış.")

    def _kul_gorev_ata(self):
        """RBAC: Kullanıcıya modül bazlı görev ata"""
        k=self._sec_kul()
        if not k: return
        win=tk.Toplevel(self.root); win.title(f"Görev Ata — {k}"); win.geometry("440x420"); win.configure(bg=C_WHITE); win.grab_set(); win.protocol("WM_DELETE_WINDOW",win.destroy)
        tk.Label(win,text=f"📋 {k} — Görev Ataması",font=("Segoe UI",14,"bold"),fg=self.gc("pri"),bg=C_WHITE).pack(pady=14)
        tk.Label(win,text="Erişebileceği modülleri seçin.\n'Hepsi' seçilirse tüm modüllere erişir.",
            font=("Segoe UI",9),fg="#666",bg=C_WHITE).pack(pady=(0,10))
        # Mevcut görevleri oku
        try:
            with db_baglan() as c:
                mevcut=c.execute("SELECT gorevler FROM Kullanicilar WHERE k_adi=?",(k,)).fetchone()
            gorev_str=mevcut[0] if mevcut and mevcut[0] else "*"
        except Exception: gorev_str="*"
        moduller=["İhale","Tahsis","Ceza","Islah","Şikayet","Kapasite","Veri Kayıt","Muhtar","Ajanda","Personel","İstatistik","Evrak","AI","Mevzuat"]
        hepsi_var=tk.BooleanVar(value=(gorev_str=="*"))
        cb_vars={}
        ttk.Checkbutton(win,text="✅ Hepsi (Tüm modüllere erişim)",variable=hepsi_var).pack(anchor="w",padx=30,pady=(0,8))
        ff=tk.Frame(win,bg=C_WHITE); ff.pack(fill="both",expand=True,padx=30)
        for i,m in enumerate(moduller):
            v=tk.BooleanVar(value=(gorev_str=="*" or m in gorev_str))
            ttk.Checkbutton(ff,text=m,variable=v).grid(row=i//3,column=i%3,sticky="w",padx=8,pady=3)
            cb_vars[m]=v
        def _kaydet():
            if hepsi_var.get():
                g="*"
            else:
                secilen=[m for m,v in cb_vars.items() if v.get()]
                g=",".join(secilen) if secilen else "*"
            with db_baglan() as c:
                c.execute("UPDATE Kullanicilar SET gorevler=? WHERE k_adi=?",(g,k))
            db_log(self.u_id,"Görev Ata",f"{k}: {g}"); win.destroy(); self._kul_yenile()
            messagebox.showinfo("Tamam",f"'{k}' görevleri güncellendi:\n{g}")
        MBtn(win,"💾 Kaydet",command=_kaydet,color=self.gc("acc"),width=24).pack(pady=14)

    def _kul_ekle(self):
        try:
            ka=self.kul_g["ka"].get().strip(); ad=self.kul_g["ad"].get().strip(); si=self.kul_g["si"].get()
            if not all([ka,ad,si]): raise ValueError("Tüm alanlar zorunlu.")
            ok,msg=strong_pw(si)
            if not ok: raise ValueError(msg)
            with db_baglan() as c: c.execute("INSERT INTO Kullanicilar(k_adi,sifre,yetki,ad,unvan)VALUES(?,?,?,?,?)",(ka,hash_pw(si),self.kul_g["y"].get(),ad,self.kul_g["un"].get()))
            db_log(self.u_id,"Kullanıcı Ekle",ka); self._kul_yenile()
            for e in self.kul_g.values():
                if isinstance(e,ttk.Entry): e.delete(0,tk.END)
            messagebox.showinfo("Tamam",f"'{ka}' oluşturuldu.")
        except ValueError as e: messagebox.showerror("Hata",str(e))
        except sqlite3.IntegrityError: messagebox.showerror("Hata","K.Adı alınmış.")

    def _sifre_onayla(self):
        sel=self.tv_st.selection()
        if not sel: return
        r=self.tv_st.item(sel[0])["values"]; sid=r[0]; ku=r[1]
        y=tkinter.simpledialog.askstring("Onay",f"'{ku}' için yeni şifre:",show="*")
        if y:
            ok,msg=strong_pw(y)
            if not ok: messagebox.showerror("Hata",msg); return
            with db_baglan() as c:
                c.execute("UPDATE Kullanicilar SET sifre=? WHERE k_adi=?",(hash_pw(y),ku))
                c.execute("UPDATE Sifre_Talepleri SET durum='Onaylandı' WHERE id=?",(sid,))
            self._kul_yenile()

    def _kayit_onayla(self):
        sel=self.tv_kt.selection()
        if not sel: return
        r=self.tv_kt.item(sel[0])["values"]; kid=r[0]
        y=tkinter.simpledialog.askstring("Onay",f"'{r[3]}' için şifre:",show="*")
        if y:
            ok,msg=strong_pw(y)
            if not ok: messagebox.showerror("Hata",msg); return
            try:
                with db_baglan() as c:
                    c.execute("INSERT INTO Kullanicilar(k_adi,sifre,yetki,ad,unvan)VALUES(?,?,?,?,?)",(r[3],hash_pw(y),"Uzman",r[1],r[2]))
                    c.execute("UPDATE Kayit_Talepleri SET durum='Onaylandı' WHERE id=?",(kid,))
                messagebox.showinfo("Tamam","Kullanıcı eklendi.")
            except sqlite3.IntegrityError: messagebox.showerror("Hata","K.Adı alınmış.")
            self._kul_yenile()

    def _silme_onayla(self):
        sel=self.tv_silt.selection()
        if not sel: return
        r=self.tv_silt.item(sel[0])["values"]; sid=r[0]; tablo=r[1]; ozet=r[2]
        if messagebox.askyesno("Silme Onayı",f"'{ozet}' kalıcı olarak silinecek.\nOnaylıyor musunuz?"):
            # Tablodan sil — whitelist ile SQL injection engelle
            GUVENLI_TABLOLAR={"Rapor_Gecmisi":"rapor_no","Ihaleler":"id","Idari_Cezalar":"id",
                       "Islah_Amenajman":"id","Tahsisler":"id","Sikayetler":"id","Veri_Kayit":"id",
                       "Islah_Projeler":"id","Muhtarlar":"id","Ihale_Yerleri":"id"}
            if tablo not in GUVENLI_TABLOLAR:
                messagebox.showerror("Güvenlik","Geçersiz tablo adı: "+tablo); return
            pk=GUVENLI_TABLOLAR[tablo]
            try:
                with db_baglan() as c:
                    kayit_id=c.execute("SELECT kayit_id FROM Silme_Talepleri WHERE id=?",(sid,)).fetchone()[0]
                    # Tablo ve pk whitelist'ten geldiği için güvenli
                    c.execute(f"DELETE FROM {tablo} WHERE {pk}=?",(kayit_id,))
                    c.execute("UPDATE Silme_Talepleri SET durum='Onaylandı' WHERE id=?",(sid,))
                db_log(self.u_id,"Silme Onay",ozet); self._kul_yenile()
            except Exception as e: messagebox.showerror("Hata",str(e))

    def _silme_reddet(self):
        sel=self.tv_silt.selection()
        if not sel: return
        sid=self.tv_silt.item(sel[0])["values"][0]
        with db_baglan() as c: c.execute("UPDATE Silme_Talepleri SET durum='Reddedildi' WHERE id=?",(sid,))
        self._kul_yenile()

    def _ilet_okundu(self):
        sel=self.tv_ilet.selection()
        if not sel: return
        iid=self.tv_ilet.item(sel[0])["values"][0]
        with db_baglan() as c: c.execute("UPDATE Iletisim_Formu SET durum='Okundu' WHERE id=?",(iid,))
        self._kul_yenile()

    # ═══ KÖY KARTI ═════════════════════════════════════════════════════════════
    def _koy_gecmis(self):
        koy=tkinter.simpledialog.askstring("Köy Kartı","Köy adını girin:")
        if not koy: return
        self._son_islem_kaydet("Köy Kartı",koy)
        win=tk.Toplevel(self.root); win.title(f"📂 {koy} — Köy Kartı"); win.geometry("960x700"); win.configure(bg="#F0F4F2")
        # Üst başlık bandı
        hdr=tk.Frame(win,bg=self.gc("pri"),height=60); hdr.pack(fill="x"); hdr.pack_propagate(False)
        tk.Label(hdr,text=f"🏘️ {koy.upper()} KÖYÜ",font=("Segoe UI",18,"bold"),fg=C_WHITE,bg=self.gc("pri")).pack(side="left",padx=20,pady=10)
        # İlçe ve alan bilgisi
        try:
            with db_baglan() as c:
                mera=c.execute("SELECT ilce,alan,yem FROM Mera_Varligi WHERE LOWER(koy)=LOWER(?)",(koy,)).fetchone()
        except Exception: mera=None
        if mera:
            tk.Label(hdr,text=f"İlçe: {mera[0]}  |  Mera: {mera[1]:.0f} da",font=("Segoe UI",11),fg="#A9DFBF",bg=self.gc("pri")).pack(side="right",padx=20)
        # Mini kartlar
        kf=tk.Frame(win,bg="#F0F4F2"); kf.pack(fill="x",padx=16,pady=10)
        sayilar={}
        try:
            with db_baglan() as c:
                sayilar["rapor"]=c.execute("SELECT COUNT(*) FROM Rapor_Gecmisi WHERE LOWER(koy)=LOWER(?)",(koy,)).fetchone()[0]
                sayilar["ihale"]=c.execute("SELECT COUNT(*) FROM Ihaleler WHERE LOWER(koy)=LOWER(?)",(koy,)).fetchone()[0]
                sayilar["islah"]=c.execute("SELECT COUNT(*) FROM Islah_Amenajman WHERE LOWER(koy)=LOWER(?)",(koy,)).fetchone()[0]
                sayilar["tahsis"]=c.execute("SELECT COUNT(*) FROM Tahsisler WHERE LOWER(koy)=LOWER(?)",(koy,)).fetchone()[0]
                sayilar["sikayet"]=c.execute("SELECT COUNT(*) FROM Sikayetler WHERE LOWER(koy)=LOWER(?)",(koy,)).fetchone()[0]
                sayilar["ceza"]=c.execute("SELECT COUNT(*) FROM Idari_Cezalar WHERE LOWER(mera_koy)=LOWER(?)",(koy,)).fetchone()[0]
        except Exception: pass
        for bas,key,ikon,renk in [("Rapor","rapor","📄",C_INFO),("İhale","ihale","⚖️",C_WARN),
            ("Islah","islah","🌱","#2D8C55"),("Tahsis","tahsis","📋","#8E44AD"),
            ("Şikayet","sikayet","🚨",C_DANGER),("Ceza","ceza","💰","#C0392B")]:
            cf=tk.Frame(kf,bg=C_WHITE,highlightbackground="#D0DDD8",highlightthickness=1,padx=14,pady=8)
            cf.pack(side="left",fill="both",expand=True,padx=4)
            tk.Label(cf,text=ikon,font=("Segoe UI",18),bg=C_WHITE).pack()
            tk.Label(cf,text=str(sayilar.get(key,0)),font=("Segoe UI",20,"bold"),fg=renk,bg=C_WHITE).pack()
            tk.Label(cf,text=bas,font=("Segoe UI",8),fg="#666",bg=C_WHITE).pack()
        # Detay notebook
        nb3=ttk.Notebook(win); nb3.pack(fill="both",expand=True,padx=10,pady=6)
        bolumler=[
            ("📄 Raporlar","SELECT rapor_no,islem_tarihi,talep_eden,tc,duzenleyen FROM Rapor_Gecmisi WHERE LOWER(koy)=LOWER(?)",[("no",90,"No"),("t",100,"Tarih"),("talep",150,"Talep"),("tc",110,"TC"),("per",130,"Personel")]),
            ("⚖️ İhaleler","SELECT id,tarih,ad_soyad,bedel,durum FROM Ihaleler WHERE LOWER(koy)=LOWER(?)",[("id",40,"#"),("t",90,"Tarih"),("ad",150,"Başvuran"),("bedel",90,"Bedel"),("d",120,"Durum")]),
            ("🌱 Islah","SELECT id,dilekce_tarihi,talep_eden,talep_alani,durum FROM Islah_Amenajman WHERE LOWER(koy)=LOWER(?)",[("id",40,"#"),("t",90,"Tarih"),("talep",140,"Talep"),("alan",80,"Alan"),("d",110,"Durum")]),
            ("📋 Tahsis","SELECT id,basvuru_t,kurum,amac,asama,durum FROM Tahsisler WHERE LOWER(koy)=LOWER(?)",[("id",40,"#"),("t",90,"Tarih"),("k",130,"Kurum"),("a",130,"Amaç"),("as",120,"Aşama"),("d",90,"Durum")]),
            ("🚨 Şikayet","SELECT id,tarih,sikayet_eden,tur,durum FROM Sikayetler WHERE LOWER(koy)=LOWER(?)",[("id",40,"#"),("t",90,"Tarih"),("s",130,"Şikayet Eden"),("tur",120,"Tür"),("d",110,"Durum")]),
            ("💰 Ceza","SELECT id,tarih,ad_soyad,konu,ipc_tutari FROM Idari_Cezalar WHERE LOWER(mera_koy)=LOWER(?)",[("id",40,"#"),("t",90,"Tarih"),("ad",130,"Ad Soyad"),("konu",180,"Konu"),("tutar",90,"Tutar")]),
        ]
        for tab_ad,sql,cols in bolumler:
            tf=ttk.Frame(nb3); nb3.add(tf,text=f"  {tab_ad}  ")
            tv=self._tv(tf,cols,12)
            try:
                with db_baglan() as c:
                    for r in c.execute(sql,(koy,)).fetchall(): tv.insert("","end",values=r)
            except Exception: pass
        # Son işlemler zaman çizelgesi
        tlf=ttk.Frame(nb3); nb3.add(tlf,text="  ⏱ Zaman Çizelgesi  ")
        tv_zaman=self._tv(tlf,[("t",140,"Tarih"),("k",120,"Personel"),("i",140,"İşlem"),("d",400,"Detay")],14)
        try:
            with db_baglan() as c:
                for r in c.execute("SELECT tarih,kul,islem,detay FROM Loglar WHERE LOWER(detay) LIKE ? ORDER BY id DESC LIMIT 30",(f"%{koy.lower()}%",)).fetchall():
                    tv_zaman.insert("","end",values=r)
        except Exception: pass

# ─── GİRİŞ NOKTASI ───────────────────────────────────────────────────────────
if __name__ == "__main__":
    # Global exception handler — beklenmeyen hatalar loglanır
    def _global_hata(exc_type, exc_value, exc_tb):
        if exc_type == KeyboardInterrupt: sys.exit(0)
        hata_metni = "".join(traceback.format_exception(exc_type, exc_value, exc_tb))
        logging.critical(f"BEKLENMEYEN HATA:\n{hata_metni}")
        # Crash recovery — DB bütünlüğünü koru
        try:
            if DB_PATH and os.path.exists(DB_PATH):
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                crash_yedek = BACKUP_DIR / f"Miras_crash_{ts}.db"
                BACKUP_DIR.mkdir(exist_ok=True)
                shutil.copy2(DB_PATH, crash_yedek)
                logging.info(f"Crash yedek: {crash_yedek}")
        except Exception: pass
        try:
            messagebox.showerror("❌ Beklenmeyen Hata",
                f"Program beklenmeyen bir hata ile karşılaştı.\n\n"
                f"{exc_type.__name__}: {exc_value}\n\n"
                f"Verileriniz otomatik yedeklendi.\n"
                f"Detaylar log dosyasında.")
        except Exception: pass
    sys.excepthook = _global_hata

    root = tk.Tk()
    # Tkinter içi hataları da yakala
    def _tk_hata(exc, val, tb):
        _global_hata(exc, val, tb)
    root.report_callback_exception = _tk_hata
    app  = MirasApp(root)
    root.mainloop()
